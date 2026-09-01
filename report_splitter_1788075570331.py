#!/usr/bin/env python3
"""
 Hilton High School — Report Card Tools v3
Tabs: Generate | Pay Codes | Defaulters
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import sys
import csv
import json
import re
import random
import string
import difflib
import subprocess
import shutil
import tempfile
import time
import zipfile
import urllib.request
import urllib.error
from datetime import datetime
from pathlib import Path
import base64
import io
import tarfile

try:
    from hilton_ui import ModernCard
except ImportError:
    ModernCard = None


# ═══════════════════════════════════════════════════════════════════════════════
# SILENT AUTO-INSTALL (Windows-safe — no console flash)
# ═══════════════════════════════════════════════════════════════════════════════

def _ensure(import_name: str, pip_name: str):
    try:
        __import__(import_name)
    except ImportError:
        try:
            kw: dict = {
                "stdout": subprocess.DEVNULL,
                "stderr": subprocess.DEVNULL,
            }
            if sys.platform == "win32":
                kw["creationflags"] = 0x08000000  # CREATE_NO_WINDOW
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install",
                 pip_name, "--quiet", "--no-warn-script-location"],
                **kw
            )
        except Exception:
            pass  # We'll surface a proper error when the import is attempted


_ensure("fitz",    "PyMuPDF")
_ensure("PIL",     "Pillow")
_ensure("openpyxl","openpyxl")
_ensure("docx",    "python-docx")

try:
    import fitz
    from PIL import Image, ImageTk
except ImportError as _exc:
    # Show a user-friendly dialog before giving up
    try:
        import tkinter as _tk, tkinter.messagebox as _mb
        _r = _tk.Tk(); _r.withdraw()
        _mb.showerror("Missing component",
            f"A required component could not be installed:\n{_exc}\n\n"
            "Please run 'Card Splitter Setup.bat' first, then try again.")
        _r.destroy()
    except Exception:
        pass
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════════
# THEME
# ═══════════════════════════════════════════════════════════════════════════════

PRIMARY      = "#1a6b3a"
PRIMARY_DARK = "#114428"
PRIMARY_PALE = "#e8f5ee"
ACCENT       = "#2563eb"
RED          = "#dc2626"
RED_PALE     = "#fef2f2"
AMBER        = "#b45309"
BG           = "#f1f5f9"
CARD         = "#ffffff"
BORDER       = "#e2e8f0"
TEXT         = "#1e293b"
TEXT_MUTED   = "#64748b"
BTN_BG       = "#e8edf3"
ENTRY_BG     = "#ffffff"
PANEL_TINT   = "#dde8f0"
CANVAS_BG    = "#c8d4e0"

FONT      = ("Segoe UI", 9)
FONT_BOLD = ("Segoe UI", 9,  "bold")
FONT_LG   = ("Segoe UI", 11, "bold")
FONT_SM   = ("Segoe UI", 8)

THEMES = {
    "light": {
        "PRIMARY": "#1a6b3a", "PRIMARY_DARK": "#114428",
        "PRIMARY_PALE": "#e8f5ee", "ACCENT": "#2563eb",
        "RED": "#dc2626", "RED_PALE": "#fef2f2", "AMBER": "#b45309",
        "BG": "#f1f5f9", "CARD": "#ffffff", "BORDER": "#e2e8f0",
        "TEXT": "#1e293b", "TEXT_MUTED": "#64748b", "BTN_BG": "#e8edf3",
        "ENTRY_BG": "#ffffff", "PANEL_TINT": "#dde8f0",
        "CANVAS_BG": "#c8d4e0",
    },
    "dark": {
        "PRIMARY": "#34d399", "PRIMARY_DARK": "#064e3b",
        "PRIMARY_PALE": "#12352a", "ACCENT": "#60a5fa",
        "RED": "#f87171", "RED_PALE": "#3b1f24", "AMBER": "#fbbf24",
        "BG": "#111827", "CARD": "#1f2937", "BORDER": "#374151",
        "TEXT": "#f3f4f6", "TEXT_MUTED": "#9ca3af", "BTN_BG": "#374151",
        "ENTRY_BG": "#111827", "PANEL_TINT": "#263b42",
        "CANVAS_BG": "#0f172a",
    },
    "contrast": {
        "PRIMARY": "#000000", "PRIMARY_DARK": "#000000",
        "PRIMARY_PALE": "#ffff00", "ACCENT": "#0000ee",
        "RED": "#b00000", "RED_PALE": "#fff0f0", "AMBER": "#7a3e00",
        "BG": "#ffffff", "CARD": "#ffffff", "BORDER": "#000000",
        "TEXT": "#000000", "TEXT_MUTED": "#222222", "BTN_BG": "#ffffff",
        "ENTRY_BG": "#ffffff", "PANEL_TINT": "#ffff00",
        "CANVAS_BG": "#eeeeee",
    },
    "elegant": {
        "PRIMARY": "#6d28d9", "PRIMARY_DARK": "#3b0764",
        "PRIMARY_PALE": "#f3e8ff", "ACCENT": "#7c3aed",
        "RED": "#be123c", "RED_PALE": "#fff1f2", "AMBER": "#a16207",
        "BG": "#f5f3ff", "CARD": "#ffffff", "BORDER": "#ddd6fe",
        "TEXT": "#2e1065", "TEXT_MUTED": "#6b7280", "BTN_BG": "#ede9fe",
        "ENTRY_BG": "#ffffff", "PANEL_TINT": "#ede9fe",
        "CANVAS_BG": "#e9e5ff",
    },
    "wood": {
        "PRIMARY": "#8b5e34", "PRIMARY_DARK": "#5c3d25",
        "PRIMARY_PALE": "#f1e2cf", "ACCENT": "#2563eb",
        "RED": "#b91c1c", "RED_PALE": "#fff1f0", "AMBER": "#92400e",
        "BG": "#f3e5d0", "CARD": "#fffaf2", "BORDER": "#d6b892",
        "TEXT": "#3f2a1d", "TEXT_MUTED": "#7c6653", "BTN_BG": "#ead8c0",
        "ENTRY_BG": "#fffdf8", "PANEL_TINT": "#ead8c0",
        "CANVAS_BG": "#d8c1a5",
    },
}
THEME_NAMES = tuple(THEMES)
DEFAULT_THEME = "light"

# A compact copy of the supplied Hilton High School mark. Keeping the image
# data here means the compiled application does not need a sidecar image file.
APP_ICON_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAEAAAABACAYAAACqaXHeAAAN8UlEQVR42u2ae7CdVXnGf8/6vr33uSfkDgjEJKhBBCIVCRFjvKADFQXB2qJ0xLZYtSidUZxYp7XI4KUjXplO0Rk6taF1HKUFRAQCCAECE1KoQQKBEBJCLiQ5J8m57L2/bz394/v2yeEWk7CDl573XPY5+8xZ632f913vetazPhi3cRu3cRu3cRu3cft/anqlJ5zw/j7++6cDfPDLc5l6eI1qDQaTfubePIsbrriLZl/zDwyAWVA58QS4bxhO7gQJEGlFvOGdvUyZGSBCtAEhBCEQRxrc9O47DrqHB3X4yrknkHbXyIcbIHB0D0FzJL3+8LmdR86aV+3r6FN0ZFuMXmPxMGhdEtTEBkGM5udvv+33DIA5NSrzjgEDASRmAGcbPoA4DumQJFUyYVrK4a+rMvWoCkmVus1mofuAHxnfFNDO1pA3vO3Wg+LtQQGges48DBhSBb8P6XPAHyFCOasx2JCkgSlHVjz7xE71TU1kg+0R8O0ifGVNsvuOWVkXxgzsGOGeD9zzuw1AOn8W4bAJAD0OfB64CNGLcBn+C1ywoXtC6tfO79K0V6d2tIwF2my4NI/NfwmqNEXgxkU3v6TXF9/1WiRI0gp5bkIIPBD7WXrK+lcGgMrZ8wAj6LK4HPQJRAKYBAgq5/MLXLCh1pnwurd0ecbstNUUg2EI68uJ0q/nMc/syE3vvPW5E58Pf3fxvGI0CZGQk2NHJFEfzPinBf/7oj6/aE4OxJIz3gA2WY6MPo3117YT2x6N2eUPLRxQ8SlQgPpwzuq7h7Tt6RwVu0UEdQFfyJ19hCCCxMJfLHrO3J/79PHlaKoYzsmdfc/Rl2AOs6HWXX1Jv9N2ARBqFbBJg9+KfbFFBRPHJBlc7nVFqkZBcfl3BRgZzHnsvmEd/64e17rBdrTpxnwxOl8BPJT4uXmzi8xH+zys74K6yxFPFHwMe9dBBUAnTC+8gC7MxRZTR4NHY7I/mniVvzcxeelHWlaDB7Y0efqROjPndRCLphiBV9u6KEZ9XDhbeP0iEGTV1SSJyJzXhM7F6pbkICnap8XoYxDLXzJx7QAgPXIGREMeTwHeXvT4PemxXeKjAMow9wJ/j/0h8Psw5wPflllTMCG06fG6hnZFXO4WMWKsMwOc0DOpD6SKUajU57LimYlM6MryaNVbExevypGae2t17VkC0eQdCaGRnwH0ApFWsYsW+wvABuCrwBKb7WPdchKuUR6/J1iMOG9oIE92bMyYPqeKowHZMNWR0wc271hNEo4N1tGCrizq8WsfmbH8HUf1XxXNSdEcLtwALZH08N5cb0sFOIikkU8E5heBa5Tyll8hJNqghL+Q9V1g+5yOnQiRPbMRBQhZDuJR8KeEfuAI/Rub5JmxVVZBxPbCkFRiHuJ9wM8M66O5uBZ8/V1P9XVlUecYf9LwJ7YXY49IL10BbQGgZC+HYo4cU++jlVfrgFMW8sC6Jdxy8iIjmUeuXkPjPx7Ad2ymseQBolquaDfiMsGDQwMxZHXjSNkLhCOzHZmW5Gku9CxwA4RPxcjEPHL1sqd6j+4fqVxpcy2wy8CqXTsOcgUU63Qqdu9zGp4L5+fMDZy8kPiv26MWvAMa16x8wRjZNSvItg4TZ/RBZD2wpDkS3ayXeEagaAkT8uipMYcdff08vdHkWfYE5i5Mj+CTD23qnHjnul5+MulJLn/jCq49dc3BBsDY7sFORysguuDCFZh7rKnVxK6BBIWXnjLe/GvCpl0gA9zl3DvzphUjOJoYbedUY3SvczjykZnMmC4UKt3AnJI79ICqRqx+9bbf6Ht7lkDxESiLvmj4RWV0dcHkaSIaQohI3utYjcdWoALULZidjsUKi7G1G1gxxuDcbJ++vQwgnge8pTxA3h2t3xx5OwEYC8Xomb4oACo1qFSEI+CoYkvfi91fQCic284dVS4x4Sgi2FaMAYJMtM8BvgTqFDwG+m4alKN9Y/ltAqD0sIXBWDikgulJOCS29nnKggsbYhQxllUQRbRi77SOgMPHJb4nNEPFFvu3OD4UFBge3v0KALA68M3Vb6Wzs3CQUWY/CstoP4yGHBP3JTGx/GqVfizwjRHyHNLEsxoDI99CfAM0DbRK0seSWrheCpjIyj+7b59COGAidOkD85mUFqe2XT9+iJkXzpu2YUNervyyGFpcv2wIcqu/7aO51fyKUYNEV0es1jrj5VkepjqGhuFHhn/A/rVHin+786xb93mKA6qAKx5ZwOS+pJXeiVc9vfCsCRN0iEf5P3s4UBkIyHJC8G+Y8pIy4xE5ojwvgu/uMDMmNZk8IQtpYKpgleCToAsk/To0EgjizrP3PXg4gAq4etPxjAyW53XHwyQuCfJ1Dz3YXJ0kEhEIhqhRNgxlXZhRjvR8O3XXazlz0wy6KhV++J87uefWRt7ZJw7pNUdMa9LVGWlG2LFT1OvxqqRDl0XFdVICgsHKIPe/f9l+J3O/ARgZnAiGQJgexZXg1R9btPWWZFZ6psvqlzV6xDWF4mtQFJg9RHHp1g4e3X0SzSSQPhvJg2k0hjn3gx1TTjo5OW37TvqmHJo7l9jab558Ej21LmaDu33d5Olx3cTZE+nrHOa20w9cL9wvABY//MYiJFEDLhU6SdLihed3s2xZXR7jhdjTAFXu7JkiUuDmZ+exsTmZrGnSEEkMTdwBOsZKzsgafv+0aelx3X1ONmzBj62NPLXBGhwsSJcgbHsm5+kP38LI4SMHFvmBAHDZjB9z1Y6PYvEezEcQD1s8vX5DDirDfJFUtA6EOebotId1Q8JZg5BWuoGjDW8NpO+x/OZA7AvS4/07/NMbrht519bhZELDAeRC7mohK1528PsNwHe2nM/6/iHmTO39Y6QOweGYY9/8ptqyJx4fdJJqVP9glLmVlx6WZqq3O0WTBa9Xmi4AzwcdK+gSrA/StZJv7O0NyxZf8mxPJMzvOjSZGGrlohkTfLtsvwAIDsyo9QhIS2lvOnDlaad1fGP9+mziqocz8hyUimpFdHQFjjgqYXJH6onV8KZaDNcZz0KaAq6DHgP9G2KpYEW1km5sDude8sNB7HB0SKQi82FP3HqO3PLKArDCz7JgwjTnkSUS7wReJTgua+gHf/qhnsEd251mGQoJrnWIzs6gjhoOAef2ZMGA4XbQchUy1aNDg0P9nV1dADSbTdK0xj13D4zRTUvhtIWApXYisF8AvHFoMq4GDqtNv/mZxqYPgv+8bITTKmnoO+QQk+e43rB2D5qNGzP6B004rBKyJLlr3nSfnyb5jnqznibO6zFWmiFN2N2sE4lkeZU0DLfYUksqHX0pVsAY2v1KA/A3J97JVU8s4pnGZsD3NPLee6vJ0KSJfWHC96/efeaaNY2vNZtO6nU8XDcjdVObmHL82RWq3ey68VfplsXHNmsJyYlSODVJNAF41Hil0GON5sBAnlSpVkV9xLR63vOt6LftwWD/mKDgL2ffxtDuOiYh1aCBbVd8u/+Jm27c9cTatU1v2JCxdVvG4GAkyyLCJBJB6MRDMm0ZDkNe8cs7bL4f7bV59CV59C/yPN5QSatfNMzt7DLRKqQjjamCsRCE9gBwQGeBzxz/XMaVnn4caTUkY/0bQwTUEv9FYOW2wIWnvJe1I2yK0VfmeLUdr7HjgmgviOaCiz5b/dq3vtr85c6BUlHGVinsqdhp4757u3dr53E4juW5o6tYGqsQCuB/BjbTX2+wO6tTj9nqRsy3N2KkGR0ze2aW+bK3vC2cEveIBxo9WppSg2uP5+2RxQtnMsbkfOyWXZ4AHR3dk+TEaGKeQRpeE2P+RWC2rTg6mNgxMsJOtS4JyhFKbVD26Fsv29pTAYUMvg00Ukg3e5Bwa9s24cGBSfGoyi6AWSFJvmBzQ4z+cIyEUliVrXuThAuW/jxfqUB4XpgyNAz97doI23Y3CDwCXon0NsAqJQCMYmSrFZZecNQTMwN82Pg8zGsogCnzaQHXS/rE/fd6vaVC5HTRQuTyxawC/apdS6A9oujUKSD1Iy5FrAUCJiA1GnX+q5v8wjdP3Z7Ivh77S2Xwz78xbABXYdb/7EcR2esV+LpgA5FgCDYbwP+ooK1DO3cdkK/Pt+TlDwFhSrWAMlTWgpcVuqi2o/CVj/4Vt7/miManMBdiphXcFsZy25Yvkp6qVau3nXqaffvt5NnGxvKOV9XuRvRJrED6vLN0KSEjrVapL9/w8pPXFhiB9PQ3ADkoJYtN1aqVjs9dmp4ZI1/HHOHRjWEvU0qDwBWIb9revn0HLPlJhaRGRXKMWZIXYqHov6I9j8q09QmR2llzic2UixZX6OjiFEd+DBw6ehneUgb2NrHIgV8avmlzi+ThLdsS7lwemDoJfrUqo/87K9rmc9ufETr/56/jVTM6wfos8DVZtl485NHTXet70epUbnK7gduAJYbbaymbBodzkjTw7JrdfP/cR9vib5svRiCn0brNXQo8WHCCVjsfe3NSUgQTin0AgUP5HkAP8F6bf8fcOtzwVwyz3Gav27kNAhDzSDqSknVmK2SdZevd4GMlTTbUWu6XT4GVV2n2mAooBiqIUQPoBzYLrcLsorwn/J0F4JoznmTxiklkgznVrnStxD+jhJg2Q545GJQgoiOBQFJqXI7gYEQgOicHJyHEBo1YpTYqJ+/jjdc+28F7VNbwmduO4dxF8/nZ/SvJqqaVuGT0omTMebd8WEolOLmKixDloj4yQhc93HLocpbP/n0BYAwQL5jxxd77LVnbl8AL7LcY3LiN27iN27iN27jt1f4Pv+MO5vMBbMgAAAAASUVORK5CYII="
)

MAX_REPORT_STORAGE = 5 * 1024 * 1024 * 1024
STORAGE_WARNING_BYTES = int(MAX_REPORT_STORAGE * 0.8)
CLASS_OPTIONS = ("Senior One", "Senior Two", "Senior Three",
                 "Senior Four", "Senior Five", "Senior Six")
DEFAULT_UPLOAD_PASSWORD = "Hilt0n.High.o3"
DEFAULT_NETLIFY_TOKEN = "nfp_E4ZR9rQEy9Wc2szLqcvff2gWPNXmCWcqf93a"
DEFAULT_NETLIFY_SITE_ID = "bdfc8f16-e305-46ce-a0eb-87186a403e8d"

# Complete website bundle: extracted only when a deployment folder is compiled.
EMBEDDED_WEBSITE_ARCHIVE = (
"""
    H4sIAOkCi2oC/+y4U5Qm0fMl+pXRZVeX3WXbtl1dZpdt27Zt27a7bNs2vunffx5mXubel3tnrYt91lmRuc9DZmTEPici6ejp6AUU
    DNwkTAyMTRwA/7eA4b/jf2UZGJhZ/sf1f3hGBiZGJgChG+B/A5wdnQwc/j0e8P9NMHEQWjuZW5vwMrJzsDOwcTKzc9CxcP77/Izs
    sID/H/+vh6uJoaO5k4mek4m1nZWBkwm9oYGxmQmdnY3Z/7X6Z2P5L40zsrMy/M+WgYGFiYWdiRHAyMrEwsj4j2dg+qd/Vjbmf/pn
    +N+pfwdnG5v/ow3w/2z9/6EIV5ATh4f9+R+tw0tKiCgBACBx/5nQoP8YxK38zH83NubCwgrmRkYAwK/44rh66UEBrGG936JzUlB4
    uLhMaFDmguf7Zc6/JCiL1erOE4okqUmtJZREcPG0uZ0U7esjgP42+/yE0AFc7aeXeYQBSa81lEP52Z8DWy8HW92TW2ADssKs2fnu
    fVXtLkSJKWxSmSjfSVVcyw58ucN97x4iEo+hRRyWtz09rLrLuOjQ0QmYqpMu49KHmzWWMDYJpYsV2uS5iahs9BI9mB37D5m7EbM7
    AVeGBhDQ4uNKwydB+sNmdxiS1evka4SjsoX5z9ckuihLFheWl8xbe3l+cPNxeXn+eFh28ca1jI3fPSorqcLZ+MmNwbyrTioKxeyC
    zf3KrC7ZHSOzuwhqcd5x5YX0RZ/8Iur+e+qzH1lgYnJ+M/knX1nuuqXwIjn+G2q7ta1YJ9xPpnnD8v6rIHbOqH8mIOymo0r1aNZS
    OpWCSKou+kbJGBRBydOfVpd9V766PC/+8IMMexN6QuHigTVuEmx8QzxGOVqyKNiM/QsAoZGEkiwA4EEOAPgFAQDv/yi/0385+08Z
    5/oAAFcaAIBpl9WhyPNvAdJQWlwEAPzPyOsuTfzHwNhJaDr+syD/NeXJhpb/2VYPFQ0nJQNXQjsHW1NzKxNCJ3c7E8L/ii5HdLGy
    Zs0GA9znCqUHeESlE+7vBP9tIGhxCylCAgYiGRI9CZ3RDOknHdnneX4BiIc+JK2UVrSfuYBWMzYObAMjsUQCDdRAK/TP5Cg4Dlho
    u7eMt9Q2LwszNHYFmIvANUSSTwagihrjPOcvoQyG4l0CyJCRBz91PnbmHSQrHnqehgfZnphggkqA9K1zDOtM4DtN2140QVv3vvq6
    m+1QzVRRRF1K9C4EPQff3W8+/tu/5T/PiwI6EFeDf5G/9zzSythfsCBrntQHfQH+JqOwR9RCYn2guSbqyXxeS58Lv8/U1bwsU7/0
    uS71yJB6IRjb5q38FHlY6/g83oJfQTMrNJelY25nPrdSVauv8dLsnWHhx/J+/+E0OpBpsCUsJVIbjgxL0BjP0uGBJKdJm1uoKMt2
    +9BnXzv6qVIh/Ayazu6/EfM4MUifpJUJ8UaifMij2jP/x+t8baBUZFhYLEBWfFQl01+sVdJLcmNZ4Akrio7Hik5QzKTo6FZ4xpd1
    9WnPF+fy/XZ79/jRE4yjLEbIAMfQjpKE/oRHdZAVhI5yDzFrC3tXqPRC818GACRFBFXc1i+/3POs11QxYYEjWI41zS6OGRsqVKow
    cQBQIcIgPgnqgJ5Qvqf8XPh3fz6cz5yJk8ZCZO8yAixkNBUBCMF3Qtk/EEFxOUEqEsJaVDCSFZKsrZxb257VrCdYBx9GNc2zfK0V
    Si/k5Yvd7h+piaonNY3P65bTplOM4g84RcduQoI4L64dx0PNK6B+oLCgcYKAKWSAC3QfJOjuoAmhhgLoFJZUzS96EnCMMI74VM+B
    +OypcZDGvc6wfkA8FEe2e8U9mbFfO8wYLiFdLjN9bgCg6ugYQcHFFUqgX/BIGhtAqGtzduQPEpiqxoDk9frx1EyWhTuYjVRnE72H
    6O6mXtW7gwaJFQijMpP4wNrq3PDkmciN6weAi0PeZfj5PWVHiIo2dhQN1wyPyMyCW7uDIAXQ92vj1rQDRd7562f0u1XHnBECjoFg
    mxsU3MxvUnn+dgUAiTtRLOr04A5+KafClQ1t1+YHYmYEpv8Lz7IY2RiXSM4uIB5hNmQhpSg/FR6wuYy3LJwdQg2BDEjVdYTA5mTY
    pEVsnZLF5hAr/j39C6S9qxmA1cEUFxmYPeHWU/I1+fF+y2O2npsVku7mYvX2OnFhhC30+QuT3x1a/hksy3/BgwsbyQuL8TIvfw8a
    q4Qk1f2UEHQQ3tivypZVu6Pn4TUyG6GByd/9a8kuqYaHN6A7yp8JcwSZh2dJiwkAADcG34jDANScWUAVTsUzRYNA4fTpB/Jezxmo
    FL39cvqc2mAMuTm5y4UjFWCYmpJVVKUzpeC2DnpK+t0KI2buvC4Mx4HEFsqKeeZ5zBSVKABw4Qhj595/Dx0S871bhYCjR9lvqvfW
    zph5PLuNvkAsmPoDiAH45Q3q3+XXjrZqIVK4CMWrYe2EE7xlEQbYPt6s5pQCd59AhcEFdvftOFCnJMuiSon5zdEWBiOhBGAuavYN
    ZHKuXhuzI5/2PUq+zjw8dpFfOs0FCIjhL5CF3uuDWnJ9JkbUPFkY5RGwOAehEYx+DeoXwu0VJM+92gq3FIi/EhaGn8YfeCCpIphR
    sFMs14eGs4X8fYUsXQel5IS7RvM6gCBQz2b4+Vj9VOVGII4Ebv4Xuq90kjjbnGf7D7VerIdHUIs0AgU39arJIkLW62Ntaa3m+yAI
    TnM4U677ogk7RFYskGi0Dh+aTyFJ99c5jduDbCT+yMV3w0IM7aA573lbFxZnmnPW6iD1KolaB7w9dEDuRO2oAi/bza6ltyet9qdM
    b1Ipl87HWEiwcbVZ8m8E5gKe2fK0EqZj3qrck4z+219XQA6mi5+Kzo7mGuypffeK8Tq5GpYql7GfVWAikFBJ8kt6gK7tsLc6HaWX
    Aqk4R2w8kVg62Qli08L1lNGck2iSaLCaX26FhMyA71Xl+ymOkZEC7y9rmNu6TgOwh+Z9U6+v5m1RMwiMGa+LMzP2nCBI96AFnWmj
    Ld/XRtvvjvIkXodYRZVgnHVUmx7C2g8n3r127UaiAbGr0JmmZ1Np1BFZHLChYt5Bjy2uhaJSUA+ePRzf1cVNdkgytkXOn6I5yfIf
    YjdWPccc52Hof59SPucyDqdScry6yw4yIC4ZN4Adi5t38Dbk5sGIwym3jC6TseumbBZqF9K1Z01zSgyflZvfGb+r+ZVa7C2JyvRB
    pbEZAwPgtDsLDLvjkkPvGjfoNlYRktDdhe0pqbePTyF3+X46yMnKmd04gllzbpaWgZkUzL8b4KNfSIe4ODgaVrna9bQnUhl61BHG
    XFUcOOYg7EwH+49khyfmIqMcWPKtP3bpYvWLHx2oCPFMdI/tFZJdO/HV5B5fzvh6+YQy/UEjZznnjLWSxGnKue8xaCSoVWxcLnT9
    HCZRG3YZOeIcFzRoGrnPu+HO0u/x4cI5qobj7SxvOnBKNnxdJ34rCzgC+wPMlK7XOhC2GfTAdX6M0gwjcLdR7ny2VM9rfme+lvcy
    kubQPVih+81D0yrKT+WSYiCyvla79Zmnb02x4UZCclsedbvgicaNFjG8uNJwtThr130LL8huAdbVs7FerRhFWIqSusyMcngM7sAO
    iiT3Qo7wc/hFK+2hLlGrF5o7iP3403T0nbA7jP0x6cnn5Gv5EH7mYrzDhn0wxNzcrav5Rn9rsvCR6X4hmimWLvOOZKrQPHy2Ksxv
    cM3dLilL02RkmQdP/aJko8mIK0x+WMschtcahcp0yADiJ0G4uxuL7m/5leorcBcepcZv/RagLGUFbI6dv7OpuBtV9cxlIjtQsO90
    /qcdlZAcbUipI14gya2TvRmUtRXSJDz2X0WprHoGBfsCeJIEyC8rDVF3NFien9LLC35huFxRRATyC6ygawIuaqqaG1HmmzXP7Ia7
    xjjkB/qihvpOPtUn0KjVtr8OcRKyx8JDrW3sSRVdqfLCQPL70XRmUkJb7S0eymTs65eMWPpAhl8d1sXg2OKiGqf22Hiu4BE58XtL
    Nz+r+F3ruKPVG7Lf8mwQwDgwM+dHdjnxDMHxPN4yau3A1go111aovv+uFPuya/NP6c+P17FTIHH2T/DlQNY34zsfkU2Wof5E65AZ
    hhctZfu4Xzqg+l30rIW7thr76xh0TWNz9HkHTczs9w/25jIqN6WuoBtLHyarupLctcHYL8ZTZG3wEFSBcH04xigCIgJPc9wqBFBF
    bEzjwAjGP0JskJJW711wfDFcE36NP+ZK7Ct+XRWyJkLru1Lo/b5fGXhyIk0Ol6kpvTlny3hbG685kChtMPu7XpseV1hqyHsyYbTA
    pUYCJjeGsFVnAjfhj7mjuYadnb5IhmZ0BcPyt5xiKByAB7F8dm3y+I4DcCVT8sO+X8VPFYeDDI2O2DubtKFTOWyiLPUSnsSvLZ3L
    lGaAQt0cX0U/hCRWOTD67R45Z6V6+FOvMD8LNlSW9lPFQ1jMUwUb1DRgAZOcggg6zp5KZuKlGkoNN5Mu9qImhYKSbv0J3+gz7NW8
    tApezaR2srDRdOt7KMzMdGP6HN5riugu/+2nxxCh898Gy41FnLiEDPTMeGa0chKHhcVWa5gsB1diMiYK53kURVa6rez3+S5ULH3K
    p4zs5W0b4rQXHQoRe6xD7H3KZkIKYuZJTkz6LA+mK+wRwukOnjOssJjddJ1xlKkV/mvKEx83blOGdFQIv7/F4CGdmZps0yXnpjnT
    jb9DsFkEmR1CfjYFsK14R1lJ4MnIClxZCI6V4AkanCsgbdKag8KNTT5XGIHoCmmVBhvkt7DYwnNaqjEL6aukBFeZynSZMf5hYmKA
    ngGfKRfR/zs+Po61sWfRTdi6JTSNx+D2ZtuSpDwKmNJptAzh7ITVvJ3KmWr4wCZY+wghQ48rOcxofNbkjiPTEVTpJsm70Apor3j+
    GUbiD5/iORkuMJmIipSJHdbB9uQtf5jRmZQalPSBDpYRuei4lBK3cY21iNNRJcmaD96YbkmH4wMqvNXJV272VLUMBjtrR4aV/IeW
    nQksaU3AxubYtOa4IsWjsUY+KK4DDlXdVCUlnUYp4BcPqrDDOQDCZCGZxF+1WV1WJEd51I4ZTBAekh0O2akEO2ccRLmtP4um3vKW
    pgbgwAR6lKUsBU4UV7FNlXAB84Iyf+bQpMYFJ5CqjE4VaP6MVQCmU3vLUhlYnGN0MJGBTNwAakwEUAg1vB0KJjb2CyZpXznvV4Rx
    RI7ZZ3A2DUxPt69qbNcynMcKwMaGnpkROTxur6oxhRQ91OLowGAlxpSGXM75UFQdTwnf7JMRCrMIdA7VNx0RhM7RDh7TZ4K0Rvjx
    sFvVxQi/CWPhJ5Coij6lZQD119laZ/adDlU8XHWmW8ioDFpKIqq54QTbFNK1tggkTdBuznhqABwrKitd6p8vnEFxjrOCUzSQZ8ob
    zltWPWBuG62yrIqopg6HDSfO8gCois3itroIGKgkf6trGxA0+61aXZk3CzqhNSWpcBJUxTsYyJ+bQcT1Nx8VFSyCoLjbL9Y6o3xP
    qlaflwObpWT2MQOStBM24TwSHU7yqnG6nuw1B3BIqihEoJtbOtl/n5cXRSlxXMAuvJOOrUzrK3bBmdXJjrezgtLBIPVCuTezUyZN
    VFUVhYxB/7ZqrSRwxfrvKoSKALwbUjbh8Ivszfb+hPCVdkRUujCam0KqgTVmvpj+xAL8gCEy3DRCAigJZOdxFlnJok4wELq4oqfY
    qzskrq+vuM2Dim+0VB8nc5Z7AhaRAsAxoJhIubAFQKEWBYZcqufs6tDbuBUhKuDr5ZFAmN8xnP0ocuXAp2jA+xpR7G1LaPzbJ+2E
    kZDK5BnV4uPTI2SMlQkL5alkw+38bbksw4TCppJoyy3JDOrEOFhqz6kyKImwxnC7B7InF4kJrvotN3rTFpzpbTD0oeHM+DSzWFiV
    5K75yk8ESMs56kDCL/h+m9GI2LTD/TY0wBDTRwCvajDPxeKx/iVyty9jxlDV5ClHCDtOqJqMhnKYAxVr2+tdGr4QSk+g3VW3U2pS
    5UckntfmGpkwDqlIxl5FIhFhOZUvsxdCHsT9ris6hDU0ZyQDQFVapooAhTvMSotXqjBMXk+eIhd1dWEYlTQrQpjWh2TWFuQpF0Gv
    7uxET2zL5Brl6hd1SVS4EBJYCfTv88tPB2euriCaofsdQo3cXmmGiawOKSRXP7bfYH7iB3ez1VhRwSQJjQK9sCYDLz4amf1WhUKj
    TuHcbcbpT8Wwmr9G6qAWsRRFBGvDXUl97Vm0bTg83gXpVHebiNpAqQiNO7U2UhVqL0xDCCkTZHRY04ilGbQ4NRXEPoV3I942HYok
    nxHS/Ed9wXS7EoUUOoU6W+7wdug8pLEA5fRUZyVqvAnO28gY3wocNHVe1lpOTBmliouDSI1BlPYP5mqmYCJUQGdHp0o9VIxamlwc
    0yYRHAYNZhoLNCgIDYV9S1BHSpVCyHzImYMb9WHRQjhHQgdkBZydEUcBKWgCickkuSweQoy3mdmgg2pTiHPj5lhJlp+A2433JHcO
    GwSRM4WsDphT05/SRPfMXT9GJiRi6FP4ePlqRdhDOOY4IkgWELimSBrjxhLQ9NKaMmY5Vg1CDg5BZuHU3ZBxcHXEul6oSSQxtCya
    zmxznesrBWEUEUKX7GGpzTp87SFh6GIOY97c2ptkZVulh3FQpUAwj1CQijAQ9DotfVtIwHpxR0Rdge3q0sFgdrejHIe5IrsJLbZa
    thRA4MFhWaWc1TRG9VdmQeUoRWN2O1glWQlqWr5hWoGOvi2i9Qq7czIN2K/ELrfQqdB6Y4Ed+FlyULAY9d2fqrghsTFUsdQxCA94
    mIHcQbuc1Sninp0n62DWx8gxbCccMOKVNemLkGlEZotLu2+swmhllNpp/Zd+ZFcFIkI8mkUd4Rq23ew8Se3DM3CRojXnSCSAMsq5
    EHQI9uazL8m/dRFcjNbHgABokEp8yDMUn+cTn+JhGUVx8yp6sYqpTf6nljeCr8LZlP4qXdM8dHq1zEFqva9j1+2Huvmuq0q6H1BE
    xs8mqPME35/yvQf7PHJ547JLtNtPIx+6X++towyssUF5nKNj5fxR+R+P8h9B8t+B8qlI7lF5PhdbiD2PzraTaX1gnth5kLYQlfTv
    BbNuI+LXSPv6NWpGoKi7KQ4EB1NkpqGg/B+OvYIGKcnwBniPrj95DpbgsfLnRHvfi/XuOYf7cC7czThx5E0dIdaDCiHwu4JWvq4z
    9b46dto/dGx7PY7SHA/SATIlT9mu23dV8t9B8r4zjXzvhS/0WcI2hAKs8PBko5K5MUG/hvAeVabu1Wt2IXofUW6rZHc1GB3R7eye
    8rvNn9Jzv8dzH8M7S5zD801QgjNQguCdNvdTfY9Fxy0aJ4gIXV1eSJs+qX2/I21pVrP7WClGpfXeuG/cb/JDzVMDOdM0HVuZk7u4
    PiOAV9LPLmU1JPp2wf5BFI14JJ/fX6O93wPqQi5+5ueSYcFKOib99+WtN179J7xmWsKsOHF+oTVVurc/epH3262UhfGEC3VmT4mr
    Z24eUoF3JyEaIHaTjIuA8nz1xD4EJ7ysOTOeW8D2Z4lrzze8rcfKzY6n26Zux6LLmuMMQrYMOSZ9juc5Hia9dhrvS2POF2eu++gS
    qgy5vrhx7s2Msi7bQsXlKHNeiJ73ng7UKKcmhAxMPEupTz23+I7e2z+XjmSUPtds9DFK9MMQul72jqT13rWGhxRjkR/OJQKWKHyn
    8HKBnASYQoSEZgzCGovz6cZaHe7Dy4ktZDTc7lQ2joNOAUuw/M+fmDxNFkw3MKI3smNbSiRnhbn239I9j2v6fqAkBlMgbOI7C5O6
    byl670upLe1s2zqmxMTN6PN7V5gVpdv2G28/cjlASEVcrzy6bk1qgWe5vooT/bRIw5tPiSFbqYpqzg+Y/HsHuTYA1AD99lv17uNm
    4B2UT3shpyehmUudWQn6wSui1/7ojs+JHB6A6pg1cUzGrrXn9zCHUQiSGp7POUTv83ZoDJHyywShHA5XMrN7UCgKx0ObeMnBxPXX
    PP1HgCgUVMVPnQZJPK5B75thX+DZzuuRnvtZ8mDORXkkh+qU/BULrCjgipj/M6XS1mND2vcaU4/P/PLT5ym1917eay3VKwsQWSQr
    EoVKJwWwD5sn3/y6hvJ9pdD9W1Lysq6ikpx34pLh4v6m7nNVPPuRUvsRVopgI4nAYMcmAjNutzn/8r1CIM81qYUT7t9vk1VA4LOk
    97qPEzEmxakYVolP73Wm03uXsOLMJckCYzAxpePKE2mu9ZL7sDmbNv1CnXehjnxZQxM99Bf037t19gDFfd7d0IlzJ8MzKwJKI9Rl
    ZXJf2nPvN2d9EITCH4Qm3Fa61X33S32M1uTXe2oG2rOadLJVuioqE+U9n8V9+gl8rts5AahtfehNlx4pVChq2AYBJ8DHC+CHL7DY
    c8ezuMIVT3YCsEI5nU6tAr5btFL4se66/UG3963u83SE+ezmKqf3EgtcNvNev7JAd4clAgW/alT8sy8ZqveZnNq9N/3i/9zBt7xU
    9RuvM/tSLmHsLvfNUa9H2R9PSGI8TJQI3Anmhc/tkXcHeHDu+BCuWEcUziXO/13B//aW+3YZW2nBUHiVNHbOBRu5I8GwmxhGIav/
    xIksieSPeA7HfO17K8X+sbp9fBh12FxneaFBlmsKFQKNiZ9giWucJjP3nFgpnqQ4abwiTOu/lRLXdqmO7czUBLvMUk4F1bbOf01u
    k2xiUNPBZxdmkSOr45JVr0EE8gCZiE0kl3j57JTDU9fYFu+MPAcz2eI2KkAWNRIm7dSVVy4kV76Uls0JylL4ay1jqpno3DwpDwSa
    6zTKBW/Uz8e1wxUIhqiWyD1ztTBOJUux9JXakfu1pvcvz9+mpaFg12tP5G+y5NcRRFRVySaJcYWGp+Y2m36X+KzREni83rfy3Pzb
    Tmb5v2aB6fzP4sDbUdPhu/Hx4O6Sq5oO/30Ecbs0Xd4ZjIwoyoUgatMGafHRh/f9zh3Hz5uehzfgfe5XewVIuqI+VQ4SQCBVw33U
    K4dBEXFkX1to+3SEBukQ1mRIwWDtfVyJaVbnGflFJtrEp8lLwRF2SQ76lPdHoPCi1bhXfLezHJr6z/Y1qJlAkfZz9dHqc9A3N+QK
    SGuxGR24eFCOS/wsiYZsOTOdddWnku1sOc+ra+myCsp1K+WJS7u2iiXV4YBACN2Pd8+ej97P2y/ELG1etfKZUh6dJpuwhTeWUZIC
    1AJoM1g4EcHhskwO3xPP656HmGzb1xHbj4BZz5mdj/XazVxx1TS8RIpkGMWqnzTKjX/M0GaadfwuPVa3si12tgObYNE9hpusakKN
    s4iofOqZjvAetoLKFYuHYJ0n/LR7PISMBjCikY16kHcb5l7ZC3Oe8lE1cBbp4Agj/o5FAoQ2wPqGN667enPBmNnFAKrxyhDgCWdj
    i3XmD0SAK1i78GnOsZw5UhO5CU2jXRjnSqwDqxjoNOS+STIW6Zn8mI8Jf07wJtxClzUa5+pOIz5BAIswvOZeGRorSVonN3jSbvLL
    FGOI4XMoRwM3rkXJF5UcFdhK/PRfU+Mw9+q0tBUpQ2iwO6hkUg90ghDJEhSJocORpDHZfoQJHMWmypattZ21ndJEjPqYF50B/veO
    k+6vFwK/6cr3LAOgcT/vnJNLjihm5IGp0knWRw+EhkgYqL0dNDtycvd5jWvXwdWO607n90vvLfC96iv7zWFNrvSB7wVbAjZooghh
    KeY4LQmbSNM0g5wCdkna56MaeNbbfX/z1TyLfKIQ1WFjwm4N+6J5vxVFlfriDEpj6FCIsYDFIbgVrFmxljP2JvrbUt678mTm6riH
    TOhaQKDV4qWkuXxQM61CgFAmg3ewZl03laJCICpMNisobALZYYSlQGNpay9MktrmmXTaMO0lUtXlBj2SrJlWNCQTymAfxW8UMovm
    LhJ4wArc8/SpU72rw7n4MW4hqiMbwNm8VVGBd5ov7WofGcTVSSq+cbFtPat96LK0lwo8sezI+T5j/x4W+CMhZCzGYcUL4VCx/+j1
    nK25OY4dKfpbNXGQhKu3U1hgokYEZfmTiJAQnCHyQ8XbJs0Yk0zJJc+11uX52SiwZVgFaZ8h12GGSJsElk8ERJ49AnHOCLuBmLKp
    KKQ+UQOGwtW1ksoOABqDiTe1loUHZ2g2adhuGikEfRiQE3F9ftkzY3HD5HgjFDeGvkB6AstUquGpqwroP/9V84cqbENkKBlaNTdh
    6C+FK0uAOYnsrT+YIGiAH02Yx6LMMmqDIhvxK7guslK4exmA45yrLRAsdFxFkHYdYh3jrx8dBdQP9ZQppRAKw2KBypYTDYiYfrJf
    4ioizQUzfFAl/VbsShLpCUxQUMxQ6oWnND1cgr1Ddyh6H7E9Fy/bmZgtvRPOy+tWuuyYS6+HrxMyas060S643ksKEevLSdQlSUkV
    5SMlxaYOe4VqHt9Qvqcrvo+hN4bitLdWA/VVa9giSs6c6UeE4eZkVaT2xU+KOVlBjsAriN63H/Re9EsaTxIlycZimNBCci1GHZlV
    TCgBdfHG8VNHMzOBYz4EH0igY7L6cw7gLK4TmxmKCvsq6kXeESyouUuaP4dtmJeSdgWFBMQGdxEWYdbaWzZUMxnlE1uovmCKjSwp
    opNWZfew/P9A06mWKbOZTbsgGcPMtSysZkHJbilGlptRaB2bmP7aE/fZb9554uXfWuNbOelAE3JmfHK2qdbLmcqoKSD5jZdjSbre
    nyXvUt5aGOie9YMczRWDj7fFXX1nbsK693mevetV9cUJAVE5j/hEsogqr5Ut/qUEZrIqcmtrq6Wxw8NiK+vi9+WytyZ+mqAA6hgz
    qP94Hvo9mhaIkF9bODeJWWNVGGtVnSY0LpcUdcgPbThDHGNKdnQ+tJ/LZSoYvItiQp0OU4SjDBd7+iBmFByV7absYPYQiCxD9Mna
    RpDKMCRpMkhCyYTYIpwnh0qypgvezWBuef4PU1JxjLgqboo8fWvQImJoWDIY4+1pDzBYyrNBsmHYpEYZRJt7LEE/HsRc7sZoFMud
    x1EgyrGrDX1b+bocCKe4ZcmIqhw2jmd4eUx8W3XwmC38pLGs3sQ4ULlLDyLWwKKEUrI4oUxHUJar5WKcbV2BuN977F/73TfTiNoN
    ZzPbXeXkS9/FoR4TqxuRMhwdTi+O/Bdr9TUlidmTqTHX7rOBL/xOtlbPGaYnf9zRStHtOTRqyP+YTn+a+XwXz7rPLHz/Ndsx2E5x
    vSdiD4eiwlF3ZUeNwxI90GqnOG0FfhC0XQ3svCXJb3Kqh+MU2yCfc4GDgvJwkRUSmQvYhFXaARKZOcMDtTepRe1AEYv0D8FVQ/uZ
    kdKFNejU6i0cZmTLCpKm5BgRh2SKIEmYaIt0kyLJkrLNT6Ug2zS8/OAmBBgM7ycUkQUYyPm0ceXNPtTlbrlrf3/eA91YT4UyuSe5
    drqvclx5b5u4+S1s2bte/FsFrxz5Doz4PpT5dT4OuHNbNuhSgAeYevY73pWhtbpyau6kkU2NjRs1McSiHPw3v3uemH3vO6sDhE+T
    FBvg4EgaKlKWDLHDpg8PRYxEbK9Il68ZgisTkaAOUSz2QfC9GXKx4MG5wEymuUwsGstnWz6hEr7ewPePiW/BVpFwQZDaMrDjHaEn
    JCCh26V/LoRLYwnu7BubWNfOD3XYbDrMVCsQQnEiYrPNiJXaq6wgijMsXi+wUoaBgjIHNZDULaBhuLqAeu6HSaJTdFeipN2hAydP
    snFNRoPSs2nM6WWXMqn3Z783f9TiK4nRHtEVERo9S0zyP1D0vFJ6FqXjJlnM5lATtyzxlDie0E253nXa9tt+H1t6cLVy9yYKj1Hq
    cPFxuH6Xh44iFi8tmflOZAMfCbxHmXP5SMVqtFRYe4BQep/C/N8BqrozmZyGjedL7UJZurqHoRvPnr6nqblfZ8BRkZ/PKro+xluG
    HJzBA2wnMzRCwbqTp8YmwHs975PYr1Zfh7rHr7Mo4lZNZarnLhmwu1iw61gUmZc2QvHHhGV9VLdB3U5ukHrD1qpT+HESFYAYpOoD
    JYCsCZKE1BI5IY7YVbT+AaXhwiRt69RNH0oyjdQImQf41uv0svXC/oEZXqA5E57zNhXrcdt9w8nH9X4au/bV6XtC3fvaLt+I7kLO
    wTdHeI8JXCztuZN9pQAeQuH3rLukZNPi9bhKca46+NokvbNvP2XQfxfI66GG/x07E4tJHEzceuNbpXjIoPkzS2O3YzlqYF4RH4ic
    gj2pmy11XwHvHnZ4y9GwvDnQtgaOkZ7AGYzUZxOEFER6PDVjFxeN7Q/g6YN/gus8ydGafcboVj0VgBYR+UsEpOnbGyKrf4zXZon9
    QXAf7S7Tyozj5uAfvzkij+1f+ERqgORs+sWipnd3lMtvHYZtKgHpvAifYLS2v0Tvu5Q8OFNWEl13JhDW9T7pe/dw3eUH/YuEGa1s
    BHnTVcPllX3nu1uE+1YSznmj9eGs6woxIUem9L0UeIgJrO59C36ALb+gSI6ctEHvwPe9qPXaO/HeLgVm9wK7QmNTbZSZVZyTZdA9
    rj+Ddzy+bnqPZStnvw5hYhlrDocwsMgIB9PDSNMasQGYLlBmFbX4uZ/RvQ/Aj1rgEQ4mdhAXONsKIyv9QkbUw9dgxnGfRYPFVKma
    Ap05V1kBikef/+XdYVH2FI44qwFSITvF8yItsewI6cLAKQqUQyAOAaIzL6fW7sv2Ws8evG0t+cd1fJBRXeDg8LbK88u+9UnpzneX
    72s9gWD7C4ziwW4JRXkt7l+CT/B3ar2vXp/JVp+TBf79541cjzvsWa91x5wvKP79mI2SaN4DEuGNkmnc74IbzOuTNyd8o0slzy4O
    lkYJn64BKyWZnN0+U+2WVYs4C3HkryQT5D4JUu9/dRaqhj4iAxExCtOk3CY1gvw8ta477quGzFjoL0oS1x9XyNveOyDj4AHG6TAi
    sAtCxL+d9KvW2twDVXY79EfGKdfK8n/XdyjeJInjBwo9mPkNiha4gaoyWGdGAHZ2yg3+BnHlGVat53xIcV/K6k8aEWIJCsid14jI
    Kewte/yYSBaAn1qrsuVIS7cLhL3pgqr1PuA2fKp3bnZ9Oz8yLe3xcqtjru28dnNteD/vfQz+o6w8XDJyfmjMx79S04Thf3DAZ/G9
    8n2f+6G+80Zybk3GLK5VJUd3ncvPx5k2Y9b7XJDq/VzqVW517wY6XBO/JQQ9sJ+uUBD9Y7xCZsfTbqkMyA58LgXO7QAjqG6Tg0Mr
    bkXNxKjQQn226k1X9ANT4/eE8P3AG1p0OFnUkhXdtjJvvG/nhwIDWD1rLrQ8TUQmHlNaNytHVNF5T4RSBUbkrKCffXgV2Zj3PSII
    PuoxFVgDVVWCkxzQL7ayT8shUnzvY337xfmzm25mVVUlxip9kmOOPT9OZpYiEfQ+xW3/teFek9f46xiejo+85Ogdy0JNW1Joa0iw
    7mycY9GbQ4Bhh7mQiARPvExy7derlSLzrWcPlDSq8vWWH/NY9qVfW71AuxIt5KR7rtiBMCvqOr5T/WPW5njNGTSWiyOwr3fb1xUh
    6gaL1aN6bk5O2YNL/aPINpxM7BN0smAzZTMQ/3ETWRnDG0TLHJC1coIxGL/4pr5ZWD4t5VpCAWUEO1jIOMFx6/bzDAtQWJFtlqLz
    +F8oGRWQJhQEQgPjeFKE/swsTFBifj+mVokXYaw3HJV1OKWmrfifcMM0aNTKXFzp+b2AC/QzwLET3xm5vrfsNg1c5rNvr5X88skn
    YxQY3T2jnLem0qyNXq95XFctqPKCqpRRopCkK7wbt5uO9hkFFv/pRTwEJmZQGLicRk0x57tyAkcHt9na3tdmec91ca8lvfeqihI6
    xKnfLCqRYezzx6V2DoGckZw1jCCzpsyBP0/zQ+m/Sm+A4TbOdvM5AkgVAkPuyatutj0rmhwcM56Ys87uqpeBefE6GduIqbwvFUQF
    8m5z0nqas/zLa9yWlEotUVllN8fM7zh6X28KUxURLP1G0EvXj8+FPdJBFZBg8wJY08anD1XgMOry99dyALYFY8gQghBGiCVw/0yX
    ZT2f2/ne15Qbzh4PX2Yj1+j5KoxKLHJSmWQ93ql8jxLhrHtK29HpPwkdBptOoNJBcfCExYfgdW2jSxMiLrjwFwg8KXUAM5NTxdYG
    BWnWcdPt1gJ8Ydby86TfkIPtRWLJTS0vFCVBSb8YA2pgaA0uhX8yI+LMDmiwoICLPzQwDaBDo5E3wnAqqZkawtDoC5I3SYbjjoG5
    RqMGYgvhjwWNsDDmR/ACZDjZm2ngbtqBPWwUda2ThR2CqQgY5bUnXRaxOR72wE/csyJ+9g1G/TX6i4oSJKPgxOei6TtmvoPNu357
    dNsDKxB7Yg08gmxu1425/rn9ieZzirc4edX76CU0G8WiCGOFlZ6VEQZO3eCKCp9XCp8jktw3e73b0GeyAUg1cpzGRuhfjxzs4GBc
    CuCiCgX+OxOA3zUpCDHTM1Ynz9KhkX6aJZpxHdMjMwOKjrq0kqLkzjK8iVB0RUSlAec/G3oucHMH+nhfiLnfnzlkXmJS1TPX4CIk
    urSaZl/LYvOifsZF2Qg+gVuPbN10eP09sxcq+WOqkCkGY3dlpTW9kbzc2G5f5/WqmdvS6tNk2EclzvyDiGajEqO6prLbETMW89jp
    d8uqczmC7ER4GaPqHk/6CN+SVH85mEmVjrAP6vuXDRx7dh+xi/jzQIQoPmimag2MymChEWjtj2JiQKGNkjokcyjzMFY6c8PTun0+
    ONfv+fOv0xriTIYKvEtCxCCDRk2sAlgT5XEvp4AOGJFuJNFh1XtyAakkP0s/0kBQmLCFCXNKLIvKtkTuQAi60iEttc0UuNDfAB6G
    6+6XbRk+GMXpLJVuW1PV6xjB5/JZHwxK03qLWMVZheERv3FtDOrKcyq2pYbhUe/L0d492sjsY+0NvJevc8rWree/3Yj8BzFmuj6v
    fHRUh31l54SB3PpE4I6zOcFdEfMsdtOUT3706h7fLaWlS8nSqYi9Vzs+O2+ln4FUA8hS1SQOxHUY9Qg+pPKXH0HwUdNoKxA9E1Y5
    MuSTqqnTFx4pMR0ebHvbnXB2DAdSJSM6M7wcQwq0nXwYaor05atd27s7SFFIkjFa2h6v98ZcOpRmj9DLz5NjCVzdXwPDW2tZcIL6
    0ouyCqGiXhetwJOVniNezhl0IjZ0hwixdlnuJZteFYRoKFxXin8hXpv9el2xboPyJ/oBv1Wx01nXqi0YlFCeWfii2Y0Yp3XBB5Y5
    VnQjr5SJvIoEu9cQb0+ZZeCkDEm1iWXTwlv94j0KRxxKZkjisHSo6wD8AROF0PwjHoQ6SrQiwxEjHosEeQ6+9jB2je8wvpq8UtRy
    bgxDlOT3UfoqSDC1AIPwOSNSGBYImvl6sz6Lf6UVg3jYJXggJQmMh9CNtHhIe2V/GJGIob5jQ7wzH9iIBFk8h1/YOr51Z9a0aeQw
    cZDCivOMpqpRVDGe9p5mJuU7gsjmy9uhCAWqKax3aMmAQ3uWTMbwPscJS7pahXdZuax0ndBuDUBDwOoSkok4Lk6LP13lWtL7LVrv
    RZ0+p/YEFVyxLwZBPcZ9iedob/Jto+frQx2vAWo/oiiksVzZfqEn9+AqHt/1MOuHcoA7+fMAHzmx9w5z4c1GifkmHkPj5bfgJLfk
    OVHf1o/MaYu1YY1MzRXnk5ls39f5wOkiWOSMfnilDcdombTrLLArRv7mlY2GreXjDa7u6c8za9fO+WyYj+wVnurftksEJhfH1wtn
    LMWLLLJS41vtdb8SjzalAT82Gz0hsSWEh8p1nJX/xu6hGEU3oXR21+zY/9j9TylyXP+MDphc7+/sYDqGuIN5QH7wv0XfA7s12Ss8
    upeSq6QK+TXfYr1CHnxSNLz8x0z2kv295VDfjNEU04HczIVOWrVKqHxF69zYzk8ELNlKSbrcunOJRW5MAYYKP2uBJYl1WiDcqK4f
    m93N7nUxWVghpBqMRhqPeaBBR5bK+GrjyJAtE2a08MCFj9s23WzBsVJVXETJYA4vUJH9FCPSveq55H63W338j9bBBLpM6bShMws3
    e/rlkIlWltPow9e7/lufFX9VJYsCI07y+w9jzhdVVMyr+EKKx5PnO9veK9tIeuKUuqW3r6+aaaLLR52nktbtB9ueB4rpQK7O2vHJ
    VxqTi6fYWehYM0MaAMeJMZGd8ugifaCxrDqeeGJHp1sOhfMPLqtOKP9CpHc+WvIFi8AwZ9n3LsWdPa7cEeRgh9gjB81RU/t5uYut
    nqeS7XSWh49rz5v2H614/5AjgGcXyd2fT+z2mi+8wKUv136mVqedOd04nXhlnGuKiG0Mbriageel7SgN0J/nKzY2JVJxDstZ+N5W
    aavOCB7LH1PdA/troEqbWbLiNbaIzPTLNuXnO0OfbCsZ2k+LJFgIBTbZgdWNauDuqC+c0TDGJDNiuzn+M7n8Wzb79/dv+uw2Ff+i
    CFV013VVs5woOGO+jLivr/MWYpM/3Xc4x3alLSyApGluFyOxFKuQCCGdhw7rDdVjLATEhfavUz/PcjHBuT0EB1NXV8i06LAVQ6Ho
    RHsWh+KZ/omSMlxo3IaDRL7vA+JuLimv13mHBlddZYKwWj43p4837p3b452PIjO+C3PMk8kjsrLp4xPnOUnVnbdFR8IqrqGftu4f
    v2m0Lz85NDPJgrlq7sNHhp16HejelGOYvJdZFgasO63OK5DDW7oeb8dL+S2/ZfWArd9vrV9/5/b/OM/LX5z19487ucpT9Vw02+bI
    5yBglgJ+ZU0ipNo0gWmHSFr5rcQoPYQJcmXOJDS36wnI87008psrbRToECm0yKlSr9dyLOfzaZ1F6n54DjUrs0EPCio9r7Bb9YAC
    mqw7ma3/ecWX4C+giu7yJ9kBBzT4qzvS+1x+vSYjAMLiaq9fkLkdCSf9fiYJLkRTv2YIL6MMDZuBfhFOkwczZYyV/4bLzzRMkhw2
    X815UoM2JucrffYcA/Ey66CFP0el0edIxB/M/dZ1x/Nt60Yzgj3C+XRDtjSjy7ftDFwNe6lk6nBqzrEmiO5S/nup9RuCKMx/tOAG
    oc8/yPYrRr37Irh2YirBrdliQ01Jaloka0pS/iKji8BbjMyYKBej1glwnJTdAEozROzHxRGH1R+1Gq8yHR0eee9jw49BktrGrCCU
    PkSEMSbuq5UfxUWbkCMIQEtiSNIi8ksLwOE7N2+VJ8HD1cmepsvw3CFa9Hd3cHt7xbtpo7/2OhdstaKyESqeJEm/dBWnDZX6acPd
    9XR9cOa1TGFD27zcqEIkEyOP73M+nL3W9TPDwogJcfHFNHF8iBYPT28L5owS8zhU/a5Y7zEeUed7AXjnmXujUq2WbSeKmyjJujWt
    S2WE6t8vgtCx3OyBN32jH/d2urnaZ9F8voAi05qb+R3d+176vUsgYFpGAm6cB9JARHuXoM73mrHjeTLxvoCA9/o3O2snFyIqzK90
    y7/nd9t1Ni5NAYuoO0tjy+zL9ot+Zp2Wy4cyv+u3dm4Zb//wl9vJ8F9Fz03nkUWTAPnNAm4c5yE8wXFNKfP26HNrSJc0/wvy5zUe
    zZWSs6bvslLeJ+B3CT1nbyScE+8xeoMgM0cBe9d9zWDtRwewKiyzOcBv37mjmWxFlqJuM1tiDIbQXmmBkM9SBI4YmgcjmnLJCpf4
    73VEiWiWDFV4eLG5C+OJqufwWCjtXp0yFM5MSr6Xdpqyw3kXkikKXNElrNwMtXOdHymlhAC0H7O4M6EG4bVHNdUZl0r97RK6NO5q
    OryhfiUI63oU2a0zDA9+ACQ6khilw/R5TQR3boOr11FO9lZkuN3iC+lUjfOfsTPVWi/mFb2c+yQDc+cG5mxLV5WbKbkUHVTyRJ2u
    V/RaPu2W8AuONcbq9+37gpxJUadK+ddV+V9XTTUPfvcqM5RPS0VauATam+lr90p49ezCnQPqn+PyXDyvCTfvNbkdV+W5KqrTew7+
    VNzx/gxlqoSh8dLIGOxly2js6aDDsAfGFyjYslaAC/XI7J+o2BYDpOxpyygkRkziOk+VR/azRPr7hKWX2qAQ9/BLjjqxoeY6BPRK
    aSP8d6O9h+rfJxuxnqbIJJLras7FP2NS38H4PyOAD7W+18mF3/s26CRCQ/xATN/q9HhGfC6hLEFwvefeb2lRsKYcZkUX9RlUOSYV
    L74nswZUaqM5urRBUjkXwhofr6N8LkKKReQhHRHaKFzIjIXn13z94iLlv4qP1cqONuabQfS7xKBomC0aIeRGC1IMRJYSpVkU/0bJ
    IaoQ6aWMyTb44aIzJYVey9YRtToIHFG2SsXBOlK46wEiC2RzojBcsmdhmpypa4eyngRfs2S5ImL3aCGGOibHdTuYR2ohgnYh/Tly
    pePcSkjSmDO6jFHrQl1efxrZCBwGCQiEiL1BbPHktRtsz+sDf7UlffOXiYQee9KH6nQ/y7P7ezNwRZ6GizRTnlyeV/sY95h32MJ1
    nkr3vgLK52hFT9CgBO29aIwpHlurovu8RvfdSO/Rval80k0jlb7i4vjQDdtoi0V2YSiwMjW4nl0wyHtJ73OA/140d5uupsiHjHjh
    DWF4IgvORkQuGGx9cynepurM/ABTaubrj3RIJnh5MOhTOnwWUQZoB+UR5VMQehUFmKYhwgA8T0xy3ETZ33rDx2heBEGjPIHXidrP
    g3J+hEVUlQfY7VGJf+3E9oPvu2/PcXT2IndEGo9L9tW1gprjwo/ONOjrToKYPNJuDYSKrGxAbhurZeZ6HSI+Ci1Dtnx6u86RTXLV
    ZNHXNVrjKSvDD9w4UCXZdpCfA00Y9iSYbAbU9HTsbSo7RW1nLDTMJEQgZqzxUt2GzmV1/v98hCv37I+e+rF3U2GL2T2nYopLHpce
    kJUiB0lFlhDF/nR5EYLYHp3s7hRWQBxbMkf5+Rfxcz0yh02e4qppbZwzszGIivvctKLlfMG+KvovxDnwWO+DuXNBaYeKmgI0aq3a
    RWxNi8tpCtMAqp77DSkCAlzRBcM4oWCb+Rxa9jWk9oNf784M6C8zS6sOFe9KkUyg3d3V+9GUpE7vs7bxfbzi9dIKnATKu1HtOwHb
    ghnUqD9P5bvuxXOBkdk3jxhkyxSTqrQL5uib3ZJfJ94zjfxXqBYcU9jFbPUPGpeBRlUesoma3rNfNz3ex5h0WUageY1xOm/8n521
    3zs9+zs7vpJGWbj0CJVrLGRWrxbwKpsa7OfVjgUyaDAA192Zv23I3Mwwwg5KQlUw0Wj2RUFUDhB9nFQC5LKrc9qV6iKmf8zrhEw9
    gL8upkqdNdJurucpfPbFu85dt2kXshVESz93hV90V9sdmy/Ltr/Egbc7Pl8rwOlZj5n2dH2Utqs45hVlJ17dJ251n89SvW2X1mqW
    vIkYUSR4pJnm9Q6G9Tg1fow9wUwt5trBLnML8ugyiDovhUl8ngoyFwpKTdQw6EPFKiZ8EyeYovZ+2jfZMMTgPmL6s5sWOTAe2jXX
    amgiVSMMR/JZuApS8Ed/eAQ2onl/m6lobUhySvl+Rx2IvToNF8LUX+CoB/v6+87InH60cE/9IH67So84YuTsDmHM3H/KzEvl/DAz
    FMgpOLjr3b2G/xWJ1dxebm9ubBuiIQFhlds0nPL/+joEP1W8zk3d70m9N7OZK7gdVy1nx1Wns7EzQQCq/9ZQ9NYuEwoHA44nDDNU
    Wt0l5oPFc1Lp7qzPcOrX6g9dJ5sOXyev9r2KcGcpc6atC+un6w81373Knr0BH+/vFl7nCc6uk5xQ0ikEjGnjF69rTxim6n0b1I/E
    yeaTbsn3Dd8Dy4JBRR7tI/NISlmQi4/tRqKHK0y6Z6j//HvW57vJCOuUXMLa+3cu9T68fN+54jrpJMl4Y/ZhjamCcW4mC+ThaBzK
    i0NREubl6ARWossFF0EzKkBfU2L3JcELVqljKMqjCc0svkA+LD4RhTCOhw7LjFjRrsGR/FJ1xgtBQdJ7viEfM6G/8dOiOsR/kHyi
    73mr7H2v5b+Z32P0TJ6iy9kwOMlpm412vcPpOevk3zc/yXE/uSL4Qm59eUsj6/3YuvE66cTn0WYLnUyfv6ohs0xcK8RaJTrmSatn
    +oXSR24sNzNEB0JTounYa6yL8jcI9nyAkCEcVdkhbRAdzuQQBZx9hmugL4M77QnKaJtL0mhozM7TUfEF3euUhQY/VHoN3r+eg7O8
    dF9GREF/rR8pQCZVOptljneQ1eH8g/q8P1RDyeAxEPkHZKzIbh9rvMkFmbJxTMzm8nFBk/LWci1DNniBlWm5Ub8Id43ZEhN22XVn
    SzjDlSkzbIp1MAcIVmVL5XmHRZMdMB5nLAjtNGzNAI0cKtvQ53oovWgorMMhmoNbLZaIbbE1Euols9QO8RrEr/KXmiAxSUY/gBm1
    csgyAClcw1QevZbqWoqT3rWsBElnLB/uTmCra2Dvjgzh5Lt/xWfN88bzFt31a9snjelcznCl2mN6U0Zdiro3LaXt5XoQCdQ++7Tr
    ef/pu8BoX0u9UqD6I/x0qXSr/kPVd28y9y3qFFLaOspZjSlQpWQz5A8FBcXCmrVcjfW9PvpCUYWSkPISuNHLsP33lDy8HhDx+7P1
    MxPo1TalSCdBXrweEmEJCCAauPCncdEZS99WpyHhjh/dpLQoaknmkwWtZRKCC0voxDcPl0pucBhWOxkWilKMXyRqr2W6yuJG1LmW
    Ia90IiUSs8GgCF0gl9t3Kh8HaxKNKqKkCEIiPGzk6Haoesg+c7B3M2idew6bbwj/EBobKaIEl6uMafd4v3qt9D2b7Hn58Pk0bi6r
    yF4y3JtK9Fg3k00XzbmfeWm56TqRpfecxknTjHws5H0JV5427ox2qsT4uzU+5KBu38oOcyoYn1AtAT7Bneh9Ts8s0V178XM78VjS
    OuW9qKwII8RTXmJo/zYLYgoUNCYLzCsUaXy1LA1tXFkeltqKGQ7UoUVpAq7PI10jFbN37TZ/dZmr2rkiuj+dMVsuMEii4Q8mpQHf
    sDGTvbohCZOCmmXyuwJZjdSeACkhIho2OIUjbpb+XEsFZZiySMnhHyOF6eziLu2iHvHUuWKrIFNIFMi/FxOcwWMonCIsjee7h0z6
    jMRmh3BxTZqI8ghEU3AYaJ5+EvtxSAnSCuFUzZDjjBb68P3ENXSxckqeXVyaxTUVFGssYqQQzJtRZBUHf+qZCFV2x64hyxo5UZSJ
    z6Fpf5mIm5FZqITE0YqF0K0fBPmYmbkyxQ+CfV74PqXynxnlkvdcdggfuzQcGCREoajPGTrWga60PnHxawarUoxblE7ADF/eP3LT
    F13PN5kDd/SeRr9fc71iGDkIWeCDouL8RTu1nyN1MZdUMCaFHeCHakYECeEESzWawiUqMpRGToD53MDL3O8CMaVbtjEqBxxOiOW+
    iQKaGXa5sGQmIVoklkZyaJisrQ6fdIffK0w75dFMsm7lCoRypjzpfkIrL9mamfZdWqbWsvKmymXNOq2Gx1yQxFrPV8UkQz8BURr+
    qq1adJjWrVHJuPh4IKfvZPXIFxbNp3hr/oTilYlrTnzmcSjfm7neM/wfI72ed9qs5TwKti5wusHrHEqQWmQVLqountvSTjYmnpU8
    1qpJ0XbzdDeITk72ofG1dPLXb4tmO9/y32cri6ybq2lq7TfMKaPS6ib3Vl939+cF3BM9pNeONoG7+mtk9MtyioGBwsZICZTBQ/P2
    pxpJZI3GdDZnXBLBGiTYGJV8IND7U1JYviEedVLsWTC3RH2oNQbR5r9ZyFOtXyfTt1DS0kPjAkPxe+SPmU1+Fcvgpf/soPpjS8kM
    HYbkx9SWFsBs0BE67pWib92LBEOXs6WvBlWZmUFVc5znxyc9gxAqCYtqfUXqqN9gTC0PkxAGL/volyTox+LvZ+oxA+c3IK1zPsrJ
    9XGNv6LcmHx+XjcdeXQ+D9vG4A9qLTdFCKuVH6FhzON1G6eUSHQrFHEQpSq0omkL0p3MURQhDaPah8xzjVIxOUohZm7ru0zV88Td
    cwzBD3wn+AjP5dXSnglbE85NruyowTGV493MUAjGKgSjfLGgTuhyjdm5HUn1fUv1vl75vqp9XzUDx4C2Ozo0/kkht48u4kelssDl
    TkCNiVECbS2PMbqQhnFVUWouipa5JqnFujQjDLU14/O15PsVee3+nwLxfZLyv7HzjY11cEG3J8mJzca27ca2TpzGtm2bjdPYtm03
    jW1bjZPb573vH7jf7w+Y2bPWXnvNzJc9hV8pgg+FQI5zFc5lIsk+wKqFkB0ofqZuOQwiouJAT/YeoXWlU4RXE8aJHmA9k7LCDZYq
    O8qjXqz64bp2yKmDVQPTc/ZtIPp0RvLqctcicCwPli2Ad4lvkbaVZIHPlSyw6XKzYUBbflU5GoYSXifSmtguWTJ/9nOj1P8+1P55
    xv8r93anf3dO+q8Anem3SBA6ncSm+/imMYIrgxXP8dJvspPQ+93n591gu1vP3Ldy/7+K/o+NuR9N3bhUBPSHpy5oIkIB7CScCH9z
    OhQmAoObPwp0HssnhdKnXklMFF5tiqMcoE1KgNNSGepwdJiwYKpVcR37wT9nSoPgdvJCiWBlzHmP2N04ULaY3M/QQ5HnUh0u41nu
    RBLEwMxorEdZpipowUYZfvmNXvWfRZTXFsWBFO0CnQj45KEDiXVV5hQGQ8YcPSDrFmeGynV4IyDBqUyaKi1VmeKuBQX4e90JSdUg
    lQ0cXWYVulH2Z1Y8Q1fqKRrc7KFLwisxSxA1E1iRvlX4rv+mFVoCll+/lVVpIDOpVoDzF2Sp9fjbrbnDD8glRjQfADlZUuEcZ9vu
    9w4HoYfYiIGBfFoVFOKW5pjpwpxqt87hgKP0CZvFmr/m3qtSAYfKHR6Mn8WCTxdMH2Szn6u5/ruy/60mTESMGqvfw/AZWNZDlKJR
    KWzf1OetevCYsrO2PN8+/0nP/zZc8GsmM8HLJUfEx17gWDkZJG5X+ytr9XCw0nNVq6ZQLZF/aNH2GO+OE4W0n2N6/8Civq2OS8c1
    ffH5Tf+tV7/7bue659Y2zmDYQ0rUHOTCjta396pLzJ7ryPLt38LmH8eYW8a0mY3frh917Oxd9OE4zJcPzgNOMvS7k2plr+8b9JOi
    R/NZbdjKEEqmVGrDX6rAjohDo0i+EmJisKOxStk354c0plKXgiOIxaqVLXxFHg1XOna2op7vguvXDa6+rVzVYFUIuZ9Vp+6fWV/X
    9u8j51/H9j0XT/5fPeJZ+AfkmaJkx7SkVBQml8mnVtovLozRyZOK3Q5vua/6PTeyPSftue+Orai/S3r0TrdMCgjxbcN+fC+jqo/2
    +11bw6gODyObDcds63H8/gIK/3itIljJR5pdyQ2lpnMzk+FUzwRYHkp+m+Uo1391kQTAIhB725sJE9LAPVO60AnshqpF14qG36UR
    1+WRmpOGmyudOTMZrgUlqff+xau1yvzIjuDp42dONt/1s2fyW2sUjCuDT8GB2IbqotUdtfcd/Xj6/ruanrHRtLFkTwpjmjIl955m
    dmX92zO40nHJ9WD3E0HXJcFT1gKdq5S4sBS6davqaqPAZ0Zrbq6VDIBYmLbjunuSq+sUTNw359oq23Si0k56lbTx2M0fS3CfJ5e7
    uvoaOZMSoCSpsLdVN0an/+Etqacw40AC6+A2mdh89IAz+/LcreVgXSxWVK5aJoNkatXJXSusyfht/uUiEJqox+PNFoCk0FVRrth2
    znHtdzrp/0Zg7/80mIaUFr3x9N3gPud1hs/Oe4nK//2b36e5/5qzHd2cejtkljweFMC1Va4KJX2wmFe3JK1rrznBE6bTV9IhDngZ
    D0vlf9I+u5aFrjV6hG/FTabGQ8ZDvtMD/TBt75k1J7Hp7OSGUv059k1Qsm0QG++p7bcAauLQWeuPj5LK8DnPHy2CrqBHXzyDmORj
    v6XRi1lxZ8kjEbDA3WpHhylGguLlnis9q6TYTWEUNxg+H8f1y12e2dd0A7gyL0kVIvrT4t5HiZtfYzEUbrZ8uTNrw7Denb2juY9a
    RIMnLnPq0xgos7D0xSPrxcFu5q0sAZCVWLW+H929jxMam75PfEg1jFuPM+G7LEBYVHnVzc+mUv7jeYFT1V+P0D078AIddUWJpZS8
    yf0gsR3U3yCvtnK3huue1xp3QvyB/RDrabsl0WT1nZIzgSdngYryFJbpVRFqw8ITbwq7K++bnWl7b8WSNRRJRC50dq+nRMUMvzsq
    36ozJf5FWYOd5oLTpq3Jt0Hrj5CoZOSffRPWcYmAwORjuSr6E2c6w7USekdwBFO4Xep1MXdmTIZC4xPNQJEdVCwwoQS2GMMgdUMh
    jAhSbk1UUqGcJyoM8CJpWJYFW3t8z2KHMBh8A8hX/lygBxTAXDaOWC9eaTPgtD6EvEZbDQwk1DlkA6eCclgVjR2ckRloaYQrQuya
    TvyTnYAQLrqxEM1jhoRbBRmVux13GBZjU7nGXcR/A/fQzD3pzTRdzUwGv+OHFuXAyTzKR7vN3FTe8IMejA2Gg1r+MwPGUc0EyhEA
    XI0Edo6BQoq4Jm4dv7LEhkZNGX7nJkYTzCKfh8k5/aok2XQjI9jh8G78Bd4Z+g+ZzanoH1i6RDSoMDiiTVC8FxC/nY2tvAJ0zKIA
    Hx+OGySECaL2nhaIGkDTTFACplbIZvOctUDAR+xAgVDyoxyCjiic4/ABy03sqShucalYlR3kBFWpKxlkzSgSKCEJIL5+oJ5EDPEf
    tzlruxI2uH7ESZAkflaYorQKpi805RCcGLPIBLFIRDBJPm2QSKMbHaowCAw4U3kYPO2yx8mWUyEKrg4XSR0jXlVm7XK5eEoOLr9t
    tcjMLhqQjZqZsoMffrmlZGUg1YQCgr3wdQKXxNEOeUGguBBKGyJ+RixV/d6/dyCqVLE1GpF2J29YdoIQDl8Ohj/YsPi7yrQd03UB
    /6gM+4zWJCCV9xFIcwt3RSzs9hDVgIn/3oyDiQKTqhNZ9EsYaNegvTMwODueCEZpImqyAOhXaQ1hc3YA0oYq29/4JgfCiBjk2aKS
    TaFCwJw4W5u5cuCkRMeYRGA2UBUzS+wjmwSAoZjMT6Ft0JhKYcwaQLCxC1ul80Q2bxZY4rqSVQBmcMqh0Y0CWL1EoVZIYrCjYIX0
    VFMSGkOVwE40zPKUeIM7OlIvVjcjIWGgGG/tlgoXMmARhBbVQRFyrp0nWVqLbr9UMq/VKglpTl3UYawCEqdM5W5xsKVkfWFKnzxN
    UA4Z3fGDaIMWxsMaUuzGcEChGzp6PHl2uDwHFyR12tdA0vHIEDrXNhJg2yACkhRAVeXSQG2Y9nRiPMJaZBOvUxhIQwLjFF8SiCc2
    JV+dXGEkK8Wz2AhDbD3DwyWfvBg5hmHrRFlBq5nIQaBeKRczTj4E414UwjvgutORhHadr8nIMJNOtoAUYsgDzIA4YeOdRRMl5k7g
    XUDlIZA/K7NQuaN0kpmklRAKzNvpSlzV8xLOAwpphbCDxfSJ8pkH/15NHdUMuqYcRLcfLAiM9OneQ1J2BCeZ0BLJy05QqcDj0Tot
    AktEIw1Ui55QVlOFyjMIKHP4SJHgsFAehcH5HUJsMdAXnvNKzk8SIoDHJgq1qTGdUe7CNAjc0NJolNoEGACh0X7TFTwL1TjkIMoQ
    R9CYi84Tm0oJZX2HsmrkmgbmtXBRIunijsKOlMwYBqsjARYMDcDxHCOpp0Jodw5WjvkOEMxVPBbl2FKkk9x48isok2nk5p5CRcEN
    c+72V2FQA4DOzk8ttqJV+yEaOdVTVIilujQw5EfjcmSNllEqdDuoUzcxKPm0wpRDSawVWs0TP+Sp5SrE07J4h7Rn6Ffi79Yf+/Mx
    iJHqYNLmVDHYdgsqjeeaXTnkM0GRFSdgOFJWIASE9EeQmIMJkcxwHbG5KggpBXFKmIQcvQ+BkVmTqMisQInEPARA2kQjT8HraJOA
    yQgDYu/T7SfM+l1M4YxSk66MyhIM035vg2Vdfg7IUwKiKoNQISgNwPmckFWAIt8Pf+3GRR1HpdIULWjUB61X7S3xFkYJ19P+PDyK
    tkO1diKBXDugR0EKxfp5sKB+JbrSxVtW9Q0j8f6mVQXQv6yM4FfycUiZrKVYFyZe0R4tjLBhLhWB9JAzQo2ZC+Vs5BAYG8vu53Gv
    cP4sAIFtso/kTSDUcARBIMMZi/YnJ2Wsq6ggv8RwNJ8M8RRmnyO8oTF2clIp9hDVg32NFIV4MC6MmRKUcvy7fmY9nY07hRb+Jqng
    xDxIgJj7cSbFZDC2zWbdYyd/2KI7Lq7cDdtpcBTTqMqQJUULz1h+gr4hsCgDdqC0dh9b5YGzShmvsgURmHeO5tbiCOnhUxF4UYcf
    mS3k5/sF8XAR4nnwsHoR5gpxl8D93SeE08LER2peFaF8+FyHN6cHpraSVP9e3O31NTRvWrYGM3fNt+2BkM+2CZlzFZpFBQY13Xdn
    mOBVNsxDITDKuTx5F6jEadD4B7m61xvf1YdXcgahFA1Ibqfw2/Davfp7jmyBLQ0Hq4lgIKySOSXbNxL3Mm2YDdfyii3FtXNkFMW2
    sGEjJPcFymDUQfGUOZqlqXjOKQRjJo5taLw+M2wpw+VmSbHGHCQ63CIY5GHt8U14PUkS8Xk0WJa/DxJXhDFCQ1CEv8yfR3ZbrpbX
    ZYDhNP0+ZDdVzCEojrDvf12Lwvm66F2M0UawwWb3o6sOa+JocKc8v2PYjiEZG7iYvPZ58y400sUrg3PAz9uxwCOt7gNTis4GBHcP
    RGi5c86GdpQ1eDnobFuypFGvz8100tCNCGC0V1lMCMDGjvDyC+P96cFQu8zYBCYFDdj5p2mCT6QYY0iR4SbvLd1Tx2ux72rCUpku
    HoFgiIfRMA0ZS+5HdAq2NggT4vjnfR275802K6CbPiw9xh9PBLEfxTdzFkAIpyDVU2JHoiE7Grqy81FsnUFLFM39n9Zx4fN7BMEk
    2EB0qqe6+rn1jTtHLbPThevvTmso1TiOo9a5/F+76Ix1CzaJTvNTEQkB/HVUfUD+1JrRvvw+BZi4JM5Ax9h1fEaxZcYGGzwkKlxK
    Hp6yuhq+hf59BGpr/dUC+BxafoG/CbHai7hcYrySxLQZ1HgjtXeTEiSZLMMZQkDC8Z1oqaTHWE4zPR5VIJZtTX3TLQFeFPLLb1KK
    uBFJxlqjRU1/C5j0AJRIVfo05d1zWkXNngPcajN5Pf0fUYppH5ISXsyIm4Hz4xICSZiCV6jUobBcV2HKooQ/Am10y8IMmWV8Yjrc
    SFp5jGdIcIZ4oDqpcltKHU5Dpj0CHWt/SA3MHJSR/V7VEtQVPFLmCVTUIa6yb7uVGLRGcsfP1Z/QZAjOGKPJbpl6RmX6kE3geOML
    u3ITktCSAzKeoc0d8Xb3pwPnHVlTW5lmBx3wCvA4n55Jr6G/7UNDK16k1lkgMn1UdyqkFqQnyCVlZERFWJOaN4gIGihAyQDREUtN
    jvDOHz2iZq+y06vVUs6EZtTHGCCZ0wgDVB32STwT+PrGwNP4akQHmX2Gy35jZlgBdqbUBa0E+gobmJN1ItEF9GK5Bq5G3vsHGpLA
    0I3rhYqDbvNgY1qDqhcVARNHFO3xEp9Vrmas4COIy1LyyconK4nCsCY3gJs+zR0KpiRoJ+UJCQFK7yCGGDFvigNbVW9eWzc8uNk2
    p3wpflMB7IWWP3zizjg4xFGEvjnrwnd75awAE1RmB3jsTmxwPjxTRkal3L5VnmCMGYkhdTDAYwhQRsjwrl8qLn0XIQtwUAVF+Mlb
    54LFmuV1UcW/UtDtdvP8fs5iecBv18l1joAYq8SJAIG8l7FhLMjGHSZ43+EhScOo/zdFobhKmmZ6/Ia5IsMuG8zYnDbL95jjePt1
    cqgtpcnvgFlHsDoCo8CENGBqM7Ir6X1j9v15U6tcDJa06ilMEeuXJuIkAbL+cPz7iaz/sWvvdnHFanvZ0bjTkkBDwqRSS7SMl4zT
    BCC+aySM5MsN+4OzTXtWgQSZ2Fg7H8G6dmeuHIbeGehOih4WpEEcPMyqB30ec0+tPsF0BMSNWCBgvMFd7qYYCrus2oCbCXZqnBvn
    /C2FzbfXps/S4DSk+WKEoRlGWhcVZmro8RkYY8vcYcvfXz4cQgNGlQcEhvdB6O70RxymxK1yCmhy4nAWMfDl1mBAjAtbOhPCrfKd
    8za7ueSnqfsC5T0MPRuUJki431wJQMWQgFoARG12TJUb0MltIgSIUACdElahznhUHKIO/ihu0H/eTy40IRUgZA4Q+t53hB4gHDkX
    jKFwA0UtPQbbhOUQpMxtu/C8P+9sXmenub3h85PhoGRFM0sgGyEpGgWDbkkCIiMxh8OImaeRZSyGdWw1AlCfz7l+pAneFtEoMpiz
    62aWzNoVGoHBiGgEMDNDs5gYFfwim81ShJPdzOjeha62GkDcg91Gu8qRtFct9E0ukb8HyaANB6PDMO4lswQL4rWOHJH3/HIygapr
    VM7Qw9MhbuBgQhsTogsDQ/dgX1BAOg20gbkQw4eJfMEuOuLQ9fiO2f3AgoN5xq4zq8NoiLzFXxH2WNb1IVEMDrkUr1HGOu5p+431
    TrmQ3Ylo2fCJlXE0z4HlaFeRM4rNmbkw7D7d5lZzp7mC4kSIV4ZsZCV9OKTNzAUPswnhKtumumf1RrznRaV0e8+aarSDZRndqdox
    xy9kQyL0vGOniLEa7GYEdicUbGd87LpmB5Ov3+wA4J53ACA2YwPMwC7wUQHqxqnr0EggwPgpoD0gzVcQUJSc/pAO6OIFCOuheh4A
    pldCy9pUJ+7HyRou31Ng/J+Pb87OsbyLhIgsKafH+3BMTIQ9czw9rej1NcSLE11xk2BGpg5spvYdWE3X27jrgnUBB/0l/gzJpvaD
    QWJXV0dKjlE3prkfTaKweGbZVJCKVX2sA4PQbS4+oBLlplOEnFcXFR6XRbM+xe5ttbNfGMH7vG9T2cxTf6dxgYNBtYAN7KIdISTN
    w7JNa7LgZEcDsoDh9LHhqKuhgJQzkgvTVqkbA3rjRX6qB/YVlCitJ2WzlUK4KhWM0TAvH/V31YvSXe9i2nR0HkeWqwfKA+wndc70
    sOxNDz3pTdiJHwuKmtOmyCJACy2/wPyFWGR5dvrTPUwX7b8aEkc2Aat1NMx949DG3Dq3PUksfKi6CPk3Z9xLIpKypVUZUtGexVN9
    f6KxXcMd1exTwGkq9iiFd/YHEGUAgwWofT8hxCPD2Mqg6Xj6aRl3e4d5jvr5U644ON2hV7poTYmgl8yQqxABJQgpiWQB3E/qFXAW
    iA71/fkPIZCsOJzG54h0YcvKJ17NzrSIxdgDuZ6iKcgtUE8DiSbEwd/0I/p37QDJkLhYmAYJQaF7yuE5nWxZGLMY/RsQPGnDf8oQ
    Srvy66BxganIUV9kfWgAPIsqAgcqkjgwgPMt8pnX+tg7+akGETiitT9RfRxOu3zmqshmXUhyoSkBXgBDzTErOua5igNzf7ZEKc6t
    kKeAG3+ZJvcQ2BG3dg6gZrywEDxjcNiU+Q2N5d82a3KA2RdcnxXRf2MQqvV3AUENbFXquZ7rPL4Gbb7uHoeeyU5MqLyZDdtQ0p1z
    1SIJKdQjSGHZhLQOvRaCOHtC5IfU54HaKv8+G0HMmzEMOIQNIxhHWJlQsPn+vHbRYDDGLN43+ZZs5kIVjB8j3lgUsiNjMXNPilBt
    u1OUweGlNNfcK+M5UPvecyZEXCO+wX4lcrN7YvZCULVhZ9PsnOCQZCggbP47MBQUT8NPZ5a9P9cKGSRCUJBM+7FRzdFgYO2q8mMK
    crrV2+YwFztSMGFTFDke7JQx1wyFsmMRQ4p0FqGud+cZMrchJNQVwOyEn+pOwIDIgcx5RxRHNQknQiCWED4A43I2UiHaAQ1mJKKn
    5dDbcAAuAMd+KUjgsqUZHJKKyxmbPgH1fgdEXVBRB0IPpZqMeRSDpd92mFeTKzceW4jsM2w+oXIMR5BgVjEOhY+Psq6vY77T1Oq5
    +BUV75GNXSQxYhf6EZe9O1yk7dPYjVDPdH6766CSacCgsrQxhL0QNOARmj+qfLAstOa2ST2T1/8QMB/ouLQV2Iya6XN5L+Bb9AiJ
    tx4b6wSrBhwqWwmW5dlGQ+lIdRzMxY7yq06N4zqQJzxojlWy7Fatq7/QpNmwKf3jwBZV1kRf2vw8eORbg4PMrt2jG7VtafvoSSQp
    9B6Tyxk34Khnp7iEeDYkBIx6qM831FsC2YPhkf6CY4L6CYYRrs5QUiM8OGzwIlLvhFQYPf6+kOrFG8SpNNMYK5Ds2KkksbuainO5
    iW29bgPUqJUPtl5OtvVgG/lLUtifeyYxsf9YvMuMyxttGy2XTZ18yivm6DzHUZdQGUSvxE42IUaiicXARCq+HoTSec67dL1grSux
    iI5P8BVirRmSFI+KAIL1Izd3C4Jd6Rsfl2GDJsMV0axZpczNOAh4sSTkiae9WitJHn1xNrt5q7tZG5HkRCy2hotLQnXJQT9zjFOI
    549nUxRf/uW5fRxOVXruAwUoGaBOPR3tUVFkMOIckx9QojuxmeSPkvJM+dgcyXB8/oNduMOOPcMp+K45KrkR6zmI03G73F/c6cny
    hh3uwpD2UkLr6xUumVC2QlYj6BvzxLHHheAZHzxv7eFzJNl7c+Ep4B3uihVJw9cnmp8svEzE1x8de+FovegMO0ZsPmJgsIlwg61d
    UiTffwuMLpXcO9pzcXfSkUXHhUifabuBNkdXIgWfpsErRtZts4m/Hk0ksClml+jfn1XpXHLqHX5DZprZntMyyepB5X381mg8eOdh
    zfU5TVEH4D43i9iOMNw8RGSjwklplN1ivrOLJ8om83R055+Zi8CxgbBkr9nXPOMX8DTkJ/LDlgHDvA/Nt4lVNMYFOBHMd6yQuCXP
    MW0Q2wkEcvD+5PRRDJYZPky7JUf+henHur8iIKfcghPvT459fI4OMtXfzzql0fHp23nUpXGyloP/vhmUsvS3/+eU0a626U7qgDxV
    nYyj5518p3F5phUnr8gYdmG3D1uoVTQy/vlhKxCs+5qsNueEAv/rwSKemw2/Ve9RGMgH09GytGrAaiaIqE6x3xx4F2N8jDU8TVWn
    UFgJHZ0PsTNQYdfydzPtizvaP10togtT2afenKx09uE3WY4F/9C87a0woL3+WwdPWFYMCUYcTNX8WlRY8wxWZ7NtBFdOH26y3/Jp
    K+fR1FpYFN/QawJeCgvlpX080fNtWB8SleB0O0Gc299f52pZfywpOM7G/eeRg8kASgMJZWMU/EF6ALYbIdIdBE0S2NeEtweB57YX
    YGpaWRdJPNDYKp8HH4q1gOQH/szeRZEaqXz9dz+UuHhI92+oZpxsIFd3MzPG/miDhH69xp8AKvApCihDzjYzDgsUWFLf0lDro1cG
    /fV9Y4KfIBoePqFQYLBwHLDOzPwIRRMkXla+Pz4OC9mer93CRWLHGzgc0M4jw039dqKHFidKyML8+0w4An49UU9x5eGvh97oqGYI
    vGQ0Fidu6qNbt3UnabuMRWsdijJPsoqV6ugUSdgYaMZzsVkZN+Zv7eUtBH/QKjlCZipbWhc1dKdipsUvsMsOY8ezbBH07IUUO0vx
    c0p0TWMY2PZstBPRFcCEmoseeGdbqMqUHpcqYstz326859am7FKP1DhYWLtGo+dLZ0aGFqszNKIuOhgJdR6UBhUcTD7HehiHT52B
    wqpEIGQEjvAqWMOi1XiQ2Zoan7KYgjqVyrhTe4WpwmdRvZbvTts4cAl7vg7KU/r79IUlNZE5EtAjRBIQFoO3q7u953hbNpHGzLdj
    dnS1ukhwdbSQzBzOwYcNy8aC+ujLREoC5K6khaEnJbX6YLB96C7vQ5FwoG1Xp0xa7T449wioqEmtVmN5rKCEkZ9X2xJcRf4I3C1B
    1yRjWc1EXoaJ7Da2i+S9E2MmrkOG1FsYLUtmEJDyUEOvddMKAmIRJaVR17WpTSqwOJC/CMntISRpEaj+adJG5Q4REivXq1axMltU
    N/C9yTNSqwCPmZr5oTJfoUmXf79xY5PBwZoqCWzwGuvvAjnew1bNmyijQygbuchF176c+//4jea6pyV0IMeINEPcu7XRcXnbEffX
    4tLTwQGYXlGrHp2cfGjVPLaK7+Pxh0cM5NzSc+2szuasqFHLr4vqSFG6UIVZGLSWaNt6Q0Reh95Fe4YJYS1tEjJsqM/VOc/Nw8Ob
    x52h6B6X1u0Q0qoT+Cgc/PMPwZCaCDoNm9eIzMkz3dc7WwgK1BmKA+yl3reKNpqLxdb3W19N439Tvx14acj24AvJGPAK4pvpbgVV
    EA4S0s0OhKRG7Y/LN7GBAUe88Mf+gUwy9xzak6GbPPL8LVv07aet4OuXHiK9KrSGBh7VrmZ0ZzvOQEiAgccuBAsZFXciCxodOFlf
    Bloyi1Sf0GoBAvu39X1oP5+qXfqDKyQGU7WwimmkyqA/r4XqrKVD/eWT+EVGnmc8G452yz6tSuz8jHp9WMvxHL7ewiR8NAqxchbo
    W66zqk3dcDk8Gq6YiF6VrV//jHAFWiBH5UAIJ2OsGykikYUCiKL8UxZnsH/VMZoJpxVG4m6aLS/f5dbisiCfH1kIpGorI4KyUS88
    2DwMDS2gblkSiO9EnWEcR4shIvrt10PGT3QpzsnPjai20zDkn3ubhShOdSTcmvisWKwF86wKrQ7z/ul9sApxjaOSi1stvo/9I1DD
    hYZlsdCM4vsW5MbgAEYWIYXWNj9XdLgqhpYVo+jRAXUD2/phrFaEFnDraIaWNtM6TE0dm/phhuDqvjWjaAY7HBvIJSwdG8dJA7q6
    /yJ0oHQCm4exdDZYV/4nx15q//9P8f+QIoTqDaylbVzI77UTfWz3gTsNPRs/IPZ45V+UIrAFdwkLjxj64qH0Pfl9LTd2ctwQB8H/
    eQJCSP2hssY5iG8DbFR7Dad9Tj3J9hbM8F8VHuQXF0wtjkf0MGYJAdn3offfkigCMpSH/mi3oUIpN2Nq6PgE/FccIaxfUy6iTZk2
    /QkfJHg+6Z5JeRxLsd++1TaKp8sH8r/T44f+qx2sVBJV8frcCZrs8GKmYgwWUyPuQsq8hgzyYdyvlUEA/z882aHzzYLxkYa5msuu
    KC4MTqQBiuqSrlUzPzKInFXbh7G4gv6B3YMI4pmac/M4pNvmz1ms7IlWVhie3+FnJxZcVewYDof4j6DRwcITrVIzD7ImWCaf203B
    R7Zes7hhVlD+XAMFf3P929JOzzA05H8cvyFDmEKKv5/Ed++41na+PW1tXV1e4nKvlw9XRHRccq/vBQddCq3tdfx2zxDSW370VUQJ
    cG7J3sR5+p73JOgxZfj0MvsV2gJNwgpNmkLEnfbM6H5XnfnvjYX9V2lH3xTgrCKS5cjmilCUvETueXHb42VxcOHwddq+Mw0CaZ0Q
    RqKUBYHdi7DVsfB/+ZzZmyNdAxuitl1HD4ycBpi7P+BoIN2eM3m5VipbvtY86frd6/DBIuYzJl5dOVBxgWHX2WIvYRER/wf7PDQD
    bAyHpA5HxQn6ymvT+8pzRyhTD/leISv3ffxLGbIIbU4cgNlOZZIYaOsw+L9BkkZkaQ6is1H1UAmNrGTmpgUVA/D67+uTahlyjRoJ
    L7I/PtegGUlMqRcSoAp/hieZMcUrbfxTZNj/CNBAys8P46E7lPnqBL8EBgwa2uMD7/qK3gQZjqUyaSTtawnPlBiMpW8SDCnMnqOK
    QUDoP6Ev9818IhDj+vjIGYVhDxP03sgydXnVudT8HOKmPEYxZ/JerURkbrMLi7d7rGAQwP7Hig5C6UyrJFus/QDMX02Jq9HXFHSe
    NGQcAwM8phpwS6fJBIC5Q2gw5n9fiW1ZG5THUkDIODWN9T79hU5jwEdgJYLAgfHwh8JrT8Sz9AxD6/N6vTjIA8qs2x59Of+D0Ypm
    wVCyE6D1QgLNo5xoqhDAtxAuaJw8btXWN1797pRaU6s7cnBxBTWRh8CKLbULA2XYZmltC911qne6RVe/TEuFzEKpgex6/bgr/QEh
    +OAUgCgDDgNDDCVEc/pWAtEFBcWnRQBYNeB5dGjOHgcjSPLDgB8GlxIL6G95ESS6+g8erluFLUBMfI9bNdLLq4o66Gnc6RB4/Di9
    nMVzGUt3tyVOLtXs0rWTpp9iibWP1MdzLGCOlurl5GDivvm4BHw4AugawnGh5EdNI/YQK1LNM1LJl1eA18sCYSo0q5uWbaRwbEv1
    luSAiVSdoi0eQfX/ILOnc3b+BqNlnYj8DmigaYQpvDEMyFmJnD3hi0exBxY2mvidtWM3XHXbPw58617IBnCyhbNiCKkqgr1+uPcc
    d19firlF4IBDbhhyEYYNKwTh+mIWOT4tnuWdx6myOsuZo1HAWbFNjE/FePi8MQj4/HciGn68wquItbiFwTO7CKWMGOwwQCze7Mmv
    anPVYYvLwlCTNrBSjQWPd7hS/zH1kDWr3QxCE0jitP1U8WO52fP9Ud+fog8uoW9cK8G51mvwL6UtajtaUHkYtSuB8TQouDC5SdMt
    R3u40ZINvn6iRKvVewDZ+Z/YGKBGlW7BDKDY0pcKrEQxUFJ+gxFI5TQk82Tk4tYQBp9EcVa/AuBuCQdhTeBCiSHon6e/9XKGZQTS
    iJGSqqNqLi1WqBo/3tz33vwZFK3yejgyi7qr4mozJYOGAAFYp+jDACtGrn8opszr7ShvLAyI0yGYkZo4+vw9jOnq7DiycxLO/XEM
    kLiey8lVgnjvoKTru2vXf7Wzocckej6w/6VBbQS78rgR7rp1v9aqYnV2QOd+9JWtctWDhz7Kw912nu1f325/aDg1Ae+7W1FvHpWg
    BruDK8mLq4zGtEADJ+wxuYQb4CD2naiYmGqK/fXsMKpusKlCzPSrv4GhCJFh3DZFK6C9LVE0OoE1vZ+Iy8SZGt3RbOKbtKTEDnF1
    CYJLVO7z+LYWQ3AmN3cxPSkSmyIpmvEBdrYRZu5d0rdeLHCyJO3XbBmzWcQeptceyH5DNNG/FZm6Mc1db4RfS8cf4d96Dzj0rIN4
    R1CCsT1Wn7jeMXU2VP6TM6ZxHTjYERBiIT/kYgyJzBEYnHtocTy6AotFYwpgtqCZMgBfknA30qLXgVb//rmO/CUHys2anq3syFh0
    JQA/zn5bKNk4zl4E8n/FpnWZLRaycSeobW4ur6+eh6o3B3Wa5EDoVYC1fwn7nN/3PHG4b/oeyuBVwP1bmysUJfKwz/Dgsv51zMF/
    ZDgxh9s7kPRtTseA//irU+u/K7tVFsK4awIpTr6A5JjQXumpAuLvmPZjkAyD0xMDYjELT/BAy37XI+96tJUkuaCMrmQc7YBlcdM7
    juojHRI7uCK2RCdiughl6tHB7g7iGULZDk8yPboKzmIq1vWjYh1qsSeyK3ksLpRL5B1tFbgu/fHViexSrkWvBj5Fgl8xRsM3x1mz
    O4ivE+hKZ+jcYQ/VgVcn57GZtf3adf18ReT1SUdT8CvM9Ju6DFWlAF3IWR9o+5iUde2qvttS9OJUCEN7JBxCctJtbQE0FP85NKts
    I5IdYLR8vKss5IDDYMm8Z8YqyC/deyP4o+kkPwQbok8knmdLbGzKiZNUDWQgIbP2qCczafWnKcwLCwnrGU//PveH4/J6DhP2NFGX
    SzY6IAQF0NeqjEj0+d8doplzqpsoWDTiKca93eRsf5B9XQl+HjkiV+zTgpLH42iICJnWf1Alh4W+pWTxIrjzGFubY3wzGQ92pESg
    hwdkSjvXDW1eP3T+eDuq+tXkjvY0R3cc2KVrKXHUV3BOI+f3JH8t8EMEG66o6HPljcM9pG+q6uxKydWspoJ3ZNXRPjqO5Dnt69V8
    y/WL1fCrKFam0fu6UxfcXJQYaoVEJ/JxN666rwNUj6NlGmIlcu15PSgzPz5D8HXY+3HR67E1644jj44Uxx5xcSUnigINE8HSpNep
    /NNvPzWkvwwZAP7NyRA7avn37WTvq8piB1fxv86uAVIFpxwj1Q1QCUKRaLnAml9VmJm+fp8JhirDmhTcZcEPrYAKgLFdL0nJiIxY
    cFgAd/qF57Mp+Da+/ZrTqiugjuiceUFbIasRH2UFYq9MJDWKRjb/z34iG2kCHIh5tpy/RKlo6FXmL/yG7HseYv2fiHy2r7sn3dwp
    aNYIoVxyzCHB6JlN+NDPhUQ1z15jgnUBYFMBpvEvCEz+L3hrQxi/pyqqzvJu98EIPQ3mWKkqSxdZXSfjPt8UvyaMIFGHUF7qnc3T
    Vwl4rEEgEGZEXgUVPAXv9vGN7fYtHZPXy/UPPY+46cFwaNiTQPH1fClVCJ6QmkHMjmGG/3j86VDSGpVHCuGJqqk2rU4Qc9wtX4u/
    W+0ueOP/se/v97TOZ9i2e6CZlZnu6OOkG+pgW3RamHrA+7FoPHS+/XvT+/qn0jSGVwPv2ubm5ZsfigPj4FGiU6JWfoF4Z5TjU6z9
    68dTVylXnuL1/SK2b65IArb44pvAKPV9Tsn1rd7Xbujs13Vvhofg1Mwj+bertg8Dx5RN9XL9lptHUbSN7WN3G8fz/9z0+03J5rji
    YpGsBpJpg/sGaf5q781DeK3L07b/I17PbZa+83LnOa482ZUBOPRptZzN4BvxonOMDdLg19c58qfFhMpyJ6lLA3+U68+FZU76Z20Y
    3p3uFvwyK7eAIwomb9CMEsj7wusxNLzE5kq9ErRbzNMW5f9F9OPZWfDJaq+0qqPugsOzw8upzX/9l8pZx9UfCjYItg2ARqptfvT/
    dBr5sNFM+wMKq4eL8YI5AAqlS3XW3YzPbHYtIUPF+IuV09sH8hvXV+W1dz9Tr7QVfvPw6rLFBIIRTHtCOAlLZoD1A6NI5I3T7+6n
    l0r9j060rMHBwnEXgdwrfAJGOhyZP/EjLoZMrMENnlWZyeLr5l6rk1d7AWaUUM5O6x6jHNfywm46j+OabXHIev2zH3d6vTeVtfz8
    +WbHLZrTqvb3qzVkj+WQbfJeLxORPNmGDmyFIgNd/YP3Y5w+uoT1o11+dHD24P9AJMeINHJ9euy2st5Zz9CpJeLD9fUZR4brxmwg
    0CPv5EQ4tuoCWWCzv2c/3FP4P6wTWvTcegrefAi+rAuAcf0E+6kWE4HPAIuDmhdnoxRElt2Ugyb3W6v3T2igzrB9z0nNy5Zq51ZJ
    Y45bBixiFjO4rljI8BSrJyi0e1iaNKpEeK9JFs6OsKvJa7Rzjpc/of9tw3UcFoLBEpUFYcFVbRaLZN7L8DlD32uy90uxt+MG9vKK
    8ElPnd0xM+2DHXiAL1qNpf9U4k5c4F6H4ixmpLJ2UCiIGRGcTBKPs9kiDHfyIJV2evpY7lOREPF9df/w6/Sfz2XBZdmw9Nyq9nFX
    fqYEhO1CwLGamScEYjB7KNnFeVX8zVW/9Bty93uK9112//JxskPHq+hhw0p/cVpNo57OjgWEzmYSgJqqaVUwzwdBQTviTN3nf4op
    pC8VaPAWIslDlEhXYMvjruBAvQ67Rmc/adyqFczmlZaR9bXenxU/RUqlhMxpQpRAQNyCLrXEZCYN29rVCG9EuPfRa885/5f+XHvb
    0uN0HIEUDnKnCL729kES6HAqKG4SskGi+ZOeRId2z78pHY8sBuZSzC1tsP9GWUVW7BkHtxhBdhVQoOE+8koGQ18li/bZ2tsDzllc
    oI/vM+SCrO7j6fOOzLLa7Hsk0Uu1/WfJsXk4PAFjfDY9JzlmuQiM60HLZRdRFAyaMoE57F56113rle+9zlaOXOb4wtjG4iq2lVZo
    UEXtWqFm1itaimJcmuvmEPa626mJ2ZnEv9nd3lxcHa4Msn+XppxPkzhIKZFdIUhAFeHB9+sRJPgFX9vm1bb3ytIdZTeET+hxovJS
    itnC8G8kdYv64yDuyxoUiNCPgmBjEFz1UkjExx1DNK5fP5wcXd3npXI7ZlBmvXznShCeX3rCwigHW4FX03PLey9g7X3bUpqL7zzv
    s0D34/FX752/z+gtWufx32++vSdpIAaSgsvXmvMS182QMShHOSu3nfzwODP09yn4HxYCtiXHbGWy3zUw6MDlVImKKEdd9ioXi7G9
    SjZabHdwAWOAa2cryvt98Z63K2OokH6hK515MSBKAwkJ/L3j8jD118ub/juE4LNsuCT7eqfBNdtjQgShz00/mq466XPsTspP0GqY
    ratlv6qZJ+7wbwcnG2pBsswZqPm4l/T18bsD43VoRVYJvkULCc+FP5PmF7vxS6ZxPeEZsZjkQmJiQGIp1eEvQe9rAnNLdXX42uNV
    797ToiiCUqrjIRQKOhwMsHuqH7ektEHgQQM/2eGicFCVanGUlZQXvClgkpzxDQywKRVn5RRHZUoVPAa33jx7BLdUN7/v+zgMeuBd
    aZy7+5lXWFcogLoNc6hqFldX8LOaVL1aeyztBIln/ef0dZ2NQa8LlMNDiupccXWx4r8kauUiXFOwxfKEMpHOAs7t+SvYlSYmOspF
    5YuKjLBnNOMOYMofS3OWOErgpk2e4jb359hH0k0ipaW1ujNL6r8+en0TEI3EEPbkkkxxAioiTdM5t+xfeWX5/N/+Vm49aFxHpms0
    ctc0Re3WyDkWlEqLAZDDCx3+pKImwpWLtBA+lblLNRq4na/Y+ca2xHc2LEaG//lr9Fg+ymnoyyakSJsG4OaBbJ/Ee/uezXCo7rky
    e6ADw95ZnwKcgFtmF7PpMlob6mrF8Lzaz+3YHf46dO912rQC0M4Gt4ubbzo4X+CTDNRxF9IKY36Abf39ax2edUaBmClmlMZLapsp
    w8HLWRSzo6H+9HWg6OffSzcXj0YJU2UM0l9WRUE1V75H9r+AH41rNuKqJkWc4AX+ZGeKmSz8K/gwrvP5PPu1Khuf7NmmK1+RPe4Y
    OpCMpO6tSW4Feq2wsj5rlCjdSs6eudXYyfi+94XRcp1oT2qTG6OKaBogZLtNO4MRa84F4nMTcr/8gD95nubm8xWB2Q0pVY/mhIrc
    ElFlXSrOdpkmXyiUwXLXdzi/nI7Wvwtr7f6IFp/L8PJ28V25djeGmVs2weCk9kahpNj6/OuuSFeZDa/DuUSpkMfiJUvPgouZdBR0
    7/Z9l++I6HbszbQciWuJV8su0rfsXIkm0AZufhP7fFTAM7pZCQRpJHSKl8UDpi65i57bdjjce+NNQs9Oso0d/93ahMoEY3RSk8rJ
    nIIajwkXct7pXKnJd7R636ngs8g2YMHqXEyFRMyKr6BoFLnM5iPC7crITHwhKs4WWmObfYWyShXOtvOHQr9MAcbPPRtNnxY6Z6MY
    PAwhRJTrV/DGkYIxx4GOc/6uP+/vAx78jOcTJdIEeI52atLunzGR15+3uT33sf7PhkQsgr1tnLj1xsm8zBhb0Ro7Rx/7ENBcK65k
    7SjlSAIjZRUj3Na2D4hBs/h+F6eKdHg+s6qE89fMyShKoGKWkvnx8cyKRGkcXd6nesUfFCHUKEKIMZMAG7nbW0W9J+ve25Iqbm9T
    0eNvirK8nlqK8wqtVi1LYXFfhgxadOXEUOBtNYohchUybBRMpFOv2EdtvRkTVaEyVNxWlQ9+SThkOWqgjZaAD8pBTsMc3CYVddKI
    KGmcYIcQdcOKiKA8csfoBlxkVIBwkzr7qCxbX5UKpBBksqM0c5geTGBEekaE0cPS0K1F9Vgni1Il6My4GCuZMuC4wmzX/vPvfu/X
    CLJdzO/Lx4nfIZqtxj/AH/vCBT9Pi040mou06nkISy1piUFWmXmkWgM5MbwZAnuGaQ9oU8grDddmS8uq5qsoHPL7xkp7VXwBp7qP
    w+canAkHp0dq4IFsNQFMfi871n7708lOdipPf2w8fdljSSl3pTkcy7vPKbxIZ9GFf7mH4e2zB2r0aSpjGyOkGyLI0QlxG+Oj/mTe
    0KSQQkAEA0bRbr2NN1n2mrBmDGFjkqKd/rL60Ncw9ML7Cq+T1yxFqw5giYnMo0VxJCUOzVmSvzSeF8KUBjdGHY7mpEpzVjigngCo
    a2hYb5fRojYbb1dMQBW47stGCkw2ECcdEWS1VdQskgtvJDbTXuf8jf+6JOu9ifZv4gRfl5PrYSkvKVhyI71R/HqoZcFtGffgpeHN
    AN96G8L0VEdRVlXMsvroScv5fgLo68lpO3mfL03+OUQtbKH+wgUHXERZlXzrR/Z/udFOxuriDNXw2ZuH1v/IFIRtOHBVYI83gWzi
    mMC5mkBzT87mApq1JEvjCaGiRkSkxOKqHO03FqIEMc9vvL9PTLBYz6+9AQoPYmQbSC3BgWN8kHiTOx4+4SeEvC0N36/9rOr/kYBN
    9DLEYe+53mkdWG+npsJgt8chzWdpZrAD+QOFHNcaLUuse511K3Le3cVUXCwyhgUpICIChg1P3Qj5Dk4wQ85IrQUnC8PEYLqm2xuJ
    gE89qv1qpLD35o9g+9GIBHY2IvgTJ+7747ZviI278wa1Y4zrlrRYIYyU/ANat/kdLlfvcSfRvOKcnCSSx1kyUVYp8ULVQPUXMBFF
    bDF6qf/9UP8LPqkIKgk893kqp/fttXcN3100+1tRoqZUirYrg5FHAa4FzsvVUa2YxCL1RT6ZtS11h6tFeor09ZWNHsgUwdDRUgOH
    VG5OrOL7CObnOVH3Y7T+65oLLiAOFzMRJ2x6HF6UypBvhHf4PyeAhQxLPmP6er3Wf6Tz97mMxctRIFwsYsRJiwDoYIej+OqlAcGr
    Jnr59Cdsy5rb2y2rKQaHEJrQrsHLKkpmDjN+rnCLouGiOTHGsnDWr1Zx6lFNvsLjqd6vViJ1nXH4Oy5qlM4HaPNIX4D0u15WrDZP
    Z+jqZTVRRSyqmlnKnhIRozZ0nXZ7wr9bzIBnCM/vJCLz+W7/8iyBugR0zYT9UgOpdd/lGx++UuLu0OlndZ0tf+3sJFLm6URHA3eX
    4i6ef6fU/cKYm9JG1LlDhYtUKv5oJUxffMUA/84itxgXZPOSZerJaMJfVa9ViK3B3RB3zH+d9q/fuK+QM8DBE0tVFbaOSMCODl+2
    +nXYNS5yQbTAnquJwJj/ct81wxF8npvV9zy87n7rfls1sIVJM+JO1U4htgQ3LrtVHtKah3je3uSp6jtysFkZxl4/hTJTjwVExlcP
    O9/MvwgbZoigk/8TzGhWsuvyrqOq7dIfHf7rs9qangtQr/8ix6A2893zEFWNzg9eHY7q80xuzkdsXNVzNRp5Z9Y/Ibuqv94vl4g9
    thWdRzOpuL+6fOPIIoUwrIBV5MykwITJDr7t16fcr420G3gm38sRasuigHFcxDmPQ+3V1vMNiTWj9j/fW481+8GFpj8/Wm0jFWwc
    V6XfAmWUhGBQOTOH6EpQ8VR9PunoZp/H0348HvKpzUcz8eL1FVWhJzPP8mI/PYi20t2UEq6HuNLmhxzmbF69q6y2DUqeb79X9z5P
    bz9baAXhbDLqlGiXghXLopGqpsnJ9H23r6yqra20gsm/3GsWdyd1YkSUz46h3Ct8FlKyKIQpwKynp5+2t4mzXcvSEfiYIdsv7j7e
    9GvzevlZ6z/aef0+KPh+YXCbzXeFH1vas3wOprybse5GnkROU8lS5ve7NYfIqihOXNSMxoP/WfgJNUOE7NRRQxQVjRR0NC39fELV
    63TI9DWXaDxllqG3TVbCM7sC52hNhF/ELLIxMdMU/GnbJM4mEcba2L0fx2eOjuDyLdzadjVqTQvqFzsJrRB3R5ig9/yb/ssZF/6P
    YYPbKvLcRqQsststCBgS6gUtdy67NVbl9ZquXCObyYlfRiZ6U0zXvE7Pbipc29uP9te9Z5XLuk9RGp8LM3qa+GUW36tLP+hdOX/i
    oumYzyXTaTW2Zu86wOwcKFe1AOxXxaQyH9wu0gWdDpgRVjyM22BdVSvutTIZQIhqBecH7oTfMzuujoo8B/G2cSFI/mZ4d7rleq73
    f4jrE0BMoA0xc8lT/OE/qo/X1RY7WxCpHGmUR5z3vXZeu9vb2eGFZl2YkVC2oEJSS8To1rHXl94QyfyFCZoww2BpwNoBEodRmAaZ
    N+Hq6JuTmKW+EAo4mZyzgTtMJlQLYaEDCwzAKvdZ91r4Jrb9NvV9XiqhrXIpXu6OTO4fnBVU89ZN8+CBoJzUfQWBttLcNnnQP4Wg
    LtGEGV6sd1t7GtqnMtRUpTYr2nfYXrbP+j96bz2FLrLg/Ag5fYUBTLH58oHbD6M1akmTyFmXMYxRXFRhUxdICDUhFcuabG8VQwvu
    7E0thVrTcxSAUqksM8E52f+kzGEvn7pEoKdN0U1zyP9hV0+5tzM3dv4uXYVWYBWTuHokTSrfzp37ORQvSDK7gy8/aZlqFrKDcITl
    cPXiHfF9bsqF9Ffb4uWR2fzpZYudkjpailDKgXSXLDe6UHU2fEzUYvkghm5rR0lV2FCCyKK3VMDmSpX5Nv0QHNlOtT+8CJCUG+Da
    VyrgfXP84z0r9+Wd4QCLJ2eb/rtWaxO3fnx/CLbSkoAir43KOx8bOxRAtm69dap8+O+c4M+6gBSPHdNvClyZJ2xZrpSoFK5h0zut
    +k8z+m/V7vjagc33oGSFyhMnj3KY/j6llbU/FBRlKEi1kPXwc2aYPgcoHOScFHPaI7E1M8iUTC/t0b8ecwm7lVKQfpKxN1eQoxXp
    H5T/tGoawcChXOAxSmOvjxMWOBs6XyAnD0CWoeQYCbFKynyd0svRgzeQIJPPxCwArwEjIhZ8ze2Fb0HQXu5vlO16kk4in4FUnJl6
    WwvjdrVEycGtRIvNr2GndP8O+2OhM2RVKZVCvvkzxL73xPBpu9IREcUcvlz1mxxciuG8my5Bo2WOXdUg2AEmx1gwZE8VIrsFTAIc
    9qLHJp7gTZRG90tYVgLmN/UeHc/pDsqzkVXza6cYJTLnumxam0yM6t74hl95DBxVsNYN3+/+tx0H4x//8vXqXWe76CpQB+ksaJdt
    fuzG//jKUfxcnfV3/C69w8pWbuWmy/VLyJXBsABFar/SDC4UGb8AMUCpU7+gSs4quuDErZHPrnMJIHifrPG+giyYLyUDW+iXaIkx
    WlPu86u4+4e7sfHl85/4OpjIOg24xo5mp8yjCva1jYCUyGrCp3ToA8vOFnMpemsjIFR+1tlM8Y/EBI4xvSPqcZsfinTqCnJrLRw/
    1WhsYqvZDF4/Xn4W0lUsMPDi0Y36NO9+nEEQOVw5+93Vd6aySVUnnCRmIiRyyLTmgqaL4vYWbeygoLKNGxrvvDNN57UnwiYoeXyA
    s75b59tvPUxI4acb76smRxNwwkCJCNPrKwJxpEm2ORA4agzHeGR55b7epO30nD24DfpS+ehVN98qSR6qmPLhkHixjBEO9Sn957us
    36GG/xsRFeIIiFYKXA6YlG0rbFDHJp6kRrGgsWcd0ldQZGMgjRidSEsTfhvTbe6CVoV/O+vx+rH9/Kox9bQ/N8jlAqeQwzDZUG4M
    /LfqovBFubFlqrSrB/XRxhCDHMPQ0Zrxx9UmJ1f3c2+EzPvA8RAliVG6Cnd4bdymd//cz0wXrJVpy0h4ZtNCU8bu4Dae0sJzsFYW
    jFRkunRu/lgKHtXlIno4jZpMrfEZRr7YjD4tUgk11AH9tDIRHVEBKfqW8XbTSz+P1YvK/y6JBGmRhUJf8XdYzpNLNyr63bUQfd2m
    6ft/vJ5ZDVJCO2/KYfbjQFCgaunq2swOjMcKr7Zsv5gN/5vyjeSsty8bKjQQl5z47rGl5bXTq3RkeW2T3Lue2AT/Jtd+d1Djok0n
    jxU6rOepyrxlXyDME3WvUkKg37KsC7CV7P5JMiJmnqgQlgW10n2Rx+2s/9e71nd1kNAHddKBV2tcFWoYJciVBh/e+1CbRkKEXD2h
    j5Y+EVE+ZCT0hXDv2e+p22CWtaDMFNtG3gUqnDdpeWCXRvH5b4r1AohetjC7IyWigH2g2akaP4Q4fXycrSg7LpofL54TsSvxLsnZ
    /fOOLtfzyC5idGNjY2BA5NexGEb+VNl5FMwMjhJ91vjcT2ObcVPKy2J0ZPv6tkb/E+wfTyPIPReY9lWZLyjQwS/djOOlj+RDtLOq
    xlW1oTTyulFLfhEGCNNH7SZZdT3S/z0yDGDaSpfQnLsIBxV5BAid1BHTKGQ9XeSiQPrJ5/G12Vzg7ee1zywiYTpUnag/iWxUQXmo
    AUI9GH4iWFUSZGIDXDlzuWudeFIJsJyWar4YPC077WvG/vPZsjN5mGda5fp8WoGq8pwXShmNJHLBDjFLCssqv6haW81GpXKvNP7z
    LzQ64HECrpGjwAQHQ7wEgf8OtfMYyqbdTQFkrG5Sp2yrVNOq3qZj2p2xptbLF+E6rKKMre9w9ae2+0yL7wgliBKJL+3bZMMvOUpO
    liGjaswytlSp0vl6lpK/LOB4+2YbxLzAAde323umt+Bej3vvFDdEKq0khngZfVkD9SzcY92bm2b5tnkGL1LmEgp6AHHQJqfpx/+V
    jSbp4EqzMcn/AV9AoL/EMkYPJh+DcU5UhV3WWeBk7vaDm/lg3zR0f+N3//xZz/mNA4en4nDHoUYJRa8zmfrb9t9pmS3UkBKElFmp
    VQVka1S19RlNDLMGBx+qqp/hTX/2god927mhOjTo88bqwcyypVauOYGmornAQY0SCMoSHcPAN5P1M05bvuLyXSpwTAIfrcZ2FMoC
    qYPxoNjr4rzzdqSJJAnXwQegjF7tbA5USEBRSUgBNaQcmipz0s39hect/fEbf2ZQoG+ClWmoEKMfDod1XTtnvK+D+DzPQwgioiKd
    bn9z1GxWiJSowajC29/11ee84Je/dvOqzXa7zm7i7nBS+yAwSk59qOu6NiZJTVdjMj8/IAYRYkQUxCi6PR36TYxtwjqVIIUoSHDm
    mXtYG40+sU41MrNJzKScllUUygYLezfGUvT3Jp0919+y9oKX/No/v+trczuTacTRNWRFMqkjlFQjJEJUoAIGk7KUcaQ2koFEhIqM
    OAMk0liNCeHVr/yeH/3hx6A5vGsp7aSQxpOkFLuIXUgHkgKWVBgNU7DUZCY6ND/4Pd89X8AIDLkGISB6iwZtqETSqtMoHvqw+wr7
    qBLVlFEnon4mpt46ZQRCBIlRJiWKzeb6gUc89IJXXfnMLEEvhS/H/aSTpTBMCwtz3W5XGUpSVZUKmI21iQoHT7fsP0qMtIM7j+Iv
    3/rp177+79bHSS3ddLBjY9KUEm2epD3HVgUhhNp7b8mmtqOR+30AEEGMSgRVFWk5ayEAI60AAIAASURBVFsT+ZsR26HvKYZVsMNZ
    +/bmiSmFyLqymgYzBZBlidXcN1Q22uvOD8cTIrApDhyd/s4b3nz1lx/xcz/9Xb0BhlOI5gEmQlUhFKEEgkDIRFVR68eVBEHRdWHq
    NdaDoihVmVALnv+cx+zeufx7f/A3dZUV2bKPTogBA2UhBsAIRN6SR5yk1s/33A887fLQIGUAGjUwORjEBg4AKKoIgiju/4DLIP8Y
    fO1ctxIlIwHRgkgBCqwAhJQJwYlPMykG6Utf/H1L8zAClrrr8ujBjNF4HUBZlkHLQXdOxtH7KA7OJWItOG5sahDsP4BX/NIfXHfb
    sWmTOdclkxw6utrrL5gQmtiQBtHKWZPneWzg60oNJ4ayDGgLWJZal1mCtnvjp/rS2MY9YPusnBrM+ucAA6HCFZecH5tJajhUdWpt
    agxzUKmj+BhjkiR104CNcqq04HVp3BTv+cDnnvNzf/KZq9UVOLRSRjIgA9ZWZkVJhEOkEEz0rJRYZQQPmxYu68YgFANp7DCkwtO/
    9/LX/9YvnnNaHicHU644NrFpsjQlUV+VRZ5QrNtk0Oj4sotOT1MkBtZAQtVOk4uATOsQbZhSC5sy9u3CxeeeXbi0HJW9bC7UMOR8
    FcXH1OQMoz6y+NxEleHeXfZNf3bl8jw0+JSDr2pwaoDQYGG+N9xY73QLa7msKgGYEoKrK4W6prFrq/HDHyt/4tm/ddOtQy89toOg
    iVfOe51GvbSqpcZkWVcihaCq0Vhxtjn7rJ3WQQFiMJ/YFbxnttp2z/kmwDZhnWJYQpFg53Iy3ysMUWacJYZGRmSErcEfnclgqomS
    R+179Bv0vvjVO17zW3/0+3/yhXGNQElkFoKSzJbxoEIC5thKErd3I7WFbpPZrJnU1XicUgwVHnq/uV9+6bMfcJ89FFZ6eex3bD0Z
    GpZuJ5tOhtYBsel3U+bwA9//ndaAjBAHZ9mC7CwXBClUASGFEGLC+IHv+U7Uo93z/elwpZe41HCRJb1Od7w5diZJmPx009Ho4gsW
    f/Hnn3nGLoQQWGsDWEoQESrkKXYsLna73SCqBJcYY5wqJuOqU3SDh4q97saDb3jjXx9dDeQWInJBouBIs5E0JSixioveanCWrLOU
    pUq8uWdvb/sG+NbC9vk6xYgBDCwuYDBXiG8sTOufxRoIcbYG3HpCU1RC403UovJWuJv3dt1yx/o//LcP/NO/fCDCRSKhNo8jBYRE
    ScAmRAWTEMgCFAFlNqrodYrMdTNruikQccn5c6959c/s2a0b69dHfyzLJMZJay4do48xNk1zwQXnXnG/HbXIxA89ylqnDLEQS96S
    EIEZhjgBqS8T4P6XLi33pB4e6Cc+To5pucnwG8P1uW6vmZa9PO3lstD3V77sJ77tiq4FrDS5zWI0igwMMiDF3Hzf2aSuaxHBlprC
    YDAPcF03giTPFw8enrp8x+ZUFbYtkLW+HkqAMiQ1VKhkKgnBQhrQqKwPnnnW4B7ugOM709tq7t982CasUwwRbRpYh+UdczPfG1ED
    akULTnzdltA4kUvTjkvyJjBcxxVz4wpH1yeRnLabbsqY2e8JqTJIo1jLrb0zIKoKxmTsh5veWpSTen1t0xEyp6ftxJvf9Mvf/V0P
    j3Ec/KQo0hh9kJjnnSzv+Ejf/ZTvqTyU2SVFRMdrrsiBFGqDsgASEUISxIroeFrv3Yu9p80lttEw6ndd5hCbusjddLK5PN/15cqu
    peQP3/Cyi8/mcjqpy43cJQwbAzkHEOAQIpyxMar3kZm3hAltWZbj8XhhYV5VyzokRT+aXE16gl+OHzGl2SStGANDCom1xAl0dNEl
    Z57q87+N/zlsE9Yphk3Iqwhh3+k7mYnABLa2bYYwK7OceIcnaOZsNZmGxoOprEPTECed/mCHwioY2n5LO4EKBowGg5gn7m5aA0XP
    ZR0HQpqn3W6XGao0rZBYXHnl9z39h35wWlWTclp0O0marm2O1jan3cGO8y/aNSpReniko8DCWR2TBlwreULDaBjecGDn0h4l6TTg
    Ed/xbZzENOMYm+l0nOWphfY7PNq4becy/dEbX3zmXkjwc4WbyzuGbFUGw4BB5aMyogLKMaqzuctSRQRgLVuGanQJBalX19eEqQ6e
    E4etHZrWeXBmmQMJwSuEWYiFjbKhbre49OLTAGx5rOk3Urza1sM69djuEp5iKEGt1hGnn3ma4GtRJagYNbP3klmrTqEMigAMVHxl
    nDBUyVJiYmym0ypJ8pkqlALKpAAxqZdQO5JenptWMpNcK/NSNiFL7XBz3Ot2neXR2Odd13iYBLfeiQMHDuV51xb9jfEkhKbIO4Ub
    rG4cefozXpWlZZb4M886fXPDQ+3puxbzzLkMvV5vYW5vt9tPM5vlOOe8ncs70s0x7veQK5o/e1eINKl1cXnPcGOcMGq/vjSQP/r9
    l+zbicl4tNBNDBhqJIJh2KKJGPnxfDrgNkgKZFPHLJWvGKmqFkVWxWp17XC322fqGGeDau2bZOZj2MocMrfifYAiMsGwQhoi8k0Y
    9PJ+91Sf/m38T2KbsE4lhFQQrbNN0F2n7QgShYTIxNgafrLAQC1av1SAVCTW8/2crN2cjOsSLk1cgqLTC56gTGBASGnLM0agZWpC
    N0tnu7sEAryqSTXC2zwSh7JCWrgIfOD91//t2//lwOHy6EYWTX86bgKQ5RmYjqys71zaORkeyLKlQ8cOH11dsbTQ7czdfvvImqaR
    SdRo+XoRhNCwbUwyqpsROO8Uu6fSi9RNB8WRlZGlNHdhcS55/etesmcnGj/puGjA5bDK0h6ANCUlVFKa9IRMO9QSTPAhxsiMJvp6
    OLIFGavGSYwYTYYxsZZTRG4HJoQDoCAlBFAggiEihCiNtVzX6sygaVAkbfJo/lfP4Tb+j2KbsL4BTnjTAJjZ2KnetcHN/2N1WWkJ
    Ayc/S+NL53JimpvrgWpjJYGrxWv7qq2UOM0kDZQ0cpjGEKoxQAuLe3yMG+NjvcIQaiIAEWBQKzDFLDYK2FmXZgqIghnEUKlDDKnJ
    hJJhSUTmfe+59i1vfdeBI0M1blxRXsyX3nbSjJwZbq6L4bm5JRUyPFgfVnODc0eblct2rg9rIDM22bK6yYWMciAOdTB5d6n2MiqN
    Sfu+4qqCS7sJpNOJL3/5c847E5tl2XOhSPPJcNRJF2Bmv7gXJIlV8FTKjHM2zhmNKtGT5cwmma/KvJdG9UVRjIbT1OWdvDPxI2MA
    SmUrOGUItkSSg3gxBmBRLpJ+9E0377ZpsgJb4nx80mm6O/4Hz/I2/kOxTVhfDzO9TIUSiBFJBRohCqHWWlDVkSOy0SMGpCmIZ3HB
    8atbVFrNYm0lMSHA7GOCJE6bMCrs/J7luRg20mS+ClkkIhiBwtTQwOpIidQKhSkHlxABmS1Kr1Gsc11QsDQyGkNkpYQSkkhoQJyQ
    dptmWoujFFUDrjeL3LDUKWUMW1X2bf/4sbe9/QOrG8puUDYLYFKwD8rM0Xttmq4pjDHNuIxMjjrGpNOxGi58M3WJAaAQgxQk0EiI
    RCqi1g7KcRPVWJdBOXEmxKAa69C4pDhw50p4wFI/z2NjFa7TdWES2RpieI/IkIQE0bALgBhiTGzMQyxETMy4Nj4zQXxTT01OfROs
    hqZnI8y0Vm/zfl372IT5zqAaTXxd9gdFKY2yDZT7mG5OJaV8x3y364C76Avq1oljAgF0cok3EgCY7SrWKcU2YX1DCJhoS068TaoU
    aMQbhiXXXtPOwTJUEAVCwgqltqlHhpmII5TACnDrngAmgMEGCRsEoFdg5+L8qFZWGJOEdigBAQTEmaoBQS0Hy7Dc+KkPkdJsYFzK
    UKPCpIZsoxRDzZQkibOQ1GVMjZeqEvQyNJWtxRrbWV2Tv/iLf/rQJ762PjHTuj+NRpHknU6EhBBElMgTjKpEL00V8yIlIkhUUJSo
    GhmiTInLoohEr6pQNaTE7S2fgkmCVE1QnSisYeccN9qsbJSv+/0//tM/rR58v3N/+Pu+84GXL2vkrIO6AgmSDIYxDTEipDYVwDjt
    FrYOVmJCxownQ7JNExvDxrrExpTEiAqiFQoEMxqNkiw1xkwnTWKy/iAbTtZsTp58pEI4JRZE7uZs6W7ToAKA27S6bbTODjt0e2j0
    mwPbhPX1oKTEqlAiA4DQZi3GgBO2ikhQFR+axpmE2cYYjINhJcy20qCAWGLLotgS750VhAHAEBlS1BWyFGefefEXv3IY0YINwwqH
    rWxFQAC8Qb1A2kt4ftA5dPDoJNjM5I3Q5ua0lxXGOk6LxvtKyiQlJ1qVwxSjGI4Cu6qA1EKy4pobJ+98xwff8c7PDUdpVpweOSv9
    lBMC1aUfNz7SFpSEmExB1lBZr0NbTwgBq7BYArPdKMdM1hqy1jJDRUIMMUYbaoA5SSwRwFGEKJAlUJhoxcIaeu//5B3vfM+vX3Hp
    eT/w1CdcdOHpZ+y2gx4mNWKsOrkhyibl1Oa9xq9trK55TeGSotdtpmWaoaobmIyBhoJCAAda1BiJNLMKhWjIit5kY7Op67STRFN6
    jZErso6gEuq5hdxw62W/FUjRcfdDEJ2w+TopjWfdJq5Tim3C+gZQkq2ilWEgKgwAothEUZ8k1rJzqYIYCmPaatNMet2yASCiMTRM
    FhDamiM5nmsQtQwAA9zn0ss++4U7jOEo2mYjM7N1av1agtGg0+n3P/Upz37WZZ/4xJG/fMu7v/DVm5AOdi70pWnGo3G0SLrdjMl7
    H4US15Fm3MkW5gZn5hmuvw2//Tv/9WMf+2J/cIZxuyrLgXIVjYZSB5E6xirJLFErPBUUnjiQNipVkQdDgThAvSKINqQCoD9IY4ze
    e4kgMZaT1Fl1SfCsSAkJNBGkBBOFQ0VC2u1268qPp2HQ39HPlq+5ZfUrr/2zlMonPOYhP/i9T7jvpYMEWV2phfbyZBTi/GBufjAX
    sTD2tq5HMZSgpN1rihpjbF1oE5CDmCIxG8NVmxBb0/iYZHmW5nVYU6sSPMizMRqjwi8vDYBWA94CAAmOMxdkpp9//DTNTlabzm9z
    1inDNmF9A2wVOLQtXvBWucNaK2LEB7UQUYnBWsvWBA9pTfJUjSVmsBGJkRCA436cW6+iDDjApgmC4j73uYjpHc5xiFFVWLdcstpB
    IRKopGJX7rg9p8se/x0773e/Z1576+a73/+Zd73rg7HhXmeeMldKDHUjgSS6qMmgMzcdrX3q45s/e/NnvnrDtZNKF0972HAyBsdx
    spmlpTOqTVn6wBK73bT2Q9U6hopQpUmTugpxMzarix3b68SFgVteLJbmi/m5vNfNsxR1teYSSpOuIB1P49Fjk4OHh+sbcvDwZFwm
    46mrmoy4a12PKNPYacqiXPUutS5LR+Oxj411nKT9Kmb/+rEvv/9jnz9tuf+ERz3kad/56NN388YmXNcQsuFwEjGkdKAcO91cgncu
    hTplgoGxQY2IVmjC5nq9PD/XaDmuJlHIkG3GjcnIsEZ4kJBEFZDo7l07EE82SD3+FoFZXngcx6WRW6G/bZw6bBPW10NrnQwiAlRP
    qHqToKob56yAEY1xEGDSwFg4BwBgBI8Y4RgMlpnFjeBuFRMSqCh54jw02HcGJUkMqCDKsxxyVvDauqsssRmOJu3Iwvw8P3i+f7/7
    Pe6Zz3jce9/1yY9/8iuf/sK16gZzO3b7mmlqFjqL6xsrSb586HB924GbPFhtp9qMrugEbjqpI5TTybGC48J8HqqJr9ZTjPo96vVc
    nsSFOXPWGQsXnXfent357uXUmWliJgalxilCBZkApTGVwotOVB1MV7kPPc1rP6B/bC3efmB06+2bBw4ODx87duTIcHUdO+fvs7ou
    TJYobWLlbJJ3BnXlk3QwHk8SJHcc0T/6y3f99d+861GPeNATnvCY8y9cHpfc6S9PqtQkWYheQZPJZlHkEok0TQwLNSFOg06N+h0L
    vclwv6R6xr6zDx6snCtiIKEYFERkiKHQCEN2767dOJ6i37VGpXf5WGZdF20/v01YpxLbhPUNQMRtoZ2Op3MKMFTFJEyKzWnIrVWL
    OsAyJrXmKRlCdKSiAmIgqrFb9uk8875qC/iNj+smMUHrSP1ul/t9PbI2VEpAbdGXWbnNMhU2ghtj1yuNKRg+xEkVYpYunbaLfvw/
    PfTZz3roBz6x+qd/9Y9fvvFWDT3TDI5Np92l3tpoJS/SPLFcN0ICY8tyDBJGnaV+YNTRukxu5bC2qzN+9INPv/ziHZdcev7iQub9
    MDSrIgeJRk5ro7UTT2iMCWQig4ijAamqCAQGuq5wolnUBKY7v5SevZA9/JJUsCPI7mnpR+P8r9/6mRtumqyPxNhFm8579GNjQnA+
    EFNfyXqG5v1hU/7zB69/14dvWBjM7TvtnCoWcNnU67QhG4p+ZyE0Xr1Y5NZCMIm6mWWbg0EMG0d2Lnaf9bPPuf9D9j3zmW8+eHB9
    eWnXoaOrmQHUGWQSmWPIXLpneSdzuyAd7rbycVJge7x61Z41MnDb24WnENuE9fXASkKOttQIgBNvxXkv817f96GPvecDH4lU7Nh9
    5vzy7t27d515xvLCfGcwYGYIiBVMEELk2ZoIbSUVDBAlnoyFi1CyZBw6/U5cGVqXhq0Clra1fmVAInFJ5thkMgnoWkpNIlpa1O3W
    oIF7yIMWL77vT12/H3/3t//2mQ9c35nrr/p1OyBVvzpcS60r8kTj2PCUzdTQhsW6s6sXntN7yAPPP//sxdN3YSE50HFj0S80o5FB
    lWfRuggtoTUoQj1i1KggImYG+xAAQImZiQ3UKsPCiBgyiRoXicGOKQmpVnnnpS+8X1nPHVurv3L94c998cANt9w5KgdpsmcaBqJd
    kU7tLZCTycmqhxw4Wk780VIxrkZpOtft98aj0tTibMK+MRKckGpjqOwUfs8O98KX//wDHnBGRbjqy35zeMy5/traxtxgRxWOMAvD
    kooFpQaDASxDNAjF2fwacLxsuHWKIBAGZDbiwLK9znZKsU1Y3wDtuODdZV4IVe1dVjzk2x/5kc995cMfv7r50h0hJgJLEUQ6GAzm
    5vpz8/1LLrnwggvPMRZ5xoNeMT/n+l04ByYYwIB86DhkVWjARA550Te2JuOiRMAAltSqWszYT6ahsf2laYA2PqHYy+e9pxSGLMYT
    uA4M8JXrP//Fr34iunys5F3lIxLNO8WcbeDKypkx0Z0abjz3fHrgAxbve58LzzitUySR/AFtjuXpEeINA8mTCARQ1Ha0YtZOYBgD
    WIFRZVWBJSFhaISSBtIAUDs33iazbFgBBRtGbgrV2pmsyPjMM3tPftJFdx6ST3724Kc/f/MdB4vN8ZKXHXmy0yMZl40aKoqChIdN
    6aUu5jtlNRWJYiJMw9Ak0TApc2uJJ2edvfh9P/SkJz/pihyYejiLz336I6GeZulCNWUIk6Sd1I7HVQJm9qedtlTkiLFkYwzo+OCV
    bAmNTsuqyJO6iaJMQJrapvFpYiHbhaxTiW3C+sY4vtQCCCi2Eks25Qiowctf8TOjV735gx+6Sm0iIWXKCeaOw+WRtXFRxA998l96
    vbxpxp1OKrGUWCVWFxf6Z5+177zzz963e8cD7nPJ0hIrZcSIHrt2nbv/trg2rDh1AJO26nOz0W1QTAv7lWu/2M1+IEcnBZppVDFJ
    AQB5gbUGz33+r355/53WLEaV1JBLXOqAcWX8Zu4aG48V5siOxY0XvuhJne7B/mCtSA84UxtpiJSyyphV8ARohWgEON4+Y8AoLIi3
    BsRZACUGSQSRKlPbVN36LghOpFcKELSUMMmLnoAaT8G7fTsXv/+7dj35Cefdebjznn+76YMf+eywnC86Z7hOf+pNrGuFWpd4pbKu
    yib2etZlXmVSNxPnuMibs8/Y/dIXvejSy1B75FCpp3mSE/GXvnB1mqYxIssyX3vjnAQYBVOIftztFCBYwz40ABMpERMMt3qohG6e
    VaUmzhgDH1BPfVE479WYba46ldgmrG+Au3hn0uxWVMBHZQslrI/xohf92O0HJzfdeCgqYmQ2bmFxaVxOj66Ni+4eL4DpDyfeUN9Y
    lah3Hg13Hl75xKeOGo3wb+31C5OYbndhcfmcW245VtZJlhWNVCABLGbbvC1jhqYeD7o0nSJNQlUpq8sKxIDhBB/8xHW/82dvuXOz
    QbHPaxEllak1ivH6nV03LpL1QXLwQZcPnvzEC+5zURd6gzGrpJuQEZqaII7A1pSwig6A1tiGIFCZJUHUHoC26COESICqnNhq2YJs
    qUjJzDMVgAgxS9P4EhycyVLAwRBXlKx1cMvc+YOLL9jxQ99/2cc+ffRjn7zm1v2WZSnSYjRF0ySsRZb2VKaT6dBwNd91ZKcPvO+F
    P/oD33mfixc6Fgkh+GkI0bA68KjBDdfe7tLzovJsI1xIgqTOJSaIlsu7+xES1DubbtXRGcLQmQzhZLPq9jJfq1c1lghMAhIls+2m
    eiqxTVj/Q1Bq14ZZaVbscNYMq3GedU0X0xpv+uPn/szz3vSlq291eT6d4sDBY73BnLGdJJ1fWx/2+3NKQTR4H6HeAIZhiZXUGN0Y
    +YjpodXR7YdvU02rGmCB4X+nZSIE6XdTPxqxwMAGRV5gbRPTGm/5u3/9l3/7+MENCcnAmX5TajfPCiuT8sjSwrifH7v8AvquJ9z/
    4nPJxVtjOBaq1YTFWVhWotBmcxLYci5kt0ZVaWs7pe2UbU1jUARUKSrNvg5bKSuobRRAFCAxW9/Vln6UkBVW1Ve154jEJM5wDCvB
    V8puUpulwb4f+YHLHv+Y8z/5yaMf+vCtX7vh9lG9lCX7EnCsbDdJ0rxownA6Xn3I/S57+St+aJChsGBp6rrqdPpQIHpR3HGbqGY+
    gAxPy81OpxPFiCgZ9bEBheU9OyNxHUzQwDBERG1O2JYYCd25DEBETBNLFibydDwqej3ZpqtTim3C+ga4a0eoHVUnAGVTZUmmiujh
    FPD4pZc887+87PVfvnZ15+4LEzMSoSLvj0clkR1PKmstiNkYJkukgPgQNVDu+pWfcpITY1xLYo3XYNpiGQlUSGk2uqVgUj+eZpys
    r5a9nXnawbERKsJ/ft4rDx7ejNRb3Hne6ng6Xt3s567HTbl2y6C489suHDztuy+/6ELi+OUiOdJNp5vrh+f6fZKExKlGoCKqyVSs
    sJFZpM3j2uCoDZfagpTOxJeh2PqardCqNfBiNaRMylZBSqw4ESGCgwlV0sDAWVKKiFEgiUWaJCGEhUFe+vW1tY8nWP7Ox575iAef
    ddsd2R/+16uPrbmN0XRc94IZTDcMJZzY+S9ffccLX/gXf/IHP66A987CaIPpJHTmXe1x+53rWWdpvYzWReEaZL2voxIpREK/yHfu
    Pt0lIOS0NTCigEaogtsqlaKqyrxIYGVtY3V+rp85M6030qTb6mds45Rg+9B/XZC2HuYMhpKCoAYwQlokvfY2DBqMsc4h34XXXPn8
    n3/Zn37uC1ddctllB48eC5XarCg6vcoHADFqCF5VmbSVEibrKjFTsVYiQ7xE1ZpTmySm8RWOhy8nRrHJ2cJoeexouW9PPm7wZ295
    /7s+8ImDxyZZZ/do2GweXe/PF5DNpU6cHPjqkx55zqMe9YBLL83mBxvl9DrLt0m9Uvp6ftCppqVK644amcUwM5MFQQKItya7cVL5
    DLNVO4WQ0NZsGM0+2frrMKmYmX7eLL3CSWk1KWLdUGKMsWwdEMQ3VR1EmyxLhxurCtPLu0FDOT2SmM4l513wm7/8uI99cv1Dn9j/
    tZsOe3NmwKLnflWHctrcfvvKT/zk7/76q154wZlk1awd08VFq0BQHFkZRbEKRK3zLpXNhiC1Nlc2qlFIr7v5yLvey9EP9562ZB2l
    SdYpiqKgIkOawjEUcC4HoQai7Y+j8yF00znR7S7hqcQ2YX0DCOJsFpoY2NovU6orH0JjGINOHmOEV0fm3H30+tf+1E/+7G+ur381
    z1LYzEscj4fGFaKGybjUERmIhiDio4+1TY3JxGVErDZSWY6YbTWus9wAIHVQZm2r3AwkPsROvmtj5A4cwWtf9/fv+8hn0sWdkneH
    lVnctWdtbTXUh3fPH038V3/62Zc+7hFz/cHBEA+Yuu67icY6T3rEneH6KEmLaBTsiUhBXh1iXmpjHYDmOE9tieG1YBYAYo8fm5OP
    E0Cz1RbGXeKyWZTa5oaF1RCqWFIEW3bGZTZhIXjfZFknRKlDdFZ7C3kTwubGV3qdlad999kPeMjev//nr3zks186eGzJJBdAu2me
    C8kNN9753Odd+Ru/8pJzT+8tLdGkgXWAwf7b7qiCV5tEU5OplSXJuoayRlhZh6PyXe/61PvQNNOxV2knVhhErG1syyRnnXlaWY7r
    anjWmXt371y47fabH/HIhz71ux+5mJ3qK/L/39gmrG+I41uw1C6/tigyBzgAzXSSpA6GNFRFku/eiTe+7iXP+qkXTMpp1l2umqbX
    HdRRIKqQIAqBqmG1icuTjMvo2cGHOoQmtSmATidtElFpoFbVAXar0cYMEMxoPHnHez999W9+drPUwfI5Y6GqLvMMo81be/lGavbv
    W159+QuecvZgncNXVQ8nydRAiRWk1bQSaKfT9xKBoAhRYYiIiSgz5CJ8azbDyq0wBauclNa1usNb69skfJeajshsyHLWH1SSWVKJ
    rUArqIXh1AIcPJqm8TDGGJCNAcQmT2yMfjxescSDfmpwYHXztn07zn/Bcy579KPdX7zlS1dfexUl+xoZjOus099xZLj+opf8+stf
    8twnPeY0VYiAHW697SDBAiDSaTnqFN1mWgZxMaZJWrDjjdGGBaVuMbSDoKQkHtKuNigoXH39Wq/XLafprQduWpxLE1v/7AseUORb
    LVDdnh49NdgmrK8LJSbbftB+4sRVunWjJlmn/ScZpwCH+tJz0itf/PzX/Pqf15XNbHc4nOaDbkTjY8PklEyRFOJ5PJp2u10JNRPy
    pKPc3Vyv5hd2VtNjQk2S5sPNJrVda9OmHvd7nbqGeI9QpkXyoU9eZZOU8s64RtWEIoORg9betHvp0Pc/9ZwnPOIsF77U1KOleZ5W
    pVIToe1cAaeWgSCBgBMNelWotPP7BmxOCqkw4xnZOh5tNR3HN+/iPUndgbaOUCstf9LXkEIAiTqzG+NWUSseP6qqYEJiiVQ11kGa
    hW4nTG/XcOz+Z59+1ovPffdHDv3Du79w25E9SM8fhiTJdq6UK7/663+4duj7nvH0+zceZHDtl250ZjGSCVIbY6IEldqR12CsTUA2
    wpGztXHiC1JmKtvC4lbNLiHXPbxZ9ou9sUolxjPO2nX6riSDWOLjG1p3vVJOfDzbwGoP+HaR/l7Fdj7+P4e2XYi7GUFtWUOBUKTi
    6/jIh577spc+LzFNqEbdPC8nU2Z2zjlnQTKeDqP6brcI2nQ7aZa44fp4PKwWBgvVZKoyLVKpy43dS/OsnkX63XT/rdcZjmykO9db
    H42z3lKj6eaogsZeFuP0tpRvOuf0lZf83AO+97vmU/rKYn8lz8Yrq3eAPCgqycmPtopPevIDW8980uOe78yT/nC+h8fW57fc3k9+
    wROf3Ho9uccHqcy0E1mNC0laOnMU/muLxc3f/dj+r7/yiQ+4XFB/tZNOJJSWO+QWXv+Hf/u6P/zoNGA0RZoO6io2ZSA11uQqbJwV
    0jQzITSTSZmmedHpN3H2h8jWSRQSYUTiaZAk79eeuv2Fsiyf8Z9+aLFnfVgn6D1fFcf/nuMKW9sh2H8AtiOsrwtSmeUAX++rTo4g
    lE1mjUnwxCfuObz+Xa//g79JnMvU+IaiqnJInWnEg0uT2uloWk8aa5N+f47JqFY+rOzZwRsbBwskOqGBWxiN1iaT5uLzd47H62p4
    ZbTRX15Y25ykhjPLWh3rDqp+7/ZHfXv/R37kQd3ihmp0y3wXo+F6jt5gbrGJa6f6IP6vQ0hqkTpIYoztGgTvw1qqurdvXvpT93vH
    vx549/u+ulrv6BRnr2w0iuJt7/vULSvre3afOQ7RpgWECNTE3MemKDprq6M84TQrTA2G8WWUqrFZO12mQgGIoHYJhzKbxdjANzD+
    vvc7+2EPXSbAmUQpgCz0biGVHF9Q1xMfMwEg3s4e70VsE9bXg5wQ/J5dggLh2fNxEb7j+QwAtpzUIVhr1zarp//wfWtf/cGf/FOS
    7TaaVUHrqkqcy/K0rus6gF1InAteRWRaTUkm+06b+8Hvf8QjH37R0TuGV332K1/72u1HjsrGuB6u3pZ1+1UlSWfuyMZmx3K3yzJd
    tXywZ44840euePR3zNfVF5e6m7EerRw+vLywnCdmY3U1z7+1cxKXJE2op35q2TnOSMSIT6Q8rVf9xPefd8bi4pv+5prNOuv1dtau
    f2BjdOyjV+3ddaQSmU1nNMy2yFw3erJJnhQpNEQpfV1nxi3N9zerTSWBMqEdDBVSMLFjo75xzvt65Uee/gyJENHM5goFwt0TwK3/
    6iz+aq+IqDPZ0m3ca9gmrK8P2pLAmvW/ZoKbW1wm/+4bfOOtZdGykweD7JnPeMjhQ6v/8s7PM2w36SNSU2qSsaqWzbToZKGqE1fU
    VbCWO3m6vn7oUx//4H/+oYvOXh58x4MeHiqowQc+vvrVG27/x3e+36V5JVnRcV03HW5cs2uwtmdh9YU//e1n7J0U+FI3P4jxhsF4
    z8451WZl49ZOr1DRb93RbFaEGBPnwBxCU8fSsS/ymLJON69xfu07H3af+13y1Bf98jv3r69TcZamTEn/zpVxgKYIxqaEJHobI9fB
    O9dtmkZimRecG5purpWrwXRzJVbYmXpfOzgmJgbJrENcf+ADLnjg/ZdYkTiSaMiInPwONbtITlYnvctfcKoP4f9t2D6g3xDEMNxK
    JoDoxHPrqfXvHlEcm+Drrkt8PbTAC5/35Cc8+n4UN7WedGxqotVoEpeLhhi9Mcb7mlhBoW4m0HjH/sNf/Px6x0Ea5CmKFN/+8MWH
    PvS+k2lZ176ajrupNtOb+907dy3t/5UrH3vhOcfm8+uTeMtSUcfxulNF8KPx2vySJVuCwqk+gP8bh15BTXBKibHOWiKN5Ktmfbh5
    R6dXJ3xbL/3azoWbfvd3nnbGvpV6fE2/01T1eDiZsjEuScgaVY0xIkq/UxSpFT92piIdjTfvTJ3vdhJABAxlhYEm0BSakjqjiOXE
    oXr2s34wsygSwEP8TOidZxYV7VXBxy8GggXs8WfCdj54L2ObsL4eeEu0r70wzV2Yie7xUeTpaHOzm3Trss7YctT5Aj/77Kc88TFX
    FEljJGQm45AkpstixAsp17XPsgSMaVkzZ6HJ/vLP/3E8mvXONye45tr651/ya2WwadYZ5FRvXL/QPfTwB/Gvv+aJqft06r6s4bpB
    UQ6PHlqY22klq6b14nxvVK5WcQN3Dwe+lUCKOZenNcJm46eBiFyS2KKbdrrlZGST0sdrs+KLveITb/jtJzzsMm+m184VYdDrpGke
    QePpqI5TtnWWhdHwoJSrCTb7WeN0fXkxueLyCxbnB4pEkQocNIGkkIJiYWLqwCT1Ix963/tdmpOAFBLgGCTE6kjd1rM5/iAYViKl
    48/bzmD3OrZTwm+Ama2TArjnN0shHNcyZoUPvtctfBOKrCMamONw2py+p/i5n3ny/lv/8MZb1o2ZL6vALjdkkyRBoMQVIQTv/fzS
    cpgOJ3W84caD+28Ll19ma4/1zfCKK3+79mnRL+pqrW82s84dD39Q8qNPv2hxsD/RY8349oVOUo+rQb9bbm5ywkWabW5uuMwC95S1
    fuuAAa09CdLEJS6pOTZNILWpGoXkXRubodc6T8eI9ate/Lhf/92PfvqaW+D31bFnqAuDvLAS67I6Nt/LlxfzbjF/+X0ueNAD7/ew
    hy4ePYbv/cHfAXXbeToBMSwpszCjkTDasZD8yA8/eTLGQhfVZNzLuhpOmto/6bnVI20vldaRtR3WAG0HWPcytglrBiLEqKpqLQMI
    Qaxl3wTLRlW991meEjCdlEU3j1Fn67LHN2ZbpwqFcwaAcw4QBqB1r0BE2LHD/sEbn/tfXvpXX7jmYJF1q3LS6RbeC6Jl4hjVJel0
    6g0yJir9+F/e89HTz320S/ELv/jGzXHHi6nCRic/asurf+z7L/uORy/2O7doc4vgaGIhEWxzL2pSEg5Rg7OFxJZl3ak+tP+b54WY
    oYLg1RiAEsCJGnJu3FRkA9FmIpFIfUW/+DMPec3vff5TV98GPkeoo+w2J6N+Pz73J/6fJz/m7DjFrp1gAx+ggqu/JN67yNZlSYw1
    SBARghDIUOj24qWX7L70ksQpDOAsRJUdtWf5OFShCgUMgwgSEaMaQ2wAgm/UJrTNWfcitlPCGVTBTN779p9lWU4mZZJYIjVWNfq6
    qiTGPE+hEAnMIGq9vBAFqiACMUQgM8EVtMPlBGUEy7AJnveCp+/czcaOs0yk9k0VoE7hFAZqFVbUBbYNu39574e/dgue98K33Xxr
    xXYBFGO8w/C1P/p9pz/i/rJ38UBhb095xZkaQBAXyQU2kSVyUAIpcyw4FvQtvqkbWSILSIyKFdhINjIDgEQWBQmBqHa0lpvbuuba
    n3nW5Q+/r0lxW6yPdAtnnAXpv733XZMRzj4DgwwJ0M8hgr9929tgUyYb6qaqJ8aKscos3U5iTD0a3vbDT39iSkhYKj8yjsj6EKeg
    cDxqJQIbGAtrEWNUBRs4R5AYm6BBXELb01j3LrYJa4bNzRERsiwdDkfex16v0+nk2m7Akea9IskcWxJEhSrBB6lCrKN40aCIQFQI
    EJkj20gUqZ3qMRFJhFmfTJIcF15kX/eGlxS9yvvVxPJcd0FhFRZqAUPqlFgpBEZJyS+8/Pc/94UjUZeGG5tFMV3o3fHER/W+8zGD
    C88c5nwLxTuMji2ToUI0D8gC2cAQbqUdLElKkn9L5ySR0BjxJkRuQMFGTiInIlZr5ooRBCliH7HPYi2tpPb688+8/WeeddbuxdsW
    +qvjzYNZkpVjc8tNR/7iT986GqIuo4H3in99z9Vfue6aqI016iwSK5CqKjfqet371RCOPujB59/3irnN6VBQJo6ZoajJBUGIGoLU
    QeomTn0sg1RBKpuoog6xjFrBBnKREmklDk/1Ufy/CtuENUO/31NFWVb9fs85U9e+qpqmCZNaSs+TWho1AVxFN/UE68QyrIE1aigS
    KsHUY9SgETSKBlzDVeomkk58Mqpd3ukQY+yxsIzf+K1fnlvokGpVTo4rc5IylElJSCMhmPy2Q+uwPZu6bneT/bWPf8TcM37g4uW5
    FcbtdXm7+jUmD0DZCLEQZo+tv4iUv3UHGlooSWBprETWmRezgjUYjQYeUI0ppMOxw2BjSmcPafX5ue51V770sYhf6mYbvho6kyd2
    4eMf/9JrX/tXxKaJzhDe+74P28RkuZNYQkOWpjH4LHe9QeLj2vyifcELf8ICWWJFleA8UEpTBg+2xibWpdY55zJjEzaOjREVMBln
    jTVsmAw1vp5MR9/K7xffjNgmrBMQ0aLIiDCdViKSZYnAJHkSLHu2ZUSlkAQNYSIYNRi3D49pREMQB0pQATVmzzW4Ea4jvOCOw7j5
    AG66tRqO0enjad/zfV6a2pfKtbAnhFaYhVvbVEJZoze/lPd5de2Lc4Mbv+2+0x976tl7ukcSWg1xROyTJCEiERHxQo1yoxROej8X
    UA3y+BYnrciIBCHacl0WUAQ1gJASi1OxUAewcjRmHOtbzti1sWPxphc+/9tYr8nMai/rNmVGvOPfPnTNm9768Whw1Vfi/jvGUVHV
    Q2sRmorENnW0lomb2q8/4IEXnXWGaSJy24khqbwRzS0PnJ2vxVaKGmjAHhxgIkyEa9QKJQJXBZQhRhjj8qzofUtHuN+E+NYucNzr
    aOvuzjnnzB133Plf//ytN9029Jp2Op0QpKzrJMmiwPuY50VZ+8mkHE/Ksqxa62OQWEugwO2Is1qARI0qFZ3O4WOHd+/dUzZ+fb2c
    n98NUH9QjEINjSDDamdaLmqU2CVZ1Uya8YGlHes7l255xYu+x4Wv8Oh2k1XC3lBCbKKKKoiVKah6QJkaIIJA2srPfwvPNLTg2TGx
    Ao4EJgWJUOsqlrTqpiAP8oxI0EEn31i9ZceSvd99u9/1pL3v++CRo4fzNDutjnWeLvzeH789HZz77vd+fGXFWpeUdZPmBCHHmYEv
    J5XIxun7lp7+I09rKvQyjMbo5gkzJiUEpEx5sqUM1i54t6a5gLPcKCRA1bFxjcD74KtqbtDd3n++F7FNWDO0FXRjmIim06ppmt27
    d//0c372Ja/84+uvO6g0siap6yZNGzIOYNGRiEpE0BzaIWOMc8ZEwYRVgAySop2fJhCwMpzOLVxWVnYyLhcXB6PxsGnK3BmQB1mS
    lACjVsEqiUKUHZsm7Rxb3nXnL/3yo5x+aqer67IMWgcFU06qMUYmNRaktVIDBEJsB7VbAfiTvNe/JUHKRgAwSarKkULgCEApETio
    ZRCoJqqZPCOyZLHs5NaVzR3WNd//fY+4+bZrptUoxilbPrpZ7Vg677df/w+h7hp7mkmbQV+qcWXhNBR5mkbdNGlz2r5dTaM3XDfp
    JE6r+szT+/MLsAZiyFhInGlb6Jaectt+mdYQQergHBQoSxDb3qC7XcK6d7FNWDOowhj2PjhniyITATN2Lhe/dOULX/Pbb/7c578i
    MNbNe3W+jmASMMiwTQgM2CjiA+B94qwoVBIVC7WipKxKsEW/jCRRg9JoFPN8QCKpdbWPSqLsFRJn64mGFBzH1qwU2c0veO5DC3c9
    +5uN6xVJHBsVIiILIqIAUpXgfeUS/DshAdlSzmv/eTfy4pNmtI773GxlkNoKy/DdvvFkKb4Wx+WP75770F02l9rgUbdkS+/2Q086
    DXf/iVtRJwMkDCFSsgIymihIIEwB1IBElQU26cyvDjdcJ/TsOMT9z/h/Lv6t3/za2oYVXcryhdIn3iDp5Kpal1JXIbW5gsbjcVIk
    ZLj28ctfven5z38pN96PJ4tz82U5UZJIDJtNplUvTUmjMphhEqRpkqbOOmaYEIJVUmnqamNpqfuUJz/ue576+DzZLrvfm9gmrBlU
    AcBa207WAIhRGXT6brzu13/s+c//3RtvWZ16EcNqbS2BE6PEwQfWJDFGYQQ+dV0WV00rldjtd6OID8Em+bhprDohYSvM1qiXyneQ
    +rFy2vMUGzvlWBMS0sRK5lB17LHU3vzTP36fS07fSOSOpQWO4806+ghDRIqgCjBIRQDjOqLC2LLYmt3hANDOBJwgghMxF7UKfDNF
    ly1TLyYh5XanhGGhltRsidm36isiYG0l5wEj3KaxgIUaoOUgJW1AMwFlKHPLOIASz/RtSGbqCLNfm48TE9Ru6dK0A7mh1VsWYgUL
    WMERsy0EBRMlrfm0sillYvpZXY1TJ3N84PIz5n7+p+776l/5iOjDQuzXWtucG+yPXvJ8R2gSryQGpqg9yqi1GjOpcisuVcl6iyOJ
    tXERhtNeQC+kqEQsYtPUrsNs47HJUMfRWiM1BsWyH08pIjPpZRed9rSnPK5I2lbBdhnrXsM2Yc1ANOOpkz5DBrDic+ve8NoXvfaN
    //0d//bZbndwcGWzt7g4rkubpM6QRI6qJEpsoxqNnBd9NU2g6fp0DUyDZEfCwYcaKkSBENuxaCFLoFiW7KK1tXFqnDESTCidrpcb
    V/3Aj579wPvETnazleFotJn4Tt4dRJkaCB3/XamVVJ/x1L+PiVhPnnWXf3fvHCc4AKZ9tRMR2T2GBjQzyBCSVt+dlU9Y57SBG83a
    +TOpVhLZ0n3HXZzTjosZbClkKd3TzxVCANoBUjqpJNTSmXAr3YVW4HRqOBUPFxuildzdduYeeeRDe5/6/BF1A2UzLscmKdM896Vv
    ajGJAwJMYwwsWULKxllliK9iHWUKzshlkTqVd7GBAXFi2PmJH9d1Eyh1aSKwPkbXJBL8XKf3wMvP/+UrfyQxWtfjLO2d6kv7/yps
    E9YJ0PEy6klQNXkCa/G85z91WE7e/9Ev9bLFWHqjKYdESWJTRVVnU8NGoq9jKakp/THXkbkzzGi8OayOZibv2MQoIBEsQvBkSzIs
    vJBZE733kzhtNDKTWjTOHb70CnniY+d3LF+v4bYsL0OIxhWhImvFaj0zCKRWKfRknOS6CgaItPULk61YJmwla20SZ6FtZNQW+zWa
    Rim2Fl/tei9pW7xvuUxoNsG/ZT8BJfJAfZwKtVVRPjmaU4BDnE1v2NnvpnarcH3XXJWOu9i0/zz+3wBlc7dTtrXvCYAhQsEgWPYs
    ZCMMRzFHd+/pPPZJi1/42o0bQ+LsHIe+xJQ1DxFZJ1P4WeDXKCKrkgpqVbYp2dRmHeIYGh+qaQ6X5kXdhFpqMsFYSZAUSYfJ+Tp0
    OkabcacYX37Zmb/9Oz9iBMZQZrvfyptR34zYJqy7426hVuI4CFSROvzqq59e/Fb3gx+5dnNSd/L5oCwSWWKEEkdiG4JPUk6SBqYM
    8dh041jHNb2B8aNRRoklQ4YECJa9pWCMAXUrzThQGpKUO3nR76aDrssL//jHn7dz6VA9vblIxuq9MSZLi+GxYd6JoNjezqbdYVQ+
    IeCsLMRQVmIoUZuLkbT/a0u5oY2DaIs42g8MlCMBmggJq4AkEgDlGYNsVbhmwqEws2AqbLXNZItrWjISgGdq9CQg4bZlKdy+6uz4
    ntAgnf0I6HHOkrtW37Z0EnU26b7FUyfqaAaROVr1KpmBFfIhrOVpcdFF5z7hiaf9y78eGVcDm+ye+iDlBD44WzWxUsCBLDlrAAWs
    BkMNfBWbWDVskLLNEjZBxI+TohhNSgM1xBrgyxCDGg2aaGxWL3/AGa9/3TNIMJ1MFvqd/88odRv/q9gmrHvAyZylEYAkhk2OALzk
    RU9ZmN/9t3/7vtCUJMalxhZF3TSqKhpAIbW2PHbn5ZfNP+KhF2n9tYTvvPjcfkKVUSEyZBJYIylpIpKIIZ233YSiTSpnK+JACIgB
    SOZ6q+X0AHOdUd5UgRTRTpIiggLNHB0YameClsRQPu4YiDY+Oe5gqscTLgeSrVvopPd+CkCY6U3gxEu1VNVORpBagLjVOxaz9Zoi
    7DGb/2rFl0FtXnkXEeQ2xmMAM9u/46noieeTcXJcwgqQMivhuMhy6yQ2+0qZqZORaEvZ0UA6zDnEaKiCWenlCw950O6PfvTa8WbM
    jdfY5ElOHINGSoMoGZ+wuHYAJUK9iCbkMmeUY6TaQ5A4GBjjZRqtT5PM+yDRJNZaaJG7enLw/LPnfvPXfsoC48nm4qDr69qZZNsp
    +t7FNmHNcLfAimahgcbQQFWYvFCMrp7g+T/9wHJz+p73f2Zlo2HJs24/gOrGqyNmMFOWZdXmdCkrHv2oR9nwVdbrep1SZSwkQpmy
    EaNqfbQTQ4KpNRTZjZmn4BIqCEaFtU4cfMelvgmJKYh1c7razYsY0EZDrcJNW4o+XiBSzHIooVkPzojIyXUiPSmpmuWGWxV38gCD
    3BYR6MwFZ5b6bTUNt2IimsVxUDDN2CrOppKOu7/PUlc5qXSmypGUdMt7hmav0KobiBJIZfZ5andbuD07ZisPbn/AFqvNojMlA7Jg
    33hxJoFJiTybCnFE9sBlF+/8we/buTHcHWJ66PDIkEfwG5NRqY0Xo7ULnkONJvhGygZhfbxJVKhd4GQ+ml5oUpUMnAYJ7AyxiyEW
    SV4kblwekbr8tgec89u/8cw8xXBjdcfcIPoyTfPQBGO+tZfPv9mwTVgncE81LGEKxtrxpGabdVKyPdRTvOwl397r0Qc/evWNt6yG
    yjIcw7FCiI8ON/bt3nX7gVvf/Ob3YnLxk75jfn5+KcSjNl0TbkSyqIYF0tQmjkljnswjSgylSkXkiWCJmUxd+cxaCgbe2IzZxcaP
    aiVQDjXHu3ZKwMyY6zjaEtNWl9C0aaCFcjsRPitsA7Ma/ImsLbDCaMBWmNb2C4FWPwUzyuM4yyKVgVb6qX2FCARQ6x+B9tIi4Zlq
    BWnbHFSCKLOyELGygFmMELcSPawiHFkhEIZEoN06YpW2L0kQVqhuzUmAAaszYVgGKazxMYCisBeJ5NiYSuRgUzePe9TAJBvgKtQN
    EyWUQAdRNGqi0hE4UVIScC2kRTZ/8FBz9FhnOt1z4FB+9TXr1163cnRUuaJbR/KRY2BAqnrdYHjx+buufNkzFzogIO0UsQnOunoy
    TovOzK1oG/cStgnrG4A4TqbTXmcxCrMiczCEusbznvNI0emxY0cm9WaWLlt2ATpt6u7c/JG1UZEubnj6k7d84tYDg6f/0J7FpTmD
    NUJFVDGBwY4aMjWpBF8ZSizDGCJNSVQiIKGb5XVdKyNJkmm1ScFnRe59VDKA3aqgb80l3KVQfbKnKVrLHFJRYlJRMCm1cg4nRTRQ
    MKkFxKiQqhIJQETHvem3Xk2UGihE2/6jVeRC4JaqZhX9mXMOqZ11HGcOWkFIMOspcjvCPrOMbqtR7bMY0Mw7lgACM04sSjIiTgyX
    tXPmAWp1y66GLaKPYqLEWmJIrHHsYqwSXVta4rX1231dLw16vgphik7WDyKAi5xHGGFV1OAaylr29nb37uvfb23UO3THcP9NNx05
    psXCGaPGh6BZbm2akVSM8f3vf+bLX/qs03YAEU017nUKKE8nk6LTnU4nedE91Zfw/1XYJqwZ9P+j0mDIdQunSgyogADLngwmlbzg
    uU8899xz/uCP//HOIytBOv35XVXpG19p4ky+ePvBaq44791fuvOqOz/wipc8+qy9NlbXpthIuLIJGj9tZZHZGhUr4lQdS0JKrel7
    iLWxImhExDmrhOB9G36FSCRkTGJIva9VLRuSUBvLlo1qDNGrqrXWWvaqQkrqhVomatv/sMZFNSQc1SCykiFxIIV6AwUZIY5RhKgt
    sAPYyh9bVgrCokhEnYIFytQQBYKwWgKTWImOyREZiUIUjUWQUIXA1jADpFEiEInIGHjvVQlKRIbJMVsiZhWVqagHJLHWkJMAAMwW
    gI+h1ahiQ1GCKhSNRu8cSSwJqbUsEYhFggIw9XrVIYtEMS1TcZnpcKN5nqwdPZp3uvlcb3XtUF4wG2lqUnVBua75ox/44t/8/fW1
    PV+S/nBqnctZy1A1JpZp0lx0wc5fffVP7FyEAwx8VuStZViRdyAo8q5uh1f3KrYJ6+uDZ6nNXWZ/vIH0MzeJ9ZMef55Nf/i/vPQN
    vbyzcujOtNuHQ+mbw8dW9px+UeL8kZU4PTJ+0Sv++8te9Kgrzr8CfEeMB5vJWrfbafyk6HTKqoosBLCqIpCy6vH6jCgAb0CQv7GQ
    nDxD7qs6cQUbG2LjoxjD1hpAOeEYI3y0xmVJzhDvfTUO6hIlELPhSKyWZ8Olde2ZwMYkbMFOxQiUVZ1j9U0IIlAylsmqagzCDKXj
    iWFUAkgVQeEB0yqjM+xWRZ9FAkgVBAURE1IJALkiNT42sQ4CsYbYElQlaG67AiXlKAQxEkiFVGPiUgUTEaJCVEXQ9kWJEpsqSQjB
    +xrMiU3Y2uBtG76Rhpno/vEeqm7FbK2LNwWAV9fWF3cvSWyOHb5leddgZfVokfcJc7B7JuXyW/7uY//20fFYz/Lap3wemoW6tpG7
    GRKjzeTwL/6XF+1cRF3WWc6s+u96gttsdS9jm7C+LpSUHNAWaNDWoRlGSQW+qccuX3jMt5/xh2+88kU//1tn7b1wbTgdDjfnFgcJ
    wtE7b0lclmc70SSjsnjtb33th5524Xd+x0NsvH7H/Kge7k+Nq6sS2UIkEIJSJNMwogJQIkkwU/UTaAChHZ4s0pwkSmgsojEgIgk+
    hOCMMTCqNjYaKmVYZ7u9LJ80W/UsCoKg6kWDasiLFAAgEiVGUY0AIiSqCovMuotKECJm2rKAnsn+MshAoxDANZRa+waNbQ0+KgVO
    Gms5xsoHheQhpqRFknRCOSFQSokSGaXYqHiJEmMEWZuaxBom0QgiUWWrjTWGVKNqBIXUZdaCjE4nI9WoiIYoS5y1tqqq8XrZKxa3
    5iTaNoLorIlgZvMfxBFQDjBTgckHOw6NxqEc7di9c3047M+fNdx0nJ1784G593z4wNs/cLTEBdnShYeONK6Q1EoekyxBqA8v73Wv
    /r1XnbmPDNDNUtLZrAlUjkekuu1of29jm7C+AfR4z4viVt5IUKva9IuiDGNrFy65cPCmP37Ny176J8fqjUEnH64OB/MLuc3GU59Q
    J0TDYlc2Om9666237x/9xI/e/8j615a6ZPWQoXGpBIiyCDzYKwKpspKSkrQNJp4FeOpIoRqbpgFJklhwbCcqnEtdWoTG+GCYc2O7
    kLRs0ExIuSPMBqQMy1COhlVZJuOJmmBZhCJpVBZLbGwcN0M2ibFQVQ1RNTKFVqf4eFex7e4JDBAJNYhYGWowSw89IEFiE2sfAzgv
    On3ETjlN6pBq7CRJak0m0TQ1VGziOkmWxaAgKaP3dRWlJBKXUOqMNoEBkA9xHFE3qGJdhTjNs8Q6Yoj3tZ/Uhn3mkuXBIFQEkLJX
    8kAjHEGtL45TslAzW/FhAZpIXPkNb+L8jrm1jUmRLJZ+aeKXVtd2/clffulL19smubxqdpWjNEkLA1OPN+fTTjM+fNaZ+a+++mcu
    Ps8Q0NR1N01JzGx4Y2b6Rdts9R+BbcL6Bmhrv20Fm5hUFcpQ1HXM87xr06PDYwuD5d65eNlLfvR3f/8tX7llfXFp98baMEYt8s50
    WrKzbAYBDrb/7k/tv+GOf33ZCx+Z91PUq6ltDdLj1lYd0BqmQ0ENUwBqEgO1pBaSAszEeYfZqI9V1VQB0aSJptnKJDB12M7HMIhV
    T7VnqMdZQZxCyccYBXVQJTCpUiBEQohSR9QiU6HGsBBqrz1G7ShaG60LrDVipVKDhdAqJWwZyBKBAlO5dYQsoKqqYCUbyZFNKTF1
    cOOpDU2iugiz2CRzEb1mataONYcOTA4fnIyG4pvG2gQU2XibxO7ALS7nO3f3l3rZsjUJGpgGPIWbmKQEjyCjqZYUa0OS2MyyahPq
    Eh7RGQvyoBpcR1MLBQDaalcrZjP9ALTdsoyhOdLv9GKMWT6YTLpK+6pw3mtf986v3jyo9Zykc+H6ZsiKzsAm8E2aZqgO9tPN3/+9
    /7JjGdNJOd9J4azOaIoA0767KKnSNl3d+9gmrG+EE8PWogCICKSETt4fT8YmMTsGy9NanaVve2D+ypf/+C9c+YYj6ytGTeLyNM3h
    zOZ0aJjKUpayHeVUbrrt8K/+5gee9YxLr7jkQvD+RNYMTQADJLMBclKFENplPVEWjgwQSyqwm/WQUytRaw+bLBfd+Yj82FCtWVDM
    +aa3sYEDB8a33nro9ju+PByOJ5OJRIQQQhARyGwMVFqteudsXpj+IF/esbBz5/JgPrnv/R8s1Tri1NEkT6rUTCyNSUeQCuSBMPPa
    o3aGHswFNEItFC2bK9IQu0mxb1RKQEfM3ObIHTxcXnfD4etvuvbQuoynZjQM1dSI7zozl9i+NUXTDIkjqI66GXWkNGbrHard3SQh
    X3Tc/FJn996F085cPOusM3fs6DDGBlMjU6EqkdpyY00gRJEGJJEh3C5a29n6UHupk5K23t1kxApjkHeqiTduroqLwZ578NjuX37V
    Px1e2Zt07jMZdcsRzw8WLOt4tJI5aLNx+m5+1ZUvXx4gtzBiqrrspnloPNvZvJWC2hmOrdHYbWvCexPbhPUNoNpOezO1xSwozxZc
    XLfTh6IJUB+S1PngL7/IvfUvX/izP/eGa782SYvFg6sbpjD5gvNhPF9k1TgM7Dkynbt9f/HG37/xe39w31Oe+CBbfzHlQyoJKKhk
    Aii3QwCe0IDqE7uBmoDI9rpRDVwvzxYVO9c3+qvrydqGO3Bnff2NR66+9nNHVjbZ2KSTR5WmhjMdQzDGGGOttcY55oTA02mjFYUR
    NYeDb4TIJEmZpFXTvO+csxfuf5+zLrno9NOWpZdvZuaYM+swm6AxUy3cAG3SwyS5CXPQ2Lq3RwhgY+xF2TEd7rvzIB08TF/4ysGP
    feqq1c1qsLho3c49O0/fvTxYumhpx45dc4NFaxPfxLqu0zQjMqpa1+V4PN7Y2NjY2JiW64dWbi6lWtk0N67o5PNNXR/qdtaX5vML
    zttx2YV7Lz6/vziY5OZo7jYMb8SwmSS5MINFiQUplFkcKRsVQpwJWiiLGoVDTLUSS93xZLGx5372Gvdrv/v2qA+ozRlhulgL8jRG
    PTwZreYpW2d2zbsXveD7Lj6/yHNAYcTkWTIcrg8Gg5OH83UWLdeAAPnxktY2/vexTVhfF23V5i7budQOjkffGE6Y4Biu45qmts6X
    frLUmfuD1z3/F1/2lo9/dv+OpdM2pmNEiY1vVBMuypK6yR5Vd+cK/fU/HL7ppsMv/s+XGs5h18BjoAYJRQtAOSjVkRhApA50jmQ5
    oFM2Og3s7EJoBld/6ej73vexm28agZaapoiaeV0yZkd/0Dn9zNPOv/jcfafvethDHsAGzMzMBANAIqkaa1giJmMcuGPlq9dcf/XV
    11x/3c3HVg45t/umW+S2/bf+3ds+m5jNS85bePzjLnnwg86r42FjVg02jIwNGtZoIKRs1EaQUgzggNzLvI+7Qtz3lrd88jOfO3j0
    KDqD084850k/+aQnPew7HpYnSAnWwDKEAIEQHM30O6MCEXE24QUSBIJPEIHa4/DB4Wc+85WPffTTt968/8jR6drq+KrP3qnx8KAz
    edhDznjCY++zb+9FZXNn0JJ0EzoCSsAzBMoMUcST9O4F7XS+2ihZ4N0+XvieD6299b99uaHLp36n8KLJ5mKcktW6ngz6zJj4Zvyy
    l73qvhctdzOIYDKp+t1sPKoGg/mmqZzL2pfWLRGJu8hkbONeAkX5D1l1mklKEWrgmc/+nS/fIA3PR4pgRiOOLaKJLGJ8JIE6ApII
    QCMRACusJJFVKbDGHEc/975fKdobjhGggui2DKz+Y3e1SO9Jmu6eXFUpCmQaqsT2RiWe96I/uOrqA66700syrio2zpmEletq3Ctc
    Va0VeUyw/+EXlz/1ny/fufsw+OYiX4/1JDML5VSUDBKJbhphoQux3jsdnrE5mf/8dQdvvPXwrbfcOdyYVpMATZeX9u7dc+Z9Lr3v
    6aefuXPnzsW5+YXFjrUQgXFo2hFPAABvbcBEgaoykyokwhgwYzzG2nDthluv2xhO3vdvn/3y1dfnnX5Vj5rm4Fnn5ldc0X/Afecu
    v7DH/o6FpHJ+09Sl0WBztz6ubbFYytzE7zx4LH37f7/66mtWykn64Ic86vGPe8Jll102NzcwBBHECAsww5pWfUakLdkRRS9KbMiC
    YAggkEIAL2CDtj/JQAA2hvXKytFDh++8+uovfuYznz50+I4QK2bat3f3FZfvfOzDl3bvaPJiEvWooQ0fVlirTsdAgsaYJclkMmLm
    JE/KcpJ05sd+MC0v+9d3ZP/6b8PbjtnaLdScCXd9M3C2QFOxDBd6kfXob7zmRQ+6784uwd6NiO4qnLg1Z9uaKMHA/selhEJQadcQ
    hIhMO01LKhAPDko37cePPvOVHjsbSYLxYDoxzXu3S1iUDWkUQBxBRAxpLyvf9qcvPus0EBAFiYFoMASFor37lI6/jpwk0vgfeDtu
    E9a9ejhVEUrxZWVNmvzSr/zDRz9+w8qGWdi5b1yVZAAK0+mkSFLfhMJ1EI7s7Nw56Nz44pc85IyzhtPRNXt3zE9WbZ7sBvfHVbVR
    ra5uTg4eLm+4bnTN1dX+/T7tny6UWcOLi4tXXH7pwx7y4Esuumh+YCQiMYAi1LAGIgLVJDWRoIAqogopyLCZDSfMjlsQMcTtiFWt
    wsRlhEYcPor3f+jqj37yU7ffedNoeofyoUFn40FX7Pje77z/3rnQ5/HOgasnwwbsKVU7d3id3/nea9/9/mtB+x768O9+7nNe0O+b
    jBHUqzSZyRjG+5mCMKBBvIooiW1NZtol5y1yjQKNEIFzUEXjfQiNtZxlWbvV3QRxluuAq796ywc+9NFrr735yJFjG2u3dJLbH/WI
    s7/t2y7euydZWPBzvRp6rKwOJ7aROE2dA0AwzLapfYXBmM94298f/tB75w8d3O3mdzXWTWVaw1iz2IxDbiXjqQmHX33lc5/4mH3S
    aCchczeR1NnteveLsd01+g+9gbcJ617D/08JCyCChyh41CAyXvmqd3z4Yzd4LYRJSdj6qpoWrkOwCaWhniZUhnDD4s4bn/XjD3j4
    Q/eGstRy4fprhp/55OGjx6ojK6ujchrQCCxTx7DNU/vUp333U57ylJ3Ledv4cgwFfK3WkqpaQ+0CoA/eGaMRRMSgdjqiXcQBIzSR
    DDNIoBpFoKQAUww6nZZz8z0lTKZIC9xxIP7zO9/5rve8azQ8pjqen8N9Llv+nqc8fHEQet1iY0jLy2d8+qqr/uhP/ypId26w78pX
    vvbMfWcwYBmMGqgsE7OD2sarV8CZhE0bMUVEqERobKIyWbIwbMFKW57vJ4+K65akukIVUZDmAFDWsAluueXwe9733n/6l3/y4mNd
    dTK9zyX7Hvvoyy65oDD2YKdYkXjEUJNaNxlxrLNuvmeKfb/9pk9/5DOVr6+o/I6aTU3BZOyj+pLnujk3q4vd+nnPetrTvusiJ7Cs
    ZE7+de4O+XfSDNuEdS/fX9uEde8ielCC1dGxbm956FF5/N0/HP6zN/2ThwiidTGGJjW5oSTWAlXRemG+Hk0+uzC/cfbp6erhwxtH
    uJud05SL5TQVdd1Bvmv33Bnn7r700ovOP+e0Ky7d4QwANA1EYAyMmVmlKxAjlEGEKFE1OjbkoyFmttTW7kUAJtqKsbZa77q1hIOw
    tRkdYvDeuAwGVYM7DtY33nrbm9/6lmtvuqbTk7QjaR5WV9cR5ogsU6MIT3rC45770z/Tz3opgyI0qrUNOMwuBU1htiK+9ldFbCe8
    +K4KMwpsmWdDRUmUFETcKglaC7aQBqowKaCoa58WDsCkjEcn+sVrbvzS575w7Ze/fMdtNyZ2sm+vnZ8vzz47OX1fvrzc2bVrz8Lc
    GUW669brV//ibz72qevMsdGyTc6M1B35ihLOOllZlnOd+Y1jd/bt6GUv+tGnPf7MQQpW+HLiusXxW7JtCJ58328T1n80YW0X3e9N
    kMIS6jLO94oalbOZEM4+e5expOCoyhzJGiJt6qqeVDt2dC+4aHl17aYzzrhs5cj+O/dPdy4+yCzU/c7S3NyeXTtPu/TSyy+77KKd
    u8lZsIEjhAq+AjMyB3IQQRAACIIoUAKb9jPGGAPAJIY1ikRSVSXVSGQME0AiMYaZ+qhIYLaWHRlAIdOAEGw3gwBlnRfp+eek55x3
    /mMe/+rX/9nfv/cjnz62Wck0n2oosv7G+tG5Xnjus/+fH33a4xo/tRSbaV2kGUggQOCoJEjIMAG+ARsQwRkARhUxwksrpEVMWzvN
    AM/kuWJEIDDDMlhauUACJ0CDZhSN1TSxqCuJvpMVezp27pEXPfFRFw2HP/KJj33yQx/64E03X7f/MF1zs9R+vQ6He/Pr1t2SZj2X
    dm++2Ro6L+8srW9Oekvc5/5wOh6OxoZ4Y2Nt5/Li6Oja9dde33nqmb6EaSrX62wpf7UnvJ2n/9a2JvrWwjZh3cuIEVlqGsSyGqdZ
    ttng9a//r8wcfQCTCBtiUoaUnQ495MHnv/IVT2AgVigSbK6qhLh3jx0OkReIEcxIEjhC9IhNCIqssFB4L3UjzJxkbIHYKpozmKFA
    A6ggCHxdd/NUAzFzG5eJmpbjvI8AMRtjyDAMHACvKEdlt8hN1yIwfIPEIjMYD8VyTaBO79k/8YPf9YM/+LJX/fmXbzgyaVCWSZYW
    io37XvHwqde+SzY3VpfnFvykdHkHsLVvyCbsEIGNIfpdBAURqA3vWrnT9p4nCGa5XisgQwQh216lDLDAMLQtdAk4Q5KZWUFOrCCw
    bbUsfAxmfsDf9eSHPv7xD60bVA3++b+975Of/ux1N940LbNQ6fjIxKVlki2NVsrlHaZq6uH6bWlnhxECuNftTDc3h5trRdp757s/
    8NQnPug+F86n/UyqiouT9a3kpOdt2vo/gW3CupdhGaPNMhlQmuZVxJevPrZ//4H+/NnBB5tY8TAGTLCkp+3pft/THtEhGJ3A5E55
    YZlUbfBY6MOYoAhRlMQYJIYAYwGMx8O8k7rUWJDX6CGq5EWdse1asxfEgCwBAOJUGd5wjDARRGgaeA9VDPqGgAjUARrBDFX4gE4/
    P1quJGzm0sIkHCcTowZph60pmCKwMcbOPi4459xPffbGQf90qjODWqpNh6TnaG19de/8ogjIFQAiQZLEE4Jgc4yNTbzn4zesrgxv
    v/32AwduX11dLauJiAAYDteTJMm2kOd5URRJkl180aW7d+0464y9u3d1FxaQWQDwIVoDgiCKiDAoSTK23TJGkTJxLBKCV2eSPLEJ
    o5vgJ/7T4579Y4+rGtx2+9qx1ZWIOK3KqpTTT78UJrtp//7XvvEvJ+U0zxfHTTUebfR6ST2tvWRptvvNb3vP7/7WD5cKzvH/svff
    8ZJd13kg+q219z6pqm7snNBohEYkEkGAmQQJ5kxKVKCtYMn2G8lRb2SPxx5L82w/R439ZAWPsiVRFE0KpEiKmWACQIAAkXMDDXSj
    G51vqqoT9t5rvT/OqXtvA4wympTHWL/6Vd++t+I5e6+zwre+z52W8TwTuNDmtvL8zPMZs+cd1nNsPsTBdF7GUWp7SyX+7b/9z9NT
    c3UTDCeJy3zdkIhqzVrPz0+/6LIiiTFlyymHBtoSH5CyCSv1qTR1lhPiGIMYOLABo5f2avWlr4yxxKmCQXAGTQAIxkAjVpb0aFkf
    P3bssccP3H3fwydOnDp27NhwOGyaJoTAzNbasizzPO/1etPT02efffZll122d+/e2fliqRdnBhsIOOlHfZelvUFo1EjLe4wQMNPH
    qRG0jrPZIIy9iXZp6fjObRmLjitsmN08qnziHBus1KgFwxr3P/TUF7745dvvuuvE8UVLySTjY+ectc5ZJtKNG7YSq6p6H8bjxRhP
    AgAnt9/xIJMljWlCu8/acs21V17zoivO2rVxeoCqNr0CBlgeeceSOy59zNPUgEGejFjDBFXRJGEoyhJ9h7275y45f44IK+NQ9CyA
    WnDJxbv/9E/Sxx4PPtZEMUtpXC4yGZP2h+XK7ffse+hJbN8hhQGfBgNlPA+2+v7a8w7rOTWCy00Moao57eNDH7p5YSFWDVFiJcJX
    qsEaB6O1oPrFX/gZARI2rCZ4qIJSKEE0NAiapGMoIzDIWBgIC0UFkEWy6tC0vOYEFcSIaow77th/01e/9ti+JxcWVsajUjyETaRu
    elCkp1q0nKGkJk8VguGyLi+Ontx/z5duvNdaS5ZN7mY2FG960yve+qaXJA5lQGrIMjLGyZPjuQ1FI0gMuFIXUm3U6PKWeddUh3s5
    9zIsLZezU/mxo5iZw823jz90w8fuvvf+qimtZeeyRLKMy4svOufSSy/du/f8PXv2FIVRRZqCCKKIEQCMwWgUn3xy//6Dh2+/+/59
    +596+tCxqrL7953Y9/BnPvCHX+oXvSy3O3dtuebay6996Qu273QKLANplhAQFCKOFN6DCNaSClSQGKggTSCKsqqKDKyRSBO2pbev
    fOlVBx+/r5fNnByfZLLKBEgdAxk39vxPfvn/+we/+78NJTBbamKapFUzLpIsRLWGFetVY/l5bfozZ887rOfSlFSkETZ5kR86hk/+
    xdeYZwI4c2mRuvGw7iW9plzsOf/yl1y5a4dzwGiEnoN1ACGKRoKyUziiXIEGIggsnJBNGAScWEGvDyKcWsBTT514cv+Bhx7cd/CJ
    A4cOHvZNbJogXkSkX/R3nL1r2/btZ5+/J+sV09PTg8EgSRIiUlUR8d6HEJqmGY/Hp06dOnDgwGOPPfb0kWP1iA6PF/7Lr7//93/3
    T849d9d1r3j5y6+9dvd2LJeY3licXBq5rJen6GWGohhFnuej4WKaqrPcBPT7+WOPV488evgDH/zUw48eWB6XGzfNbp6ZveD8Pde9
    8pVbN01delEvs/AeMSLLoQrvYQy0rVhJV7nbMGXO2nrui64+923veEUTMVrB4adOPfrIU/ff+/hDDz729OFj5Sg7euT+b9xx9+/9
    Pven3NRcvm3Hpm1bN+ap27Rh86ZNW6YHMxln9bg+eODAk/sfv/Kyi158zTk+YLxS9/pJkRuQlOVynqdNvZKlG9/8xtd+8AP3BV+l
    LmlCzZZCVFJ1NmmiW1wOX7zpkde99HwLlE2tHc00W2MBeF85t1rbkudncc6cPe+wnmMTRtOQcfjgB7/2xJNLyDYmaVo1tbGw1pbl
    aJAQZPTud75ukIBR5j1HoqLiQ7CmMEyHDvlbb7939/kX2SzNe5wkCUdUY4xXUI3x2P7j37j7nnvuuWtx8RQjqjSGNU85TUyW4+zd
    m6564eUvf9m1F+ydNwaN7wS9sI5SlRlEIMpC6IARzpwNXBUEVYUbb3zgznsevuOuu08tjR+4++StX/nA+7d/9WUvvvx973vtzh3o
    T+cVxh5Fg8WIkZqs8jYptmjE8ggz0zh2FL/4v/2fx08uV432pweXXnX+295y3StfduEgBUVYRtVIAMPCJVACCD5qE6OqZpmzCSw6
    3AMAa3TKkgJzObbMzF11yZx5zwvqBsMV3HDDFw4cOPzIvieOn1g4+fT41LHmiYeXvXmAp+PyaOzLJLWzKaa0EUuhSPwXb/7Gvznr
    H+/elVpYRlgZLg160y7pB2nStDfyza7dyXWvu/Zjn7hdXeHrkJpUojcEo+x9unISH/vQ7a+48nxrkCZFklDd+CYGNrBrMmWn2bca
    iHje/nvseYf1XJoAIcIk5r4H8aE/u7GY2r5YRTUq8E0ZpwfTVbVsnVx5yblXXTZtFAwhRCIQR6NqmGLEbTd9/d/+h/+STW1QdkRR
    Wnr2yIhW1Qhc2dTMKDInUYpe/6qrLn3pi6+6YO+eqT76AxgDKJTa9KpxhkDCnYYhA9r+DJC1QuCI0Mo/JGzSwr7zjRe99Q0XLZfv
    3Pdkc+OX7/7C5286dPDIpz/7uTtu//TFL9jwy7/094M0xEWEkEujN+Om6RXGGhsFLoFLMC6Xs5yuu/4VP/rjP3zWDviAnoWFRxxX
    4zobzAsgAokaVQH0CgZs609VINq5VFFE3/hQW8upTZOEY9QYtEjdoMD/629dFzzGFZoGoxGePDi86+4H7993/74jj2jgRlJWF+sm
    NqE3SDZunn/Fqy/tz6YBCFIW7Dh1lQRrbJBo4YI0JMm7f/glH//0551NuYEGY8BGWYJa9EJj7rr90G23jK5/RS+xBMA5F0IIIdg0
    m4RXz9ezzrg977CeW6NGLDM+9vFblofEBUBc+3Fa5EokWmZFABb/xk//ZJYgIbBmEqjVBbUk6seQbNe2+Ve95IUPP3ak8lLVXmLF
    NvRy1y+KtLA792yanpvZsmXL9u1bd+/atXnzVGYhAtUW2YTgoZDEsgGMJsFHAEQGxjBIqJ3KpWpUsePMOWMME4IoKRGhXm6EMDOd
    XHZxcv5FV7/7PVd/7aa7P/UXn3j4wa/vXCQfGkUiDNG+YAqmyGfD4uLB2X5pMtQNtm7EP/8/fh7KV159flmjqsNUaiWOlldWZot+
    0ZvzdYwQa52zFAJ5H6pGiSjPWRWtJBkruOOeTzpm1FYYx4gzEG0kSohk2PV7HHNMzWD7rv5LX3Z1DFejQWJRljh2Qn2opufy2TmQ
    QVQwoZGmqhvK2bnCwyvEcOIl9NPcA7v34Iqrdt5+99Esm24CWZuQoql9lhRG03IsH/vYl697yRsjUFax3zeJs3WjIYZJGauLs54P
    qs6cPe+wnktTwDlz023HP3/j7S6dLb0Eq3kvG1UrqZ2OOm7qY1detevyywqOYAHUGAsfPKsaYtVoEn3xi/dedvneIFCCEsiCCcbA
    0JqoYBugMHVaxTEGETGUGIaxEOE2q1JFkhjAtB6tFfCKClUqBjkAaUf2CNTK5EQUg6QcroSmNEkam2bjxqn3vOOyd77psltuuvPC
    vbszmwW1HtAI7z2sRI1Z4YwtIcpKrLj2igvYoFYMUnioR2lNMjOzBRHjZUmn7agEIrIMbMBkvUdVYbmEtTCTL9gC0IyjkwtS5MwG
    sYlsNM8sk639yDpnTEBLFBFJolEyFgSBFQwy5FtJOXMZBHFUjdiayvsin0ryqbKpXMIQJjZGYcgquBE4g7/+k2/9ys/+y6nZad8Q
    wQn5SBACGTa95LZv3PGlL59z/avP7xWmrgGIS5wlhBiModMXwvNe64zY8w7rOTVFFfDhGz5z5PiQ0o0m4+hX+kV+ajgynDB8b6A/
    9MOvEyBjiSUbg9rBc0zYMFk2CL5RJFnPCCG2EEq0KZ4qeRKoJjHCGE4MR0FVBWNML7UARNBUIIJNO7aWuu6Gcpk7Sob2YaJoGhB1
    yRcAnWA1y6WYZ7lNYkTIUnVQBvkKL7r0itlZrKyMXOHYwDlPvNx4lONm4/w0N0MTaTqBVTR1aYu0HA9Nr6iqsU0Kx64KkAqm4AOn
    cM9Dxx64/6Ennnji0KGny7KCcoxaFAUAEYnRM3N/UMzPz89PD1505d6955y15+zc9kxTY2kIl7DLBgxEhfc1GUqcJZCAFNH0adhU
    MWqWZ5bMODQQzdI+E1JOLDhGNVo44VQMAkgiEq7GIUJdz1172fzFF+1+4uDYmIESGvFcUBUqosakS6Pm2Ac+/KdXvOAXt2xOJcIH
    nySpAFHVtAj89bPPzzOOngF73mF9B+sK1h2BZFfB5o7ofXIhbTnOCRF4dF/4xp1PwuQuS1dCKYSVldFMf8A+sCxffvH261++yyL4
    UBdZTyNENLEJQUvfWGJnEyUmoGliqwzPBDJAG28pO4Bc11wDkGdWFWWNlvOJLYxFHXHyJA4fOnLi1MLKaHxqYenpI4eOHTs2Gg0B
    ONdJNgwGg9nZuanp6ZmZuS1btu3cuXPLxnTztGFCFUyjTepygsZI/T6sYOlUPb2h54FRDR9Kl3LjY5FkvqrVixEsnAobZ22S503T
    pL1+Ay6yaS+49Y6nH7jv8UNPnrzngYefOnF0eTyy1mZZ5r1XVedcjJGWlieHXFRJjyjtO2BYPv/FzzbVKEvT3bt3XXzhRRdfuPe8
    c/Zs255aBgDjUscIQIgKjpa4jpVNXAKjIC+ByLAhgH0tzrimDElqjYGO0XKRQg0aZLlNGYtVabL8ve96w7/5939sTBZgQ9RekS/X
    y4QaMt40PX3fw08+/NgRl581N40caYji67LXCaa2/NEgleeB72fInndY386o04YShVdAWvW59k+izFqPyzQtRDCsQj5l64B/9s9/
    fXlUcJ6txHGjkuUDQpDRUm4bJys//rb35oCKqCIygoCiGEOqsEgsG5WOXM5ZAwIRYiu+AlakrFAGBAJYCwW8IEQEQdVgZYTHHj/0
    1Ztu+9pt3zhxclHJaXflF2IFVCiAWk15NE2TuMzaLAoFD6hhThh6zu6tr3n1Nde9/OpNmwfSSJqQMYgCIvRmUgC+8lnuQGalCrZX
    hGoEDg7u6cOndmydC4yFUemKPABLK/jSF+684b99at8jB7OsZ9KsrJbzRHds6G/YMLdj59adO7du3rJxerqfJIkitrPtMerS0vJT
    Bw8//vgTh48dP7myxEXqvT7wxLH7Hj1pPva1xCTGmEsvvui1r3n5ZZdvyTKkGfKcAOMDkpijgQDEMMZyy1ER4QzXYzBZKKQCSMlS
    XZZpLwchBqmraqYoInDdi89//8bswPGlECJx1tSUFTPRDxVhZVQnuvmX//3vfuiDv7wwkumeGG5MyoSgaoWgYIIwCSmoTeDpr9Zw
    /v/o9rzD+k7WCeUIOg7JrlZhiJqySl0CBjPynvWCr97yxIlTIVI/QJSEUxtCYImFsy6uXHru9pe/6KJqpez3HCX54rAa9DOyRhUR
    YNc1mZoA38AmkypIS28wEUAVwAcAIIvxCI8+duyOO+96bN+TTx85fuzEyeFKZVwqymUdCZiZmrro/PPmZqZ37tq2Y+eW+Y0zWe5U
    1YvPs97Ro8f3Pbr/scefOnrk1OLCyuLicHm0fN+DDzy+f98f/s4Hp/v5i66++A1vfPXFl2xxKaBgg3E5Jpf4iLIsp+dmTy35lJ1V
    JslYM+swDpL18n2Hmv/4a7/62L4njj81ym1vZnq61+u99FXXvusdr53JkRgkCVwCY0EE4k7Eoj3ELTOX6mW+QRWxMAQMlpb0vnsf
    ufmmWx+4f99oWLFz37j99rvvvJ04GivTM/mGjbNbt23YvHHzri1n796x/ew9g/4AXlHXMITEYXEZeQ4ifOmWJ5/Y/9CxEwf/+k/8
    0IYN014AwDnO2JGoBNo8j7e/7brf/P1PUjJVSVLVAiNRtNcbQIy1g4NHD9/wsTt//F1XCCQ0VZbkIYZW2xWdcs7zQKwzZc87rO9k
    nXSobck8gMgwrFCBsznYR2nYJuMSi0P80R99fGU4tv3pyge17JiaUJMEmwrEv/e9P9TvI1QJk6nreqqfVQFqcc/9+48ePdrvT/Wn
    Zqampufme70cowDDYEIdMRzi1KnFxcXF0bipAh86dPyJJw4cOnj46JETSwvLiEisq8bD+bmZbVu27ti+9eJL9l5w/jm7d2/dsQUy
    hgWMBa8DZEWFKrBj2xuu2sYWUFQVnj6Mw8ePPLj/odu+fv8T9y9WJ4tP/tmTf/7h37z8mvNe/+Zr3vCac4MP8zNFuyHnp6fC6IlY
    WjI5yCI0WVpYxkMPn/jgDTd+6sab1Un05c6dG1/3ipe95fWvOeesIhIQYQSWYQxAiLEV9gAbMGNVI4MZzEgNBoyNs4gNJKerd+39
    qbfsXRnigYf373ts//4nDz351KEjTx9fKf3Jw8PDB449eO+TSVGcGn8saCSORZENBoM0zeGdb1Bk008fPhJ81YSVjZuL5eGRs6+5
    8iXXXD5tmevGZQnBiIgP4px7w+uu+f0/+eTJ4bJJc1ViskQGGps6kEGSDv78Y59/yxuuyAjT+Qw0tKqOk+Vins8Hz5w977C+sykR
    ugq1ANRS+KqCLURFNC6PfN7rfeYjd95x5+NItrBj8crMIQTDaqjxzfjCPduvu27X8iLmZowKxsNyJk1DgLX49Oe+csMNH1XVojdN
    RBK1DtGwiyoiIhKIiKxJEutcfnJhSOScTRkUmpA6u2Pn9rN37njZS1987QvP2bwRwSNxAKEuMVyIG6ZPv9grCMoEMLV8eLECEwqL
    PTuwfduWCy7e8pM/9qqlo/j4x5763OduPnTs8EMP77vrgVv/79/if/XLvzB35c6Tp5bzwZQvq5TMVNGXWCjJeLhQN6pKmXU333iz
    aTA3M/03f/bn33L9XguMF5vYNFmRKCNUHdcVMYhguYuwYgjM1jgAaImbRZRUrbIhNYBWMYpmhq65/OxrXni2KHwEMU4t4oEHD957
    30MPPrzv0Sef3DK9cRx9VQ0b748vnGK1UKdiEZfTNHU2zZ2NqM67YG826OWOLeDFqyYhBJcklgHFxo14+cuuvOEvbpNYJa5HbIit
    942oabwOpqcPHjr2gQ9+9ad//GUKhEbTNFVV1g41wqvF9+fzwefanndY38507d4yggCMSGpJiYDgYVMj4HI8IsWf/tkXknxbY9Mq
    1kqwNqnHS0VGHGvSlXe+80cYyFNA4aswMzMzGo57vWJUY/P0jnN27l1YWBiuVGVTOZcYtiowxjpmWBIRDTF6jTK6+vwLtmzadOFF
    ey++8LzNW+ay1GQp9QuU4zAYIDbQujEmaWVPB7NmVEcyZCwZIrR4iOhjjIaNtQk56ljeAQBWMAsgYHYWP/EzO37sZ374oceHN3z8
    U5/+7BeXR9W//nf/+Vf+9T/bvWtq3GC6N+1Hlbq5oaDfy0zfim0sp+fvmv1bP/qOiy66ctee6akBkogoYWbGKlBBg8S8YILomiYt
    CaKqkiNpdQ+1xVx1/QZFA1WoIUrs5HzEgKaBsUgMNs1icNnOl1+z07rrveDAcZADMZqmLqtx9MFw4lzuywBYZ3h+gzUW/R5KH/x4
    aIo+nAOgSgCSlEdNgLM/8dfedts3Hj50rGRXiARrbBU0z3pMdGJxaSov/uwjX3jP219mC7BMypqAEay5quebhGfAnndY38G6ZacE
    apUThBAA1xZdAtSDBlMzv/7bn3n6WB0wHzTUTa2SMBsSJQ2W6r3nbnvtdRdUY8z2EDySzIKgUhkq4PFj773up//6dXWNw4ebkwun
    mjpUVcVsg4qIWGuLfj43N7dhg5vqgwJYwAbGwhCkLfqQpAMofCNVklpjAYEaL3AuMxHqERtVIjJgNo6NA+ARYggxxlYCjNrsN0ji
    EgCNr2ryu852P/kzb33HD73+I//tk9e/8tW7dvXLEVwKX0WmxHtfabTSzORwaSCk1lQ/8kOvJnRcoGE8NE5gbRSNMIlJfOu4BMxs
    yDGYwEqyClwiIoEooiACYgEmAmmIEoMBDBMbg6yASFdcT3KISggqQrs3c9WgbpQSzudnnKWW3dSZZLiiRUGMjktHfN0rCtGQWEct
    tlYVJMZGw/bss3DB3u2Hju4nqYkyIo5RIwjkhJOqqZei3vCR2//WT77QwI5Go16RQVlBp1GQUlzDzj1vz4U977C+g3U4JgUpEUy3
    GClCjU1RRpRqHj+0/Pt//LG0uPzUidr2rQICq0ENsfpxlsd3veN1gwI2IkbYlkgv1IOpIsRy0MuVEANYsXNbcv45WwCUFWxbkKZW
    UrULSIyCbYdo9z7WEluuGAaPqibLEpMNBFIiEqtmqKUy5AjkYNv0RFRqLyJSZIkFO+vUthGCRI3B11nqVEsEw2z7NvGA62O6737h
    7/+wURigbqLLjCLxQdnJIIkky+pPNv5ko0oyzopsNEY9CrMz1rpMwxjRW2sLiNcyIwty2nJdT0RGSVHXnpmNMcbAEgBRqKi31BIm
    qzHRGBAkgoOCyHhoDMpknYOAY0f3isyh3yeQA2qFFxUiqsd+ejAFteOVFWbOsqSfZdLUXmOW9qEwhmMMysEaDQhNsK+57tq77z50
    crm0SRJELXHdBGYU/Zly4eTc3ODDH/nsa179wvN3Q5TadJDUrHoroQgI0/NCqs+lPV8d/I42GcVVQAxgW2SBoPYIgZy19KcfulHM
    /PLQp70Bm8S6HpGJPiRMiFU/5ze9/lLfwJq29AVfN+yMwBP7ql5RidZqv4cih298Wfs0AzGIQQaGYQ2sgbNgC6VK4AVBTLSJcc4S
    oWpilmU+UgAILNCgEVDDRJBWckoBAzjiXmIHWRK9kMICLFFCrdFbSJEawdjLSkQ0bEnZhxhj5eBJvSOU5Wh+1kTF/oP7s36PDEMr
    I8EwZS5hUOLsqeEo7aM3Y8fiy6ZRTmCzGKX0nsk1QkGNwokaH6jx8AGiSDNnnRFF49EEBGFRw8hUEyBlJAwbJQZpFI0h8Vo51ixh
    ER8BBXxEUCRpR28v2sQYoigpLJui6EFlvLJU9IuslwMCjewcU4e5ZQNjLREYEmJlCG+4fs+2rQPGmKXRGPI8tzZR4rLyJu0vjvzJ
    pebPPvL5JiDvFzrpc9Kaaob8FdMb+H+CPe+wvq2Rioau4d5WejxBCQwYX0vjBfc+Un/2c3d7P3B5ztYE4RiJlSCaGGQs/+DnfzZ3
    KBLEUFon0Vcut01TVVXJxC5LlNGJWpGYlF3GkYIaUdbOWRLAKoSAqOw8uAZ5Za8c2vKTMQQ4phYnlsJmZBNwCpOAjAqiZ/UqUaJA
    wIBRYYkEsUyWmSkSAuBLLMFohBLgK/gl9E2RqC2IjIytGStqduBMjo9OJf3Cci81U0an1BcaM0aaZXbsG48GLEjTQJnXVM2A3dQY
    VEkSiBvFuEEjUANhlA2aiFGFOkINIjCsMKoRCE00IVhRp0gt9yz3DDKAHQxBCCHlaNWnQN/CAWBU6svYBLLCeUAvYCqiCJpGSvLB
    tLJRYrE2Go0Ql2QKRIUSfPDMJmhgBItgFH/nb/91jqOigEo1Ho/TJGNrKt9Yl2W9mWEpX7z59oURqoioVDU1CKPxGAARFMHH+kxP
    RLdhOP1Pg6p/PiX8diard8S0FtgzNJTSeMqNxe/+wZ8PR7l1cyujkmwgk6hqliShWkqs37lz0wsv380CRZNaLsfDvOiNRytJlmSm
    GDVVmqQhtqIRBCCEEFSIKLW5QoOGGFuPZCxZhgkCEWHm1ECApomWOHVUV5qmREDTgATOGVWUjSYpBbFEYAIIwWNcA4LBwEZB8CBi
    NhZIY4SXsSQzS9VwU7a5pamam3XjsbjEjMeLg6nUpA4IirT0w2K69/TxE86lSpxqDumpN7XEJMksSBAblRgMDIhQRdRCYJdaHD6B
    Rx7e/+CDjxx48qnjx0+dOH5qYWGpqhoAeZ7Oz89u3DS3adOGLVs3zc/Pv+YVV/dyOAeJiAJnOUlhYFRBEAZYFcpErcv3JrVkjAML
    gDablo5FRwS+K4prO8bIZJoIFaQWx46d2rRpblQu5rkzhhQ0anDeOcW2LcXTp06kyUyWFAvLy0mRO+fAvDwcW1ecWix/87c/+k9+
    4e0RHAXj8UqvP/DBB++T1CQm+UEv4f+n2fMO6zsYTZSTiA3iRCgBgEkt8m88VN108wNJelYTnGqV53ZcRmut+HJQ0Hjp2Hv+1k9u
    noURJMx1vZIXBaDsbAAsUjbMcIYg0qpzwVibAgC8iGE2ZGG7c+QDNCK3CIFjFE44scgS430MpaaJ1QARWAI7iODUSRw9Wd96z8ML
    S+OTJ0+urIzquo4xqhATFUUGwBmenZvesWP7OeeevXt3MbuhCFL0srlaYACJde1DkuYMDPozEIyHi2I9Fz2wNqHuTW+1JpeybCpR
    NVkKDS76MBqP3CBj5MbABxw6gm/c99TX7/3Gk08dOHroKQlajqu69sw2cRlAQXpJbx5AreGpE/WBY08RH2RmJvfv/tMNg8H09FS+
    edPUpZeec+01l+3ZPZUlyBwsMUmH1GQCK5gdFEGoDkGBxNrEgAxCaFEcndgFM7UD4UFgDQKwNGw2bJoTIMuLshz28p6AxMeNc+bd
    73zdr//On8VofcNFmtWNT5OkaRpSmHwwrOuPfeqLP/LDbzlnp0nzvq+r2o+SJCHlOtSpzZ9PYp5be95hfTvjSb9/7f8KBQJMFTQY
    /Mf//AdNTEl4VDdFP23qik0moc6c17B01vapd7zlglAhcSBoq7awvLLsxfSnZ0Y+QnNDIAULhbqTugIAqE24pd8UEYCZ4SyMQbmC
    ogfnuKxQV8h7YGuaAAl4+uj44FPHDhw8/OADj973wMPHjy0EWKUsCIsIETEzEamSqgJijGFGjDGKt9b2+/2iSC+8eM9Vl1108Tl7
    9p6Xp1la1SlbqGJxsZzu5cVgY+PDUoUs7asOmbQqV3pEeWZjRFlDBVkv7U2npWBpOd5zx5Of/tRNd9718LBsvFUfxkUms9NTO8/Z
    MTMzNz09OzszPz09nWW5cykze1+vrKwsLy+tDJeWlpaGw+b4qXphceXI8ZMnTh3/+h13/tqv/9bWLfMX7t3zgksu2LJxw1m7dm7f
    smF2BqkFK0JENUaWw6VWFEE6xHlk+LbnGsFsrOloIQQoGyyvNNbpqKqcIybN0t7KsOoXrpeapsZ73nH1xz/zxYNPN+PxKB9s8t4b
    5rKpZqb6MZTgzGvx/j/95D/6hbcASNNsuHwyTRJLJigHUfe8v3pO7XmH9R2MJlfl9q6bAYQx1tz09aN33Lkvzc+uSp9lSZ7b0Ynl
    3lQ+KlemZrLjTz319372p3s50IAQV5aXB1P9cTnOB9MZbACaYHo5his66JFJwMqqHaFC04AFTDAOq5doVQQBUow8fAMyyHpYrvD1
    O47cdOvtt339juMLi6OyNM6xNU1oYChPXaphy/TMli1bNm/eOD8/3+/309QZY5588sm6rhcWFp5++umjR4+urCwsNsdHy72TR05+
    9qNfmyvm88y+/e2vfd9PXDn2MAb92byuEIeh6Gc6RmxcQmlZDXscCwjFqmmW4KatzY+NMW7we3/0kS98/kuHHjuyZX5rCkvV0guv
    vPQd77r+ulftTbNWxLWbMbIWDIRJ/9QHEHU/NwFeoIyHHx5/7OOfvP22e8uy7z1/7Zb9d9z+VJpnyrGqR2Titu2bLr304vP2nHXe
    WWft2DK/aTOYUXlIA5fAGXglmxgH+IiDR+JDD+6/756HDx544v57vm4T/+4feuPP/o33EaRuRrnpT/XSpo5pitGomp3OXvGyF7z/
    Q1/p5YPl5ZU0zxCRWOsl+kasSZJs48c++ZV3vPX1u7c7pDI1NVU3pXNp6vryfNn9ubbnHdZ3ZR1ymVRJBQjKHnj/Bz7Z62+pPWW9
    BODxysr0YFCHOktpNDq+Z8+m17z6ynKE2QIatJ/3ALYuI9hxg5/7u/9YaWrQn770wvPnN8xs3Lhxaro/PzM7O2vSFAEUBcxggvdY
    WcGJE0snjp8aDodHjh9bWlk+fuzU00dPHj9+aml5HLyQseOqtomb2zi/bcfWiy/ee9GlF27btrWfmC1TSA0MI7b0eAxrQRYhXNyC
    91VRVVhckJMLp4Yr/pZb7rvja/cfP7R08sToD//4Qx/86Ad37938Q+9908UXnLN1g104XgfNpgaIVWSPfpJZeApDSGWdGovjS82v
    /eZH7nv44KP79xHk3HN3X3PFFde/8qUvuGR7vw8fkaUQ9YiihIxUCBQ0qDiyUEYUkmhAbJwBgTRJbFnjonOKK37x3Y1/9+HDOPDU
    8tJy9ZVbbjn41OHHn3x8OBplPXfs1MJnbrzxox9fMlXI0iRLHDtlErZqWJWCb6oQNQSoWEgKTVUcROenpo8ef/zOu+5ZqcZMfiod
    DJfK3BUGBorpvhtWcv11L/njD37aEPI0ZVBTNVmRrYxGLjHgfFTXUtkPfuhTv/S/vzU2HBCUwMQ+rNF7PW/PlT3vsL6DGbBCVYlI
    wYGgoizCH//k4zffdH82tasJ4qiMTRwNh5s2bZIwKjL15eKP/ej75uZgg4c4gkpE5cskLzwgipMnR4tL1bA89pVbHotKzGwdW8uA
    xhhjjBIaAESGmQnM3CIoY6TKh0okZIlNE1sUbuvc/Nzs1Dl7dl1yyUWXXXrx/Cx0lVA8ggXWwlAHbWhLZapwDBFVARMNCgwK3rJl
    Qwi4+pqtKV9/712nvnjjzXc/8NCTh5+6474Hb7333isuuuhv/tSPv+SFW4fLCAGJ4YzdqOaF0veLtOgZV2TDEQ7uf+zLn/vC0jDu
    Omvzu975xne++aUzfUhADD4YaxwFgaprKQnb6UYlQQxsHUBkJ79ssRgSjPpBQkxWAOOwexf2nDUFmnrT696ujNEYB55aefTRxx7f
    v//48ZPj8VBis7S8MFwa+hgdW2VE75tQT/dn2qSQKcmy3qA/O+jPFGl2/nnnbN82+4pXX+isMpKxb6animYM61AOq3zKOtXzzx+8
    /CVXf/4rhwa9TaNhyYBGISJjsxi8V52a2nLjl29973uuv3RvNh6PB0UGUFOFPLPPj0E/t/b9cVgdRKWdxFuFqawXdWNlIWHtJkVa
    ChQApMwKAitDIKZj4STSVqwAekapHZWILIkC0vqACFK1QfCBP/lYXmwYjapiam5x8dRgMD2XbqjqJks4+NHGjYPXve6KLEGauOiD
    MZaty5E1IRoHZ/BHv/ern/38I/fc//i9D+07evLUysqKqIaoja9U1VqTZIWqAiRQFfIKUgHFfmHPPv+8Sy+95NJLLjj3nLM3bkJi
    EQMSOyHDUZBAEYnIWBYBMQQQFZBYIoVG9Y4dc6eeEAUSkTKyFGVYWR5XV1y+ce8Fb3HpWx7e7//9f/yNe+67f9++ff/gH/yDj374
    DzZvSocVIhTWQHhmdosfn1peWSqrqt/Lrrz8wje/8eVn77nobe+4wvv282hmFVYVXmFjJ0cPjaiDqCpbk9ikxaP72E0+tyfVWssK
    JlEJUSIzG3aiKMsmyxIBpgtcePbgwt2Xk7mcCUqdcDQB1kKkYw1zDnXdsRgqgxURoDhZhBGJRV2HXurgowBpChCcMz401mYK/M2f
    /dGbbvlX45VjbHvOpONy2C/6IUhdlkmvGNWli/pv/v2v/d6v/UIvnyGN3vssy1qOjW9upM9YaX+ZfaUdsFkBapmQ0Im/sXalBCVR
    iJIClr79u2g36A+09CTtwvmmj2wfddqrfX+KdWfcYbGAlQFSQpxMugp1dIy0npdRGd05ECFVaseOjY3GECtAVAchltQw+Qg4qupx
    muZnDklMiuHJsj+Tg3U0Wsn6/Qa2Drjhow8cP7EyXPEz8xvH9TBNk7Jq0iSLGg0kxuZNb37doI8ILFbLM9lUXfs0cXUtzhkoEkGa
    4N2vP/9dbzo/MKqI8RijUSzL0nsvE1iNqkqEQC2bot+fmTJTBfoGoYYIkgRsIIK6hlG0hHZCoFawIko7q62sUZSIEjYANxIR1XKq
    nowlKGKjlogZCJAQB6nNewliNUiyqsRFu9yv/7u/e9/DJz52wwfe9a43bduclgGUIFo7gq/EF1VMojXEiaNG1dD4H/7DHycgRLRz
    NyQtYUxCUTR4ti4o2LAxcAk3DYKHF8QIl4KoHV1CjKgDUofQgA07ZiIbFOJBhDRNuGXPIrgEqt0NCsMILTTewzCSFAyAkGTARDqo
    leJgAhEMEAh+hOncSYOBS2MEHETV5iYAVYgwvHEGl+6du+O+w8MQnJnOssyKxMZPZRkYVROF7KP7jx54Cjs3YKowFlSHhhND7WAk
    nsE/2nJeSLfvdQ30/72tZdJ2xDXCtc81ClYhhQMpQwgwWvuxTfqhCsYmCo+1EHzNyejpDpS1m4zAmlsSAQuBwd2DfxBE0GfcYa0d
    B21d/zO/KStEuSOmW3usKIHUsDIpCCwEjyYxGQupdP06dmf285OiP53DY2l8Ynp+9vhwBDu9sIjf/L//pA4Dl9jKN1FCXkxVlUDZ
    OjYynp9zP/3XXmq62RPyIVp2vkHiOET4yheFg8KPAQFbZAZFH7N9w9qXDgcIw92BWjUBogcT0gwAfBOqOpIxLrXEqDwwiU0UiORC
    BDOI0USQINhWRt1YNmQgiuURUoc0IYkdJ0oMqj44Z8gQBHmCcYWihysu3nDJuT/XKwiCWEVbGBUJsTJpz4/8VJqmpiBjDZEFWCsV
    C09pZmJLVkjcjKE+ZoO0ieAEChxfwvJyVdW+HPtRWT748KMnTy6cOLnQxKhCIYSsyDdunJ/fMLVj57bZ6Wki2jA3tWtHQsDiIooc
    bGAJTDDUNSsUKGsxjlPbrbEoGlWhYoyZ9E+gUIkQFYlYHI3mZ6darhsCQoBLEKJX9gCXTeOSQRAMCvy1H33rfb/0H00AM0DMv7F9
    pVSpvbNJal0jXmM0iSuyDcPFp/7Lb/3pv/s/31uWKHrMEg2TdIwNz44/Vgvy/z2bfq2qv8bDpQAE2naEFUY1qLQ0zhORV5mEC0qr
    4VKc+Erp2uEds9K3cEvP8KzUOjHRM5wDf19rWASsc0y8fjsK4Vl6uWvHvw2tfZTUcIsGZIMYYytjeeZMCRQRJE7Pzy6OVgb9uQb4
    2CfvGJUcyaVFvw5lkCCqIQSB71mf8dKPvO01fQY1iGE0W/QgjAjvQRbOwjoXWw7MHgDE2DHXtPlte2iIIFFFA1qiJQNmJlUyNBQm
    YwEJSaBULcwYVdU0WTIVWpYuRuNx6BAeefjAUwefPnF8eWVlZbiyEmOdpGZ2pr9p89z8hqmrrrhk11m9YDFsYAw0SAhNkWcshQM1
    5bAaL07NzqUpgawlJAmZiFDFnjGqSCJcE6ILNaFkU1fV8VMrTdOzbCHsmExCihAxqlWFZ7gHH12l4AT3P6w33rJEdK4AAIAASURB
    VHjjTTd/7ciRY0RMMFXVELUBEweVGGOSJNbaxsdoTB18nrosceVomeFfdOWl173yJde/9gUJQQ0soF2+DiGxKTzqOgaBWGMMGwYD
    VKNWJQipKpExbNkYAvppcape6qfTo6rMbJYkdHLh2NzsIKKOIOMEaKCpY7zipdvP2bXt4UM1hcwY50XZiogqlNmWdQCKT3z6q+98
    y5uuvGxgFWy58suJ668Lo9bvhLUrdDef+pfeU8/8lUDR1hPa4isRtfBj1f/hxxrPeIRymq0Hi09Mup+lPdBQC0wKDGtBF4mA1LVP
    FFHnyAdhc8YnSyOj9rWBS3rF8eXRsQX73z7yKZPNjUvbNwVDDME3jUpIXISU8Cff/eZXhRHmp+CVETwCYJ0reGUlCJvegKLCa+NA
    1bgu0oIit19TVdmAmYjAhgAHOAAhiES1xgRGJQgAwIYyBzTASsNLi9VXvviJA08+/ei+g8eOLoVKYyTfIHjJ8yLGSFBmABLhidRY
    HY9X2Mi2ndt27tq2cW72BZddcvXVVyUFLJkIKPWn5vtABELt6xgjBfSLKZsaVPANTDA9m9XGUpGApSl9E3yawKgzUBhDiAExQMQk
    AViucfvtj9x88zduve2+qg5VVXnvIZSm6Ya5qdnZ2R07dszOzmzatClNrEjIsqSu68NPHy8D73v8yYNPPjEeDw2jqcItX/vGnXfe
    /au/wfOzUzt3bj9nz+5zzj777LN2bt06NTXgysOY1JiUAR9Ra3c9YLJtctZmhUHR8k0vj3xRTJWAzfLjJ1c2zA6mZzc1GBtYgU/J
    VrFMTBoiYoM3v+5Vd/2nj5LdUPSmR+U4CikTM8XgVbnI5wD7xx/+zDUvebeXdk7ItxxqwCQn1Mk9PSPmWqWB/57X8ypacO0VJ28o
    omBLxNK2jTT+j67m8/2IsEiJlScHik4/Jfxthq2UVinVeXERm2ZyXVfGVFU9w+5KCezQS4qhH0ZOp6ayf/kr//XEUghilIqyIZB1
    ToMf93JnKKIavu+9757ro5fBl5KkFhRCs8RwtadiempxVJ0aod/LJXENYtJLFb47SEoqpGSgYKCu0dIGQkHEXU3aoCEMKywvxWNP
    Lzzy8P67v3H/gw8+evLYySzJnHN5kljXy5KY5XZ2dnp2tgcqs8QWRT9xWR3i4tLKwqnl0biJcdPKqD7w5MKBJx7N+4OPfuzuNP3w
    Beef9zM/8WPbN6VzU2xTnFpcLgbWOWecOgAooWmILITEZo57S6WKC+qCSShJjAJ1KFk0M/mp5ZVoKe0NRo188Qt3/fnHv/zYo0+p
    SIIqy3jPnq0veMElV1z+gj3nnLVxDqmDj/C15gVZQBEF3gBezi6XMTv9CiEsLuHAwaP7njh4970PPfTI408fXzh8qtp35NFP3fyw
    kkuyXr8/NUjzS3efe9bmLWefvXvT5pkkg7HIMmQ5ygpE4Fa6MfiqKeuq8iE0Qvfcc8+Rg0/ff989RuXX/tN/2L7diBZkGhVlNgkR
    a2M5qQTXv+qK9//pl/cfrTU2QmBrokobuFhXjMfjQb7xa3c+9tXb/bVXOx/Kwhp0+rVCnbp9u7DajWCUdF2ctZp5/CUW9TdJ3BTw
    oUUdcwhqDP0Aak7PtX0fU0Jl1qgAlJ91efnWT6KORPfxx45tuXpTFGVo68YI5oxHuKR1VbosC2IjuaVl3HzrAy7donHA1K+qihgJ
    kyW2FHy1snnW/szPvEY8YkA9rsX7bMrZIqt8ZdIiwP7Of/2jP//E53bvOX/nzp07dm65+KK9F1+0x5muAKwGEqGx1RMEFKpoapw4
    gYcf3n/PPffte2L/Y4cOjOoqBDhOASteGLxpdkOeZhdceN5LX3L1FZft2bgRxIgCQ8jzroGkgGjHm67AuEKMGI7wtdse+8gNH3/8
    8SdY8fhjT/6Tf/qvWKuXXnPJL/zDvzYzO9MIlppRlnDjxxaUitiiB0IA1REqiKECyoRDiGMPBB+m8ulI6E/PBOBP//zzv/df/3R5
    yUpMg9ezdmz8Z//ob23bNrNxo7WEEKAKFVTjeqpIJW9ibETqECtrYJyxxP2ZfvDjEMJUll94/sZzztn8+te+EAmagIf31V+99c6b
    b73z0ceeGI/GGhSJ3Hr4G7dHVY2goCSqIUodpRlM5aphlWarTZSUaWl53O/3Q+Mza5rx6Lbb73zzxhd6H9OBs4yyWimyoirrxJki
    M3OCt77hZf/1g59fGJ+C69m0qMdjQIlM9DGKWSnZUPJr/+UPX3D5T+eURQlmAvvtEoiJ8BI64SVSWr1mTxzW92DPnMXo+sSTCM77
    KCLKpC1hbvwfXpv6zBfdu+Bh8l9d7b6vi4pJsJbhMwAlERKgoxcS8L33PnrNCzexkrK2WsnGmHYo/oxamrlRXYrkSYpf/Me/GkNR
    NjRuxsWgADOpUd9YltCsGBq/8fWvF6DXQ9NgMJuPx6irmGYpjA1aLCzj8f3+ySddVeGhh46Ox4/a9KbIDTs455xzhtBOzwBgSF3X
    ZVk2TQMoMwOgiAw5SZpbzM8Mztq99YLzzzrvvLO2bJrde8EWFUAmkqvteB13ubVGRIW2JUSFEoxB0sNUije86pwfftvfG1e45Za7
    br71ns9/+e5xwKdvvOWuBx669pqrXvfGV116yVwD5K5geCAGYOSxElaiQ5qkGSJrTEwU9QCSYjoCx07invuWf/O3f/vhfQ/Mzc/M
    FdnVL7zih9/zprPPzi0jaWXK6mCMSTrK07Qalb1e7kwaJMmyaQIA8RoqicYlxuUCNaACEEVVa8/SVXvTF5537d/50WvHI5w6NT5x
    4tTS8sqho8cOHzl64MATCwsL3tciMqrKsixFxy3ATZWsSdIiz7IsSZIrLrsgzZKLLrhw++ZNaWIvu/jcIgVnRkPDlkgZSqysQZGA
    I978uqs++8UvD58aBcpEJEZNjW1Zn/tTM2XlTTJ91/2Pffmrh9/66m0UhdRQ55sAnNZGmaByWp8V5FsU579HW3NJUVHWPkQBuEMS
    qvyPHmN9vyIsEv72PBvUktQ9K6wlCKsq7X/yiASAQCBiFRAznWFnBQBBQpr2pMEXv3Tk4QePBsykWU/YVNXQ2sSK0QhrEGS8YY5/
    5L3XuRQgBI3DMfV7eYhhebRc9KY8eGoaP/5jP/GSl779sceefvDBh30YjarhSrkSEVVjqCVIVI1EygbRe5eYQT8lSoylXq+Ym5vb
    MDV7/rbzzjt7z3nnbZjbsKY6QwoiIRMms3OQqALDbFm6LNxMekPKEEGWoGmksNybRojIObzulZe/7rrL/+7/G3/+F/d+9hOfvu/u
    +xYXRp/7ws0/9td++N3vehESOJjIURViYAtLiY6GK4Wx8FBLErrN9vRR3PTVB37tN/6IrZuf2XD1FZf9xI//yIXnU6iRpa0ojhIR
    MhsjmtIzc5KYXi8XQdOEJEkJaItcxaDH8BGqGiBE1BiQY5dkxK34K4SZBzPYPF1gV1HHGO2FYjGJaxAn3Idt+cCsFpG0fTpEkSRg
    wNfIU1RjZaKmLFUaa7M8y5uqSfI+hIcLo/5M75wBNs3zsUWz0ISmqjUiMc77KoRQowzRxMrPzm7+vd/74Gtf/Pd7roAK2qojAAi6
    rlxHW6i0Ppn7nsOfjntrojkkUAKYBDBKUEEIEgOp6WpaLf7szG+aM2hnymG1cGoAhrB543yy//g4BHZ2VJbz/Zl6VLvTu4SYaFhN
    eoXSypdOGGZ5/2OHoyAQUiCEQMzOmaqJmXtu2qh13WRZEqOKiHOmnWJp6291A+/x+c/eVY0y2xtUTVBEw2zZqI+ZSxjjhOWNr3/p
    hnkATYDkecaCGAFwLy8EkWEM6PIXJJdevClik8hlUKiAIkLTYRrbN1VF45VU2cAYbsWZ25EaZyABljvF5lYfrGskMgMJgCiIIiAx
    1oKgcUKQYzqQJCksg4DUskzIVSyzNj4mLnN4+9suffvrLn1i39I//sX/z8Li+A9/70M3fubL//Sf/C8Xnl8snApzc6iBWoZVXHDJ
    bCx7FFI2aWJnGRiN6ttu/dpv/vpvVxVdfPHFP/M3f/ZFL9whNXwJR97CIaK97ABgwBlHEw8igElsUARVzpIsSxWUIe1AP9w9DCKA
    EkM0KqkYNWwAiBGKkWEZhiKUxJBtZ4+jqiHy2sIh22o0YGCpwwAAcAks4HJo9NaIyzNIkKjOOUgAuf5Uz0d4xS/8w59+70/+M5fs
    HJX1oJjT4IP31jIjWmedzYaj0eF65ZOfeOgtb7ggLZgo8b52jpu6TtI0hmhaUcnnaKuhKx600igRygAT4AX33f9QVLBxCKjr2tku
    GPymZq2tmwYk1poowmyYWVVFyBiICAwzuPFN4pIf1IjkGY+wRJBlllihkZTTLGmaitl8y8vJ6ehgIW3pQ1ZW4nAZWzZiVNX9LFUg
    BLSJ0nNiWZbUtVfVLEswcbhexJg0Km666finPnWzS7e6ZObkqcO9qV4TmhCiZasxCmKe2re89fXa0c9IW2MjBoGjABKtjQxLWHM3
    La7TwIQaCohoCMFamySkShN6LABoc9+u9GECgCgqSgKLFhJBGJYwBsa1SHFWcC0IHqnptkZ3ptshboP2WhujEMGAjSEYaxgrftxz
    RaXhovOn//gP/v1f/MXXfvv3//Dg/n3/+f/61V/+pV+cneqXY3hGKH1TDckMkiyrFpZ7hYPa4RCz/fQtr3/l8NSiUvKGN71+40Ye
    VU1ukqyAryU2lbFZCIitrPXEB4UAMojaXqYgkdomjQFM6DYiJmR1LVChrSa3eNEmKLEyW55MwkgMta+jVOysY8NMBE6IhUVVVYhU
    iQwEiW2Pg8YQvURjyVhmkwARbKgTRzMKAqNt3G7ZOnjNq1/22a8+lpleYl09qrIsI4pVU9vUhCDOpuNy+OnP3Paed1xQeRBr6tIW
    zQpV40y7QL7Zmv/LrOfVOAsTDBAUPiJJsP/xg0wuRiVYY4yIfKvLe4xRhK21ZgJaVtUQwrO32A+WLPD74bA2b9kg8oSqU3Bqk2al
    zl32nZ/ZFSkhBFU7HDc+IAZQ11jhGKNb3Y7PhYUQ8jwH0DSBma1la3i5QlD88Z98LITEpUVVRmsdseaF9bUvXK8ZD6HxvPMv2L4j
    rWp1aaoSQlQCtcGfMYlRryISa45WFcRsmIwGgVYSlU2aGANqKTMFiIImtlAaEFFER9QcgSbWgBAZw2k7OhEFPiDN0SiatuXIACME
    hIBQggEiOAdnYRiOYQnBICGwYQEaoI5VFO9sOnDFqCrniny4PJ6Zyn/qJ6+99LJtv/Hr/+X1r71m+0YiRV0hdzA+yaiPJBlWJ/qb
    bO2HlV+Z6m8ohz4j9zd+9O1eQAaNL3sZh6jLTczSJCoxdXU0IUTFqMTyCsalBCFRiq3A16TUI4QWAqoKBtjAGFgHw4gRgx6mBkhT
    ACSRQgAB4lt6Bpsa226s1tVXddUy2KgqEazpNuJoOTprkoxMOxSOTiqxDmqMIWbt6thd6hYAcrj62ms++fkHDRGilGU5nRVRamtB
    iL6p8jR1xext39j3gQ898Z537LagIACxdS5IsC3OVZ6x7fkvSeygPEF1nfZqMUIEDz70KDgRUWbDxDEEfAt302K1nDGkEkNs+fWJ
    aMLG9leF1PSMOywFzjn3LB++yqZQUpFADKJnlZ+UoauldwZplyF27V4KnpYWMdvH9CCrfWWMsal7Lj+nIk1TZsSoAKzlEESYYfGF
    z+2/456Htux4wbGTofJNb2rQyHg8HBsywYSoIUuzX/m/fmJlhLkeKQB2TF1zXqXN0WJRZJZbWBWrgIKAHJMkiak1lDGokI+RmZ21
    IiCgzZYEqGOXLbKBoGcY7e+Xxzh4aGX/EwePHl84fmJhZVSdOrm4uDyqa++DhBBDrEE11DPDWps6aw05w9by9q1bijzdsnHDhRft
    fcEl5w2KnJCPGmQWjBwAMxcZDctm795tv/Vb/5IV4jFe9M45MOqhz6m/XAeT8DguDrgyCaKglzsLxArOAQ6Vb7EBRGwDoWpQlnj8
    8SPfuOvehx585PDTJ1aGVRBDnCyv1KosakRNVIgSwELSUne1PtcwjCFniEmyxJF6JriE5memzzrrrHPP2bN1y4bLL9mWp0izTjh2
    EufCa+YYAnihEISNkqj3cWaQqMBH+AgRgEEWZEAuaemndZKuAgBhYTnOTJlXvXrnb/3u9icPDLOCeoN+iMGrOmeapkqdCyGwsZxM
    /8Zvv/9Vr/wnW7dgeSj9gi27EEqDQLDdh1sz/kvUsNBBy9aHa7w6eXP0CE4cXxQzy2wjUYjx2zgda22MXokkRhVxho0xaZomE85U
    IuraNsz6gyPNOcMOi8AG55+3B+rZQAi1r/IklyC0HsK/HgQMBuIaN3H7d0II7uabHz/vfXsAhNAQpa1WsHnu/H7rpFQ1SSyAsizF
    9FY8fuv3P2TzwbGFxci9JEl9bBS235tVjSvDFccSjXvRS//3vXu3bp6f2nPW7i0b5s7ZvX33ztmNcyhymJQNFSfGpUs5McyQSGA2
    FmSsiZCoDYyxJmVn2iJsiG1QgKbG0jKOH8PTh48dPXp8cVg+sv/JleF4YWFhZWWlDl6gRKRkVCmKikCUmJ2SIXJAwqYQCarK3A4X
    whAby/c+eDjLMoMjwd8hIlmW7di6bffubeft3nbW9rlXvWRrWmQnViTLk9zi2OLi5pkZthjMOwRUARoYwdR1yOf63g/VxCRlZSws
    npjJZ1xmhXBqGN2gIOCWO5Zu//qjJxdGd99338nRUtmUde1VJ+TyYCI1rq9K7cRuG1oTkYA1gGCYiEEC8kAtwoqhj6SpIciKPHV0
    dNcD95LeS4rUpv2iNz3Tm52b2rxxetuO+U2bZ6cGyfl7t05NIcsBgk26RCdJzVLsmj2UdiUtASLQCrU2DcYlFhfi0RPHjx49uri8
    dPjY8QcffLIe54eeWuj3twyHy+PYFP2cOamb2ihSZ2MTF5eXN2/eeGrpwI1ffei977lAiYXb7d4WHkEtjGX9Eqe/lM/6pjB5BTPu
    fWA/EQevtrAkpvE+dXYiCPlMY2bvRZnbpULUFSicW3uAqrT88T9AxPwZj7AsY/MWm2VJiFFgYozGGQnfmdmsFdHtAmdlJv7sZ2/8
    2Z/e08SYZRmT9QJ+7g6bKojgvU+SBEBZ1nmei8Xv/P5t+w+eSrINsRGhmCQ0GkfjXFUH1ciWXJotrCzPzG07cGi8/4nlu+87QepZ
    RszNVJFs27Jh11nbN26aveLqS/LCFj3jEjJiSVxKzjmQY+8TGAPFaByXlssTx5eOHV+++64Hjh5beOrg0ZMLKyrGJbm1CciKzaNQ
    jEZ1atLw0oh2kTEMGCwCjaQKITiyUUWjkBCAVkfaBlb0V0oNjWfO0zT1jTywb7jvwCNf/Mo9zfik1fKSi3a/7e3XX/Pii/IC0zMz
    I4++gy/DeETZlEkzI9TMzc0tjZenp6fDcLiyMgLgikEZSQ0WVlBF88E//PxH/vzzw+Uc6FdlFDYlbDRTxrQwDhNCR1pf1QHttqVJ
    OtgJbjuCVZAQibKqapT23jIbY6wxJjEGpKoQGpW+HvOJcR0PPqX6mHXiTDAmrgxPzUznGzZOb5yf2bBxZtOG+cFUzzkmjoCCiYi8
    94vLoxPHFxeXhkvL9cpyefzk8nBlHESttc4Zsq6KUaKdymzRnynL2qbpbDY3rEckpGr6eVaujIqiyPLeQlXmvcENH//UVVecc845
    DkAZ6tQSAVG8JYNvMnb3l63Jrg+ylJVABnfdeV+a5ivjdlC3HcCmDkzx7O0mE8JbImushDo0viim1kVYq5VEEpEfVIJ45nFYDOcw
    Mzu9fNyLcGKt97VB67e/6fWEod/k96Lm0FNHnnoqbtkibBgQFSb7nDE6xhjbtogxNB5XMcY8Tx87NP7QDZ8spjeOxk1vZrC04qtY
    s3ECAyWFOqdJkUbtLw9jL9vETkbagGpDLL5eOiWHTw7veeSQSU786u/e6lJJ8uCcGmKjidOeIWs01M24Hflq9UVVtawak7jGxxiF
    09k0Tdla733dhFBVICYiZmvaH4gsmaZpRBWAaiAy1rAxhtlWvm2mS9tKY4oAyEBiMNaQSwAIs/e1qAY1o1HcsmGPif4bDy099tRn
    q1/5s5nZ9LXXv+hv//QrEeAMZ9OsDo1ZXmpOWaJE82ahSW0v+mSlRj9Nh2P8/h/d/pkb7zhw5NTC8hLbvF9sROBGlpyzRnKNpBEx
    GoA1WoiQInctV4dSqwbdhVnwk03SriZoKwWPtMhFJIg2UcQHoCv5GUfOqGUiJKQ2aowSQtOk/Z1lpINHcOjpMXSk+pSIAGocRL1I
    gG2RfUYEEmlcesOJtXMu35wZoxp9qOtaKMmj6rCkfpolmVQhjKthGRpn08SlBOO9V9XB3NTBE0dD4h9+9PAnP/2Fn/9fXh+BtoKn
    iMzoyCW+66mPb2nawRpUIcRok2kFgEceeSQIjE1CiGQSZ1PV+K1eplU2ISJRZeaoaozZsWNHq48JgABRJQL9QIERZ9Zhcds2Cxj0
    MxxtNCJJk2pUpq5lAcGEwEEErMRd2YIUMCxroFAhbmLo9fM773vozTsvLuPYgQ1Z+Of4G3SL3piiyLyPf/L+D4xX6qF3aZqeOHGs
    yGeDemsyBYiNc73F0dHKl6lNXJYP65AkiZdInJCxbPokhGir6KQ0NulF8mVTlk00pBDWxmrkQZZ5nxoAhqMGZlYmr3WMRI6tUy9x
    pfKqNRk2NmdQOyesQlFVpS38BcN2UqR3rFCJ3vsgjUv6ohP0rkaRqIgkZAwLYghBoMYYUQMDcunUfO+p4ydnelNC04ulCYHCCn3w
    hps/89kvvfDyc374XW94wSXzJ5cROJmZmx2NNTOJNHDU3zB/zrjGf7vhti/e+I0HHjpEbnZY2sHseVVVnzpV9YuedXkQOFtIgIQo
    QsYYBhyBmMQHIuIWg64CElVVUsskBIoqLfkKkyVDRE3j22NFbEi5Ra4b5hB8hPgYOo49Y4mtujTARh8QwWSYqH05Zm7qIBxVI6lS
    aOtKDCDNElUVRVlHESEQc0aWRx5JkvgYTq6sJEkSmYkwGAxiUA1+WI57U4NRNRITk8TmubFib7rppp/5G693DolJFU0jIWWn1DaU
    TTu4oQRus8LvNdmaEFAQYNqBXFAkEwhHTq6UjeEkr2uxrGlifRPWFLYBpYg19hsiYiXxUTMwKWW57Dprjs2zvSh/EwBq5ybPOCmW
    +ef//JfO0EsTgVgZMUS+/faDTz09rMWoSSqpbG6jBkAJSi3zDkGQKNpKhZA6VssKhhBESY0xjYy27pp5wRVnG1JHsLDSgO1zA4Uj
    YiJYawDUdZMk7sSJk3/xqS+ORs43LI3vZYURJlUWVQSVEGJtrSFjhFQoqgugQFFZGUKqrDAEBphabywCIQhDHWlC5Jg5CJRJ2UQQ
    2ApRFBVlYhvBoobIMafMqSpJhDWO2gOmDJ0M9oJiiEyWyIgAQkyWyKkykQHBEEiFKDBHw8IUCVERqR0bUKPMIBtFo5ckcVAipkgE
    Y4MiiqlD8uAjRz7+6duePGzP3rvrk595ZDiS4XKd2jxGMSYpa/oPv/LB++4/sn//onGzEq3lVBplocQkkJZfyAYhgLj1MF0XvgV1
    tq35NmNRBVHHDxW7dULaUZgoVNWyIYVGkRhIlQESkRAsGYkKYUOWyUKZyRqTQKgbi2Ey3DbNEIkCVMiQtSAnMKTMZJls8F5FAGUi
    a61jNmQiqNGgpGTZJE7IgImIESOpWKg1pBqJNUTPAPtxhqHUy02Fiy88P0tRjgKTpVZWhJQmjQFMBpa/13UrjKBggTVA9N4PjaUx
    kkcO4o8/9MWVup/mG1RB0tYRWjY6YmUlIWoABRzAXmtjyXtv2flSZwbZ0tKDf+/nf2jn5sxNOq1MTERQ4m8+Wvf9mFU8811Cjc7Y
    veee/dkvP8hm2ke1iQvimVgJRhmA0IRdhiAIDEAdlGmdb69inTDuuOcRa66TaMhQs6xJ9pyB3ZkRglRV1e8XSZJUVbN588Z/+S/+
    6fElfPDPbv3Upz9/cuEYkBu1MFmWFGUTgpIIrDXGOCFpQlP7ZiqfJmWJOhkECYwOgAdDpLat94pqjFFUMpcIohfpCECIGCIMB0RV
    EVkFNxAxMZW17xTdjeF1kBkmCCuJCESZyHYlUiGPjodNWgm/9cUOAARWECvkNHYNWXel5wjn69jrbxsuL/3F5++89RtPPv308TzP
    p6dnLFzdxGGt9z104PiSyXOHZENklvjsM2NamAomdLLozj4wuTRH6sSTAWknt9rRLKW1R7bmvWcmNmyog151gkBQm1h0wCKBksQo
    3rfHyjARM4wR1RBi7aNLs26uUFu6WxGASLWt5zC3T4xQEiLW6aIfVUIIIQaVwAoiYtLEsQYvsTYkiYM1DAjHUYbw4muvfNUrXjo1
    QF0hSdLUtRC82HFYEVar7/Q9TvN3h64dKBSwSp6mpfpA+OMPfmRYq8umx5XXKM6ClQyxn+wpVggpSEgRCUwcY1TVJE3Eoy5Hc7Pp
    +efOfJPRkx+ofR8cFozBFVe+QH7nBmKBqLNpiA0pk7bQdgZa8PNa+j0RfVh7HWZmax968LGnDuOsrUYBoQb2OROqHI+rosistQCc
    M03ThCD9ginD3/m5a37u56656+6Tn/7Ml++88+ETJ5fLephxLpyoJk3QUAUFJbaXORoOl9ixtQknq5UKFUSAo0aNAQpur/DWWOYo
    LVERyHI7jmtgmUlClA4w2UKZ25DVJGkeozYhxDoAQqzGGGYiE5WIWdo6RaQ6xBBiyLLWqRPURhjAtqEZuiheQUIQUDQdbwRrRzvd
    SZECAMRAm6rM85xIH3vssfkNs97XosFYF5xl5oMnjlOaBmZr7bgu7Tc/NQKKWOMnWLtQr9LHTSg62jmAZ7SjBOjcnDLUKLEqRKT7
    1sq0XA7TNHUuUaPGkGGnqhJUVZmtguoQQlUDSF3WG+SN17a9076QStvCU2scsSoQo3TVLmIm2yyNicgapCzGWkOeEKB1bMbOatpD
    kfHc/GDP2TsuuviC3Tu2XHvZVtexucGHxhoLsA+1s7atiHTrXLvv/j3vLyiD1k212dpHD3zxC7ewnbFZ2oyDMaQqQUSJT4+BaLVB
    b4xpkaKqSoy6GW3fMTuVP1fb6zmz7wfjqCrO3WM2zPcOL3tr0hhb0ivtAI5Am8m3zp7XiIFWX0IUSFxKGn1tP/zhz/6vf+f6Kmhe
    uJZ75bk5ENYCYObWc/V6uSrGVWkzIzC+DufuGVz+v76TgfvuKz/+yS/eeeejyyt+YWmBPFLOncuiNr6OM4NeQIhRYmhDG7LEQmwM
    VCBQEUSNIYR2FI4URK3TYVUVL7VI24hhgxa8Z9COmkgUEa9gm1qH1DEkqMQYYvTWWiKJEqMEVTWWXOZSdrHxSgw1E/5IVhhQe+SF
    IBNWfFkr+U5Q12sBF6HX6x0/fnxqqq+g2dl5giEgqoTQkLPELNCsKERDVAmkxGvOZd2pBCMQVgF3so5GnGUdYRSAST+FAT6djJxB
    YpJEJAbRjqmuzaoIc7Mbg8TQ+BjFq1oSY4w1HEJggI1JiAyTiIiM67qMIkSkzMRqicm2/VZwi22IEikyIjPSNE+siRygMcag4i00
    S6nIkzSzF19w+Xnn77rqyovO3tPLHeIE7GoJVRUVsZ8lWT+J0tRV3Xai11CHf9nAhRURGkQtMwjKNC5rkwwOPTkqKzhX1E3N1hqD
    WEWAg0Qyq1hrWt2ArBJjdM4JYqiblK1JeOeO7ZXHcwp2fA7szHO6sw2C/hQuumjP0VseY6a6jtY56OpEKEEtCAwvtP7U0Xr2DFX1
    DXr5/Oc+f9vf+dvXp0yCIOKNzZ6TMDVJbAvCKopMFd6HJLG9LA+oAZ+n1qdUVg2Z5PKL8osueONo/MbxGAeerG+97e6bbr59//6D
    EAx6M0sry40ywNZa5xzYQiTE2JQlG+estS5hZhEJQWKMNjVrbCdRfBRSMWyMs8wd5gUxikQSZWJWEQkKEpI2UbQAkZEotp3rIQoS
    vfexjkHVcqrEUG4b0crthKyscnp3P7SFNCDS6dyYLe+4YjyK/d4swSwtLRVFMRo1/X4vMerVk5XGj5ViHUbMHGPMiiS0U0V0Wluq
    84CT11+NsNpYTmkVxbLKbs4kFqDJhOmE61GhkWP0qmptYhmqUVUV1Aw5RmXO+1liSL2vJTREsZc4X4/rsmbSNLNJZg1INPgwBoRB
    okGVRAILAeJDMJYtsTXEjq3lPPWZpStedNHWjXNn7zlr584dG+aR50gdjEVmuxHJ1SBGSFUjQXuZUVDjh0RkiFxiDRvtLhSnVdnl
    e69aGyJm5S6XZaXcGbrllvuY+3VAWZcu73nvjaHEZvW47p7WVgbUgrA6N8ItSoiESRzj/HN3mzNeQ/+e7Qw7LO00wcnghVdd9KWv
    PaQSEcHOoSVDAQks1AAKbklHn/2RFCAfVSKBeivL1b59ctmFXNU+T9aYrP/7ra7rLMvQjpDGWFUSYlMUSVVXzmpm0yRLqjqIGMe0
    YQoywJa59KrLX/RT73vRsWPY/8RTh46cvOOB/aeWxsePH19cXK6rGmSMsY6olycx1iGqeFEYInLEltSIaTdeW0dxIs65NE3L0ZCI
    SLoyVsbsnLOO6mZEFAEmIpVWb51UKU+LWGusjIIBMkKJMWRdiCpEutYzF5CsElKLcueztJtF50mkM7GODbUuRzPzc2VZTs1ssNZG
    HZW1gBVMJCQhpmnaNA1bil5SZ4x077L+Hl0Lft2JRcvk0VXQuOO701ZNAyq8vkQwUW0AlFVYIhE5AkN9CDF6IppKepEhEqipNXrE
    OqGQpCaW5SC1+RRbS5Z9mtGgV/T6gzTbaC3yrLM8z/Mkdc5kWZYX2czU9MzMzNSU6/dRFEgIsUHCYEaMCAEqsBaJRVOBDSzDOjAg
    ihC8j42YwImNogTJbAZwCM3yaHkwmH6mt6J1x+S7NFICMSkg43pskzTL3MkhbvzSHT5a4sQ5ItKyGg96PWYbtHaTeLbNaSYlGHHO
    aYiG2DJpqBq/vPus7Zn5PmnhfPd25ullVNtxyksvOjsxofSVox4JddDeDtLWlQ0Z8Vt9JCJ2NquqOJ1Pffwvvnzh3leRM0L+OTyg
    7SBhXXsAeZ4CEEkMkJoEitCoIhSJAVFQEBBjgCAxtjeD+VlcuHeHDzt+ki9TQghYHuLIkebAgQMHDh5aWFhYWFwuy3p5OB6NRrVv
    VJXJEiszN9GrqjGmreBmCRWFnZnakef5zNT0/Pz8ls2bt27dumWLm5qGy7o1VHkcPz5+6uCxg08dXzhZ3Xf/YwunqpMnR1XpJRrh
    pJ1DqZUiQUkjIXLUFmcPCLcepPNX7eQPAKHTBgyALiAqBlOjcTUel1NTUz6IgmOMibUSIoNzLhwck2NlJpESFuk6V7XmvJ61Q7uo
    opuGVyjJau2ZlIEap0uWAK2PI6ZgSFlIJZjQJIbyzFWL+/s951Kn6ntTybbtm/actW1+w9TFF52/Zevc9m3Ik5YPA8bATWbw6fR7
    AmIEm9WIDt4jVGhiOd1PGAFgMNSBOxEMpEX3UGnQeGHmJEkSmwhXgmDZMrNAYozWpFODnnZf9TnYXyBRqHE2qlHgltsO7Xv8uOo0
    24w1KikZgHlUjs1E2VUm9YGOlgYwxoYQmZm0Acrpvtt7ztnPmNH+q2Dfh+FnITIEbN1aOFMPx8M0HTQh0vqxhInDAtqZZ7M6/Qqg
    Kw0SWZeNl5cc86c++5Wf+slXbZlPqlAW9rlhbGia0E7kOOeI4H303hdFBoXp0DkkkaCeiEk8WHNrYBkIopAIJussVkbiLGcOgxls
    GiSX7DlX+VxmhAhVaNvTMVjlVFFqWWhgDCS0GqVIEvgGzCCFSMesYA2Mg6IEagWnzva32rO37NYX7oaC6OVViXGJusKRo7jrzoe/
    dsvXH9q3H5oGMqJgsFF4YlKORDzJy4RaqZ012hElIWWlSYJCAiBqrEM1u2GqqioVgpEiz0MIzjkSJaJ6XKsSO85tWlUVGW6bgkYV
    YKOC05CRba9QGQIN3UgMCSsmsOG2Xdhd/zFReQGBta27iWg0IFZRlcxhenqwYda+6C2vuuSCc849b1eawhpMD5ClCAGJQ4tI1QhD
    yC1AotJ1qnmVM0uVWIkI1A0QEimRyZ3AMeAUlcCLRIJhNgSjghhhbYoIiWBjEsdQwKOJtS1IVUVjUIFS4jKAx+Mqz9fP/68RJX+v
    6YKIRGmMYeuSccPjFdz8tfu8d0DiGxn7sUttlmVEtDIcbpzf5Jtq3bM7skxA2xTeEosP/VR37tx43jlWA+ivmBDs96GGZUCoo073
    6W1vfs2HP3pbHerE9NtBjAmB2UT+TKG6ukmk+w24hUCXvs570z6sNI298UsPv/3tew2s6tpEp/cRgHPG+9giqr57c86ucpgAsNZY
    azo0che5R7ItdYyYdi92nWnDYGZuWXCncgWaFlFkDKOjaEfLlyK0FvO3IBxZdxLYAa5LfxIHtFvUgFYLnwqVFOTayLQ9em0wEoFe
    hiIDEbZtw+WX7X3f+/aGgJtvPfSlr95xzwMPLa6MF4YNq7X5bBNJlLK8d2JhYXZ2djgc9rKsLMssdQCieBUQw9qkrUm1kAub0rhe
    oZbJFFLHMTEZYiUVBTklpaCN+EAGXkruDKRoWwEqysa0UG/DKhIkVImhPDWj8bJKcEZTZwlomkYkWgfRUhFb4JJLXJIkeeKstWli
    N2zYuH3r1s2bN2zdsmXnzu3btxdzUyBZ2/GT8hys6SAVXfjYKuNORApV1+M5CHF9vLXGGgIAxDAJwU2KO6RQYjIMhZABm3WcxQkS
    soAQZ935ArWXhSzLoq7SgnUahS2w5HvFjhpjRAyzWS7rNM9ZcNvtDyr3RK0SsixTEgi8xKmpqcZXk7YvQ2lSQ4wgQUTqUmmaLLHD
    lRPXX/cGg44C/6+UnXGHVddNmiepoVrww+++/qMf+/K4WUmLQgSCVXJ36ViS1/KRtiTfdRIViBKtdc6lZTMcVXrjl297z7v2OvQI
    XTSUJElLvBejOmf0e6oFfBvrQAWylvCv06vuUMIth2QXP0dQmNSVJ9iy7sHto1sWqO6rmgkF92RQo5PraNMlapG0a9+FSXmtkrNu
    YVuCsgBQUgOIUWskGPPyl22//rXbo7zta7ef/OCfffLOe/eNqqWmpqyYGi+d2L55y8mTC5ZNCDHL8vF4lGbOGEeWVDX60KgwiAyY
    rZlMArT/CkQVdWhI0YLvmUBMba+A1QEaY/TBW4Z1rKo+NKkzPpQSGjWaWzZOmDxreMU1F2/ftumcs86anurH6BNjd+7YdtbZpLZt
    Eky+McF2bGJdU6aVa+y4DyIMeyLfifISoF2vgdrh+7ZIRgQx3fFeJU1YOwOn32MdfwuxKp9WlaM1KgdtJ4pafzPBk5MyiVnry61r
    vJ5u7Tt9j6tWKcRgjKlCSLN8pcSnPvPEiVNlpOlVRfSWP4SVOvaT0+jIuyo7EJKkCI0X76dS9Geya6+9nBV1GWzxHHINPgd2xh1W
    miZV1ZjUJsxn7cCm+UzAlR+B0tWItNucOinLAqBIq4z9yiAQiQJ1bITg0v699z9+860nr7l63nEHIRFpBUIwHtdF8V3wbX13Jt1y
    nNDZEgGmq7wotZq9a2AhBXECuNPY2TCJrbDuldaTFHY+rNs3sqbJ2zLan0ZIoVh77urbTvTKWyBo93dmOAcrIh6J5Ze9aP6FV77v
    6El88Sv3ffQTNz762NO9pL9w5Iksne71p04tDr3Khrnti0snYahFt7IaIhhi4q601yFYSYnYAEpq80QQVVugAFQlqKhoPS6NpSxL
    sjypq3JluOScmxkkHBeNLPdSPn/vOa982YuvedGlWzbDERIHCBzBWRhFjCAFW/g2j6a1JJoCAHivbNB2GlNDhomB2MnUJx3oVCcN
    RuouF6QT5tJV70FrfqxFKj/j/vTFAJU1bL6252CCXAWtetZJU1PBMLRaCVonYrreeP1zvkdj5to3Ns3KgLLBH7z/hqBJEBKjoLha
    ppy8Sdu6Xb0ETkoxpBK9KhybulqemeNdO1HXKHLLf5W8Fb4/nO7GUN2UJunVEa94xQs/9Ge3xEDt2O3k8Em7lLr930EZBWoVpusb
    sqjGJgRjyCWFr0d/9P6PX/GCn6gtMtYWlBCCWMvOOZHnCp416cGvczfr+pLrIp11wL9nT7KfdtLX4/cntv43PEnxaN1T0EKT0VG/
    sZ6+7lvOX+1SztVpVSa1JnptomcBE+yuzeZ977nkbW++5Ms3HfjM5772la/ezezK5dCzPZcOTh09XkzlMUb4bnafWSWGWHtnLRFR
    CwCfDCNHIl/XsnbRZp48YHqqzwwf6tFwJXG0ccNMiM1o+fi2Wbz2dS9705tff8H5mWP4AAasQdJ6gQhLcAQhSIgqsNpp/zFPcj0D
    VWSuc9KqSiSASIghBk6LidLlOp3x1Txtlbh1NdhpOZRp7TFxoqe8es/rAAe6Ll88DYGzfj2snVZ+5oY/Pd1cV72iDv35vY4SMgdB
    yyD74MPjhx45OLPxgnEZu1XWwui15TIMk+nc09+DIiAxRke2l2fjpSNXXn6hY/iydFku/7M5rNFo2Bv0xce2bPyWN7/6T//bp9Mk
    rzvfL5NDBsBAGbCgpkus2k4MDBSKSEatA6IuD4e5Te9/4NCTB3HZuWjqJstSY7iuG2MS50yM6ypbz4mtw/itFkLQLdkwyRoAIH6z
    YSF55lJen1Ri3UKXZ/1tzZUJQcFCLcr/NPkVAZRXN1qnrtAq3iucsxl1PHSxiQ3B9hP3xtfsevHVu5488I5/9s9/5diJJdHQrIzn
    pgZNqKgVfOymgUAMGO4G/VVVYjsM1NZc8qKIakHc/qoFbQEqVVPWZWJptp+HUA1PnZyZmrrgwj3/v3/z01M5khQQaERuYQkEhOgN
    w1kC1IegURwbcklLltLamvZSq94DMBN3o3jGWMe2Q+uvO8ptLi8K4k50YjIQQwJYgXu2o9FvcUOrb4v1F4yuuTkxe9r/nu2tumhr
    Mh/TBnldOniavM53aU3wRZ4NPZRww59/Nu/P+6DsrFCtAKmFtrOrAvaTYgWfDgYSUHAu00ZVtddLf+RH3uEMbGqDr41N9Ht1omfS
    zrjDKoqiruskzSLQCHbvTl509aW33/20ryV21xPponxxgG3FPcBh0nDlDlbKQbUhAlnUYzXgPJt6/x994rJffrNzrsV5JknS/iAi
    fOYKhl2tbT0ZSPeDdNWL1W5495hnKQbpuifyai2eJlzf6wodzw7OeN1FupUZXHt96rxpy50JAsqqduwUXiUwa5IYaIgagHzzNAbn
    JX/6X//xRz7+4K/9xh9Oz84trTyd2BTGEpFICCFIFGeMMRQaDxJiMmiR4UxQYa7qIROj40tg0009Ik1tZSpnmGRcDk9O58nbXv/q
    n/7rL9s6i3ZURVspRmj0ZROaPM9Fgg/CbJ21sASwRih1Qo2rRcl2ZtqaLrNuVSnatoCxRGBudRjXrityOm+orEvcOjGn02pWp50b
    FsjqPRAYTPpNzykLgYWFnil43iWhNCEHPG37d6E6KdpZRnyPwnXee7bOOtx624mvf/1ecOYBJdFO3427umeHCBGhZ40foB2X16ix
    bsKFe3ZddEHeRO0nLDF+L5/l+2HfBz4srqoqSdOo2tQlZcUP/dBbbrvrV1c3m5AyxTYZnEyQoBs66yq8BgARBQmlb3pJlmQpRw6B
    P/e5m29/y8UvuXp3V3PhjodM5DlCk5IShXW8RasBUftfs+5e0BW3DLrFsf6va3WnZ78HJlOvpEZpdQL226zbZwRoALjjM2olU7sk
    jcCIBsaBYCUII7AQiJgQvBfhfmIkw3veeeHb3vavfvlf/ObXvn6P8MBHeN82uYlYoRKaMDszwwxmWGZj2BhjiMFhemrAzM6lzrnE
    ZVmSOGeMMePR0rbtF/WLfFwuJ6k5/9zdF1xw/qY5cI1JJCYKAYXEudQ5UXGctjvLB4iAiJlhGOAWOraWL7VTli3VfTvyrCYSEchY
    jegEZCZTX1jDzT/78NEEmLruyE9OFYGUWkbX9h4ggqx3SM+AlRHFdsh63claTfqk5QRa9yfulr92awD4plp339bYeoEwPnzDJ8a1
    lrUmfRulIm7bVo4ULBYUqCNuetZaWoM1EERf8pJroyLWY02Nsc9d8+o5sjPusIaj8fT0tIhUo+XpwVSluPzSjQY1wzOiKrNOqBcp
    oCMPYMBiNWekCEIIjXWmqoMkKkRwbnk0nuvNfuTjX7r00t29HMOy6uVpG2F0jAinn3jW730cfs3adbRudmL969CECWjSKF5dGM98
    t2f8f80xUSt1sLp/WpX607OhNfgkOjc64TMSQ6EbECQCNHbAckHSwrtEmWFMotpI4421oamz3sB7CKHvoCn+xS/97Xsf2F/WPC6j
    iMzNze3aMTc/DyLUTafc025ZUlBH8oC4LmJg6YRtCDAGjYcqsgQK1E0A1RapS6jrAFsDaIwITWOMFWiUKCAmtrYrmXsvZFq4lipI
    hVRVoKrqrBNpWsBtS2QBbZl22tSY100sMkhIeLWEflr1T7s61dqRf8YJa0vrzxhmbCEIk4RuQlNMICHt5CRajdQWVtaGe6eHZMwT
    sRydBGU6Odfr3239Ml6tiE14eIhdooyjx3HnPQ8rZtJ8KgSwSVRldbkKB1ZRMJS4G11oJdRb124JbAFrYm7DK156rgZMF3kMFT8z
    FvvBG8UzXFVTRfSSpAi+qsM4zTecGOITn7r3P/7mJ1bKvpgBTAKqmZuqWszSQiVV8GSmv8WkWABKof1lm/C3YxxWRwUv/Nav/R/n
    7EmKNLI0KbvRyPeKPBAmUE20xeAWEfCMeYjv4gg96/h8l0+nNllZxTSsaZSvlawIETUBBGOipVWFFgIMlKJ2WIquHQlQB6vkElR3
    rS3JIQ6RwQA1oKgMBUcyRm3LAAyGcgRFbYs4ahhu9Yt0RTGCAHEdXoknH1KpQ16sZpqkE7BlV7WidWOJk+03gWZ0RKEQAowydcg7
    WeeRCfoMYrPWTwRBbJVbqT0IE5heq7AzObaG1Kxt93UOqP0GmIRFk1DGyiT1MogTPUwmbVEQHXhF14pNvLpsnpGxPTsue+YKeeb6
    eXZ0vIq9WHu40urR00mdnGpfZy6NPlpDTFoOV7LBzNExSsXP/fy/ffqIP3YSNtmYFtOjcsy2jQBaV9sdc1ZALVlqYsWJM9ZVVcWi
    iUqqIbfl5Rdu+o1f/UlmNH6YuVRjNCb9gVPKrLcz7kGJEFUAWOc0SmhGs31c/9pLd2zppy4k1pSjxpq8rkKepz6MJrm3ga7uqNA2
    OEiZ1E78FyIhcNLE/i/8o39dBRCMiI7Gw16Rj8cNurPPqwwe30F6+luZ0jNv36213qotv4hCutotdUWUSROq/cLUBUeTm9LaI7E2
    JbPajuTWY7SXzbVe/ToWBG2nAjsP2F6QIeAIjsSR1oKGdkUbRIPoIA5ICCnBERwhIXEUHBpGbdAwgoFYQnsziAxv4AmRSQ06VbH1
    UClM+pydNyMouCXl6/6oa958sr8nHm11la56q2+10jpH2X2v9WioySMmVYjVhsnqce7uJ2Fud55l9UytjQF079PdJgSU33qFPHP9
    mMnttMfrNw2+J/8KUDVV6lIAzpmyLFU17/cXl8u8wBe+/Mjh48tlQ3kxrUJ1HZ3Noa69Bigp2iufssDGqETGWFv7anm8pJAsyfMk
    d/Dkl37qfW9lRTmqEusIJPRXLsI64x9ICaAQGoW6XjrVFiXmZ/Ca17yIeeSsGBaNiMLWpfrdHiBeTcGE7YmF6ktf3BeBuuai6I/K
    pbwwmFQoVrFK+GbR0pm3dvNi/Y3XRVwEy0japLjDZxkIQ7tYihVGYAVGV72OApJCM6CnWgjWy/BaIFVNISmEW7YkGFGCEgucwgEp
    tEMD6trFXKQrKjGtOpAWVwkmGMACjpEQ7GkDzDBAAjjAtGGMPNMLdMdhPaZAQQqrsAqjYCVe8+a09qmkE6Fw7esrUXsT4ggjMAKe
    kPWu+g2VLsjrvA9glKzCqDqFWYP7AgACyMMEmKAUFZMbCbgV/+bTwuMzvViw5gzXfkcMciajVTi8YWILYnLZ8WV89OOfXhn5Jijb
    BEwhNMTaMveiHb9YXXUKk1DtxyJi2Tl2lk2I5bg6aVy9bfvsFVfNiyLPUwMCov0rFFp1doYdFikQ0tSJUlNHMlnishjiaChve8uL
    tm3pxWa5l1vvvXNp48XZ/Hv6SELcRGY78zu/95EHHox1SJqIXlEQaXvKWMHtXv/LhVf/XV989fCurnnFaimNFBQZajsd6I5MXLkV
    u1KZ5Cld2Xi1eMGTOhEZhVWy7WKM3IUgq9PL7WITlsiiXa2Y1r+UrD6IonT9/jWckHbv0iaDROhGcnjdAzos7STuQ8uwRRpJI6nQ
    aXDZNT3S7jft87kLN7rnti+isSu1tJCKVvJjDV4gXezc+Tus9sVW2QEnCeDEA5LCKAwmIeDaEgIiOLbp+WkR0tqVpauL6xm/4Cmp
    clD2q9fXNpA2isQgRCHSqq6SLK8jVsYxyenDN3zlvvv3p+m0wlSNZ2fZUtQwQYS2r7M6KCDGUuNrgNI0S0xGoohVYmuR5Xe86zW+
    gSocE4DhcMj/80VY0sRSIUnCMXCoAEFqMNfn7Rvx1je+tKmOaRhZo9Y679XYtH3W6WH4t/7wyjBZFdInDoz/6E8+M5ji5ZEoEKVp
    i1ztRZ8hrPL9Ca+EVEhlAizgVjp+7Vrd5bKEQIik0kZQk2JqUPJKXhAwqWkA6zY5tVdJD6qAGohAFEZkRINgg5iR8li5BJfc0rZQ
    iBTj6mtMrt+rU0UgL4i6fv5WO78QO0RBl1PyN7n8rzPFOp/SpsDr0QhAmxC232P1RSbuYdW1tU/pkAhgPv0dW7cy2Yvt1MH/n73/
    Drclq6rG4THnWqvCDifd2JGGphu6yTkjWbItGVQQJCmCEZUkKr4mTBgxgyIgQVBAJGfJTQaxm6YTne69J+xUVSvM+f1RtffZ93Y3
    L83Lheb3uZ793Oecc8+pXbX2qlFzzjXmGIuGuAAKLVNB5vC3nNN1P+xEQhIjtSW73dfisTZPREmv4XVcl48iKOKiR42le0lSQ+2+
    JQA0EbZ0Xzlv+pb/+ICPjt0ApgxJ2RpmUm1FhPRoDgcECKkxjm1WBE+hDkjRoB709eDB4pwfvk2egxkxeUNc5sX1L8A6zoA1l8dO
    IshzshYSoRIJ0jT6yHPucKMbrMVwpCwAERVO84fttziUWMkN1vYXvQMf/PAX/vPdlw9Xss3xjNlhng8aXU4Mv+WS+bc3jloiWC7W
    tvPR+sEsBQpCCfNyVFJEQexSs7bLbX6HzJttE6gBAnX0LmA3rQuKIEjStbMJq7Ts8d0T0t2Hdnc0hDm1u91ZW6b1L65lHq7OT5U7
    WAmKxPM4cKkAHkDhv5V3ARCdMB+uJatavvl5Kd1tX7I0A+0nKJRaO+2jDjIPJee1zjRnV+0GWTpPM9Fde4tmgbuKuxzLPyfwMW/y
    3RrtvO3uH+yuJLVMITXMNgE2x8zjj//sFZd9YycvNuoaAmddJqqCRJyErumRT9KJ04qpp5UhNyjK2Iwno8sfcc59eg7GwVBqadCZ
    y+W7m5Z8K+O4h3yWsxBTFGUDGDBDY4q+Sn62McSPPv7B+/YwZEzkDecp8rWe0rUUvOsQrjy8NWko6vC3fvevPvflejBcjSBtbyNa
    lAW+g6ar33xIqzPHXTm8veXm0LPUQb3bJaFou1zmXfstVYdZlwsoicgTQtv6RmpJc9KM1VDXocbSRTAGMCSOxLbUoQVcd+J+XV0t
    EALggTDfyjPcvWkbgLRpRWIIQVuoWhRGWhsFRiCklscwR5x2VfH8tNvZSG1gRLT0WgaTLptrEad1FprnnosAjwKweCXMI6BFAyrt
    4l1L+Ijzlpv5L3TA2qaJiRAZ3iIaRIvUykjaecVu/uz57g++BrIRJTJJEZumEbYRGNX4t7d++QMf/hK7DeV+Ezgmci5vVWJkt3q1
    uIR5omsYQEpCYkubWSTH4bST9zz4IfeYzuoUGpFojNPEEikFoeP6gP+2Zue4Hp0VbIw1hpIgegHBZIUzdmPYP3zVzqPOuekNTx02
    08sdewJYrxvbU0GuyHsrAzV26qG88ecv/5ftMaYVCQTkiWLHSl8ImX43MsNreYtdtMIiJ4MGaIRG1bbgSwpWBSuRYg5CCkotuCja
    rUBDakiufjV2vuJ5XjCyaMtk2sEKAYQIJEUE0nzf3hg1pNSVrSFLgUxbmNJFpEMq0KQteZMSIQCJkJYq9W2pvhOoaE0uqE3ZusQt
    gWLnfzFP3biTPV1U/OeNeF3/UwKioG2BSNwVwDokWvzevG25g10g0HzDcIGAnSsP2qmMBtEgGUTTlQF1XgtbyDDot7NLfN2HdAUr
    Iy1mdYFfAIWUZj42ZW+FGJMGX7sAf/uKfxuu3YDtmvdMlIeoopRSgmmJ0/NS5tErhJkFyGzey3Lxfrx1ZN/a4JyH3W9jgEHPFs6l
    lKBIiZiM/c6ZvHynxnGOsJSgluBAQiw2Q6u8w3DJ+5P2r1YzffKPP3xlGKI/vNrvRZ9ijAqJKeRFZp0J0RtjQghLBzxq7YhIkJBY
    bNFrYvbpz178sj97S1aaBEStiEPdTIhBjCT4DjcYXss1L75ausM7B3YRAShKamJUZD4oLCU/AwkRB1EBGbjklVvVlJZCRCnEKqGz
    CAWxdI193b2cYjBISRPAVRObOhLZVq3LgBmWAPGKCF83RIm49SUUhRFoEmI4EjQT34aBBNRN7YMncAgBpNAAFvEzSCAWJprNagEr
    AijWzViRIKlVJ2xmDQkZUAqJusajNiyKVT1K6ufzQiBS1bquCUlSUInQyIDEaECSBJI0Nm1hPSISMKm3GcLQUFUMkIASJCIJEzKC
    CSExVBEIyUBimJJ6AzAhhlaugEJIKWmIQVKwEEaEBELUFFPwDGKoSqp9BdIoaVbV4O9CQkJ1DYJLqWuF9s04him0McZEFQW8Ylrh
    d//gb6d1Oa2yJlriXGHzrEwpsWPRSDSv2y4duo2MQ1KAUwrWaGFgpDmwd/iQB92Du86wVOQFAGvmmzPXM6b7cTRS7T6Bzt06tT0u
    BLTbXsaayvteaQ+ctHpkc3LJJVcdumrWK1eynLM88963lt/GmMxlIQS+RmFRQtJERCoJYBVTVbVoaHy6xc1PEQoQybOeCE8mvihd
    jIHYHMeSBGFefuK5sNV8+zx6kFBbRCYTEyubFJBZy4aEuU5S18lm/RCkzPLkYecyNpPpuCiLcajYOOI8JiayIYhlSiLWqWFVSJ0a
    H6WXD53NmyqQKDtukgDkGJaJGMYaSU0TA6z1IiBHyBnOgKCw1ijUJ+9F82xAJpvO6pSkyDJApKnZMawBYTyZ9QZrMUElGTZgk5JY
    k21v7RRZkWeWCZNx1esVKWpIQRSiCiZni1mI42ntA5gcyDCztS6k4Ixl4ul0aoxzNptN6yLLYkgmz3yKATKZVRFpmA8ZLDEaIkMG
    RJrA1jVefaLgBcTGmhhDkkgszGSIfRMA46yZTJs8y5TsrKrKYmjYpKQxqDE5kRNpvSbt5uZmr9ez1k6bKnO5dW57NHVZdlwhSwnO
    GlFIVE0R6l1h2SJRUnBkV0VHBn/wR+/48H99ieyGUC91zduE1nVqngl2VTtdVBfb+oiYzLEi1H6YOYTR+hDPeOojb332mmEwUtsw
    v0xk3ZUrun6M4w5Yc5m02HbeKFHbGxhDu9ccDNkbnn7GO9/5Pol5iuSDZ2taqDLGpJRUxZhrljbraE3c+q2RshXBaDy64ILzbnyT
    k0884QCRIdimRr9v2lCH6HhSaghzTrjBkjAcoGzBRElFgCBOYJIgy7lpwM4EmATHtjBM02myxmQLwpBqUMAV0WSE0qvxQXJnoGQs
    Ypw1fhy1SUYyLgMMyCWFwhQFxwQ2ZBgk8D5KCCazPiRb9DzsLBBMnoQlUGoQZ2KZApLN8ggnhCBI6vq9LAnNpuOiX8K0cV0GU0xq
    ZLlJmk2rxuU9ZqdCvX4RvHgfRUyv53yAdTytJC9KpUwom0UY1yvyvsvypDytIAJjidmNxyPv/cpwyGQkUZ5nwSOR8QLO7KSWfm/V
    mHI0GZdZ4ZgNm9R4Ug5CbCmRIWfANsuswhAbSZENLDEB1joiU1WSF257HLPSRip3xk2WlYQsiVU1zGAmVcMGZdlrYhpPJ2yssTap
    FkXO11n95bouH4hgOk69kq0zoBTizCfvo8wisRsmwlv+84I/fNkr9xw48/BmbfOh6MLpTog6ZkjX/QjMcWpXTSRKIogDOE2TP3zH
    257+c8+6u6F2vappab+6dJ3XM8A6/jEu5r6YQFpwAgHnKHMuI8Q42Vg1v/ILP5mbaZkHNuK9F5HWJktEvPetaeA1DkuMKKQCI8pk
    +j1P5UWXTX77918xrqj2NkTkJVKCCJjs97DPQFszZ3LWkLPY2kYVEQk1ME2YJSSDsUfZz9gAjNik9gqzsji0k2beHakgRGJcBEJC
    NUOW9fNiAM4Ygwa5MYUX+AQY1BFkIQJDHbGi1QudNNokSqDEWZMwrgCDokRRMBmYzNYCNWgUHrAFPJCAwcrabDYDqAmpbgCLWYMq
    YeqR9/ujGWYeyggJ01ld9pxzUHS7nq7MRzOMGzQKta5S2vbYrlAnuBLkUHlceWTaH270h+tNSD4kNiZ6gMA5tms0QKC8VgRQ2d9I
    cLJoIibKct7cwU6NKqCJmDU4vKUCa91AYRKkiXFWBwBNZJ/QX7WziEbRW+ttV4gEOK4SJh7TgDphVOGqIzNr3cpwHWCFMqFuqu/C
    ImED69p2pyAIxlm2eZ6vF8VaUnz6szu/99KXr+05NajJ+ytN8Mu3MGm3QTyvW1lcLY/1Vb0y6Esa50UzyJsnPObBbVrNCjPv2dZ5
    p//xrtl9G+P49hJ2Oy0EYS+UWmoUiZ2XfoVMqmKoY54V5tde8uY3v/2TZnjyLFCMQkTOOQAxRmOMHrVjsyupzZKSRGRQyzGRMSYD
    NGy7tHnLsw7+/V88B6HtQ4OxGqMvXH48p1PnlD+zyzMkgFJIDSwJGMgF+NqF9Zve/E4yvVZJpwlJNfX6ucbpwx50r9NPHqzkSD6l
    pLa0H/3URe/5yOe43Kg8ERmWZL3PtFrp6dOe8pCkoAyXHpq95vX/EWVI1IcoYlPY5gmPe8i+Pcht1/ZbzSKMff2b3nvlTj2uTVAy
    xhiEO93yjHvd+XSnSCnO2P7137+6CpbdIGmumhyqH3rIvc48bTVnnYx2hitrjeJTn936z/ec21tZa+otawNLRJo5Dkj1z/z0U0VU
    oyoxZxhPYTK89e1fvuCiyy665OLt6XYIwZp8z/reUw6efOMbnXbzm5925ukgoK4wKCExQKRw+WjUHDrSvPbf35vsYNr4fr8/Gm0P
    +uYnf+KhHLCSa94SgoXF4hWvec/hnRjFsTiVQDR98IPucuMb7c3Ym04VOKtr2Bw7U2yO8Jb//Njh7e2vf/3rqfEpxJR07/raTc+6
    yc3OPvPMMw6cfAIYOHy42r+3ZKD2dZlls3rm2GXHdf0A3kueM4BptZ2XhmCnXgQlOb7yMB75mF8RXiG3trnjV9b3Tme+kzXs0kBR
    0kW76bEF97aKqtorSOpDLm0/80nnPOlxdysdDGDaRzlpQoKyadUpvm2tgOM2jrNaw+7VsnY6drbbxYmIoS57NmNDzuxs46eecc5X
    vnbJZ756KPLqcLjqva+qqixLa22Mka/FvsPAMbGyJk5JmpAIJs/cmqr70hd3nv3sV/z+7/541gMAw1SFcHwB61qGAmypxaYoqBqc
    f8Elb/q3d9beqB3AlAkKirkT0vENb3DSTU+/TZOQZyY2KQH/+pb3vvtDX97xfbJ91VQYMn7at/60g+UTn/iQvIQahJC//R0fG0+L
    GDMVcRxWB/zIxzyELYJEiGQuS2TrGu/64Lmf/vyFgYbJ5M5wzk3GdJc7nt4rYMnWEf/xno8e2QoBRUousyZ34e53v/u4huTorazV
    ggsuCS96ycs2R/mR7Wm/b6BTidMT9pahOvKbv/5LZCFeix5/48rx6trwzf/+0b/461dn/QOjcRI2nCcfm9z28kvCBz/8P7GZnnn6
    iXe8/VmPf9wDb3gKj2spnM0zms5Cf5B/9FNfe+Vr3jbYc8OtnVE56Dsm8Tv71w4+5Qm3p9TpQoQoM89ve8f7z790p2kcSQ4NK4N0
    9i1uctppew1nDFGwKKqA7U38zT/8+9vf8xHXW7nwG1fc4IY32jqy7dhl1lzwjcvf+19fYg23v93N7n/fOz3wAbddWy83d/z6apa7
    rJpNB71+U9f/D4If39JIyQNFlKosc4WMmwY0cBl/8b/j81/0Mpvt25kCUfvD4ZHtQyvDtegjKWPO71tGq84tFVjmZzjnRuOrSh6f
    cer6jz7hbhnD8a5MiKCjriTAkKOj7+Hrw/hupYSttgKsomuAcA4xRsAadrPReM8aDu7DT//Uj51wwv5WNTxvHVxSUr1W+VBWZhhn
    CgMrImQS2eA1+cQie4lP+PgnLviFX/hLYkTV0XQ66A+O+/UqSJcqbl17bGepnhRKyAr0BuuiLu/tC1jX7ECiPYk2RnW2OY4f/fTn
    PODFR0Sbm2mNz33pwsQbAfvFngi3N5k+Fyt1QlCwQ1TMatQxTutkshV2A3IDzQYjH+qEACRpokYBshy9VZBbF7vXFCcbd6LQ3nGV
    zYJtuVI+xbpNErN1dnuV19VtTGozabjowStdcWS0NcGzfuaXrzhc2eJAVp4cdI+YA3sP3HR7ws9/8UvufNebzRokyBVXba6sDp/4
    lOe+/O/+Sc3G9tgmPiFgv097Z2FtUg881ov+Ka44cMll03e9/xPPe9FvbU7gCr7y8KYAk7oBAybP+gemvkfZCZNZT2hjPMtf8Y9v
    +uS5O42HKHyI5JgzBNiAjIo1yvew2xN0kJf7jEVSeOFZHXYmEMIzn/WCf3/7BxOvVrG/cfDMSw8lKk+eyvq2H6A8cc+JZ9vBCed+
    /oK/+JtXP/tnft0HrK1mk4nnzqs05pk9vqwY0l7pROrZbASIRxQyWWY3d/AXf/ma8792eDyxw+GBEMnHsL7Rm1abhMQQhrB29iSd
    a8uu42eXFSqgJHUz6xV2bdU97WmPtxZljqbxR3MUW6mxBES9JuLM93Z8F5qfRTvN76N2zarGDwYrEIx3Jhtrw63NaYq40+3WfujB
    d9+zyqMjl4tvcmctk2rq5AeIlXg5H2yZu0CMMYQ6kLIzGcEkYXK9rbHvrZzw1a9d8TO/9JdsyeT9VsSBjq4kdj0cbUvNUvPFNfdh
    fPMejV3rpM5Eb0GzFjDBGGGW1lWCp01qhJtkfDSzyiTpm2y9t3LS/1xwpQh6LqtnYwG+dN5VX734isA9awcSOClXTRTKZg2EnSiM
    BVlYl4uaKqFS58X4QCrWOUDhXJZnuQKzBimininQm9baRJOohCvZZTBQgI1lg5jYB4qaxVRG76zJe/3hLEINPK/80ov+aKfK1AyP
    bI6NK9gWkrCzc+RXX/Tce9391OBR5Mgym69sPPNnfvOCb8y2ZrnrHxD0vPdF5qBxtV/0cp7tjGZT77I1uI1DWzj/ovFjHv/ciy6J
    Bw/sGc/SYNBTQh0Du5XRFE00rhxu7tRZb9809f/0b/5l5BEYdRNab1BDmQ8upCKgSFw0kWJnGwvDyAp35XZ6wDk/eelm7VYO1DHb
    GfmM8gxI9ax0bFl3dra2dioyA1PsaWL/ksv9z/7Cn116CP1Btrk96feHVVXRt+GBefSCmS82HNVoufS73k8MJ5cZhW1ClmXDyw7j
    F5//8k+c+7Vy5QSXr1x5eMcWJVuzvT3q9/sAhESolRSP6FTA5mV4le6HACmxoOCa09ZDH3jPe9zthpnDrJqVvSwlP29EWDSYtiYA
    163z5Lswji9gKdDWe1uxEQu0KshMyLJMCULcH64kxcZaPzcoGM/6iZvf7qzV/SvcjI5kSM6Qwif1QkjEok5h0ZbFAEDUhYZmwtG5
    3KScmswmY6EhjVyfRqGWcv3dH/7qTz33Hw9NEIF6qvCBohIQmjSXcJGIFJG01YQhPQaXFnAmkN3X0dC2aI6bk64AqCAmpAgiZBCb
    k+kxOEKSN8XAq1UKRsPQ5vBGdBCw8pXzD49nCNN6WJQA3vqeD9L6xk5qLMW+VRtDZqwkV/bWY2K2UIEzkATOikguwKrNlKyTnH0r
    DMwCJaDMYRgGhsHGEBmN6mEkaGMskkAFTY211Q1jnK+Co15h+qmqQjULHpHw/F975Se+uD0KfdNbU40ONcl4bSi/95vPud89D5QA
    x4pUveDF/+f1H/r0KJZnoX+ji6+Y5GWe251BdtUZJ+Hs07IzTpSNYnvokgQQD8Ue9Hry9nTPS1/6agUy0znp9ge9pk5rK+uOU2hG
    rpfXbLZj/1Pnbb7ijZ/calAOyxQ0A+LMrxYHmsYFUIXEbBmUZtAIBTYneP7v/MVmWpm6jbFmQdxab5C2Lj99Dbc6Jb/5KXyr0/M+
    b+UccturGk68MWn2n/vl7b/7xzc3QDkYNBG9/jDFeN1uAFJCJIQWNVpxq0UD+KxKs1qUMJv51nCHSPOcKz/Js+GoJuPKQyM878V/
    9+kvX+ndsGY705gNi6SSIspsNTYt0RTKUTgIJyERgoANkUTfKw1zbJoxVAu3Qj6umdEJvelTH3v/EoAgL3tRU+KoFECBSQzYICNk
    gLneCSR/FxRHF3HIMfYhejVKWqtL4Agv+KUfe9azftfATpomNGSsm8uYya7s9lwaSSkKKWBIwW0rrDIA4xABZXPVzsTYtfd+6AuT
    5//TS3/zx05eJSLnA6wgqRgYBqTrPuNuJ1iX+uK6QdeZQjcvH8x7ATvRwcX/KFhAZS+rR+Mem30b+76xc8iVedk/8JrX/NcvPvGu
    EF95nH/BNzw7U5aZh9Yz6jQtW8EZkWUXMGUBJ4IhGAjv6ih0Opyd/oLOZelJpGvbSYtrs4y69tCi31sNtZ3NqsKVw8GKK/DcX37N
    f33y4mJ4QozNaNxsrA+2j1yxf//guT//Ez9w171GsL15ZO/GegSdf97ogx/58t4DZx2e1AQ+eHB1+/BFZ5++9hu//uzTbwQCqjHG
    Ezz3l/7681++0tr9nOchQqM7/4KrPvZfO3e47ap1JihCSiml6L1lSRBwCoHL/t5Q7/zz699185vc8IF332sspYDV1dVPffnKvafc
    fOwnJBRCYNGyQGYwmuHd7/vi5798UW/tlK0qVpXfU67Otq+87x3O+uVfeMyBEwCDacBXL2h+5hd/ZzSrhmt7tnemvXxQx9knzv3K
    VVvn7F9BYdA0Ic+z6yoZvOhgl6MX0GhcrQxLAqLvKh4hCGJwOWfZys5MbZEfGeFpz/z9i6+shAcJuaoV7pokSDtBbWDX1RJH6wsa
    Y3Z2dpwzg7KXZdnO1tZaX3t29hd//BtrA2QWDIjCsWF2864mpq7nnK5/fYTdjXT9Ginogf144QufnmXjfkEImtGq0QEBhACuhBsl
    UeK5Fc0iAF66FpKqmgLIy17tY14OVtf3ffSTn/mVF73i01/0s4RZRDTIShfibDY97JCMwgobMaYznG6rAp0kJneiTcRqWE0nwLBQ
    7Pz2d3+5dUTOcjrr7BuJNo2vmM073/PhKFDJvMfl39jWwEgS68mtb3kWsEiN2yu9+ifaikoHMU0r0suguYyhLAnFtRN47LY3M1IQ
    VbI2ixpFSal/8aXVn//ppz70wS/1evtmFYAiz8vJdHtjo/jFn33K/X9gbwgAxY2NtQQeV/jghz5tudg8vKV+2s/TztYlpxwcPPcX
    nn7TG0MEs9lkzyr2bOD5L3zSCaei6M9mzWVidjjzhw9f9d73f4gNmCEJUMeOozTGEKkgBg0eQhrd5pH0yn/69/MvCnUCHLbHo/WN
    jel0CqGMHEWQdlBe13jve97vbLF5eIuilHk2G2/e4bZn//JzH3ODG8C0SJThZjfPf/o5jxisTMiMRadCI7bNhRde+JnPfLWbmtaB
    5roUoQUkcAkuoZXiosV6WRkWkCQp5hkxSTWrmSTP8iaUW2NbFMUFF8mPP+nF2zshJaPkOjlGzMVa6ahWQYVVOO0kw0RJBAy1zgw0
    GlL2s5GzI8Pbj3jkg884Y6WtCVcViOCDn3eSzEv182+ud+jwPTyluajTUS8mOKsQvd2t1n/pF57k66tW+nbrqiMZ5UaYIfNONFWS
    uTgc81yLck4jEHRu4DKt6vWNvT7IZOqHK3s/ce5Xnv6cF33401MqMY3YntXOZsN+34eqa+5dJHek0nWuXdO4OkTN//K6qtjEGHtl
    rlrd+tZnOpec42kVJpN4ZISo+Op5s8lUkRyJapzc9753Y5WFLLGQHH0iCxc8UU4g33KczaJ7jgTU/ZUsBASxJM3ZcqKZoVQ1Iank
    g14Tzdvf8dE3vul9eXkQvAIUebaiZEMKv/7iX77vD6znjNwESdPKN7VHluH9H/xkXaPMe/v3rod6R8PWqaes3eG2uUT0LPYNeoqG
    yd/srPwOd7zpka1Ly54Vrdipy3oXfO1Sw2hC14JoHAul1ryVkxbW1pOpM8XB/ad/9nOX/MVfvTYZjDxgXdHvNU2DkCxZy8YSp4AY
    wYRLLrm8qeLejf1FVqamLnPc/a63PPUU+BqFw7B0PQtN+OEH38KY2ZFDl6+uDUOaGcdlf+097/mIKCqPrDDfRv35mBbqhYKjRG8I
    vq5CCGWZ93qFIG2NmyDcG7j/+tTW057+K1tjndZk3KCp49KtqouuGVocXlv7TttKuYJUCSFhpb9BYnLD1fRwZrbvfrebPPlJ99nZ
    hm0Lf22QJZK54nrHEL2WcX3DUM0JLFOR6gH3PeNnn/Mjk52LTz1hD0sywqTMytxVGUVJOqqE2rmKC6OTPZKkqdfrTaZVFGZXjqYe
    piyHB2dY/bkX/cFvv+xdkZH3ikM7s5lPrO4oeTlqJT7mctp6ra9ukAKi1yjo8X8buc0Z6uudM8/IYxz1B06UanHnX7wlDq9/07uq
    inMzcOAb3mD/ySe2cNyZZSvJonBL3VJdOIzFxEF5jsKd5J4sVECVGGqhlpU7r4FW/VvBIBGEmISUnIuSf/Iz50UZKoY7Y58XK+Nx
    4xspy1zhOUFqFBwLyzF6Mph5XPD1KxpvnC1CMzOm2ViztzjrJAnIGZlCA892wrDIdsa49a1vXpZ9lxdKCEms621uznxACLAG3tdB
    KqGgqpadY1rr9bRpjPJkJEprH/rY//z7Oy80OUzRm0yrftkzZDQmBohIFSGAGXUVLJeOymZSI6Z+bk4/7UBTYdBDmEYnYCBWqAPO
    uslNe+WaijUmD4l6vb2f/dxXEwCjx4p7fWtDljCrUzMTIqHolcg4V4YQFBhPR434/ko+8fjbf/rS817wsjoNq6aMUsZoFdk8uNOl
    A3cf/HwXvpVgtYvcU5Vms7pfDibbh088ODjpBPszz36Mc1hfAwiNR+ZQN7HIi04doHvw83xJXx/H9QuwGPDNpLAc6kkGnPOwWzzl
    xx8MORTrw0bFiOHk0Ck6aBspXH3Tr42wQghJ0O8PqzqA84Mn3CBEtzVq7HBvMMPXvvldv/TCfzr3S9srqyuCgc36imzpSLsqAbu3
    +3F4/LBCk6TQlKXe4AbIs1hXO3leTiv9wEe/WCne96EvAENSI019n3vdSRXUic+IkMpcpLwdpK02ZqvA2cET7ZbM5rLu7aattk1N
    7e8TzaVaqCUTi3EuJ2sq36jJvWeggNqiyFIKISRjcyF+yUtesrMFNhjPdgy03+tbg69fWDGX/cH6rPGHN4/0B2UMs9vf9maFgVTJ
    KZpJ3DMcaMLGEPv27ev3Vg8f2rFZ6aP6YKazIALj2s8yRfXMkGRZLCcqDZ18YLWZbcUYgZ66vX/68ld//jwkU4JNjNGxkRBFokgy
    DtYiCqpZWBvuvewbV6agq8OVajo+eGCl10c1GxXGGoATHCN3uNnZtyry4daRmTMD39BkmsYTD0VmqG4q5mt0yvq/LOnFY4SXllKv
    LCTBZTbv9aZV5GIlL4dfOn/r5X/7oZf+0St3qjzRmpfC5ivjWVhb35C5O0an56W8yAla1/T28YPOg4pSSmvDQT0bNbOt/Xt7oTr0
    x3/wwgN7MB5NAYgKaRSBczZGadqA9pgs8PjLFX47t8z3+gSOHRm7DG5Pb2175/CgwFOfcvdb3mK4OpxZDSZlLCVJ2ZVjSOeukLzk
    Z9sO7Q0HO5Oxy4pebzCZhu2dyide2di/M5mOZnVv5YQPffyrv/yCP37ruy6qFbPYCjuxsrZeNbwrkNJtoACp/QiXsetopPx29JNS
    VECcS3s38AN3v91k54hzWYjuQx///Ec/jWnjsmxNPTjGu9zxFiQVEBfif23kT7tS4MxiSA1goVZbrfelJZfQ9W7MLS2IlFgsiV0I
    9xmFIWUlsFXm2ke2mXG5zYyiacIWUbW6OlQlFTsap9e87j9DwrC3p/JNPZtpAqnaPB/NqnJtI19Zq3zwIVniNEO/MM0Uq0MbPTSh
    CfB1mE7i+uqJUMeUS8pFczJwDgCcM8RgayS5FDOjPCjtDz3sXmtD6fcMW1M15vKrwp/+5WsjlSGRiIBUJFrLQRIY5JAUKWHzyM7K
    YM+e1QMSQERJNAJZ5nTeLywBDhhv1cmbPRsnS+yRDH1DhgsiCEBWWu2wb/3DbRGqxSzTic12AjuSko8pJASBKa0wvnDe6Hm/+rLX
    vPHdbnBwdd+pVbQ+2UkV+8OVrZ3xbvGRFl5z7Scvi4c01CgcNIPaXlGOx0dWhry2IhqO/M5v/uLBvVkG7F3tV9W0sFzkVlUMIyV1
    zumSv8bSir7ejesXYFG776VUzZphMdBUrxb4tV99yj3udrqhkVVvhEwyNPfSXUKr3YcBqwKo67ooiqZp6iYWeZnlA2PLWRWczbK8
    XzfUhPLKI/K7f/yPz/7Fl3/9CjQET/BEkUxbjTZqTOcxF1rhp3kL99zc4KiYi7+9MmWWZc6a6fgIgEc8/L6FSxARthdfvv3O930m
    KzdStJx0WLoTDoAwIwqtklSbES/mjefxEWurku5E3aJzU8FpXqed64N25uhdBNnaOShE4JwhohhFVcHE1vgwq+tN6ypjdoi3m2ZH
    YorelOW+f33LBz7zhctnwUhyg94wBQwGRV1Py5XeFZuHxWWBTCJz6MioVwAReQ4w6sZbi9xBRAg8nVahDtZkTJk1BTNUMa1hrQWE
    YSQ6TVnmentWB4991Cln3HgtxcMhTjgrbbn33M9f+LWvXxkTyn7pk4eVrMyihJC6dMy6vFeuQE1V+VkdMlfWTSJAGOQARkhYXUXd
    AGpEZPvITjPTwg77Rd+yYcKsajJrk4br+vl2AldtebNbRUEpCYvNzeY4eaBR/M4fvP1ZP/vSKzatLfZm/fULLzsUwb2VVWIbkhpn
    5+t7WZCvczyjjm8FgEkt1JCa0FT9QiFb0+nFz3z6I+5yhwM9C4tECHlhCLHxM2NoNJ4AuHZ62fUOs65fgAVQ5grA9opBPysK40T9
    yhAveP6j7/sDZ/fymdEqI6VIGWXNtO6XJdGumKjOawzc2SmQMgs4JopJkzDBQoiUJVlrB+RWJ03+hfOPPPGZL3rVv37xih3MIhKw
    M2qgVM9qIgUkIUVEwVyeipNyq8g+r3t1Qcu3mjSqQlVFRERUVSQqYmFw8gm9fesloUkpqe29830fndVqjHOsw5IHfRQF2ATRRiQx
    g5lbnT1mUEJqammCgc1NXyUjHTQeMPACZR57P64jugYpGMPB19FXhiX5EAIMIyaQIoUYY3QuAxlAysI4G4qyMtmRxzzm7sBVuasZ
    kcXF2Js0K3/456+fVFyWay055Aancl5KnWaUm1mMnrJycOATH/9y9EgR0dTCs3JVgwYFLrn0AjaUZdZazqyNjT+wb59EhADnoKqs
    nCIRes6uhBCbMMozvPhFT8yzrZWhNE0TolHtp1TarBdCEASbmSY1nJsqoknICpRlLhpjjMwMtU00o0kKQOV12hyBrbiYTXzIcnz8
    Ex9zVnul6WUDiZRic7Ozb5wSBmUeNX6TFIkIIhpjSkkWpS4i+FnXlUakITZAEsSIJGzGEbZvPv6ZnZ/++Ve/4d8/uTlen9RrMzEz
    3xT9Hgz7GKTtS0bLaBYGunRe56VbIKXQ+Fl/kFtGNZn2swFEc6OzyeVFvvPEJ97/nB+6lSVkAKshtPIVkmcOwHA4sPmxfKvvloD9
    tzOub4A1bwxoyUtqDIzjkFu86AUPvfnN1i1thuZIaUxusrXBytaRTdWk2k04ES3FWe3RFlFPV0wgtUYMQAIT4Gq4SoppGvzOy175
    zJ/5oy+dn3YqFL1CAZcVm5vbodszRgKNqomPPkhovV1wDVX2bwmziEBt/xGTULKZMwwC1oc4uG9gEZ3joOwTk3He17nVR/zQAzdW
    kWJNGqBBEVTn+JzAipUhhv1sUGahmjVVQ8iahs87b+wTtqdSJcD2ekV/5HHo0Kaq+npSZro6cKz+pAN7YwMhjEdRBNaxdWY6ncTo
    88xV063p7ErxV/79X//qTz3znne8ww2dmZSZ5q4YjaKXweWH5TWv+1BS1DUMIUXc8tZnAOOyz1mZubx/5VWT8y64clbD5IAppinO
    vG8SKfDu93wwt1kKYoh3to+sr/fW13pEKDNkBkRkiA0ZsItCMSXiOBzgwD4856cfE6rLcpuKrKhmUuSr3gciMpaThsrXguQKGIMU
    sbLaV3jLiUjyPE9q3vX+jzUCU/Ty/vo0+GkIeeYuvgSzaUghlHkmKTkCSX3GjU9tzavruuZrv1/q2jOTc6a1rU5JUpIUUJQIdUop
    qYDYKfLRDAH9rYoj4a9f+YFn/PQLP/rpr3Hv4MQ7casCq7tUeNklMSwURBd3rpLM04t9+/aMR1vVdHpw30asp04DYTLsx3v/wNk/
    8eMPSKEuDFigsb0fOkMAPpYcdswyvj4SG653JySdFzk0AQlWjFOXM8ocf/gHT7z3vW40KKtQ70jlta4HRa5IgtQ1AAHAosrE7f7X
    Yqd/XrM0pJbVtmXpoFInmoWM8xP/+4LJU57+4p/86b+85ArUEU3E6sZGFXSWdOYBZP1y1dqCiBJESa+FW/ctYlY3YuvwQhQ9Vnu4
    5U1uIGnsbBLotPHMLKkp8/DwB92MAU2+SysAgIkMtLOGXx1g73oxmxzq9XhlUDpbJs1f/tevvuIwequ8UyECmxXe+Mb/ms1SCKHM
    SJqdWB+KfvPkk/dah6bBxkZHJHaORD0TOUPJj8+84YE//5OX7FtDafCspz26ri+r6yMqsddfC1psTem1r3/3575Yg5HlUMGDf/Ae
    oTmUm1moxrNpvbH3pAsv2X7Fa88dN/japTM2Ky4bhmjf/s7zLzhvM3rnTB6831gbHDly8R3veDNjkBTTCiqiSUlIVaPGRkKQ2gDD
    DI94+M3ve+9biD9STzc31ta3t3Z6vV4dahjAcOfdCwDILG5367My41PcDvWOscomf98HPvnO91/VKI5MOdnVzK2OKvz7Wz4+2gLU
    aRCW1M+txtEjzrkPAynBWuuT/6YfKACkpCklY9hatg7g6HpaNZOkBqbYnNqst+IT3vL2Lz/ysb/1t6/4j/UDZzVaXrm53V9fTRyF
    o3SVh7RrMNIpai84M7wwfFTVst/bPHKIKe3fM5zsXOVnR9aHhuLmne9w4+f9yo8YrtcHWeNnEGS2zUd2db2WTcjlOGwrfcfH9Quw
    OlY268JwqfXbdASVycDhV37psXe/6437edCwM+xbFq8aQarHRDrShVpHu6G0hUmrc/FYAYONqIliopRl/0STnfC1C8c/8uPPf8lL
    3/j1y5vtGciU1vSKbDCu0qSJPgloIdl8tM7/N42fF+3b2rlfzQMkkqjCZJFQMO5+p9twHDF5NiapmAyZS3vXixMOwDew1nIrPa5g
    GFbHgCEYQpbhh8+5X0qbKWzFMJtMRpkrL7r00C//6l++8wOXfeoLhz/xma3Xv+Ej//iqf29q9MuBxqqXizOzYT/e9a6n9ErkOQ5v
    Chs0TUMsZc654xTrzOo5D73vzc8c9KwOgBuduna3O561MhBGSDG4omw8po39i79+jTDGM2QZbnWL0298g7U4uXwtB+qZeAgVr/iX
    N//WH7/j/Ev9uV/xb3rb+X/9t+/7vd/7pzw76Buwai+zjObUk1ce+IBTnYWmmDsUWYmgECWTyInaKBQEABoj+nM//aMnHyitzjTO
    ijyr6yrLbUzJGOOca6Pt6ONKDw98wD1yM9u3Jxv0dXvzciKqGvnt3/3zv/qHT3z5/OYTn95+27sv/L0/eMs///M7YirLbNhUs5W+
    nU2vuO2tb3TSSagq+FAVrpeZ3rV9vnnuRNAW/owx1AXhOq2nQUI5XL3iSDXzcD38+39c/tSf+sff/v3XXrWZKx+46ogv+mtrG3sm
    1diH6XLM3nkUtTazAB1tfD030GbvvTFc5nY6OpSbcOL+4shVX73NzU/9zV9/CosMnamabcdm3u/UsUNJv3kAddy2xv/fxnFvzbmu
    Q5m1VclIrG0izSCR0qDWnX4+fOHzH/v7v/eW933489EHw8aQS1Aig7l4T/sM6XoiVFq299zorfUpNcoCZcskYCUmJcflkc1RL3P5
    yt7xuHn7+z7/3o988ok/8ohHP+qOwWNtAOPy0mI8qdYGWe5o1+6tM7n4v4hRtgrmnd+yYlG9Mi4LIRbsLIMi7nybM9dWs81ZTZyz
    YyApZmecfmpKMARjzDwVZRVS5Y4cRIHUPuYxt/m7V7/x4qs2R+Nq7/4bjDcn/cH6l8+74gW/8WfW2ul0bCwZ6knKiJIl3y/NeOvQ
    Tz/naf0C2zuTjeFgOOQqIKWkPGM2Cofoc0q3v/lZawWCl/F4lPXW/uB3f+rRj33JVdNt1aFB5vo5Uu/cz1/wqtd99fGPvYkA6yvu
    2T/xqBe8+I8h8dQ9p1z+jXG5sVH78nVv/9Cb3/fZajLdv742Hc+aure252CO6HWa4mRWbz7tuT+/dx+qWdMrLRjikbEL0MgVnJIm
    tRaASqIwvcH+wc/85ONf8nt/v1UddsV6VVc9u+KrpEqmDUcEpJFhb3vztcc98v5/+Xf/YvvrJ51wYDyu8qJ3ZLT1ilf/xyte8Z/J
    p7xX1lVirK2v7pUQ+0XY2f76ykr9tKf9RDX1e1cyqGUYaj15r2WICIB5SqgxRmXNi+G48SFhdX/vi1/GS//oH752wZGdGbns1Igs
    UUzaVI0XbQixyLMYMCfuzD0Iu+3v3U3wTk2AlCDOZnVdrw7609FWmVHhZLzzjXve/SYvev5TCou+46oerRRDgkFCijCWoPZo1vOy
    7eXuD7rVej2DrOtXhDXvGVQlAoPmXDlSGLBWvp8hZ/zczz/sMY+4R6yvIplao2zaBrk21T9azmEu6kDdpnLS1q5PiNSSWquZVWfU
    Rp96vQFxdvmhnYCh5/Urd7K/euV//OiTXvqvb/zE9gg+wgeo5gqa1qHrbQR2rYivy+gSQsNJVUSgTAmxxuoAN73RSSIz39TWGh8m
    KU1vdaszoMgyhJCAlqPKqqxiVKCSoD7Lom/wcz/35H5fVtbcNy792srqcHNnaor1aZ0HWe0NTyXeEKxYU5JomdtDV11487NP/sEH
    3Lrxfn21V9UVW0iCc04k+DBOscmYJNTqa5lhkJn9w/WeoYxx+9udUeShyENsNiF1SMLZ+j+86s2XHUIQ9DN68H3OfuaPPnyVZ5PL
    L8pJ6iZIv4+VPcGsaranCpli0O/vO3Rk0/tRjIednT3hcQ+8331vGAMk1YZN06CuQka5SvI6jjqO7BMpARbZMM/rcfWg+5z60Afe
    NTP1eHxodVhU1YxgNIpGJSEHlNYaQWHwlCfe+x53OduZ6Wh0WeNnk9n05JNvVDdG3brme5XXya2Uw5XxeOeKKy8mmq2szH7ymefc
    9lYbKwMQgiGEQNXsWn0tVUFErfNAG0FnmXMum0W2ebE5xvN/7XVPf/YLz7/oSLQrxcq+WeRGmGxGnJGizIuyyELj564lXQREauf0
    OiLlJa6yLJzimLKmCZaJ1Idm8/a3vdGLnv+Uk/eDJcZQrRQDKIegYMjcW46ueY9Ir/7d9a3ufv0CLAAJGlu3YVbiRCYyRwLgi2Fv
    35HDV/VK3bOKJz7pnj/7M0/au16wSWSo+/zmeRaRaatXbeY/lwoSUAQnNUkJKmxSzjGzyRhhR+JIFKEcDrm3ujW1xdqZ29XqpZfX
    f/+KN73whf/wqU8eEcVwyFWDPC/nNnfHBvDXNhS7POl2g68TrddoM8eckYJFDXDnO97KIKgGNoDGwTC//e3PNgwoRMQwG+OYjSVr
    wIZgQIQgaeIyf6e7nP5Lv/xTqyt2z77BZLK9Z8/epGU+OHjVVqh9GdJwPBamUpM29egB97/b3/7N84oCRFWdtsqeEYUI8qw0TthE
    RnSGHUmsK0TAw08CBRDwtKc+bjCED0cyW9WzI6urq+NZmMbsb1/xzskUBmpq+bmnPvQZj334Wafu31gdCvTweHKkrsehptwemRxG
    FmFDUWoxqE46JX/MY+/9kz/1kEEPeZbWVkoGJCK3QxIylIhr2JlyI0hJQWLVy0ovr8Z42lMfcPOzb7C6YpM0xJpZC2WDzJBlIDeQ
    ZoIoucEf/v5zHvOoHywK2bMxYObLrjikxjWIyeksNWpVeIfz8coqnXrq4Bk/+eDHPuZ2TZgYbsazq+qmcobKnK8t8SeCMaSqTRNi
    TNZySrIz1suvwl/8zSd/5Ekveu+Hv6D5yijEnWayU2+bvjY6TlTnhYFQmCXUzsSM1HQIJW2xta3DWtYFBbUjzwtHkMQoIcQy7/X7
    /eSnD3rA3X7zJc84cT8a32iclq6MgSbj4KwZjScun/f0L6yugaO8sa9v+HS1cX1LCUlbTeUushLqGNpt4oYDew8IdNTMVsr+jz7u
    FsPV1d/6k1cCmgQCm5SghpiISGjuZw7ltl1HwWBlgrCqtH7GpFC0rjtJoEW/NxpPrXAx2HPZodGBjX2zrWZ970kf+uiXP/Kxz5x1
    5inP/cVn3uLsYZ2QWTdvFZ27oC9YrOhU/JXQKUDoEtOFYQSOm8KMnYGFrOQua4JzMJZCk25189OHWW1oRgZUxFMPrJx5o1wjQoNe
    ZgrrKY4ZQpCMM0cgIpBjC4Q4zLMH3e/0s89+3t//3b998P2fUj9LTXDl4IS1TNMOKO3f3/ezQ/v3Z095ytPOefjpdYU8l8IVSasQ
    vLWZAs5UJdeuzL0f5Twr8mZ1YHKHGFCWLhlMps3Je/NHPuwer3/DfzTB799Y39q6cM/aSvA773nP2+599733u+etneNq5H/iyQ94
    3BMe8Dt/+r63f+jj1sZESHXgGPuDqOEKl9dZX+5+r1s/5WkPP2EP6rrhwsY0M5w1PvQKR6gzM2HjMydcIMQ6R226IqAh0cEAlvDU
    J/3Qr/7GH06r6Wp/pamnGUthfc4eEcgYgLEKL4XjZzz5IY/84Ye89A9f+4UvnJ+gja+8jgy7lGSYF9PJ4YP79tzvXvf5yWfeNcuw
    M97cNxw2abLWWwXyugIT3NWc+giqBFEFsTAlcuzggcuuqj//1cte8jv/XKfSp77J8snU26yMPmWF86HK8iz6JjZNmeWW8hQjkxHd
    rbsy4m7GoHYOK4SFbDsFB2stSbMzG1/x6Efe6xd+9gG9HN43vYxc1vdNyJzr981k1gxXetPpuF8Or3bTLVLOo8ZCCuJ6NY6vpvu3
    MRZbFayLPim0shnt3Mnc00IFnvDfX8fzf+0Pv/SVi9c3TgWtHNmuesM9UVLdjFzGxqiKQMiqocRRAZvPyxB6VDEeglYIARBiKCnB
    aMop+Nkoz7jsuWa2k9l0z3vc7uEPuc+d79iTgMJBvPQy8WGUu1ZVtSdgX0cAZWFV4X3IMieihjpaug9pUsdxHWsP3yRLtpfb1ZVs
    pY8E7Eya0bSpPETUgfu9fH2Y9XKIoAn4xhUjISOJDXOvNKsrtpeDkIwxSRABYijgA3Z2MBmHaeVjlJSSiLBKlmXrayuraxgOWnPq
    dnJbe3cWmNkEO6Nq2tRsHdSllAa5WV3JV3pIgpSULcEgCKYTTKZN0zTMBuREBKLW6bBHZWkGrRKLAoRGMZ5ic+x3JuPJzoRBmUWe
    58P+ynBY9obI8/Y0WoJlUCUoJ2VVXHZFLUoBUTUZy73CrvWKfg7DSKnri6wCdnZSVTVJjTEGSXJHq0O3MgAz2g0OUQZ3lq11jckE
    k0lV+WY0GkHZ2qwoisGgNxzaskRmEQWZgwESgtFkyABGlJmQpHMJWKyfupmxcbB5FWEtph7/9pbPvf5f3/6NK6c+rUSUaBfVgq5z
    bOP8HCy6blCmVhtpqbG5CeJcaayNErw0NlNyIvVknazTKnPNOQ//gac++Z7DHqqqKTLkxlLXgAVt+7nIAzCatffUUedwnHWfv4Pj
    egdY3/pQRSJ4ws4Mv/vSN7znvedOm2KwduKR7crkRVa4ys8gMXOGlFOjhkye9euUlmqMLUgt0rp2SbUd1K2Aixpp2g4GYwgSUpxZ
    K/1C73zHs8956H3ueoehUdST2frQEZrpbAI7yLI+g6Di6+Ccs4ZbBT8lpQ6zNHXZK1siFQjtso2TaEtxVFVWMLNlJkJKSAmtJH2M
    SEmM4dZOqLV6EoFq6/IMACKIAdZCBCoAtdZBmmVk3dF9vLssaiZQq0/YHjnGheNO9yeL84wRMSbnTKu0IYKUlIjarpr2ZNrzaf8k
    BIQQeqVTRZt7ovOkxrIC9ryKufspUbd/3/2XMZjPhjJTe54hIKVkjGFCiAmAtaa99lZl2xizENpuDyUJqsjzTo8xoTPEav/XWSgg
    qikFBpnW9gkaA4oi1whrEbyv6zrGuL6xUUckxrjG+z983qte87bzLryC3cp4KjbfUHVHz/NiiB4d17QFOlLMGyralWkUtsh729s7
    bFxRZj5MgZT12MSZnW6vF/LjT37kj/3YrWPA9uahkw+sM7S1B1885hNEKQKwcN8v2HSN4/sbsIRQJ7BFVPzDKz726te9a2dKxq2r
    7Y1nga1TkpSSM5xlWUxhNpvlWYFr2/ro2oY72Gqr9dYoM0mIPtTMyDND2ng/MahWe3zXO9/ip57+2NNOgkTsbI/27x3GOMssA4hR
    MpulSBI1yx0W98n8OboI/Vnb+woiIFZjyJhjr1Q7tjRaL45l+Gh301uLzzlnArRAQD32UEB3Wy6tgiW3atCu0938T4gQ49Ix29+8
    2iS2INX+V3tiR53J8glc+y2zOLHlkt8xU9EetsXHZchePpM5andHMYbmZ6hEndZx2zt+9P0riqiIBiYhIRKzNeRUumhRCYpkmHyo
    RVOW9RU8mqLo4TNfkD/+s7//wEfO3X/yaeyKw9s763sOjscJ6q65x+Wa4hrT7eawEkMXfKsIEmZlCJM6cIoxNcnobN/Q/8HvvuDW
    tzbbO9izCoakWDEkswXUtokCWk5XB1jXO9Xj6zS+7wGLCJujiXGDosTnv+R//Tf/9OJLR1XIy+H+mYfAuizzsanqqc3t6upKNR7h
    2Efdkg6ozvvqab6cNGBeI08pSWqYUToT/Ch3Uo+uWlt1977nnZ7ypEfc4CRoQgjVSmFrX0uMg96QwN5HZjbGgjrAapWYo4BUnaHW
    c/WoCILhvTK3MUG3k2CYmRHmCl0tECyCKSzd2CntwlmKu1/zXH7uWLxYmg4Vak9mOQJqv13+enHAlHbPB0tIuuwkuYi2MEeWq4dU
    y7+/DKbHANwxqNRGcKpo42ZjaEkERtrpm0/j7gEXr6SQeYhnCEACPNSriiEWUUMZcwa1miAJYLCDD946WzfeZUVQzKb40McuePOb
    P/jpz3+9Cm6wvn9nWjUpuiKva19mw6Pb8pdXnRxTPyLt2KGJWIkTTOcySNGHalja1FQa6561CKmXlaedvPHC5//YgX1YW0fGCLGS
    VK3kpQAM2wEWupz/fwHrezxUARKmBMAn8slGwdYEf/qnb/jgR780rQsf80i5LYY2zxJS1cwaP+tlhnapnm3X9KJ3Z1Fd4N2lZKlq
    KoDKsmS2dT2LMRqmtdX+zpErHYfCqUFT5PbBD7rvo8+5+5k3QDOVfp8J2NreIU1rayuqSVVhmDouT2cyCkHO2PUsUWAps8PyMqfu
    9msjoGMCCsxDiavfnAscwdLdvghA5gc/KsJa/A6uvS12OcRbnM/yMdt8bfkd2/FNDniNY5lwu/yTY1LU9vhtZroA1naFtCchIrRg
    oyvtRrut+0aHGYmRoIk6U1KAMyiJaIokSsaSsRxEqzqBbJbjk+duv+HN7/nYJ740rljMoAnsyeS9AdhGScaY1NTXsvN2zYDVGnAm
    QutjlKhbko6VUkCoBk5MnGo9uv997vELP/vwPftgCFEArQk1kxScE1hhoa7bsuoASwD5X8D6ng1VEIQ5TSY7INPvD8eVCGch4l3v
    ++ofvuwfo64E9L3mTVIfoZadI5LmGMBaCK53WzMdGxgAhMQ4E1PyKakQs3XOMVvVNN7ZKnK72i/GoyODwknyvcLGeuvRD737k57w
    gwcPYHMbe/cAQF3PisKgK/i2JVULGMy1vjW1JRtq77QWsJzDgr9PrAQSjapq2KWkbV7T0jha+s880BAiWiBF+5P2fxejLetYu5R2
    Lt1SKSoztzc8sAipxFo+BvVUIaLM1P4CdvuNdn/t6lFV0yQiBnRxnu25XRtgLXK69sgLzm17ScuYuHivxcXqPB9u0QqL0phydzTC
    LCSwtl1ORsUQM7Uq/9KCcZIkSmwNkxWgajQIlSWuOoyf+4U/+p8LrmS7PvWA6XHRi8ozH5LCZqUk+Ho2KDI6Nry6Vi5Rq7oBkkSi
    pImRqG2ARfKh52i9ZyeHLh7kzc8888ee8NibAZhW2u+TijcsjBRTY9kYylSzRQ1L5h5f/wtY38vRAhY0+WpWDHrj0ag36IOzBJ7W
    +OJXD7/29e96x3s+GXXQXzuxakyTqOz1mjCaBy48B6yWH492JS1o8awQUh9DXhQmy4JPISRtXQZFBv1ejN5CVWKMTfSzzNlBZrNU
    D3J5yIPv+8hH3X5jAyAMChCE0LRbP20CZNQSOQIpkqpClMgw79J8Fv4s1HZnECUkABZWVKir3rcLcnf9icpSubrThLDWMh11k7RH
    AK5p17pD7t2fd8p/uxw3Wr75l4+sODZSO+Z8CGYXTXSRKiZVde6aGTaLpqtru4TdJqf2lJRbwBVRYm3RR1SYWKGqpKptY1bbfN7B
    K8CARkUSANz+1LIqvHg1zGwVmHkxhj917uar//nfv/Dli0YTKgb7tsaRs4JzN5qOy0Hf5dmRrW2CGwxWY+ONpiXAarG5k/dZvprd
    620AFkDpv6eNUqfOSCA1rHBKOcv2lV+/3S1Pe+Ev/8RtbmnHo+lgYGez2cpwKBol1rlzolGVSSybbF4Qg7Tx5v+mhN/b0a5S3zRl
    kdf1tCgyUAKoaoLLBwm4/Ij857s++YZ/e9/5X98u+weNXTuyM3K9bC5ftysGKQs1dGod3Fp3NgVAhmNKABEZFUogJmMzN5vNMmdj
    jL08C35WlHnTVCYpN7Lac9PZ4bV198M/dM+HPvTOJ54AY2ApMryFMsjA0Nw7I3XGsoYAVRYRUgsATEuRAhQ6b66cR1XEAJJ0W2Ap
    JWa+RvhosSOmKCLtZtkiQrn62iXqUKb9hTY2ae/2+bR3AMHMBEqSFiej2MWO9l2wBDSi0nL6W0C5ehnr6meyGO1lLlHJdfEWi1xv
    AcFtIJkQ9AsAAIAASURBVJaSEKthA8jSn9sFFrdZYRODM8RkmVreXtdNrgnkEIEqAAYw2JnqVYeqf/rHd777XR9X7pHtK/cP70z3
    Hjh5c2dkMheSZwYMI0nrMxabaM3yDF8jYC2jlXRdtJQAMQqCGIFRyRjix4986L2e/MS7HDwAAIVtYppmpp9UQ2iKrCDIZDYpXN+6
    TNUsKu5CUP3fXcLrx1Baqpl3wVFq1WA8TAId3sIv/8qf/M/5h2ZVxvlaYDdtfJYVhm0TorO5zYtZVTFRjNE5Ywyl0IiINUykusuf
    WyI90EI6XUDCrbInCYsxKcsNN37kXLA8lbh5n3vd7ud/9gl7N5ACEJrVXg6N1Wwy6PdjbKJG51xby2BkPvpmFlZWViC0aCZbPCcB
    XCvV+lrHNZZ7l6bvmHG8uc7X9W65ruezfPwl0TsCiQrATJyShhCYOcssAdEnQGxmoIBEAYEMMW+NUm/FHNnBYBWXXonnv+iPvvrV
    ywj7k5ZKSKyJIEQCqzACi85ISQnCaBXToK313NLMtzXTeZYqRGRAbCAiQaMp82ldsSI3tmTD0WszduTLrPnFn3vqAx9wirFwFgxp
    0qgwOcFSa9rWXf68yjFnYKGjLooiAWRg/xewvmdjToQnWpi1AYAqh83tI2trG5uTWV6uK/DBD1/2Vy9/7fkXbQazwvkAZJuQQpKY
    ALbO5aJKRKJRJBpGljlWhNjYrhlg/lhGp5iOrlg6VxTqNGo5evSKMsTKIvYLmk6PZKhPPWXP2Tc55cd/9JE3vRFVM7DooE+G0ISp
    yWxEFInRB2Ncz/UApJQsu3ZXe/e2awFLr+s9/L+AJQDqulbVtgpp2KAt8MXEUNLEzoAw2tkpyr7LikktCewKzDze8Z6LPvjhz3z2
    S187/2uXnHjijWd1qbDCXiklDkoQImhb4basRAqjSoiMhE435BoAyxijmlJKIpEhzMTMYgjOHd7a3L+6ZjWGyY5N1b4Ve9fbn/XE
    H334Tc9kYzEaV3nPZJYjEoMtLOnCnOKo7mXsAlYCRLun7v9GWN+7IaQRAQBr24u361pFJh7ZumJtfY/Cjaba77urrsKfvfxN7/3o
    VybBVXWwrszLYRVS3SRj8ygoigJMKQWCxBhSSs4QS+c/uKwWpABItaOethtUMocVzbIs+ejr0MtyZ6zUtcaKpNpYcQ+4750f8fD7
    3PRMaERstLdCNULtp4O8z0BMnhSFzXz0mXUAQe1Rt91ume07NL77gPWtnMP/y/lcA2ABS1W5JEkSuC3Xa0ppbBkCtrZUuNEszhr0
    h7aK+MSntt/wr+/8zBcuqDwLCpv1hW3TQElAUSgpx7lCFUEzFjJiSdkI0bwGmjgKy/JO9NJJCUOI2h2fdnOQZ0H7ZU5xZmUKf+S0
    g8MnPe6hP/TgW+UW1iGKRAkuyxPgNRBMTmzkWiegnQTpVqkCxP8bYX0Ph1ASeAEz8k6uXDvAYk5Ramswmk57/fXpNBVFERWHJ3ju
    8175kf/6+L4TTm0izZqU99ZqryF2vYtZlllr67pCkuFw6KsG3eLDMgtzzo8/qgAB0kgNsRoULEaCZbiMHEsqHHy1bXRST6+6853O
    fu4vPOMmZ2JWwfWQFEZhGSHUjikzRjUaWrhadtf0/6eAdZ1O6dovp00J293MdpewpZISqiZUxIVPtvKmPzAJ+OwX8bwX/HYTsq0d
    74q1iGxWeZv16qbJSiMkQFsq2FVDaztGjRhSZrUkhtUKSbSNUrwaYAkzi0YSIU5MKpJSCiIYDDeq2aifB6mvevD9bvuzz3rcwXWw
    wCCKRDKOrQmKOgVidmwMOu8cxaJ99ZjLb6+5Wz48zxa/T8f3N2CBUoJXgOCgFkt6CUTwfmYsjCEi2/ioYthlVYAp8N73XfpXf/tP
    F116iLJB1RiYMu+tVY3UTSC2thNdEmbWdteolfI41gjg2KEU2TazpnJUONuT5JAchGPj+7lLYdoryaAyxmfG3/jGJz/6UQ++5z0P
    WoO6QpEjZ4wnVW60X7p58YqgrDBYBqzv33HsBH7T5PQ7+86Epgkikud5SwHxXoJ4dgmGiUoA2xO85T++9Oa3fvDCi48klNNpItcr
    ykF7xkSUNCZUQm0rcutd2pYFEijO5YyY1UFykkxJkq2UwpKyVTcHhjSlSIjWMbOKRkAstB5Nchtveub+pz75nHvd82QWGDSIvshK
    JEKr9KVIDAWCpJwN2gfqgp+6qF4sNliXZ/v7Ga3w/Q9YCrTinwRY2SVSISWxlgGZjMdENBgMAB6Np3nZ3xr59Y3MK/7tLV96zRve
    +vVLNtUOZzXZrJ9lgyBaV97ajJlndZW5UnfjG9A11rwXi4ASqAEJ1KbIClPkA8uZ9z74xhk1rNFPnFVnNYaJQXPrm5/21Cc/7ta3
    XjEMKwChtEgpOrMw+CWFXUR039+A1X1k3xvAWoyUuk5MMmgUs4DJFB//xHmve9O7z//aFWrXq4ZNNhyNm6LopZRi9NZxqGuyZDIo
    QPOiFTp8iO2+3vwdTGsRqITWI+dYwCLRGInVWRiDECvva5CWVk7fP7j/D9zuiU/8geEAwaPIUj0bDXuD6bQ2lGV5zi2vigGg8r7I
    2kbFY2ZPl6M/bomx3+dQ1X2O39+A1Zp6dYLuLdFSpOMtG+9jluUGiBF1XZd5YZ0AQVR8sk10wpg2eNs7vvJXf/fapOXhrdpl/awY
    Tme1Ma5XDirfCM1LBHNyVvfdXF0bXSW+W4uhnvTLnsszH2NIUTUllRjD2tra1vaRzLrBoDcZj1MKe9bWkWbstyymJ+xffdKTHvug
    +93IORggNL7MdwN8hVmoUB6XQUcXaY/r+B4BVlXVZVkQIQRVRZaRCqY1xOFf3vjV17zuTRd/4ypbrkXYOhpXDmMiJuuc802TQsyd
    YWie51WdZJmu34nq8by8LYAoJyAqJ7Rqg8q0dHXtV4akNSFJ0jTN1GV80gkHbnBi7zd+/nEn7yPnsNyCPp5Ug35JBFbEiBDEMJxj
    MkhIusuVX9q83pUL7epW1yDS8H04vruAdbUbg3ZFYzA3ibgOM0qLQ1JSSoq4ACxRMpQrMBnXq8MiRkCUODB7IhU1UTMyVoCdCpMZ
    3vv+89745nf8z9cuGwz2NgExUeaKmRdxRYLd5WctwdYSZi28ESkzWQrRx4YNmYwSgiDluZvNZoPVlZ2dHcOuzMqqqpkyTXXfVINM
    puPN4YBvfMODP/iD9/jhH7pNZmAYZre/UeZvx22h7epKge1M/t9lQuiaI8Tl7qDjvwauQX3p2s57N7+52q9c8zzI8p91tSUltBqP
    TdDaS79nQPja+ZOPf/qrr/3Xj1xy+bhJUvSGUbmOSjYnY1NU51xVVVDNnXXEJKlpArl8CRcwN3kngFufAOWgHEANKAoAzaHZfLuX
    u+1sJIMEeJWKqRkO7S1vdZOHPeSB977roA+E2bazBRnLZJOiaVLZMwvhDdP2kCpiBLGqEZ3PBGFeyZo7VLRMw7Zudc2qMvMr4Wv5
    d6n946iVBmApqPzuAeJxB6zoxTkGt9RHISiDk6iBAaAJVeNtZk3GXj0TW5jrilkAloqgi3HNMXArDTTnJbRcCCQgJfiEf37tR9/2
    tg9cfuW49pSXK+PaVLpSDNeramoMkdUYfdnLQlM750JTpRCNcZaddvV+lzq3uN0nLUiBOH9ThrK03Bw1jGRlVjhU1dRwKnLU1ehm
    N73hYx/1kHvd8yRn0M9gGaxi0DAlgIMWTYjOsDEk0VtrGaZpUm7dArAW7JsO6Tot3VZxQHfrr0pEprOqV0A6j4Lj/hD+ltWXCAoN
    pG1n3K5GprTZMouBJE3eR2uy3DpSxBmshQpiFJcznADJR29ctjmZ9QarXnHZZXjHu8593evfvjPVwHlCNnfNsu3ncjSeynwaVUlk
    956c+zqrbfn/zBRCA4rEEZScRYgxqHF5PzWSgvSLoSWWqkJqCicW0xQP3+JmJz/xRx9617vdKDcQaLZ476UpukZoXkzmtcHN7liO
    7xawsrR7qJ3fAC9BXrtsuO1u7DIIbVXVnHHwMTkXG4xyOEU/1IaVXbbkK36c1s5xj7C6LQpVJGVlgoFRKItpu7VayY4qxqixcBkv
    SOfHe1DbM6HtsqsatZkdTZCA//qvy17+16+86NJDKxsnj0Nvexb7/X5Rloc3jxhjhmurhw8fZibHxjlnjSFFjJE0CVjhugLDrqGO
    LIeWLWy1bqYMiU016BXe14AM+r1qulNXo43VLHfhWc/40Yc96Ma9DIjIjGisRKLLhzJXbWqqipTzPDdsOtHbzhZqYQglZt5p1FLV
    2/6e+WDTqYbPeWWAXp9qHQSFtvTLdrd0F7CYIIhtk0ArRJgCVOEMZuOYZdYW0CizZubyzGTZaBbznt2p8Bd//R9vf+fHfSymNcOU
    ESw0t5DRtnGPgWV3uN2PshVO6ET4wAC197JS200pTJplLoVmNpvkLuv1iyqmGMUZm5ksNh4xFEYLDn62ebtb3vinfvLxt791Pp2I
    MfVK3xzevGrPxonHNf0/GrAWvWKYd30sB48C8MzHMsubWbKsZQEAKraaadEnIR/0iKpaXcu4B4Fv4Irvc8BidDoeRHNvbSSCakp1
    XTGbLC+VOAraVhQmHFfAWiYHzFM8VVDtkWVuNIMQyhLbY3zyU1e9+S3vPPdL5zXRNHVU5L3enpmXWaX7D5xYVQ2IUkoxNYrArNYx
    UUoxoiUHdgpqx1jYt+U2aVsSSZFSKrI8hOBDXRRFbk2MM4kV6Wy1R6eesH7/e9/xoQ++2wl72xIsfJppSoQsz/LWUjgEZHkLVgmU
    5jRWbvurTcvCn/MwVNMyoayNc1uPg27lXncq/XEdutsyFUFhHj1w8GJtBjFQ7lRuRF1Go+l0OOwDMqlneVYmMt7D5fjqefiX173j
    vz7+hUNbM85KYTOajF3ZVxSKDFiEIXxU0kxLPiOLOWt/RKpL7Dxp95QVKaXc5kXeU5G6nqmEorRNNVGp+4WJ9Whtxd76Zmc8+hEP
    utlNNlb60IRhiVAnH6qVlcF3qULT0R2Wi4mMa6jcA6Cm8c5ljtGE7aYJudvjMlaCjzGzDSFJyhi5YWrqmBX2uF7BdwOwuqpCS7JU
    QCKbZIB5adCIGGIrwHQaB73jKzN/jYAFUBJEYZ/YZTTzqDyGAzQRX/rK6A1vfPNnPvs/o1EEr5blvp2xzhoizpSMEpQiORijCh/i
    zHUSRhZqoBnUoKPq8TzUiqAEaotqkAQlzqwjtiGEFL2znDs4juPtqzITSyennLjncY9++P3vfZb32LeBzEAU03EorHMOSUCszopS
    anNPgQWMwhBgZO7X1AZfXXmjW6mtkGZ3i3YdetcjmX8hqLThQCQERgBiq3WRPDnXg6Kp1VoyDlCAJSBNQ81siTMmI8B//09447++
    90Mf+eLOtrAbznxsQl2u9AUhiKrmx0Q0y0UGUpkn2t3CIXGs3HoyL8x0BXDGNE2TucIYG32qq6iK9WHfUl3NDpc5Z1mUtHOTM096
    5Dk/+AP3uEm/QAZIAiWUDlD4BimlomeOb4S7xCxZKqDM0Uq7bxZJaEriQ90rez7MwI01vfGURV1R4tCR8cYeLo0luFDXmTOqyib/
    Jk5o34HTP96AJdIJp2inogtDYq3UzajInagYyhQ2BnbGNA1ccVxP5yg2E0HnWrSdl5iCA3Q6a1xROMaokpWSFdjcwqte/Z+vfd07
    x1Ob5wdMth40U+MSKEhISDCiGgSV40BIc7RyUNOBV9ctqJ2pL0VAhGDzYjSZOXZ5rxeDeO9Jk0piSgf375lsH5pNttYGRQyTG556
    8uMe+cBHPOT0IkMSkMA6EBCTWANClN2La4nWTCAjuyUjXYQKu3p4uqjRdDrKdD1q3ZBW4ALSdroQoplTCkgsAE0QVZMRKIrWdfCS
    54JCYBT4+iXyute9873vO9fH3mRK02karqxnZTGZjevQEJHLs6WK1aJpoeMELGqCC8AitcuABXQCeUqiSVJKRVHEKIZMlmV15TPE
    THacrQ37O9zhrCc96YdvcVZf0UZ0aTbdHvYHDm5na9IvBjZjpHkF/ziN3boV5kpiVysBaEdAbZEtxWAdi2odZjZHQlbNil4PTcCR
    Q2H/AafacFIS3yvbhqdCcTxT2uMNWK0+UUInfWkMbCuVJo1hjKfjIsszV4YGzlrV4170PRaw5s46BJpUU+Ocs7kArU5nCHHgyhjV
    JyLGeIx3vvf8173hXV+/ZDNokZALWzVWiBOUGdappCkhLhELGWqWXHZl7rKj1KaHWVb7KKJQBpssyzKbpxQs8WS8ZViHvUxi4+up
    YZQcTjvYe9Q5D3jA/W8xXIUqsgyK5OtJryiXhFK6S+RWwG4ZsHSpx6iTTEjc1fISwITs+gRYGpDaxKsthpMSi4UyEpo62twYR7Wf
    Jg1FzyVkY7gEXH6VvOWtH3jzm99/ZEutW9seNSsr+5KyQFtvx7Ismbmpw1xfZtnnSo5uY1iuXzOLwdEOzIu0Mfq6LF01mzKk1yu8
    98MCB1f5bne6+Q897AFn3ASiSBHORUkzQynnTJFio0U+RODZ1Jdldnxb/eZ1q+Xdw6P//+hfb69O4qSqy8FgEiZVg+Fg5TOf2/nt
    3/69u9/1Dk/9iR9a7ZMBDCJQz6ppWS5p2B+PKzjuEVZneYWYYADDiIqqihdfeslNb3JDAClJbjiFaIwJTWOL4rie0G4L8TwI7oAD
    KlAGvMa6rm1mnMkZNjXRmZIMZh6TKYYb2Jngrf/5lX9763uvODzd3PFkeiYbJrU+JBExFu3D+ahZXogCtgQCbdVNWUhqX+VlTtam
    qFEU2tVBDDEz544JkqI3JIbAKZTw2oxue7sb/8gTH3bHO60REH1dZGS6DQzbKmjRUW+9lBLK0jx0gNUGFEmR+PoHWBFeAQNDatto
    sbupDJKPJrcg2ZlNrHPOlUdm6cjYvP2dn3njm99xyaWbvcF+5Z6Pmvf6s7pmgxgjk3WcpQAVtuxU01JIJeiyPFksFSjzkm7aorKz
    8DplBamUPbd15IqNtdKapmnGxsaTT9x73x+40+N/+M771iAK36Ao25rjtMyzpplaazPjAJYkIrAmbx+hx3H9kwqCdp2VzGqwDFLL
    GmjdNaqEim0WxcKhAgCc+7nxr774d6ej6WMf+7CfesZ9GJAmGtRF3lJIvq8jLFIfo7MGQJToKCPCzkgOHd75lRf82vOe97xb3PJg
    8LAEg+RcRErK5XHNga8GWK3pQmubokzsU8MMIooak1eHgXN2NvWuzMhiFsAOMw82eOvbv/rGN73rq+ddFrVnTF81F7KtDje6Ircu
    7RWi24HqnFFad3u1mSZNIUQBO5eLIEUxzvo69Pv96Gvvm7LIVFMITWGtC9GST2mHzfgWtzjlCY9/0N3vfkZG7YZ8lw22davuOQpd
    1F8UtNgT6iRXdgFLWkdIc33q5hdSQWgNkNuOubkgB4SCj43Ny5BQRc5zOryNt779k6967fvHUxOEyBUhUpNEDZPhIMFa66vasitc
    2VQRiYuiSBIXm2JK3WYIdj8zBoi7LRQGRLgr/LMSi+VWskoTU7DknWt2ti697W3PeMpTHneLmx1YH8IKHEGk7W9tm/O1sxkSUW2N
    5YBW3ksIfBwF9oQUiAolmLkYzoLJON9YOFqphiSBsggcHqO3ivd/9NLf+j9/0zREEs55+L2f/az7UZJhziQC1Mys6o4rYJkXv/jX
    jt/R2ycGkwapcuYkMaotCvroJ877t7d88PyvHb7BDe9wYB9biyRQHRurgGsaYWYRhBCdYyKEEI1ZNmzRTpahbV+9DufT3rBCpJiL
    doIIIFbTfmHYMBmADFlrc4JDgskMCIJkTAKJMxyT3uKsfQ/4wTve5MY39ZPRld+4OFSTXlmAnQiKPE8pJknGcBR1We5DzPKSycQk
    0kp9EDlnY2wgwsRMrQQImKCAY5IUATEWrS0WG1bSoiwqH5Rzk/W/cdn2h//rc5dcvD0Y7t+3r9da9TV1UzhbzZKzDEXQWjkyMVop
    Y+0UkzuNzV1eJhFAMDSf1uvDIBCp4cSaYA1iiD7UrgA4BvazFNX2wFwL/fPrPveyP33zu9791dlsPcQVURcEiQKMKIkigViVLOUM
    py2IWSEN1G1FKGjBPiEhEtE8L0MdIayRJGlZ9GNKbOGjZ0KvKEc7W4Oey6001WbO1caq3vKsk571jMf83LMfdPqpA1KfWWSUmIRZ
    mEDK87S9LekztekUzeXcWqbecZv/1vJHAagh6rwEWtCS1LAhQU2kIcUk0RgTUjDGVTVgCRn+6XWf/pM/f/X2jmsaBultb3PTu9zx
    VGayAIQMWajF0cKw3/klcbxTwgRhEpHKMkdBTHmC/ZfXffyv/+Zfa28Gg8Ef//Hzb3QaGOj3Rhkk6QBqjza/80WRXZst1XV7HM2L
    jryrTETLh9TFHkonF2MlUheqkAp7UGt5A4GJyRI7Q6g9zv3U+E1vfvt7P/ipaPaYYmVaNUWvrHxDbPfs23vRhZfkvR7BAHAuz62L
    MfqmUU2ZM6wqC8u8ud7WvPC02Faf06mSALBsLDOnpGnmTFVY//AH3+0xj7j/DU8lRwgVyhwaUIdkCm8sAE4pEZFll0Rns3owGEDp
    GAkKXMdOg+M92gdTikiSstyAYuUr54zXpJwLsirgY5+8/A9f9qqLLx0XxQkpFhJLJRaulZvEQTlJm4irbTdAWu9vQmybutotwo6J
    ustUIBFRoRTC2sq6s3k1rQD2vlbrXcnT0Y7GuHd9WI23ONWnnbL39BsceMoTH3mbW+YQVFUclDBGDJJCqSNtzfurFkHuLmBgsR1y
    XNUUhFqKEUgt0XKekYhE4GNsnM0AU0fJbA7YppEQOFn86V998JWv+7e83BvDIM8yba76sR+5z08/424QLYggre3Q9z/TXQhAUKkd
    UwI1MffR/vnL3/O2t597+ZWzsj844YThH/3+008/DQ6xabZys0Zw3gdjTJ5zSogx5bnR3cD16HEdAWu+GbQrzHi1e/YoHbRl29Vu
    gw8RAME0MQVPbHrWIkWMR9ga4wW/9VdfOe8y6wqhXOGaSKOJP+XUG06mTQghSGIGaUgpskHusujTMkYstVQsc/dlca1RJcsyFRPq
    aNSUziL6UG06rvZtuAc/4C7Pevp9+z1UEziDsoeEmOBbjmiIIYSYW5e5QjuzmN1pXE6WryeDAPWJLflYgQWsbPIEM/bqMvfZL05e
    9mf//OnPXjRYOcnmG5deenh1ZQMaAVGOgCZuq5MtTbfN7FqbUiGVXfHFFrCoo/UKtd1QMhgMQuPruo4+OTZFlje+MoUJcbqx1pcw
    iWHHUv2Ih93/0efcb99erPXaYmAqM1JtYmiccyC3Kzp6tW6MY9YvHef5nwMWt7Fe1+dBAFJIEyAZa4Om0bgeruw1cJdduX3CgbUm
    4Sef87fv/8gXDp52kysPVaT9fu44XfHEH7nXs59xN7SKEXHeYPn/BcCSBMwss4KDlrXgxb/xr+963/9EXcnLoW+29m34P/mjXzzt
    FIifrJYDxpwAYYBOKCZlWfvNdwSwFqPjY11bnbNzHFmQAOYbagCgSmTamyFJtwcaAc3xgY9svfq1b/zil78WUilUEPUnlYCdgIy1
    1nFKHhQ7ox7J2jJ5+4bY5SIe40rdnY+XZIwznJFy8oqohXWlM7EZOaoNjW54ytrjH/eAO935pgf3YjILbDTPmIg0BmddW4yZTCf9
    fv+ojtm5SOH1DrBSYkuKFDQ1SYWyOplvXOFf+y/vfs/7P7c9ImM3omZVE1fXN6bTEXMNivPuznn3THeZS65p8+SPuwom7wJWtyo0
    NHXrNpI7E31AikxqrNTNmKTau7f3kAfd4/73vuNNz8gyBhSOoKkyqsZQipFEbV6q2mNqsnL0gtWlPlAcvSX5HR/HAJZZyJRSiKmy
    hqvokyJ3a1VQEZPnuOxKPP9X//S8r49q7Y0aDIZ7pzvRaSjNkR9/wj2f9cy7QULO7v87gNU69AK1NazgCNcAz3zO33/4E5cOVk/d
    3mn6w8xPLl5fbf7iZS+45U36FOC4s5bzXkSkKKxvGxJxTYJK1zEGpY72ffRPr/0g0imLtm+GrkUPIIWPofWUz/O8LV8HRR0Ag6j4
    5KcPvfyvX/2FL1+0tvfU0URcNqyCKBlVTSkYxwSZzKa9fFWxsEr8ZnpbSlAwWeu9V1XnHCuiTxDj2ORsZ5Pt9RUHjAzv3OY2N3zm
    0x935pmDtlVXMS+sazKUjCHetcfa1Qi83g1SSBLopG5M3mdDh3fw1rd/9pWvetuhIwKzBjOomriytjqZ7ihFYyOzBwRakDgsSKHt
    g4oikLRr5jVQSxBCYIjMoW334aRinTEsvp4iVs4yA9ZEI3VR4FE//OAn//idy6xrtW+qSb80GTO3wlkwgJGos7rp9XvXcFFLH2r7
    kJDjDFXtOCol3AWstt8oAZLAiiyAFQgeX7sQv/abf3zJldPtGWX9vaO6DhF7BwdnO4dW3OjHn3DPZz/zbkjBGYfUXcDxvorjD1hA
    EjACWFWoVhcNHvUjv33ZZn5kZHr9jbqerQ0QZ1ectNf+we/80i1v7EhR16HXczEixlgUdtf48+oSetex8aqNgZf43Tj6gMeKnAml
    thu+/ThIQUsRVzsaX0VJxpAxxifJs54AswaJ8ZGPX/VHL/u7r19yxBXrTWCX99lks1ltjMnzsgkBZJWY5nQD2sWsYz/69o4iwwBS
    Cq1NTpZlzFYTJOpKv7ezc6hwsWkO5Vk8cHDlZmee/Cs/+6OlQV4gzzopHAIsSetevqsHMc8Prm/DJ1UQW3jFO95z4T+++q3/fd7h
    KuRl/6AthqPxTJBiqlZWy+BnbCUmD2WoI81IHNR0KvgUgSQkXa9l1zUFq2mxLnS3YCxQIfjccQhji1CWRmOTcXj8OQ982EPuefAA
    nAEzHGE8Hq0NSyDGFJjZUO6jzKbeuaLXc9dgCLQ7zbJQJeqqlsf5sXFMDcvMjZhAMaQmCqwrBLYKYMKnPjX9oz/9u/++4CrJh157
    YrJoIpHJtaeznaEZPfnx93n2M++ClDJjROaPhuOc1X43ACtGOAOBqFIlFCzu97Dn7vg1r6tbO3Hvxj71E6m29g74hA39P7/647c6
    +wDmFr4hwDk0zTEp4f8bYM3PrC2xL3F/cQzrEkBCwK4bBLcsqoXGe1rEa60tMycjdZLQRM7L1Z0KtkAD/MvrP/dvb/vAxd/YDMEV
    xZqvKSo704uSEkvqjH9k9/TmJfc2mutuJGWQ+Loa9Ho2y0IIIUUlSSn6UBeF9aGJvjl4wv7peJRCMOyGebrxAfOER//gg37w5pnD
    zg4yq8NBe7gw32OgRSsvvlUZhe/SECAIiPGZz0/f+G/vfdf7zh1XZnXfKXXgceWbUBPJnr2rk+mOswiNj0nZrbQRK2vb4dvFTp3M
    NEFJE0FJhIQVVrreb+0moSOROquTyeagZEYjaXxw/9r97nX3xz3qngcHWO0BQArKpNYi+Lquq6Io2BomlxRKpt3qbWotc6LdB8P8
    lpiv3XZ7h1TbEqJ2zfPHtege5oA1F4MnAEkpCXhzVOflwFi8+c3n//Or/v3rl2yVe0+5ctRkg6E6M4sjCREN7x+UNL38yY+/33Oe
    cRdKai2Jpq7lf86WOE6jA6zlvHq3+1V3SdIAhBIfvbMG/F9Ibl0vWERmAUAUlSBY3O0HnzsJA7gN0bKa+Z6zmUZD0enW/tXpS//P
    c8++2b6mQa9ADCEztvv8u1OGUgIgkPn58LIFG7C473Z1QLAIXpb+R3dr8NgNN+aY1ZqRCJKQ8DwNnItzt91GUIUowCBCAkiTRQ1S
    kKl8ZNcDmUkAO1QN/u4VH3zr2z64udWIFMauWFPWPoWM09wggTs1uOVQZ27So6wEgriWwJNaXXIFq5ISp+l0vLFn3ft6e3u0sb7H
    e3EmC7PtFTvOeXrH29/iR55wzu1us0KKporDnmV0Bp/dFSkfu/3Qofnu99cgPrX0tS791dJPj127x/Lgdu/hjmIu4LmdCASYNnjF
    qz75L69/22iixm1UntT2hA0xJwrOyKwa94p8Vk1yl7EpQioVhpA6oy1IO6tzmWlWksSqFIUTK0gcd8+8+eRDGCkzkdAYmlpbP+gB
    93zykx+4fw3VDPt78LNJr9dLMTJDRIy1MQTr8pQSGwegCRB02nt2N2pewqylS14CLFHYawCsa5EwO2o+d3OG3dj8mnBDUhdhZWbp
    6agUFbQ1rQf9QQRe9c+f/8d/fuvWlsANKzjN8iZo4kg2GUMZuTQZD7j6icfd7znPuBMldYZEVUwElOC+OWAdU8Gba1Qubz0s74zx
    YjuoHeZXf+3XuuKGLjz+tLVU6+w8WuYSQ9Aogum0EmlB4VumQ0lXatn9V6BsyAAawQ5C+MZVeOObPxnUqVoSFGyNkhAH4kR55d1n
    vnDejW5yy4MnWUMwRpEiRWLiNo+p6oYI1IZsSIa0ExsQRCAKIkcFElICtN0NadvqSIlkDrkyF8LqNEtaEiUrkzIJkxAJSFs/9e6F
    7t+5Cyh1nTDULRoQMcgBFkTGWENKkMyQBZUWN7/pDR50/3uWGX/tf74kYVpkKhzV2VE1c871yl41rcqiF2bBmZyVWZmp5RpSawUs
    3adI4NY3VLtzF+NsGT1pMlnWl9T6oUa2xqPv+nu/dN6FH/nEJwfrJ5x2oz1KbA0I7JtgOCfmGJSZCSCGzsMQhSiSdlOEjhQ1x1Zq
    /Ubn7KWWxEVQkKhKK26jUCVSWsrmCUHbaSUmNFXQoLZNtFmbZhxkRoZGwTfIlPHxL9Q///x/+PCnL9yairqBmFzZKiHLXIoeUBWx
    nKmwNQXgVA2pYVAnsYNFlMre+8FgBUR1UxsLNkqaXGZ9EjUkqqRqjTGqNsaeQZpubfTogfe5/Quf+/RHPPiMQYZQj1d7xGDnMgUx
    GZAhtgAz244YqgTAMNh0y2RpoSzdLXMuzVzDxCgx1CrZo8sRiw+iTeUWe8gtixAgBJB0m5yts30iFUKi3Udf+0o0/yhF4NioQAVs
    oAggbI8nw8Hq9gx//lcf/Jt/eNt23csGJ04iqTUKdQQHtsomQSUZoznR2Tc99W53Ojl5n1mrQmQoSGCy7TpZvHYXDOZrpnspdYVF
    JTXdxBCIAqEhJIEPSaW9B4EADSmYF/3q85ha16O210yTBEliye36lXPLOfYEMe0e7TwJWlB3jwL1pX+FlFt1awUZ8UqHtvCGN300
    qCEYFmMUChYmIStkYqTxpPrc5z9z29vfOS/hDOfGMneFBzJwmTXWBPFJkmPD2voNmxbFiNsHWld4mjesdE/OOaFpV78A3fddKYf0
    6K2OXTA+aqF9kyFEStwiP7U+TogGoglFwatD3PY2p9361neqZqNLLj5/PNlpgAMHTvCzqW/q3GVNXa2trE2mU2NNF2S0RsRQBTqH
    euoAt8NUJQJT5yZtCPNFgQSwccPN8Wxj777KNx/84Ae2d5ozbnzmyoBE4axNAgKs5emsynJXVZW1bolNesyOKi2mgOa7WrvRw3xh
    yiJhXjJdw/yrqKIMBjHIWWNaMSuRyWTL5rnNyhpMpjdu8Oo3fOTXfvvPp344mnFUSyYH25CSaFrq4G7X30LDx3RLn9I8rlElVaLh
    YHjl4atUMFwZhBBVlJjH46khkzuXOcMIFGsNY4tqkMfHnHOfpz/1UY/94TNW+1RN6n6pA8ekgcnOPcOv8XVNuPTNxuIXr0HgRXen
    eFHi3E0C2nWQCAriXR0z7XIFakso7TEXSmJQUcNOEhmDEKSuG5u5qkm9/uDSK+Lf/N3bX/uGd5vegWJ44PLNbVf0hBNBrZBtVxxI
    SYiEJZ1xw4P3uOtprMmSlQRjNUGYeDcFnt9su+DcVToWs5Pmnbx2aT4SoD40RCYzObFtKdWGyLCh5KfgTjUvqgBsyRpaIql0aX8r
    BiC2awLBclV4yUcIgCyVLlvmNptIEMClmZgvX4AnPv23G7KKjMS1ba2JWzvlZNX72ZE9a8WwaP7gd3/lljfNDODrSS+zTdUQXFH0
    Ouzh0NmWi4EacFcYV/gkgZkBJs2XlkZrss5LEWbXi8DdE2uhhUS6sE66jvpQy3anrdaAalJiH9XaIiiI0EQYiy/9d3zxS1520ZX1
    zjgMVtaaJlhXJkVe9KqqoXbBtbo0rdkBMaDSaUIpFoti1+BgkQvERVeQMjmm0c52adP+9f6hy7521hkn/N5vPf+MGyHU6BWIMdaz
    6frKEEBVN2VWzp86y5JJnbrWUVK5SASomHZV6rwNKHXn1rK258SReV3OS7TGAoi+zcREtCaCEBnXaxK2J6ACP/Pzf/Klr5yfeKUK
    a0I9Y4wx7QZrUiVmbrVwrpb2MGm79Hc1fFoGiq9DXgwIHEWNMd1aIaGUmnpkWYc94+udDPX97nunJz3hnFNOQs8iszAKTZAkICEi
    Y4+nuP6xa2nRg3rUWBiCtytM2l1JaPuwaBc5iI+p83RSOTB18MRKRMzMsEkghEu+gd/8nZd/7JP/4wb7Zt4EZKbsJyURYYVL1Drd
    KWniSBRcqB90r5v9nxc9tGeSFROi2kIbbSwZnqelC4Xl9hSu4UIgQLueM4BaBY6jwgVFEsQIUTDDGJhf/9XfIMNEwiSWwMQMs1Bo
    BgBOoESdoLBTmHlM1yWLutuwKUtYvhDzBIG5pXKzRvClV+Ct//mRCCvE814QVoISCWmeGcvMKEbb03M/9fm73u0uLofLrWFxmRpr
    2DBATGAyVe3JFm3bOUShQhCWRCKGDYOXn/MJpERHr3DaLTlTJ2g530VaXOV1qyDulisIc7JCF5gwWx+lJRTMauzbx+ecc9cbnnLT
    Sy48b7a9mTm2NpvO6tFkvLax3qRGSZSlzcPmMRSABI7odiu108XpHmPzYlx32gRCaCpirK9tIJnJjl8Z7vVe3/nO9wCDm9/iRFEY
    y1me+dCE5Ht5H7oIDYi69iXSVk/56JsGlFpgn8eA3WW3OlIEs3i60tKcaBJjjCQwk7WkpMqJXQZTbE9BDh/7xPbP/vzLLrpoW6hM
    lCt6bDNjTEoppURkjDFEpKpLx969k1mZACKhVuaftM0ArHHW5T4EgKzJppNpNavWhoPUjDcGufqtUB++x11u9svPfcpjH3XrjVWU
    GaBRY4KoIbZMltl0MPedxaVrX0vUXdRyEHfUCqa5YG475a2bQbvhS8tHgBIRGOAmRh+qIs+UZFo3xuaz8P+j7b3DbbuqsvF3jDnn
    aruccmvuTU+AEEJIAqFKIPSOIooFKcqnIipNuvTeEWkiHZEiIAJigBBCTYEAIQXSCem3nLbLWmuWMX5/rL3PvUH8Hvh+sp79nOfU
    ffZec60xR3kLfnpleMObPvKDi65T7lHWa5Jw7ozlmDqMNBklmrdOugzLatq1vfeQB9yhsEwRSdQ4iCZ7ABur81e3eWuoQuWgMrWT
    KLm16BHP2kmCzmCDGdbBWVgGA+aFz3iRxKiI4NSVRqokiedEIw9KRKqwKhaYNa82H3PqmW6Ox+endtbp6a4XkllcEuZrbsCXzvxO
    BKtaAh8csJTg26YsB+ONMBhuC16++rVvnXDCHXfuLJpYE3dNKyaywYthIriErkIVoUApMISICZsaIEpdrUw0s2/4n66OzT7c7M86
    eTblA2XQr3aRbbbElOZPqQAMm6jRsfEhGcelQxKUDrc5uveIh969Hk+vv/6WeiK9akvZW1rZGBvHiRNRAEWGEMiKI5CQPzAqPbAO
    cqD2Ip1TLYlA/X5eT6cxwJjSmUESM5kEH/CjCy9aWQ2n3v3YjbGWOZPhzFgfvSV7oCel87DVCWcdSPMPGlZsFtnUoYrmZSDxrUJV
    VytDDVuJkqJYx0JILI0Er9yKzQq88a1nv+efPzupczWL0xrGDYUMkU1JUhJVdDTPFBPNip1bVWfd5kOUqAumJCBSGCHObDltW2My
    a0zTTob93mBQru29YWC9NHtPvP1hz37Gk//8T++5e0cR/HiYa9vsrzLNjUKjihi2AGIEm99swBLSzV2TlTbbMgfoswfzaNHd4reW
    oMdMm39WC9LBq4CuBl+frhdZFShbnfDlV4aXveI9P75kD+fbOB+MmtYUzmR2PG2IQGR+acBy0N07+g970B0zBiIUYIso0XDXjeLN
    lpTOXvXsOwI+KHowd0GANmdhdiaeG7pXG8h6kIf4GIL3gbN+5srSZn1QEZWjEtiYTpecEyiCZpq/rLg17no2BKYZUo4ZFrAzzlT3
    gJkFIz3Q8mpbf2D4NX+azVrDuXJtox4u7VxZaxvp3bIiz3nRmy64ZA12oKjWm0krNRBTCoiwnbb57O5NYCVjQDkkh2SQbL5o6Rck
    X7p/d/DHzTRUAJmxBX9BlfFXudoOPJEQtOstzFdu3y17DKVexogNA5mNkODQ9m18/rMe/pF/fuEpJx7ZjG+RMK4yZlLThR5RVVF4
    pQbcAoD+Ykk+j1/S1f/ziMZQCl6cLQxnTZvqkJRLky9n5c4mDj7xmbOe8OTXB6EEEMw0JWezzafbbNoemK5sPvGBU9RFdjl4RHvQ
    a9vsbsoBwSkBK7LcqKL2jRcytp9MvjbF7z3+tV85+wej1tliy+ooLSweae2CJJNSUlUmY40zplOJ+5/LdOosAgFAYVUzgYPk0yZk
    JnOGVFrSOrT7EPdvWZQdy/q8Zz3pHW/78/vfeydHVAbLpUtpY6FwIYym9Zq1cJnRJG0dYvy1Lof/5+N//C+bqfM8gQbrgQCAzanX
    wT3YWwPoRWFMvtDfcuP+DWvMRZfe8pSnvuia6+vBtmNuWqknQXpLy9PWT6dNnlcp0v8EBSUiIjLz+nhmo00dZbLTrZzd/p0rIs2b
    jPwLXT8F9GCJsVm8Jgt2gCWREGKdpLEZF73c/OWfvSwJYEHWRHIRnMAddIAoEUUAUAtYyIGNjPXA1sYHzt7B7UNGJ+8PVihF4q7/
    TPTTK6ZnfuP8QAwY7hxHuvKJVMHO5NaWk+m07Fcb40lKYrPigh/+6A53vNPW5SJzjilZkIoaOElQA+YAJNVo2DAVyaOuYW33cgx4
    ZsHAMAwYStzhM0i7gR9vChTMQ6rONEi7FhbTr7yl/gKktFvDrlgS0YXhIISwMVrtVRVD6nrcy3ILyozt1v4BD7jDiSee9NPLLhyt
    r1p2LJY0h3SnKAnXyuEXTq/OrBm6Ze5GG7O16FqkMajl3DqnCuMcWW5DiCCBVdi6Tl/4wpm3v/3J2w8pDNvW+8zY+T4uNC+KCQfV
    g13dN88cdQZnI73Vro8Dichsetgl+hSaBDA7UoJXYmNqwYWXTp793Ldet2e6MU4m7/vEJuuF6MbjYDOGChEb45g7SinxLM/5xXXp
    KkFQEoLAKFlBBs0AWxb9FLyfjKsiDcrYjK7ftTP7/d857fnP+uNTTtgyKOEYuUWKjcY2d84HX7qecyWAFMGWbOZs9hvH1nZLSAem
    WgdSlINy/ZnMKVHXtxKiAz3t2WZFt8o/Z9GKMJ5OXJY1Qi6vPvyxc9745n9x+e5pyCdtKheGJrero3WFqarFFI2q7YQ9fmlJeMTu
    xUc8+A4kiRKz67YunkknER98oXaFVPfdLvtiBRHMAat22RwcKhnMRp5ICGDD7Mg4H9H6YFbXF6/8+S1tLE25aPPO4LE7JZtx0c7M
    FOaDSZ4DDTcfBzcpuj+WOaxh3kAiZggkgi+9cvK1b/5AePa0dCBgQUFINJ3WZVUmicY5Y4sYaVq33zv/e0ff9nY7dwwZlEQ0qrMZ
    MYxRpaQIqkpUKnhSY/8Gqh4SIAYgoxDTTRZmreIDcPV5sUMHxSzalMadIQh+vYB1oB1+UFOcIGC2KaVBr0/Q1rf9ovLeI848DRVS
    lXzo4fl973NqkVU/vfhi7rp6CkVSSuBOySgjtV1KS2p0czROgtkgv2txUzfOd65o25DlGTFPmikIZG1KQpwRWe+jiJ53/vmLw+3H
    3marwlrenIt2NbRufr0ZgAAo3Vo+Y9aJmO+TOmsfzsLIrNEGAhmbKYgMPBCUg+LMb9z0mte/d9+GjhsUveHGtFG2xpW+pV5vQJyA
    pEpE1LWxVNVa+z8kWTofCxolIzM1fWJFbOvcSi+X2Oyt3ORRD7vnM/76jx9+/2MXqpmkuoq4WaPKGHLW5DFqSsrGMdsoUTQpybwU
    /Y0ccutTfauPm6GKBJSUZbY62vUxDWZ9j1nh1aHJiIRINvMJIZgsaxJGE/zrJ7/78U+dVftexKBObMuyiTFqYrZ5VgYvvk4Lg2FI
    7S8NWJziUYctP+yBx2uMRg2Z7pKYZWQ6z+wwrwk3zxrrQenOQe1WzOTPVEgTsRIFRRRLbKO6626cnHvexWeedZ55y3s+/Y//9L6v
    f/t7n/73L196xcphR96xqGbOIAoSNVDT6SeJapLA4Nh4ayxxN8JWMEXfGmM3l9InhCTKRASFQGGUiFEHT9Z+/oyLLr78ugg7J2He
    KmCxsrUmaisqnfu8sVXTxtFoct65593lLnddWCxydsZaTWCgbdazjEejUVEsRNgA/PQafcJTnn3RT9eOOf6EwRIU2Bg1ADJrSSEa
    2rZmAhPH4JlBxBITzzep7u4n6nKrXw+2O0eizrKJbsvrsAjMDMCwhRKBrHEgtZwzWyJiIueIIAY8qHDqKbse/Yj7XHThRWure4xh
    YgjI5WUMxOoMZyqQBGKjQiEmpu7lb5K15xkkSIUMm5RENLGdm+NQR/EhKAzZtvEX/ODCjXW++92OJMBLnEwnxhomatraOdc0U2uz
    AynoXDkV807Z5tWIrgygmTQ2EUS06/kSDBGLauNbdTYJYPDxT136yte+L+qgjQ42jyrGWWKKqkwuSVIJ3UnsgjAzdx132ryDD0Qu
    ISKXF61PLiuiQCQZyxq9syk3MePapNUjd/fe9+6XPOLBt9++lDkSS5GQiMR2zfrOiA5G1RJlzN2UCcrUSXswfoNNrP+eoc/fnMTY
    GsuNnyaJotEY9rExxlASGBujkOUEZVCIDSGmUGeGmnrsHBOZum3IZkIQYOLxD2//wsc/+VXOtnspVidNVhZBEhlSYaasmYZeMcyd
    m04mXdH3yzIsOWLX4kMfdIfMMMusZ+VTbGMH65l5YXbXgCYlgqYkyROkw1V1opldjAhCCRpJI4lHCEgBJhJFwbnn733u89/+iU+e
    ddbXL3rwQ//YvPFdr7nq2pVLfvIzzhavvHrv57549vnfu/LyK/dt3XaUgssCzBDpVLWFjULE5U4lTMejzFKK3jfTrCpiapUEikSw
    pmvVoY3BMpi425mCCKw561tXXXL59YncAdTMPGARFBqNhWEiZrD1IcYgVa8vSYJPF1140T3vercsgyoVGSmiy20IoSqH02TGgVrB
    81703j3rcukV159z3kWXXnbDYUccu3tHadmmgJSELTnnmC0AYy2BNCZCd5Lnt8WBocGvd2kSzZxKZ5nOHB+/+XG+V26W9AdGQAQw
    JQMxqgZclTjt3ne500knnve97zSNN9aurddF3s/YttMJgfKiECGA8iwHk2K+5c6nLfOIwtpdHV32R+iwWyKRSUWECUo8Hre33Lx+
    +ZV7jz/xdoOKsyxPIEOcWbuxsd7v9+cvfRaz9EBKeqtodSA7IKiiaQKzsYaT8qRu2Nom+qwoVsdYH+Ed7zn7fR/8vMu3C5dBSZhn
    Lef5menmNf93PgXNjm5nMCLWR82s835SZMZSyKw3aaRhZesCP/1pf/yCv3vY9kXUG/WwNBa+S8o6QcPZYszBT12mKh2fYLMt+Otk
    3P8vAesX2uSzNRVFMkxs2FoTSZhMKxEwmkzw4ops/9oEbCIkxLZ0WWZ0Wo8ck7HZnr37Bwtb1iceztyyHy97xUe+fOYP2G0ZTeGV
    e8NhIvHRW5sROUOutNXayr7MaL+XBYn/Q8BKhx0yfOgDT7Bd94gB1sApMxkYKUrS6Ey3WwtR0hS69BUpBu8JStaQ4U6rP4G8GDUW
    5BrNR7X7wn+e+6+f+Pob3vCpL3zxnP3722mdTjz5Tk9+yr3M373iZUccddvvnHth7Z2XKsTevv3TSy655kv/edZVV+2VuOiKoc1A
    BpFSUG/YNtNaVKteCcNsyeYGGtlSxwwO2nYVKUETgiNmGBIQIUhSa77wlYuu+NleYdd1uHg+TRMCKDF5y8nHGKMak1trW99aw0Qm
    BZ2Mw7e+ef6d7njqtu1MjI3xqlFizkaNuNztH+GVr/v0jy69cb0mcj1XDn544UXf+Ma3mHtLW3YuLzCIhLjxGkUBE0OEqrGOZmks
    Du7X3Hps+CtGrO461/lfz6fgm/XnrX6XQCoUhBITkxKJITHdvqCpGfTsobvLhzz0Pjfdsufmm/fEYHJTkLS5I2uspETKzmVJ1LcN
    2+45ZfP/EHgOPujMrmk+AyFAjUkuBzQopMgq0nJjLV378xuu+dllJ5x0YlZyCGQsra5tDPt9KObT0jnG9kDxQkAH/z+Ah+jeIDOM
    MwqMxk1MUlWFEnnQ6jhFMW/5hy/8x+fPM25r3ltaWd8wmRECwNo1FQkMQVe//+IySIe+697T7DHjG3Db6vLilvX1FWviQg/15EYK
    +3Ztzx58v5Nf/6on3+VOi36CwmBYudHavrIoAKOwOmdDbK6X/gJIh7gzVeKD23m/iYB10HVCMye6BFJjOKQoqsQuqXqBcma4YGt8
    snVE1c9g2DATOVGvEq0xLsuhVk1ZR1NU9tIr65e+8gPf+/61RW+Hmmpj2lTDSk0aTzYyl/km9fKF0MRYN9u3DpjGo/F+4xzAvzRg
    Hb5r4aEPPMFCWLgLWAkppFYp5YatgaEUUhPDNCVvGJIiQcg546wokqgSVsZrkR1bjoqb9+Jb56y+/wNf/4e3f/7ss6+59hoJYUGl
    2LJ1aevO6rVv+Iv+IujmpBnjTf9wxmf+/buD4dF7908B9PvFeO2W5YXCN+uh2X/Xux73xD/9nVNP2QUgB6aTUOYus6inE0ltVZWq
    yRoi6pw9JUjqFIAZDJBBBg9mNBqTs3/+zI+fe+GNgTNRZ5RJyQiUEBmEYLlxhmtvRDKQy8pqOh2zIs+sUdG2yTDasijvec8LduyA
    Qxo4EwLIYeTxkld/+oyzfyhmKSuWm8ZP641+roX1FPZvHZi/+vMnPuphJxIjeBT5DACWglqjzJ1oFw5uyunB19CvGrD0IAztrcrJ
    zQzrVuIirIlqhRiUJGaWTCjAAhOFcPPqWm9pe1Rc9XM8/RlvHo1DXU+NKyBlFKdUhEBiqOpXrZ8IBcbcxhWWxIBkJqiiRmFlpjev
    oOjDpCysppjaUGZVhkKisTyGuW7rFv3AP79u2wImY2zpo5m2wyqH6lyUleaynLdSe/iFO6+zd9tsXXT0r9rD5Ni3gaf99ZuvuWZk
    3c42mqDCuQ3aCs15cMpzbyHQzNvx4GgFzCvBrnTf/FIFEk1o6sXFzLlmfe26rQtm9yHD17/qhUftRlsjhbg8tNPRpCodI5G1Aiub
    EfagpdmceP4C2978pjX25p1fVtBB1CgijlCoDYAPUmRWgLUNkGK4gB9dsvqGN77xH9/xGgK29DGZ1MuVM6RQAC1A0r+0gSizazU2
    WjzhyS/es8eyXWaTTX0oqlyMbozXqqpKkQs33L9ntDRc6hVUT/YBa1kvH4VCNfvvwNFS4z1PPvTtb3hcyWICKwNWPHxCUiQDsjAM
    iIolM7+zDMBJupQfPsEHuAIR+M5393z4w//x/R9cWVbbXLG4sjYxVAHs6+bQXcM9ey555rP/4PF/dNLUB/OcF/5dabNjjj32a2d+
    d2W1ZipdPhiP26IYjKYNyBT94Q037/mP//zqF8845wcXXuWl2HnozqzA+gQmy8qyBDsfIIIQE5gNGUtWAKgQkYIYdmaQbigSfexT
    31nZSAGsZFiYlHmG1lJQhNSOEUKsyqoNrWo0DGPIt95aF5rQHy7W9fjc839w2+NO2H1IsTEBMSYB//bZiz7yif/Mq+0JZd1GdsXi
    8jZVOxo3ZTW8Zd/aD3586de+eSHb5SOP2modbr5ZraMso9Yrm66y2JzhH3RN/tpFYQfyZOqy+QNAmlulWgfB1WSODumoiLPkLES1
    xhV5r60ly2jQw0MfeM/c+Ut/8sO8sJJCjLEoKpdlSTTFhAOOAjSf0XTkjA71zphBBzssqKjEzLBlJjYiHBNSIhBFlUnjv3POj0+8
    0z12bIchhMiGieekoM0EBN1/miNH52+oe9MdE5XqJiZhtggJvkVW4MIr/LOe/6afXTueNNbYQR3E5U6hgkQgUktqZieQEs1CBA4g
    DedR/8C/nCdcqkIqy/2cdSJhP8L+ww+pnv7UP3jR3z12eYB2GqqMh6VhILRtURbEdjJpTZbpQcty8LL/ghYo3XoNf0OHzgftICVK
    B71lapqkJiOmFLmOMz7SRT9pXvaaj/zT+z/h1X3hjG+cdp/7GYNBz0XPrCZGA8OJ8fXv3vCsF775hj3tcOGoNtrRtMmLQgjTZlIW
    lQpSSLlxvcJpGLGORVaf+Ce//cMfXpC4VLj/nmHljMMOGT7sQXfIuGPoQlkE0YEZoiIMYmIF+aCtl9qzybKoNAlIDA9cdW385nlX
    fuCD57z1H7/ysU+cect+Xwy3euGp93mZs1WidtCn6G+5//1OfvYzTh9vpF7OFDTsWZksLC+cceYNr3jdB9Rs3bcWe/3lENo8Y6a2
    bdZBKS8cM2KYEm0ctmvL3e5659PvfffbHjVkhQEWB7AMSdAUretsFDSlJAQwWeTwUELnFPTQx7zxllE+VVZYmwzNtlMkVoKXdv/O
    HVuuvuaG/nCLcgayRE4UlvONtfXDdh+y7+brqzIN+kS6/obX//3JJ7gU8OWvXvumt30g0MAUizfcvD8vF9jmo9Gk1xuUuZE4YZ2W
    RWpG+xzisUcd/id/9LsPfuBujfA1hsMZpZ4Q51eN3azu/v8cBzHpbyVfcxBYedPJlTbRVaoKMDNuuXn/9p1bxtMQkgwHuQBNwDk/
    vuHd//yxyy67qezt3L8/hVgMlrZNG0/mgOMLlAFD4gAFt0DcdJ9WMEGUkFnTkfKYWRNiTFBrDCk8ZIy0sTSgN776ebc/ttwygK9T
    LzeErmnbFUrcvSWDTYCozvFW3WGmbczzLAjW1uPioo0R5/xgz/Ne8Y69a6kql40drK1ODtl16J79+5i5o0XN4dqYAetm5+/WyzDD
    s3Zgdzl4XGgpZmnsqCmKeL/7nvqcZ/42A4Mcvml7hQM6XQwO3ocozjk2Tnk+0vrvaqudgepsjjULa79pRdaDrxlCUu3KeQacgH1C
    XaPsw0ec9721L3zp6//1tfPIli7Ppu04z7SXy1te++I737FXArFBniMo3vex897z4U831Df51n172qXFrUmjsdw0TUppOByur645
    Q6XjGKb9kiVsvPrVLz58t/mdP3gmlYdH9H5phnWPk3b/4xv/oGdAEaJQGyPF4JsycwyLTi1HDYG76yYpmoCb9+JHl1z/7XPP/+EP
    L7nh5hGlLWV/i8tMk5oowRijwm0bLNQgLlTqsPHWN7/gxNuXRMg40nhttTdc9ISb9+Mv/vaN197sJ01psqH30WQms0lIY9TojWpJ
    FHq9OsSVbrxy+CFbHvvohz7sQccNSrQTVAVyg5SgKs4xMZImITHktIUSKMNUca8Hv2SSljsuoRFiuVXA6rv0uN//na+c+eWfX3+D
    ySq2ZUp58CiKKkYP9dakZrxujS4MBoOevuH1Tz9kJ373sS+Z1DYrlm/au760bSvIjRtflYspWO+9qM9yTWk1d+LXVg/fuXVl3w2H
    7Og9+xl/cb/77AotJIZhzzHCQWHl/4cXwDz9mFs8yMHw49nPZoUb05wm3Pmey5zqVNd1r+z5ti3zqq19Sslx3iSVgUvAez5w9nve
    +8k82zlYOGxl3WdV1bTtLGB1UFW11PlZUgtKB2nDMykrjCHbtAGAsaQaiMU6gFwKWfRhUEXx+wqz9o63vOROxw8dUFrMA5Z0ffFO
    +9x0Wguz4Nv9tEOxmbWNaa8/IEZMEMVZZ//kDf/w4RtXh94MqiqPyYfQTib11uXdde35YCHD2bvo0MUHqZUdVE4TaXdgBjkjZnY0
    LbF3eaBvf+trjz4cSMgMRisry8vDDvDgG29cbpxNQOtDlrmDZIVmCMyDAhaATZtCnQue/GZt0LprpquICUlV5xKp+doo9AZOFT+5
    HP/0zx87/3sXtqlAtb0RmxVubXXPtq19P91LzdobXvW8e995d85gixe+9EPfOP8n1D/khtVWTT93RUqpnozyrBz0FprGJ6+DXmWp
    2bv3ml2HlINKX/WKF5xwe9xwLZ7ytFfvmVa/NGC50NzjpN3vfPPjBw6UIAIxQShJSEWWA+x9VDHWmSSYtsh6+PS/X/rpz3352htX
    gmZ10ARXlAsxUEISbYW8KBiOtDJiqiyL9SrLLW967bMfeN/tzXhSFd6gJU0qgpXRtL9UXX09nvTnL426fXUEVw6aGBJaImJyhAKx
    D8So+5hb0lA4Qqrbev8hW4cn3fG2D7r/aXe4zWFHHgajCAGZBVv44NmBYSWoKCHjScJd7/vcZHZFrgSOFYRISKCopAb+2N1bPvaR
    J33yM1d/5F8/vn9lojQALwFFjBJCXfVcim1mua39Yn9pbeWmw4/op9SurIa2JZP32WRtDDFpXvRGk0Ba9AbD8XijKNnruLBaMUmY
    SKgdhzLTo47Y8oiH3ucRDz+pn222r9TMKdC/NMki/KJgqRz4CVgO/N6cH3UgZvGtrkwGiJKbByx0Gjkym1tx09a5c5o0c5m2iWye
    RNYkaJYbwoUXT977vs+c973LmBaUS2GnMB3MRMFC3AHtQZHQBax5W0Yt1ECdKllryWiME4Fnk5KQ91WvGLT1Wi/zltaHVXzhc57y
    gPvs7ipMo34u10pQqx0xUmdwm4Nj1vrGqD/cUnu0EUWFL/7nT97z3o9ef0vL/dtFrkajW4rSuMzEIFBnuJBoFOhMbWcqbB3lGwAi
    OsfAWQrc5XQWqqrCEKZobHJGMp485U8e+PjfvzslMEFjGJQutHVmuK193qsggDXjSV32yqhKtMltn6l9/feAtekNrv9PAetWDOQD
    8KSDtEZvXWJubnKd/ECCdhiLqCDg4p+0H/345y+6+Ge37BtZWwUUDXpBLFmxRoBmOtqzbWj7GZ7yhD96yANu87o3fPq73/vhJOU1
    CnEDNZkGSTH0B0VoWolkqfDT0C8zDatFPtq+HS960dNOvuOyD4geD3jYM1uzK6JnhVln0lRKSuSp3bjbybvf9ZYnDnKYWQfdC8FR
    LgIfkQBm3LwX3/ruJd855wc/vuiKhNxHM42knMHmUdnH5DgXDWomsIlgVKxJPSNGmsmWgd7jzrtf9fLfLQxy08R27HJL3kdrzXgy
    halciVe9/gtfPOOiaRo2UmT9/t61PUvLQ01Rg5iUk0LZgxLUAAryRC1TSwiFNY7wgNNOe/QjTjt0J4YD9EokmWbs2tAaZxOcwlx0
    RfsXT33T+mRIbilRBmqUG/Ckq1yc1rc/ZPiRDzwvMT73xe9/6KNfvHmvFOXh+1aa/nAgiD6MizznZOpJU7oiz1zwUyVJxNpBQueu
    J6yGKYtRnMvquu4N+6PJWl5Y1aASDQyTamgN2mHfbFt0z/zbJx5/7PIhW6ABJopBMiBkFoy60bKglEAENghtdLkFJM3SARZokgS2
    Fr9E2+FAkJoNfTp54i6EWWjeNbo2dSLkv/3hHBhBwrrhR0VWjergslIS3vb2L5x15g/XJ1nduGqwfaNOyHOb24mfRA7WgFOaazjw
    DBemDNgZl59UKDKScDfgJ6LSt8jzAjFaG0Ozf+cO8+QnPOJxv327pmkGhToEk8hwicaAaRJHVa+E2tAma0wHc7UZJcgkeDEVMT78
    r+f98/s+F+LQo0qmjEx84Ax0WhSGxOoMYQVC7Nw3SJm9A0nZN00cBamtcyGosxVrQRGUosSJoY1+5W97ux0vfsFTj9kNuxlAqWtd
    3yoIHYTamgeR//txa0Vj6laoSykP1q6gX/InMmddbjKVGJ3KnYggpZTZAmQkIUWQ6SbpIBZjNCFKAps8AaMxfn5d+sQnz/j6Ny6s
    gzF24KXzfWXlnk/ROjU2SpoQomU1xK7DuybyIbUQGJMYKWpmK42RjAckBbGclVRoM67yZvuyfPijz89zBImFtRtTPPy3nzFJW5P2
    SJnVQDMSA1iDFnH1NkcVH/3I03OL6Edl5hgcxTCblXVsePz0ij2f+9I3z73gsqa2bCqNM/6NEhKLkAgr1Dn0vW/yQZq2G8ZZjbay
    S3E0Xq4oNyv/9I7nHn2YKfMkyVsufZuoaYMPTb9XTFsRZMp4xGNeszIpRr7KBkuR0r7VfcMqc4RUh15R+uSVWDTXGbo6ghtQ1OAz
    Mj1X+unGoGgf+Yh7/d5jTzvysDy0qHJ4YH3aujK/9LL0tL9+Q90MhJYTWWWvXIOnoBYkuY7uctTih977nFGLoLj0yrW/+uuXJ91Z
    VDtXNtaN46pfTkZTiB32htNJE1pf5RVIEsfEXUk1u9OTj2VeEWQymSwtL0fR1ZX14dKwjbVINOQcGxGR0GYm9koN032HHtL/w8c8
    5PcffVLpoB6UEDXCWJvDEDTNBDKCj9YyGVZNOgNxd8UvGGw2L86D99VbXcedtqQSRH/dHZu0jbVAMlsmmE5847LL0t+/9G1790kd
    CtjhehuqxeHPb7l+95G79u3bW2U5zQgQm5puncH6ZqtIgJkSS4frkmSYXF23/aok1Otr1xx+ePm7j7r7X/3pAwx0tHrT7qVtEGAK
    GEapPgZnMiITWhERl1sF6ugjOePonf/85U98+mw1W/ftl4XlQ8ZNPRs3zrtdXYlKOlcoBmayMJRI2Uox6PWvv+mavGf7w2p9tFHk
    PYmUUVZvjLdvHdQbN21bxvOf85S7nrqtzFGYubfCb/74HwPWZoQj/LeAJXNwL6moDykldba0GUExmsZBz06nnrNMAJ9gc1xw4eo7
    3vWR627YWFmLQr08X4jJhCTGOElM7GL0WUFNMzE29frVZDQJUfKsp8wqJkpQErKASZqggZy1Po2Z0SsrDTGMpgW1dzvp8Hf+4xNi
    RJa3DLPWxKooTrrHn3J1dMAAEFZL4lgcqTVoHW0ccoj5t08+00cUFg5IASnhO9+56jNf+Oq5P7jUU5XclogKqHwrGXfllCpHICXu
    9iSbGluWeZNGQoEMZ7aYrEwOXV6Yrl7z7re96LfulhsF0Bi2TU1Mxrzsla9wWTadjtiyJGS52b7zNmd/89su74+mdVJyLutqGWvI
    S6PGCyfteM6aARbEpNzrLdSToJL3e8vTKS644LIzv/b9r591yf49qap2L28n62yTcPEle7705bNd3pc5Lg5QgiHNSCqX6E7H7Lr3
    3U7s9ZBZLC0Wxx930ve//+PxZNwfVMSmbhpni15vYdq0IfnhwjAlP6MIk8xBRsRqMmMoRdaQZ5RiWF1ZLfOicJlJycRgJDkjjhWa
    kDQGVc3qKZ33vQv/68zvb0xo15GHFQsIzC5DlOBDSJqaxqtqnjsRYoIKAzyXIu3mR4kQgXDQIynFGUN4xktmUiMw1HGefr2hExnj
    VBwTG0ITkoJ37uST73xq1bM//NF5Ueqil0+m436vJ5GqrC+RADtHp2/qOnQN7k57o2MjzmxjYlRjTF7kIbbO0mBQaGrbpr7hhht9
    Q3e645HDalBPY+6yEBtTOMAQEUiIorHROESkaSMwmXX01rd/9VOfOSumwfo4uKJUqCgBZg7N7WTXDxpvdmq3M2ZZAklM09F0fcvW
    bT6wxKxfbhmtTYdlpWGybdm0kxse/pC7PPfZT7zLyf3SYYaj+c2FKFKhqDPhmgSS+efKtyb7zX4d3f4gBsozdid3xGBVI8pgcrkj
    g3E7amPdr7Jm3EKyPKe6xRlnXvfqt3zmQ58868ZV2VeH6DIurVffNGMg5dZZpjZOXWk2xhtlVSnszTfv75XD3nCxaVtVFY1Q6cbD
    BoaVnDVtO11cWm6btq1bRsrIn3aPOzzjbx63fSt8s5ZllIRKV0w2cNEle6/bnwJ60iGWGMIqHJW9q2Byfezv3cMn3LIX//nlS97z
    7i9/7F/P+/znv3/V1RPjdth8Rx1MTM7anNlAE1EixFnvgDrOH5eFZdbWh0FvkSibrm9sX+6H+pYHnn6nx//RHcQjzzCZjJxzzjq2
    MC988cvGk41+r8yMcc5sTNLhRw0v/cn1V197Pdu8bVOvv6BBFGKMeqnJihApMlXbIe66JEMkJdEUKYmxWcXcm7a8thF+cvE1n/v8
    V759/lWw26rBwhVXr5xz3oVFOQwRAGsn9qRkxLFkTtNdjtv9gPvcJqREVqC8+9D+wx5yr6+dee6NN91Q5pkkxCg+qojawrS+6SRG
    Za6vQiBWy0qltZqmmiZVye103CuKrVu2bexbq2yG7kqDdJM5ZWbrmAqY3CcznoQLLrr037/4lWtvGh1929uHGPulU2ONtVnmfBQi
    Yy2SzJu+KkTMxKxEpISkkLmKxMHzwU04+AEdyV9bbQuY1lrk1DTRh1AVblyPYtJDtucnnHTY7t2HXXnN5RsbK0WRM5lQe5VO9IG7
    7hJhLplAPBdM4jkMYqa4Z10uIsZANLTNJErLxFVRre7buPaaG3ztTrjDYb3Sbozqsm9FQhJY65I0ES1II6SJyItiNMXb/vGrn/nM
    N5MukumryfKiGI8nzpSsMy03M0N2K81QtJj7celmzCJDRZnX02A4R8pikxbKIjRrjic7t5m/eepjn/zEuy72qMzAiqZuc2t/zTP6
    awUs6KyYna3eJlprputEB9RVMYtWgZFmeICDwmlIYEtBaKOeCmmWVWxdgs2cW12lz37+h697y6e+cMb5e1ZiK72JR29hS4R6SWyM
    zRwbKyJJfZSpc1SVRfTR2HJhYUvTxumkzlw2V/MWA7DOlM2TxJQCFL6pewWGfbr7Kce86hWP2b0d6yvrC/0s58xoJhF5hn/9+Nl7
    xhyp7GjJc+yRKmndjLfu3L649Xbv+qfPvOc9nz7nnMtvuCHt3S/gBXJD5XLqpU2aF4XNbNuMbbeac3qxdpqDBOK2biea2NoytUrS
    lq7evox/ePOfWcAgZJaMYSgTGxAoqgIIfiIaYvRlb8tUzL5V/OETXzINi0JbVzZal2egKDoxOaK2UAvpQXKoAwmhZQpNO1paWESy
    43Ht2OV5rimK97kmY4LJ2jpsuCrbtn3XDTes++CIKoGdlUskJAxwqWu/e/pRr3jxo+Bw896fb922K8HuXcPNe/GcF772lv312kQH
    S4cmLTamdVZa75uCM1IoRSVAHaklMUYVadwr03G3PeSKKy9JkZSKtsk15YUbiiChFQrKXgzIsDFGYEYbk8yyyziGyaBv84ymGzfd
    +5TbPu/Z/+eQnVhfxfIyVCCpLZwhyEzBVEDImTeHfUHnFIuD49Q8bm1i6mfX9a87Jg8eWYbJJBqLrGAgNake1U2/v12AOuBvnvkP
    F1xwbX94GGiwth6R9RUMCiAPCnPcg4UawB6Eyex4iGJcNpmMspyYkSQYIo3Cohac22Qx/d1H3+ev/vK3BiVG05VhVWSoRITYR/ig
    USiL0p+2eOc7v/KJT34ty3a6YmFlY1QOq6iNNXlqXedpqhxpNoMDZvKQBwRfO0MKUHJ5XtdNanlpsE0DmvFGv1CD9VPvcvizn/G4
    HVsxKEFAO9kY9PoinUr5bzDHSt0o5QBTB3OL+V8a3zpby03V1k4WijvyfZjvWAmYRkymoolf/9qP/viia9fWpibvTRpVW5aDhTbJ
    tPVCQgbsrGMjklJKlgKnSZXl6xsxswtANW1Sr9fzsbWOoAFIDECUEgO2YzDBxKZZG/TBvPGIh93rBX/7YAOEST0sjDMWgWBofQ39
    BfzJU992/pW+psVZ62BuB8eqzsl4tG/LclVPx6mlKlt2tAiYuh5HCZybrMia1NbtlEmtNY6YFCxGaFPB1AgF1fXMOshQkuGU+qWn
    9PO/f8ETH/GAY0l9SRSCz1wPwKSeFmVl/vaZzyqKPPimKgpjWUkcubyCsf1vnP3dslism8jsRNXkHCQoMoUhsSDqKiCmAArOmbZp
    Qkh5UbI1o3riJQ0WF5o2Tprg+lVWDUbTdn3cttHGZJmzGV+skxmgxCoOk9+6y23ufNIRbTtZWqrG49WUBGp37+LTTrv31776dWMz
    ZnvLvtXlbVu9eNcJYm2isGFY2SgbjYX1VR4+/rGn3OOu951O1q6/9hoDrop+6wE2ZA2YlCiqJGhUjVGMdVlRGpvHyKNpaht2ZnD9
    z/d85SvfvfCim3rDrVu29/MMxthJUztriZTJgEiF5+oeoJkvuSHYmXTtjA7MOmeH04x00r3xX6+I6WRznGGXcTOdiEY26Oe5l+mk
    rntF+aAH3n3bll3nnfvddjItyyrNiMKRaBPRbeZo/lupj3S5QYiBmdiAjRrDWZ5bW8bIuRv6lljdZVdcPp2GO5x0VK8sGUSKGFpj
    bYKZtOrsYN8a3v+hL33s419e3n7MxhRBXF71J3VtnCN2SGaOypTZXguZf3Yw97frDfJoPFnoD6u8WNt/Sy/XzEyX+uHpf/OHT3z8
    PQ7bCWfhm1HuUGR2dWWlKkv8hovCuYoOzxeOuGP5478/dGYiMVMr6pjnlEACioRR3Xmm4MZb8O1zLv/YJ7709nd/+tqbwsgXyIac
    DciVqtS2MUbJXNGr+gD7tg1JlBAh0NbpGsu4X/VVJLSpzEsyRlQAEU0dIpvJQJjUMBlDYMSqiITVJ/7Jg//mz+/XNg3rdLHIfV1T
    sslT9FRVaBpcc93eCy+7UZAzlMCkhmBIHcC+CYuLyz5Ea3Lfss16hGx9Mu4tDpJNbWx8ak2GsnSZs0wgkY5b3jVtuntBORoOWe4s
    96WVKjf791z1+Med/sQ/uvNktL+fa9s2KkRwSRSGDKt59WtfPZ5sDKrBtK5923a3hMLe/va7L7zw6huu21eWwwSTRGyRNT4yVVBH
    IEJk8sSeyBM8Mxl2hp2qKoktDDkZ1xtsMy5yyuy4bqetJ1uU1ZA4m+k9U+zaH0SB2BuaPOT0O5904jaXudF4/3AwEInDXm9jrV1c
    sI/7g9O/+90fX3b5lTt37N6zb6XXK6Z1w2yJINSdUGuUrJLV1mL86Efd884nH768hAfd/7jDdh/9syuvXllZbYmTsUIqbBJZwHbg
    pRCDy8x0Og4hDAdLBlnbcJktM/XW1+Pqevv1b5zzhf/8ZpTy0CN2G5fFpJ16o5LZlCOYs4FpU5h1U2n5gFDkbCQojDgDCf7qdrmk
    08l6nhliSjGRmiIvk5ckyVrbc+V4Osmz/E532HHiHe9y5WWX7F+5WQ0pC81nYgSDmSr/3OfxVkrQxGTy3IElphBiaFufEqlkMRpn
    SmViQ5f89OLRtD3lbrdpvRoSwwTOglof8xv34rNfPOdfPnGGK7duTKMpepzlTQhF2bcuH23UmXEzERgWzJzySOcswjlmfjOemuXh
    1np9TKnesbWsR1efcvKuN77mr+5650GVITdoptNhVQXvow+D4TClSJ2BzW/sOFj58yCNFPplj+4kC0iVWBlCRsgmos7hABY37MHH
    Pnn+G9/y0a+edeFNN4eb9oeNmNn+8qTVtXFTVEMiF0Ic9IfT8QY0InoiLTKrolAMKzPMRqff48TXvPJpBuUlF1/U7/f37rtluDgI
    0YsIQMyWZh4FzKQsQeJ6kY/+7tmPf/zvntyGjcU8VYYt2LkS0dicU4DNQYz/+uq3Lr92XVGyEqvhzhJNDSkVRTHa2JhsjHq9Xq/q
    NW1dt3Ve2v2jPVlpXWGitCKJFMF7X3tn8y69msPrFJSIkrD4JlIyKTSl8cccufiC5zy2KrFYkEUEuMh6KbIorKOEaJ71vGcMqn5d
    t86WVVmpig8tMxy7Y4494UtfOjMKsy2UzLRpimKA1AVaYUqEyIgGkQFLjtQYGFV471WFmUCGXRZFp20rxGXVj4oQJYnMOqyzNnAi
    ElA0qE+7+/HHHr1dkKqqaNu2yHsEFJlVEmf4bne/28rq9CeXXjMYLtVNY1w2v3K68qqjqEZH00GvfcmL/nhxAZZRT/3tjt36yIff
    fXnr4ef+4NxkgiKIJgKxyZgdwThnJUme50RmPBoTZ/1yYdq0gFHKfGQfkGC/970LP/vZM67+2Z47nXwSEZMl6Rr+fKDl+gsMjwOf
    z8m1hK7b1V3wv84NRpJnur6xzzehLPuqzEQipsgLUvjQDooqxmjY7N5R3u++97zu+ut/fuO1XUuIlUkdqYVaUgMQKB0AWJDOfHGI
    YopNU+dFUZWVCAE2yysVk0SZ4UObNF530/V79ozufvfjSmPAZm09ZQU3Ae/74Nc/8JHPuWJrkzgZJ4SQhJlFELwWRTVLpzjNtYA7
    3AGDzAGF7XkAZUWY1ku9LDR7DPb/1V/+9gue+6htCzBAaSEx5DYzzJKQF0WIwVqnv0m9ql8StuYr+wuH0sEplhFYgUvEHeA7AV/5
    +o2vfv3H3vlPn77syr0bYzOprfDA5Etqy2mbrHNl1Q/BhxQy60KMeZ6xKEMLRyxtM1nZOrCn3GH3W175zN975J2cxUknHRJjcc65
    39q1+5C9e/fmZQeXIWYLcFJSiDEpM22/bF/8wr948AMO9+14KWNIa0khrFHZ2W6nUMZ4nK7fM/7eD68TlKRmbvDRyRhpSiEvsuHC
    4nQySSLUaSQZU/R6PkrTts64IitSEFZe6C9K0E70YS5m0pG6k2XHoMwYo03p6n94y98cfTgsicikbesyH06nMS8yY6gNrTXGvPJV
    LwJMZgsmlkTMzlomphh0+/a8bfPv/+DHZAqYrA2wxkG6K0tYOyEUMrAGuVGrEbkrUhRSYrJGnTXZeOIFTMZ2vkNkDbMKonPc6U5x
    J0QiILaF1fvd98TDDl0uS/Y+5VmPQBAhkxz7NgbrslNOOf7qq1d+ft1NTNaaDGpiSoXLJYlRxHa6PMjb8U1//dTfO+VO26oMzJMq
    46b1uXO3u8PivU6/F7npjddfFdo6cy4Fja06Lps6ZlmP2aSY2LJxJOqVghoIQ5lhXIxOtPShuPqqvf/+72f/4EfXtG25ZdvWXg+t
    og0BliVFawwYPnoyRIYUIpqM4ZlkPsCgFI1GQ2R0Piz41W6URGiL3OVFD2qYCICZG4Baw0piuStLqcxx+gNOGixvueAH58bW51kV
    WzbIM9vzbbKOVT2Z5LJOZ4qIOMbUCYJYY1U0RJntAkkTYpaZkFpryeV2777Vm/eur66097zb0fv218PFbP8a3vme73zuC98SHlC2
    EGE69rfO1JKYyHbZt3KcSeazdELS2imvg3yIw8FwPB4zhAg5a0lB2n2nnnz4s/72Dx/5sOMdgverpU0GZMh1UmFsTCIitoKD5W7+
    /4Ukmink6MxfbQZiIIUmcNeBnqs3+BAFiQ3TfLiplECpTS1MHok76d61MS75afOvnzr3tW/89H999Yc37YmCxYhKuErkghJgc9ND
    EkJi21FxQmJha9tGnC2MqEzXtw7oEfe/09Of8ug/+6NTdy0iNkgBC30cd7tDE+jHP/6hdSYlVUGW9UJMRVWFFIWFqdmxbF78/Kfc
    +15LhpCbmBEsciRrYMkZQRPQWsfCrMRXXL3+rfOvjtSXGU6YhFQQhCJIFNL6VpWIrcsqkGmbFEJmbJ+SYbjksdhbEK8SkSIMW1EQ
    k6iIBuMICkmccyZhzLL2J3/0oPvdZ1dpENtJleWZK1UyazIFVGEsE2Be+rLn80xWdJMqTkTKBpNxustdjz7//Muvv3Ev25woAyxD
    CImQOoV8o7a7R+5973tR8tPJ/hhHudPCUc4m+jYvnXXGEBOgopISNDERREIbVDhzhTUuxBR9KzJ+/OMedNhhOQEhwLGBMpGAQkgb
    Tah7xdBkeOhDb//Ty+TGm/a0PrLNWh+cscZw8u1iv5hu3LR7Z/5nT3rUrp2MNAamGREoGpPFRNuX+e6nHnPnk09Z33/L3ptubKbT
    hcFCirHIC02Y1rXLnHHWh4aMtOKJtfN5EtikTpFDC3Cv9XzL3tG3vnPeF7709Wuu2791+6HbdlQxkrVmMgnTpsmKjJiSNDFNnaMo
    dQgNW2JmKDGDDVTmdg+/asDqml8MdbdWIe70r9KMeA3MJF8Ytz3ukLvf87cuvfCS9ZWRBMOck7LLXRKf1Is0SYJISlGYTF7kSeL8
    KWeCskqkpMxiHfswDSmUVc/l/ck4Xn/99Rur43ufdvuoeNNbv/b5L33L5VtsvriyNjFZ3l3f85ppU+5mppciM+2wGZzCWOtbvzgY
    3nTjz7cuDzOrbbNRsLdx5XG/c9+n/sVvn3zC0EANJj1jTNdPUdtREWTuWqm/oOb8vxG2eC7tuOmqYng2Vk1J2XZU8mSsSxpDiokk
    QVqJysRcTiKmHjfdgjO+ctFb3v6pf3rfv/3o4usUi+PaKPcT5XUrPgpZx4ZFYUASY54bho7H6wSp8iz62nFkGW9boPvd6w5/9ZTH
    PPkP7njUrpLaSKEtC1c4BJ/yio868ghr6PKf/lSFF5a3rOxbd3k+Gk96w95osj4cZve8y22f8EfHqaA0aKZrlsmiIHGEziK5FU4J
    kmBAfP2eeOZ3fqxZxTZTQpAkmthRllvObBQhWGMzSZhOp0lTWZYpqjOmtMwSLML66p5emcUUirL0MWRF4WMSlbIsg2+D96XtTddX
    ty/bOx6/8++f/7B+htHaaLHfg3YQxflQqmumQMxLX/YCzPoam6Y8TKQxtr2yVMLuw27zlTPPJsqNKYNPhpLpxrRqoHkXrVjjPe5+
    4guef/eHP+JuRx+5a/+eq1b3Xmu0FalFpyFMkaIltjAQsJAll5nccm4oVzEpKTPy0vV7+vg/PH2pDwVS5MLOaywOhkWQfMq9t97j
    3z79/f37NrKy34ZgnCMwK9Q3mWmNrD7tLx93r7vtzAyaesVZNWyAZNiotobFj+tjdg3ue+8TTz3lxL23XLuy98YQat/6wWDQxJRU
    2dlRPYqcFrcs+tCABMRCqtqNNpyyy4thSFDKRN3lV/78K1/91jnnXX7tdavHHXdMv2+yPGsiT32dyMNFQWMZMK0hECTGROS6oNyp
    k//KAavDe1nAdALZ8xYwgDQ3Zu4Eqkw3CPApHrKluMfd7nX9z29ZXVmfjqZVVbXtNCvYOgUnY8jaDIoUO3B5p+dlNuEOBCJKqtFm
    CkiIQQDDefRutNFcf/11N94YPv/Fi7929vnKPc6qjUnjigK3UtLu7v00D1Vzg/I5BgsgiTKoqnq6vtAvylwljgzawoze+Jq/ftDp
    tz9sFzMQ/SQ3jmHbJllTzk2sIDO52k2Axv/CITLTiv1vS+CTejYAaVLvU4QhZjfxLdgYk4GskA1qo7pA+OjHLvjgh89+7/s+883v
    Xro24qy33biFugWZXNgYZ7Miy3JDBszIrAJtVdlmXIc27ljaWVDWjFYHedg6rH/3kXd+zjN+7/d/53ZHHZpLg1RHo5JlBhrYJsMJ
    mpYG7pBdu61zF/7oQlXThtDrD2xmxtNR2Svq6ereG3426G2/zTHbQpgs9kpSNOOU5XnqIHGsETFBUiJr3EZtP/mFL48jfAowxpWV
    cXmUNKnbad06W2ZZpcpEyApnbIp+WlgUNmhYN9goM5+5djhw93vgvQ8/8tBrrr1mXE+zLHM2jym1E79rx67UtLkNVTZ5+Uueeuh2
    E2NwiM5Ygp1NZOYeQmYm8ywboFw1gx6QhGIS4jitm2R6Cead/3z2v/zbWU1ayMstIbaqCjUkFuoYMBCmyXAxfPQjfztYgCNMa4zW
    9OIfXfOtc79/8ZVX710bT0Yw1Lc86HJsMpyiks0lmpDEOrgMPozqyVVn/+c/HroNSGBFxXMWF09gmgizXldZmb/slV8/46s/EMoj
    uYmk/mC4sTYqrOnnmK7/7KTbb/voB59pBBZNDOOiyBXqQ8hc7jU5ZHXdgrgOGAyHoxbfPffnX/yv7377nJ/6mFM2DGTVqLK2sQ4p
    lXnOagCjYlSsKrEYIlIRleAoZlaZvaQGFDMTSxdPudNt73PaXe9yl9sesh0MBEwZrYEaiApEYKWwJjfsAIh0Se2vvOcf5HwIQEln
    mgKbgC9lpRlzWwjr0/WiWiBg/xo+8pHz/vPL566u+6xcGDc1HPnYJlBeVAQbgqqgs97TGeKBtdPooggKQLTGiEhsFeqyrCqs+vH1
    1kZRwzZ3WX/SJB+07Pd9SN1LuDU5DgL7395QF4xCkVuShlKtOikzvc2xhz7vWU85+XZESVIKttPkVyUYYxxkjnObpVcziMD/ll6V
    HpStiWA2a6XI1I79KEbJi9JwnsBArqCkMAQvWFvHykp91dU/+8bZ3/nWd3/ksiNXNxKU86oXE1ofyTrrXJBEXT4MESTVlFJSCTHU
    O7duC7VONya5MZkJt7/droc+6NTT73vc0hD9CpJqDdGyc6YAAIqxHtkqB9GkaSJnLuutT/H+D3/9wx/7Armt02BM0fNB1KKfOb9x
    88BNn/k3j3nsb99D6rXlfj80xtlOOA1wIaIVaEqcm97Vt+DRT3x+a7aq5tGbGAhCLjNZlnnvLdngU/CNs9qrGDyJ7YYVOJLDD9t5
    j3vc+Zhjjjj22EOPPgY37MHLX/25i3/yMx8zSQZiqrxo6okjHZYmxT1///wnPugBhyI2/UwsGVXbMXrTJlaRYACQWtKZF5bOlZm6
    D03T9spqdTSpBsM/+eP7fvO7599wi0866mCeKsosEJE5KHEyXTvjKz946MPuuNjXpZKXSz5yx9EPuN/RjeLc711x9td++JNLrtu3
    b6NtVQkslpQyMlzkPsZJPQpN6yo5ZPuSgXbyjo4BIAqYQWy8CLhflvmPfyJnf+N8dgtRpG7rrCiDbyUmspwZ0QIPevA9DIEQrYE1
    C1BKMRGMAEggdr2yBEVr2ma6Zk12//scft/7HH7xZf7Nb33/VT/b3wQejcNgaUtGtgU0YoYR5wTIpt5I8GFxcTFju7a2Flr0ykXn
    XBPqut0454c3nH3Oh4s8nX763X//sQ+5zW0qCZVGv1hmxEnRWpsBmDaTqiiVfuURYXcjHQRKBKJAmTryx7yIo00NdBhguRo0IRpj
    Fyv8zdPudqeTb/Oil72h8RtsCmMrcNEG8e2sxdaB3gEm7WhDckB3kG3TNmpQFAUTxaAxpXFMVW+HxNoYk6CrkwmTyaue9y3IzjhA
    s6Y+ZnilAx4d3dGJPcigcqmdaBpbUy8vmONvf8Rb3vCnpBiNR8OCy6xz8LYicx+iuV4VAwKYudfRr6XB/3/bGOZswVnriogISjpO
    a8hSmZUKClBGEYFxjRCwvoELL7zpq2d++8ILr6iblGcVzAkbE07K1nEEB3hYIsfC3rfTssoBtG3t2PQHFZFrpsmW/Y3VPVVmdu/I
    b3fktsc95qGnn7Y1teiVIBYgwGgix2xT997baLMSAiHJXeHYKrBY4SlPPD1J+LfPfi2ZnNl6iaVdaNq0sHjUaPW6d3/g89aZ337I
    XQNABknBBinCwKmqIwYlAxQOEuoUR5DEmuWcEYPQoh0Pi6KeblRkqoVcZSyycciu3m2Pue29737Kwx94QgKmUwwq7F3F+gbe9a7P
    /PBHFyQeWleycW0d27btFaWGjdHadY977H0e/uBD60lc6hloC1CSwGw6zm23xKbjESiRhtCJWSdCUt5U0iV0Wve2Frs+xbU3xSc+
    +Tlil4UXEjlVhVqoYzEENTQxvH/LUvOed7/gyB2VYr2d1v1sObNFBOqEdorgsedmfO+CK87+5tcvvOTSshquj4Nq1RtsgXWTeoNM
    szjwn/7wy3cswHR92oQQkeUKigkxoFgd01/+9duvvsYns7gyWl3Yuhyjb+pYuNJo8pOb7nbqYW95/Z9Vzi8UakQgpUYEhcsh3PrU
    oMnLrFAWmJQ0EZkAagKRNW3AdTfg/R/69De+8T2YqmmTKxamgQRGwWSYmZVJRFSI2dZ1w8qD3tDZvK3bGCWzKhgpNQzNLEFq0vbw
    3VuPP+6oJz3ht3dvhwFUMCgBxXQyqqrKwOBXb7ofuJ9UEOczPsyKxIPc5GmWYCEFaduQF6VaBGDfBlY28Ld/9/I9e0MbS+OW2BSN
    jwJlRkphZm482+I2J5gqFFTjLP9SNWwBbhpfFFkMbds2zplev4zR+1acy1UxrwcFEOUExDnVuRv5Hwi8Bom0LUwI7cr2LfZZz3jS
    Pe965KCP1ZXVQ5YXGBJTJDJMTghISEmcO+gZ5oZjB8LY/8YxqwZo00MLCTqViTJUrQ/EJifCZVeOzj//0q+cee7aWru6FttoDPWJ
    8yQsCUWv56MXkahRJFrHZDT4Js9t9FPDKEqHFKeTkWrqFxlJUxXmIQ887bG/84BDd8AoCgdWUW1yx2AkSTEBZKQbZaRY5IVI8N4X
    ZalAE6QJyEu7OsE/feC/Pv3vZ3rt18EtbT30+hv3K2Xblst6/arlqn3b6557lzvu7OdIASZDG9Q58hJydh1gbM8GHvHHL5rKIAYj
    iUmFbTKmNRTbekpJdu3cfc+7nnqve55y9NGDpWVUJSShsPAhWWfWNsJw6P7rrGte+up3Jl722q8b7lWLBEzWV3rO9or28J30zrc/
    I2Pdtkhtva/ImcnVjWTlwqY6h+nUqTsDHm0VVpVDgqRur1LLMLaLZrEdNSHrDyPwng+c894PfJaLQ4NmoKQKaNaJgxsEa8a+vvYv
    nvKw//OkB7C0g8wZZd8GdikmUs2NARitx/oYa5Px9394yfd/cNkF379634onO7AuI8uV2/jcv760C1ikiBEppbygiBShEdmnPvvj
    V772I4tb7zDxvDreP1zqheA1ac8tJD9J/pZXvfSpj37ITovk0CAQQpUiyMHkiJwU0SD3fqajlOfsoxpDSpg0viqyBIwnuOH69owv
    f+vrXz9n72rwdqkVF2ISYjaGjFUhEYVhghERmqmUkIjE1HIRrYVG0oTCFEaQfEsSNE7ueNyR977Xne588rFHH1lt3TrrjqvOPNp+
    xWMuVxJm7rCd5AOMdupXOhtgAYkQDThN2WQmxpZyWq1HZbVlqljZwGvf8Knvfe+a/StaDbYmJZDYjGNqIMJKQtxtSPOmOEQky5wi
    NM1UIXnurLWitLo+3bp1K0OaZmqYRGL0Icsy6bR1lGdY9s1UTbvUfhZuGEIqjECxyV17u2O3PfkJjzz9tB1tHUk3hlUvSRaDMrNz
    ABDEE2DZ6kyrhw+yd52RSf9XuM+qEBEAzMwMVYQQ22htCZ+wvo7LLh+d870LLr748p/fuHfv3vVysNR6SrDsKiLnoyrYWm6aDUDy
    vCS2KYKZDXOMscjMZLxG6vuV0dSkON2+demYI7Y/7H53Pe2ex2/ZihBhzUwLExp5thsxz3v/AJJoR7KSCFVkFikhpGRzk4A2ohE8
    70Xv/c75V+TlzrUxl4PtE6GY2sVC0uiWLXl43tOf9LAHHwGByRDSJDPcRF/aBUqQgNURHv57rw5mWRN8GAPT5WU+8ujBobsXT73z
    HW9z9DHHHLboGJDOYAohNtYVAEJUa2l1DVdcs/acF7xpFKq1Onf9bUFM0zS9HI5DataGVfuGVz3tnqdudYCDTCf7e73Se59lvQhO
    c/VqA5DCddpn6hUMNX4esBjqGIZFmdH6CbkMnK1MAMYfPPEVN61mXvogEdXul0kcA7EZb192JDf81+dfnSmqDBxhWAAPYyE2JSSB
    MsgiEjYmoddz0wbf+c7NZ5zx7Ut/ckXtfWEmn/ro67ctk+t8xRKSirG2iZqYxh6P/t0X1364suHKwRJlurK2t8izIsvTtM0oHHV4
    75P/8tTxRtg6QE6dDppBBAwi0KaY5zZAQ6DMYWOEYR/TKfo9+ChsNSYfA/pF6WvkGZqAd3/gzM9/5ZxJayeNRnHKJWyeBFGlK7q6
    esGAjDEEVYrTOHbOSYJE5FyQGCOcO1dZ3li92ZrWcTPsy+n3u+ujHnn/w3dVC4MO/rQJ6dn8oPNb+lZZw0EaW/FWKbOaTTaQIYAS
    EBgGyTXjphhmjR9RbgSuUReUVfGmt551xpcvCFI2AaLG5pn3jbEz0slmG6uDRzGTIqUUAMlyp5q89wALZUlgDTEkBW8sZ9allOav
    lw68kwMHkzIrk8KoEIJDbXh8xO7+29769EO2QQUkbZXJxmiclVutJQLaIKCYWQtIE5rC5QcFrAPZltJBVfOvcBzkZjQT+tpkhwcB
    dGZmGQPqGpMa51+wdsZXzvneDy4QJZu5xscIyoqiaQNnpZINUZKCYJSZRHqlaZpaFWxcSkrKlphSDH6yvJgxT5vJvl07Fh/wgNMe
    8qD73O5IuASKiJpMpmRi0EY0Wc4JBcFCgAQVNUxmDphqA4jgOmUjgbFQwKsmonGDvMDfPf8j3/j2xWSWWvQaO0iKXH3fKOq15Upf
    9+pn3fmUHBaC1Ryu8W3PbkECPEY1/uBJb5p61+sVJ5xwzGmnn3TKKVuXFkFAStGRNYIUQEBmAdOq+MZLUQ184hgghL/4m/ee/8Pr
    isGu6HrrY58VRQxNzlHD2pYFeezv3OfPn3SvyqGZNI5jr3SxrYXYZVmCSfPNspOjtR1orzu7OmNIYM4Vmt8emkSErQMwmmLPOh73
    py+fxEE9TcZZY6g7e9bkseXc6rBojr/t4lte84TFCnE6raps1lgRQOdtYYYQAhSgJEgRrJhOcdUVN11++aW/+7j7k9Gck4FISoaz
    BLMxhevhqU9/73k/uDavDpk0pLCqanKK4g3HTEPPNm9/03NPukNhARLJDIcAa7vtKIWQDDJiiMVqG/eNzAte8OpHP+RRj37EiaVD
    SBDXKiK8DrO+TKMrbAt4IER874JrP/O5b577/Wum0lM3aGOIUGJlZlJWIceZMUYkxeTR+Tdq6mzyOt9cUiYhQ2yhLB5oMxMzh9zg
    qMMPeegDT7/v/Y4cLiAmWAdHSIgGMyxqCJGQZ45ZoQLvwQbGgM3mGomqWjaYt4c7/fZNf+CQIgBrbWexlwTCGE1RVPjUZy9+93s/
    OWnKKP2s2Dae1mpCG9rlpS1120zGbVmWimSMEYlzvmF3k/M8gZoP6w4Ief3PaQtpkmBNptE5sT1XNhsr/SJo3Pv7j73PH/3hfQ/d
    BYnIHCQmNgZAmEsdzp3YZQ6H75phZs7o7GzwaFJPi6pi4iQJCdYaIoSgzpHIrCclCaowBs4iTOFygBGjghNZFtDE+wg2mTPA+hjn
    nf/zr531nUt/cvV4imldJimSRiWFARkkpCiJrUlRlZy1TsWEkCyZ3GXGq2ryGqNGYynrwEhxbHS0vEj3uPttH/3bp590h90G8C0q
    gziV3DEMyEXBetRpxjbCEoaQ/ICW/mac/YUW6OxGFoA22jZG2+vle/bjrW/7l+9890ejJpvaHZz3SWpHLFPNWbcuu9e97i9uexwM
    jUsogWNd5szOYN9+2buiRc8sLKEswbaD9EaGEkwKasgYA9VuZB2ShKYJRbUwDt6n7NnP/eAFF+33styqK/tmNF5LwQ/zfJhnzeSG
    +9/3dq986eNyC9PBRzr9yvmuM6ffHhCD7b6gJL88gxZRY0h1Nt9lRoyy0fCHPnvBu97/H3m2NYkd1+OFpcU9q3uXt+yI3rTTOue2
    cpN3veVZJ92+6jn4ZiPL+1A2wMyklmYBqwmtcQZgTWrJOAIETQvJYQwYrUUksMLetMcvbu997ovXvvP9n7phX5OXW1qvbAuNIgSB
    NzQxsnGvOx/9ttf+aWHAqrmlFDs1BiW0pIHU2lQqo4meiuz5r/m3//qv7w2rpcUqe/KTH363ex5XbUFFgGqVyIUIZ5sUbO5EwIz1
    KX50af3lr1/w9e/+YO/KWtEfhBB8VKPOuTJFTJtojC3Ksqmn1rLLjJJEiSklKAGWyXJXrqWE5ImCNZpZstAU615ljzl618kn3+5O
    Jx571FHLWxZhDEjhaIa2Cg26V+KKWcGSEiBCpMYYw/DeWzbdTQ6BxBSjRI1l3yWIqjVgCGJUVSGDxCBjvOKMr17xxrf9C/H2q65Z
    27p9d3IIKUzGU7auLEsfo6pmudEUNnvwHemElWUmYfyrHkoatDXGaEs9V4aNyWLFhZk85UmP/J1H325xASoqMs1clqKmaIgN3EGc
    aHThWADwAdvzg54fEEjHNk6pe4udtA6iF+d4s5VOhBgQfKpKEwJCQge294I6oI245CfX/uiiK8+/4OLrr1+ZNJTEEhfGlj446woh
    tG3jY8uWrLWbaCFWVlVNiDFCyCiGtmz9lHMUPTNpN0IzOvrQXSedcMw973z8CccfevQRxEDbTjLncs5meTMBqqvrNy8tF2AZT0f9
    almkgtqDQElx/umtp66z5qaujzcG/WUF71ttirJIgpf8/TsvuOhanx9680pb9QoNqPLljf2rgx62bPEvf+UTj7/dYs9GCT6TXuW4
    Gce8tHULdnDdzAOtaiAiRtZxZlNESspM1kERkqZpE5RLm9t/+sDXP/KJb9Vhq8ewCZJ03KtcTlRyTNP1hX7zrx956aBCr5hXGL/a
    8T8GrJTEWt7ctI0hAC1w4wZe+sqPnnvOVeDlvFjat75aLWY+To0xElV82tbLbnNE75/+8U9N8oNKUmKCM3PxK5kTxYJEY6wAIiJR
    rRpnQAyvYKMGE0Vk2Cj5aOpWp3jWc//5x1dcJ65HrvIeZTGMLaeUDPte6aejy9/7jy+71ynL6lOVmZRUlbjLoMVbhgFrUiizNT+5
    cu2P/uwtPu1CqlwuSa7PetP7PvTUv/rL3ykTthewHmmCSOAMIaDqIzHqBHLwgjPOuOqj//rv11+312XDmHg6Tez6VX8pSr42nmSZ
    DRJERAnElsgISFWZuQNhibYqDaEhEsNo2gDVwmXOsp9MNPhdO7ccc+Tuxzz6IUfsXtq9A7mFc6hKkCJKXKvXbO6yLLewXReSFETM
    c3d4EWiSDrwOQtQWgKG8G9MpOqiSKLBvfZT3FsjiBz+q//bZL+/1dt+0f+LVLmzZ1jYhKXybBDoYDNbWVooiB8WDA9ZcQPLXC1jV
    oNqz98bl4dBIXVpxOnnFS5556imDQQ4A0/FaXnJhch8js7Ocdd5Lm95DCtUUARhjtOuKq6p0hhQQRmCJEENsyRKQVFlh2apK5/Xd
    tq0ztixLAE2bKDeTAElICT+7tj37Gxd845vfv/rqm4Uq5dzYAtYmgU9RVckAhuvGg02vGjiXtW3wbQTYGQtNkAhJhmAMOcsZU5iM
    +30bsOF1/fgTjn70ox54jzsfs1AhZ2QMVvhWmalwUEUz7cp5VTRF5QBu25AiZ7k7eJhM5EENMAIY2ALNDg5YXYaVII2PxuTOWB9R
    t6gKPOv57zvze5cnXiQZOO6lpMuLg1tuurIsmm1b8eEPvGLLAIjNMLf1ZFzwwLBBNutCKEAaATHEAMcI22EjkxpHRCIIAlJkoxbn
    nb//BX//ljr1uVhcndb9wZIkij7kKqVrQv2zf37Xq049uedbLTP6tXq4/2PA6mgH3Q0gIsxMhEDwwA9/Mn32M9/UNIOoW4Jo7H0F
    QFNr+zjGVWxUVLzWAZWQbWyD0SESigoiIqiIOLYzNlm5IETsQMXCxMBrYmFgYQAm2C1iYrdiAQoi//c952w7C5B7v/j/vu+7M7bz
    9vv087xxFHkLYnf53qsUTY90Ff0omD+55r5V9ePSbqdSTJi+ayfcubhoToBxv7TQlK22Bax+7ZGzzJRpom7Rge/SWt581K0091Gj
    HlXXHVzcXjmZb9mE5q08NpMvCPnq9KDqUOvq2bc91qffPHX92ZoDgacKOsccWb3/7aPMq84Rp7kl3w5eqX6Zfmjn/KoTN/6QhBfR
    tnXcdGtC4vqct58KPd/PmZzwbfnEPUYmjv5mra/uR1o8pZd1X9W254Sdu58CV7GmRVTl+Su78kcObnbUZNuVgytM+/j/sXa7a7H1
    iAEhzU8jA1M3rezsYxK7/I6pb3sK2rpJv5SpIWdbzvIY8251pHHHmudGRr1eBvh6h7o1zXz4PKekMmtYa6/7czbcm9XBuF0v423c
    kYvHh7+5wtg0d6mvr2902LPjg1vMOzLFr3nQmFEvPx+wtB74VHqLlVCwb8yRayufLZ+Wc+DW4kvb3XdvUcamjk/syGnh6nzqZNaE
    N/mfzVdNSHa6k3P4fD9eQc23ohXrZBOHTKWXDRd0mzp1UzOjxUMDOx8OSt24caxiyLUHB5C+X49XXCsTp6WNKk7/uuKi8Pax0hz2
    Hy43o7OnLvv09ULMuqTvw29HJZ65UvG+y5DMUCdhYnDYoKyYuJlpcwIs2smtLi5eOj3FmNL3XKbprfRjF+fxXL2FWd7L2pv7eS17
    1u1Mn8Lq7tKFTd2/HlrguuVNq+DPFKuLU3nt2u4vib9o3GzooudjZ1qv3CxYPaGPX5HdoNPGB1f32zmwebBbmxGDrj4pqlC8NA08
    IJm7c9X+xWd5ncU5J16eeteq9PtFK6/qA9R58VO+Ohb1n5sjfylljmm1ZL7b2oTTb6OqYwb5bC3Z524dEjiS8TFgWTOTF+mhnPDl
    QVaikS/cD40LSAsZd2549xxBm+4VVXYLjxSnfVp0+uGR7LXFM8e9yqxqsa+yf/7khfmVX5rmBE0dzZv3vcZprMhzUHh5/Ojtfwxo
    kvXoS83476tLfnza8mAURRixdtd641WBw4ePWzHsYc8Hh5c+rLgy+WhyC+Y0vyfp7jvHvihu4bglwHwzN+uAMKYyp6Gxw7TPJiEb
    Q48fkax1S3x8oWQvc/yqZU/2l7e34ASuN79y3JFlKvn5qn/UxPvXDq7dRRme++yWWWavZs1Rx7AO5k023L372PoRY/LHIxHlbXzO
    Nm0U/uI9++j52N25d6uEruK41lVhRRk/9xe6v1qfmXpy/5tOtN4lQdkXUVqj0bfLhJn0lpcZn3K+ze2SLyuPz1d+8HN6JzrRIjDS
    9d4h2pCBDwvDLEbcCbl5drBvpMXpPghzxtIfEzL29n6b22v3AnmM/anGQwOp8v1dOV272M7uduRLMSfuwc/yo/EfRjh/TvGJf09L
    cuG9pGy+v6gfMie9W1fHT8aR1XbysuWH3VecnrV26MfK/Zd25DUMj3xf8q1Nj8o3Ll5ZTY6u3PfaoXVlRL+T5++++fxZZsy+vPjh
    MbtJAp/qmgkHGoTNL7i2Z93ciV+LnLwq4gqXd6avFVNCT9EeFcwwO3A2YMqRXZ3Ot9swwKThF+/B3mtFvmtLR4T33DEgg2E0/8nG
    bHth28ZdWk3weBdZeEXC350//srw26Mnv0ZXRCXeYiVVBI4+yGqycbvjlZ0HNqxKayjc5967zJvTcFiBMHrRlrXZG7MvthoceXvW
    4fCOLR2ezn+MSCxjjZdQlvQpvPx+XrzL96DMNYrZST+LFjgeHxXawfjN0xQk+d6SlNT7LxLnzD299uyGzrefThqZuqXsDzOTqw70
    JrR1RtRGU3pT/4gL63D5Yyv7Xd9e3E3xNd8faH1hu9j3ODq1Y874wDFVLyqiEk68eZgrGC+o2i+es/Rx1xWmt+N7tCup3hXz7uHy
    aSVuZ5s8tIpLn5LS8oRCGVI4rH1bzpyVKUjDdRZtF7O8qtM9S0tL8iRDEk8FdnFr/WAv82HMoHH3gleiXhk0u7Nrh9z/uSr8x+OQ
    Nk/LDjyf6fc9quWII4eXXOoZftNhW+ZwydUm7i8T130rcvGqcF2086Nlyo7Ix2Hih86KQnbZz+R5y/JLkhckfp/Xg2PaeVt6g+mm
    2zfc7YxeVnZh/XZtpfFZa/uBq4xWWDt577326LBd5bORJYdTd3SLXlPoYM4tWrSocYR14acMr8rqKs+3IysOmvZpb7w4midskrLQ
    +1zbiB7eLWOnzU4qrLqcWNL/8s5q5AJz+q3b+ZWlQa7mNZ8Wlk/81mjd0Q+rVocMa1P1enqPSdXlD9ybPy7oMnnSbM6CAPPnl82O
    O413zQjbRUlq3fYsMrfs7r1JRcZ55e8v/HzaoiTn3XWPSzf3z5+yRukR5xfazKbPjoZj2+05udV5qnngb/1bUY1eyk45FhQ3sbPI
    HSNxXZtSkH1xrb15pcmoQkbmrM2snadmfi7mtPZ4fjneuYwW71q+8j79D3rzaGl8o6YpFvH7rg/l7rLYuv2jMGRRE99nzYN3PTIe
    Pbl9xveT8VETvweeOtHRL2/vw7BcdsiwQ/PlH5DI2TWFU9DQfY+6BbbxOlaxK95DdMHCNC99FyWs0azVgyeKK5shxztEFg5j7hlw
    s3Fv+r1dzIFzTZBuZrn7zD7EFTVuzzYOMGkwdOcC6bUh4iNfP8de/nbXKawos4syeJjJyzGm7z5PfJUcJN0ecNPKN7p3QMCGlBGx
    6aPn9uo3uOhuanr6zaomq2SFfTwr745/E077vth88n3GbM+vJypySwWHSvb+1niG8UH6/WTPqZX37wQp7RNfLBU82nTaRd6gTdS8
    Q/clluxi+dNmnj/tasbfCDx07NmZyeJWjanrX6QPjfjt3JElu3nDWBmtZ32bZ2mBPik7x78aYTYkeTga6NvKeV5+4ovZl71q3kTv
    yfq0xoGzdP0Li/aC3R7fn21uk/dyd8uG8xZ6PvNb3nf74LOXsp7f2NDLKGA4O8s79F3kiUFjKubsy/ve/kVW9c+igH0PB+7LcRjx
    ID9k1ZjHm86He358uOL+C2nft9PuK+4+86uZbDVWOfVLp4YdBzgVFXXq2Cnc/0iM15gf2ZPz3sVsHP/KLwLp3/4te3dJkCPz+jT3
    Hw8PHfvpEu+ldOE4TzNHp1osbjg4a+Eal1UHgpzOhDdrZbTfVCQI9//2JL1T7vMeXl0uFrmanssx6/Jk5Uikd0/LG+VXqeUPJDWr
    fyKikcHTHf3W3ZmU6XNzFj9y/fH9fhk+U3ZIFllVl90Lz+y92j0+60TekIlf/dxmbY093TG26qtPj4k1D5Od8p4+sGs21+nHqfxJ
    j79EjflA/fmTM7nCvI2nSND2opkpfZGpVa8+5/s2i3V9N6Opu0nEj1dNvbzexr6cWFO14v2De1u/Ch2D4wLcznm+39+xQdC4EaPn
    MA8ObXungek0h/PXXpiG7DncgX/DoSFtblrKpsDk8y7bBs7wbZK62cLnRV/T1nf2JhzltDn6fnz/rsk3zd3avLajWlwJdrJuoYhI
    LIyN+/Jg45yNF3lG5bt2t5o1qKEtd/fR3+ySy2KuuFU0nX7CsUdLa9vtbGOb08arbCk1+d3TDyand1UOKDvi9ONS2/5jR0mMYgeh
    LoKC6uhBqULRZHlX57IlFlc6WAQbT005ZxlsenrGyUtS/sBFjdc98jhW+L0kt7KgS/iunCfb3+7KZzsIfG4GdPO/cXBrE8+2Z38P
    N05tE+B7s5gX9nmK2fJiszcno3jFST18xmZfHmXUrlfBCC5tpuOVrZRtGTtP54y/L2qT+3btwsnffIJKw3mJ/KYdjE3aNCruuvF1
    sVmfSN7h2G3fN825Q2emTKebtt1qk9DFbHi7EWe8cl++jap542S1cf6jwt3J7+eVjl/oc/RZ2rnpDRsENz13ZMrmlG5fPkwZ1nm9
    z8yqq3sYjLkXzudWlUseHPGbcoDhfKpF1I9r+/KUQx/eOxDTu2x7Vtq7P7L3PuDsKZ1xrwxGBKl2qyYt6MUaMKLrCaXZ3Q8Lhp9e
    adG5/ZnTpy8/XhV7127t91PxyR93Dji1Ifymcu3LzW3msPYtTSspPtztSlJ56YOa/afevD5PP/S7MCN8bauXNUHHTx7rHjq8rblF
    7MopB27deigNoMLFnagf2RWNjx4NyOXL6azxm8Znu1sPy30c+kF6733pz0KvS626nhjf9Ei6+cq+/n+IBgxNs3+3k5Oy5TT3j9Gz
    N10ePUH5ZWenY5+OeDS7U94rWdq934UBYR7zrrSMGJl0X1JSVprxs2/s43tD0+2nT+yZOGbugPYTx3w1NWNZtxj8+8S3PzzsJp7J
    FzaOiRz3Jvfe8sjXdws9K80rjl55sNCdcntK9vEC58PXb1K7/9ya/1tl57J+xig6JPfC+DaTT0sefJqxdvhiWWuXlsfTUq0bmpy1
    Gtb+5h9ng433xNjRKj8Fvlwb/2kh7diTytXltoMb+y2SWUqRzLOh0OwYOnDDGOfFc84HsHwDuSt3mm/o9cV1JnPtnumsAgWvSZ8A
    1tYZFk2XzC5iDJ5zd0yJ6F52Ts5Nt/LiyIhD1A/FsrNS+ta5MjZ7D3rua+Xafg2LLzpedDzu8m3GKGn1pYRCtyvuVe/knZq49RQO
    658dNPXwlUs1SRfPKyfcbPfw65vMkg/ODtTP97uIV/OSg1tuXOQbI7x2cxwzu/USx+KL3UzMMip2/96mh/RUWVlCSGnkyMPolbwu
    kmL5g/lO9xOeXtrxNtA4d2t6b8un3F2LThVto3fOHGgjaGcezes4nRtyO/gQd06qpXc7S2bjlJhp9tuN046Zc5cPu0druaPbvRUL
    XCbcDrB6UL7Jrdu+JCc0dWmwmUnoEickhrLUR3ins1FKv6HFRsuXRYdMXWQ+e2y/UGfbzkVTq7t8frBv0ufFbSZu6bzZlJNR+PrZ
    8v7u7x5eSf6wDt11Y2+HYCRh68ylzbcGuDaatHJ6o45t0/bPCqy4sWgD6w73etqRDTFO4xzmrsj/MKzm+/PSik4LxBvnN5zFpVgG
    rtlVUn2TE/X1pBXFDtl3Lu3zxI42DZdZGwsdvFs9z7p70Nt04vPUdXmVgT2W/lj5bHtB6QVKL5e0+4Jzd7lxpRu7J17ft67apCy2
    xUjx8ZTEuQMXTPKgva/Z3OtH7/xWM45cnts11TLNbMKVpld+Xi3M82gW0ZD6W+9w32utCkVbMz3c9giqmy3Mvz0sI2GCeDEjxGTe
    M4pT4JJOt0b37iQPfxaUU9D/ue+tUdOd8p+2WFdzPrF9VPHE2LQ8CT+Ykno7rMrFRdz5wJnAmu93wu3asKmh0/a9Ej6NNnnJpP9e
    cpeyu+cffI+MbzGS6juBSzoK10h+cC98Nk+q+NCjvOPGqMSbp0rLb+b/eJ/V6KaNyf3wvM83nHI/9OzjI/Wae9J387azd5+0q/4+
    bFLFz4rDbz9Pce7rMeG3duaWx4sGp41tmb3pvuObriOnr5TeOvCaU1P5el++UhxGe5gdP7q8deeFvbcc+72lv2Bw5y5mbY1Xpgww
    4SgunzHzm2F8eZTFQJsZA5Y3LXC2GP4p4He+wn/t5leZRdnmTRCbOSkN2to8mnq66cfnkVUe3bI5g/KjaLRWnWZPfyzN2bNnT07O
    iNazsixScr8vXwbQnnpg2cHA/cf2J5e8WbW3/HarEcARjfyyLj3rwLKy0gkPorrGO2xvfq3nhb4zTB7db9ll85SBt63WTM2L7nfj
    yXKPnB+fu9Byn2Xb7jpSXrC6XJQ6rsPyzuvXNQsy9vVPH7WV2fmFWLF1u++WLzOXr+QFn9q1eut2Y8tO17ZGmy3vdcPahInEzLHZ
    tEfs8sxPWF1wPSTxyadDSc98eih+vI8TP5hrtr7hCuXXzQMWrbx7kju/YqfjlfNm7bosaRk9sHfnXlOuOgYWHWpmY+48wv9F7KxT
    lzgTH7/1gq++ZDxY0XHChUWmJ3+KIj5djPH6NKbCOS/mcKs82UzvUN9SK7+iITaNFkvf9Wt3qVu7V0ZOtlOstwTfVo4MKyrmbZ85
    /kazz70d5i5rYpd8KSZP+eF5e4mQ5xssPdPW8fSYjA+JVzhjPm/o37H1zN8a+6aM6XsNXXPTqrHPmbNWB5c1jEl+v+V2/qfcYSU7
    LhUJUlDJ3ZuczOT2g3oeXv9JkPl+4rlB3feFntm0w+PNrafFtrNMtlaXnQ3NFx8PCREULIm+2uO2kpnYq+riy0mP1rQQNz+y41DO
    lxe5t1svlwW+XRS6aubNXcPvXKr5fMKpW6K15dtXne5eOHdImeE2/I3vb47VKYcbmf/RYPv4/THJVZW3j32f7rR79oZ8kbKns/2m
    pYJLbofEsY7lDjciDlIu3Pj5efHCmtx7EUHdVrj5LzbdeWBWIL9Dy+yhKzeMMvGsKb1dfcrNM+8aP+zwIvGPfTVfnu2YFHchPK8y
    yzSbvW5Y0vOIkspj7UsVjU+1/bhi8vcMW7sHi0JZ22b2PPjMJ3fPiNRCp6oJV/b2f3/s1N7ZP9dZVX8aenTF9kvpTBtnm5uDuHEt
    B829tul9xRexcN/BqeOnf8mffOkQvI4vSlZWOatiypUDc7xnturUJyjtnHBkx+oDB+aPGfoikfb7at7vw2JTGZdZ3j4hs0ZcGClO
    Sdtk5uNmL+r028oTBWaxTM5g/l2ad3rTnUFDx2yQpLJjt/qItm1nzKJYOA2UHhw1Mn2UE3/U64H9VqD5YQHF5i5HptzKdYwVie8W
    ejzpfjT+a+rmYke/jG1bHnV5eT8jOSm3Yolg3969ew8yxvQQDc684brp0ZdFOQKf1HE5F/c1X9ObkdchmjO49XMkLWRuQPszO7Mb
    rLc1bsloKxhs05vXZOzK9MTy5MC5wUMqRp5KHz+r24xe4wos9/FXX/ssOxBb+PG0W3d5xdF1e95epaLeESseFZqVs80Y0+abDQyI
    bhrTytw0scFKWrfyLNNlX4IbTQvvInAV7KB+9706186+/4CQTOYVr8pzpRPf7Ys6yuvt9+zUxe+v+bd3ZU2uOjR5wiG3irV+Jz6u
    3zU6Y9PmpbKXxu1C+F1el81qO9StQ6Mn0xuFzT3B2yLaQosf7L9+yfgHpa1uNLYs3Dr0XLPOnj8Pr8v9KeqRfSvkbS6DNnipRY9s
    tHD+7TEHOTsO17wu9KoZudBGrpx1ZHlVwjE2c3jRxRdOu3dR7boxx318Osh/d/KDKvExdtq7Pf4Fl2fkTA7uGLZ9TRLbPXKUsWTC
    jXXrpEn3ZkU079nBLhOgy6dRRPWsRkUJySVj2QIOej1ixJQz/htvP7z/o/pCzets2yc5WbMuRMR6CapH5YePOP/208/rWffLX2VJ
    kp7uO0TbQN9aWNYvP7Ss8ZcLDTuuepRlKxl1orT604VC9qNtH7e/y3507UrHk85bjozZuemsFS2qekWb3IkvowewH3zsl+FysZfF
    4mncbi7LfZZMC/abkfljXUXNx5Me9IhzByMO7vpxoTAq2v7H1+wXF+0mV6yO8nr60efbst2VR9/fW97oW6HXiagfZ3ZMKN9zUCTK
    Kx2S8W2lW8nznNKMpDsvvWr8opxcBnnccTtbvXrMl2NHHzy3nmGyqig91cFnquvBXfvbcYwb9Msd02bu/LJFnYKaBc6TGQ19F1U4
    hJb/KcnryDer8MjJpy7tcHqQY5eQI7y2asyNPzJmzZz5fEsr0b2GT5gvWhpt+HY6I4vB9bPpvczXt8AkIraBtcXiPmkpQ9r1Xvx0
    zMvvP1us5A8avHNTvw19hrp1pgQNnL8wi7X5QpPL5RHr0n5nuq9u9/N8YIWLKKHAp/Hxiwt8+OXlSSNbbl/8eefVxqW5FwanNliz
    dFnK8S3T9zSw6DpsLsMyQOjMuRw9pU3a4YEpxv7LjtTcMm8TPc7N8TFlWvvOnFSGTe9qyexrOSvtQzY+mbMpHD3a+o7s4rWjbR7k
    Pn+2ZIHXk5QTyzd3nnp0UHX2LoeQbwdR3tuuviGpBYcvbd8i6MKLG7nl6h/j0wo8Iu4MeG8dNn37p5W/F3cVTev5tVnpxMJMz4+P
    7Y7M7Fkx22/I4+KN3yp6VB2vWht3vch84nzbya4/i2d75Vx1skv4GdWqEhFEeFG/NDe5dlYYYsX92KLBvCPdFY3nbN+ydWDL56YO
    PTMb0/ZIpsknJ1x4+/DHzy4VZ1KYdlXmAwsub688V7mA+hvn6/EbQya+iPKQBSN+zk2jv9FZCxo/G37s2cAXL3K/eU3+fr7a1nqe
    T1jh/IizRY5hh6gSYMtJvDd3yFdeO1Z+60rOvW3+DojoqzDIOjo2y+8eYjs5d0rZ2/QcmzmXqMipzb9ZT7xaPHiBW3bNs0Kvb2hN
    lCimO0UwPOPTS497h5QVkbwPLfdF8sx4Y0eYnhw8ddSGYL9l1q5j2w7caEZpxZi2ldFhWeP5601YNIZ/L/OC7F4htDkt2lqYGTcc
    PmtuQeeC+Wc6z0p3extZWf6Z+nJNxaUeyYVtz1mEPf9xu7p/+O3XX3KTH8ozXlbcVyQdfZDx9vLv4Ytivb/+iI7fOfy0Wfeq79fs
    Jj1/vGP3hMVFA+YJR4SH7WX9Xph1bESaSbPejm5Ss5IjsbMrXo812lTUv9ThzBba0373OiiHLSk/9eZq+Yp1bo4Ndtiw/FK2tB47
    S478Zn0r4XFxjnAG+8mp/nweb12Io2jz9Mc/bMeUP7mIXCueMKry2BK7Yz9FnX5yXy4MWjXqYYuMgAeTCp2K9wM2Tn2z/33v3a7K
    iV27OGZ8XmUblPhxqk988vvYIT+4Nee6FAnvHVxPu3BzcezzlOZth7fZspaObLpBGdr3Rujlw0Exgozex30dpp0bv2byhzedMo4e
    aW60J2bIsZcZpZ4fH/Rv3XXlueFFv9s1qa4u2jTStuf7DdKlZymrQpdv8by5fpvZ3VWmd6WhjIEjJ2xeGuI7aoCgo4es44+GVTWp
    Nd3+eDKgrZGpTeAC07YjHs0fnn5h1qmCFYHDrmxOdNuzwj+VkW8Rm7BrenoXk97vzXv6NxQ+c2+8gDGCcX5PcIqxxcm296f4TW8o
    R9fKLA8UpfB5IUGZYTasgZW8xZsD16UYy9OcF3N6j1iZvnlO96E2ZiY21zaOtAj5vcnQW/MyJXekP+dFJXxIzBw8xCe27/HUV0M3
    IcerXAfM8PVtPcpH0Puq7yrTxk8Evq4Ly9KajZlhsZAyp6vNFkvqqT+yrk+krGiS0db0Z9nU6VHf33QrnfTynmPG2EmlFbsTPj4u
    7PEgd4oCXd5p0/nw0bKwhu8qA36cni2+U/XsRtjprSce3u7uUv3VJovi26JLVsDi1ytye8VWHzK33Bh036znDO9I3shxLcc1aTnj
    quRj1tHJH/fmVx4rnLElKC0odNnsvunGfdxEJTe/lOZ8mL9iTPbIJOuZxtveWAbOP7DA3PP0rLXjKz92ffhjeZfWV2TXRiU94ueb
    vf12NbZD3Iez958V92J1OPPm/G2RV/zq1gemNH2/0y9Y2S2Libx6mTbh52pTipHgxZFqsdvqkBkP+x2euJ1V0mWM4vHy5Cq7ypsx
    eQ8FojetbJ8yja87dNoT3OX8zy+Huju/fGd/2SteuLeH5T0Bv0lh49mV3/ovfLWk7Zd91Veqq9/vXE5dt8bVlX99/6s2VU3ovsob
    V1qcYY44OumM28PEbQutV4/t1aeF6ERb5YwbVzs+mXP4sVnWXnnuw9ySxK1R3eO+F/8Y+3JuhzPl7+3GRZhcbcxt2GXP+/2B7L0r
    er2f2bykZK/9keUrZi3t6u4e/57Vw6M6dUnJx6MrgnafKB8UNibuU0TGj9dTzy2gNq0OuxCG9rxm2briYc3NrEPJ1zpyVlSX7xjz
    Oe3AyGkWPqFtrd8GW+56cHNNr/Q2a7b0aeg54Tdu98RxVk5v7HfnPcoKjwwuKp68J3FcyeKk0qRvr9PvZGa4Xo/8evfAjZ2tstd7
    c5nnf5+S7mPT9qHZy9R5pn427CfWJ6Zb24+y3RHNlw5eeDk/VLx0vIUvxcgiYelCm+DAO6Z+vFTztY3F2RaSVZZS02ZX2i1ouOFj
    rHH//k39N65qmOXB6LzReo7v4mbnj6cXtOkfwuhu/uLAUK7xBpve7VJmTDGZZpzajsN37jqu3eaxlM6BFteq7jYN2/LUx3TBKHMW
    t1uvhNN9rqWsp+28Q899sfPl2kmVyy/ZDlk5sP3OiP4972aGmPh/9bbdNDe0Oa/UZ8/AiNULzrvb93samGZUeL13h2btze5SZs7d
    5Dinr+TqsOaW+1oOHHw+78NlauqQnXmfXuTfOXw0Xjqp+D3tXfFhyZdVZf5Pj3USJk984XZupQsz6v33xOrN99okl65vWuF+uClr
    M2/+k8cjMu2djjo7RaQeikwzPXlEMDZVxotunJ1UvqbKbA1ras/T++/5TbZWdnBdcrBty0biI98XTE76dm27oNPTHW8K+nCNQ7nc
    q20uvWtme6Miv/RBzBpmBv9bk8b83gXWbZsxmIMf8Ffs+GBUmlQjsm86cHXWq9bTg75YNukoaVXgxsmjtH8T+SM7K3zM9x+Tdq/i
    d+ntcmLp5tvn+0w/HlxeM9rtyqhBV7b15s6dmjQwnTd2nkODnr6L5u38umj8ezevT6fWTneltLhukTW44+xTGUVBG8PcD70+/eRg
    xMNbMv6bM+j7NciUJ0NP/N45vIHR+oJ+PXp8u/ryQ+LYFeLAfaJl80essGiSOeBYTP5WxeR7BwNvj664e8nKImbRwozcSyG/2yjG
    PWt9+sTKEWFfhr+/x9nR8+vCHl0mfflQ1ePOp2r5ZdeyB7fHPE55u8Xhbu+CSu/3O0/QNzz/+fhyf8mxMtvpI0zPnH48vSos0PnM
    srD1fY7GZEz43GX3sR+pl2ILXV0PnXhFL4m0HnVh/aW86pwXR/KdE3JCDj4IzZ2DTjwx6NCxl5HhfIfjV/v24nNkKdu2b0216NBz
    le+94KJr9n8sHZR6YOdV38ZTs7d6LOh65krGqw/F309zFpa6KBLW7zp58qSFb9vF1qxxj1e27UzxX2qJzqTMizZva8HuZ+Rn3iTh
    ZgPab4nN5/ed47tp2Ya2h01PzL4+qsv0TYmMrjfvrtoi2Hmr54dXuyWXRvZTeH7pZt5s/FmHtOhRvl3fvjiefVPGtOk9d70/supm
    ukm/1KE3Ur6cPN0wOmb/H+afFvU9v+TJ0BETJ4+wpA1qH7v/2eXB1C4RXrnjyuaiXbaEi+7u+bB0vr+ttbfLDMaAEGPzDd0jW6Su
    6TaIe+Np6JgODW8fcBpolf6UY/Hm0fQDs4wWjO3YY5o8Lni5t3OA/VvTKWfYHZbRQu5Nm93m6p0yp0nlsRXv+x3pSom4cj/hwuyo
    C61eHM359nnHpWfLv0/90Ctr8NuHT4ZSPYdev1y2+qH0DHrvdiNr2bfmggmDB44zH3SsdRPTa+Nnf9wb2n/GWc801irTaf077/ny
    aWabSV/t7atme3ZE0dRY27XKow/2Wjht71OVFPV9rsjp553ny8c2Nj1jPbdrBz9RaMr67ZzlnRq8Wpbzc5bdj7G3xwxZezjjUGJm
    21Zzr38t/LzpQ7j/+SVFmw/9LOuxID545R/9vftFb5h9cOzY4IovDj1mxlR7ddzJHm26eHn6zqq0zGe2W2fcmTV1IXffvDE/M72L
    Wu2Ie77txUyHIRb3JuTkyELOn1mX/3WmVfKPU7dH9WOtLDWldo1YWpjfasajmfevef58/aHmR3FmPiVuXPDdXt3mzT995NzafZZN
    StCWrhfezXc6Vn3s0mtui6Hrw8vS9nv0SUxs2KrEIXKOctX8W+94RwPvikdOn3Xt645tDozlM3JKWu4PbCc70sRMHjm6rDz//uHJ
    YaO3dBiyOzLgoMue0r7WrFvZvHHPn1zLdJIIJz/+Mjm7OKJrj0EWX6Oz2ly+sXlF1GplUZuBsc5+O9esfrB/r/ItdfKaO48enRha
    ePr3o/Qen+8XN0YcmHTu/q6HHvVffjd804EDL0Y8eRFinElv2JOZMfH85FKvSQk1yW08y0+LOy7pxVm8hN15NHNoUoLlcZsl1pz0
    dB+Hxv1MzwwemrLi4KjggZ2n8BXLOnZuN2dmQCpjqHOTOEHHfiYWRgHj+g2rbvLiZcyKrGuDApv0nJw9y3zuuzn8r48YS9s94Zyw
    Xrxw0DiHGZvNU8x3nc8bf4Lf5PbFzdf3nh2dsW3FmvgiQYRjddGdK1FHmzTnNxwwltrSxCxvwarTve+a+HV+/LqDUdaBawOnjnT2
    b9M2ZlhprrFA1H+9ua9Jv56LB/wu7uU4h5Vv/LloKV2RkrJzp/9YStyn0x+m289+OSbxk2BveKeZr1u5P389c3feJF5e99j4J9Ly
    QYeS+U8VQc1rEip2eVV+yY8N3PVy/dLqD9uH+33tT5tcNq/kU1G4I6vlxgHj+ZZdi932Ke8pxo5oOrXioe0S2oOtV+2XfmjcbPGy
    MSddBfNmJt6vubpQUvmsm9S2vHRd1IEXm7zzZYKQ8Ok3zSvPBHWMqlk1b9cRWWzXDWvNrxkrRWm9ZoyyGTrqduDGUO/f7ZuOv/ey
    9Mjbewspywt2rMkOOzPaP+rZYcGMN+abDn06fTmmplyUfsJ31jjpTZsVR7pyjVL6vR409OG4HY9u2s87XuOZIR+09cDQ4MW9L/D3
    TNlg3+70IPHwZvd/nKiIlD+wauJz3nWY/6EdwQGb4mKyvVaab13rNunrfLfSkkmfvmdtqE59evfpSlnrvKYc16y9F2huSZVL93lJ
    b2SNvPEwf3Le8wKrgkxem25d33ILOmRci9xaUpWxesykh1uc8n6iXYZErm/JKO67uB/X7DfO0oOtUwJvUFbzbzU02pYWF9Fl/dXS
    /pn7er2dbT5ltFsUu6L7w6MV0lX7hTucHfZVZfWYWJoeUTYsac/TpP41z99WXPZMPj3g7f5hVx6nxB29tPXoqbwLhZKfrzJHDdtR
    0e13hWvs2wYUl1c75zyjjGvi3e55anZewZIh7skv0rsY3TzQ5tmBfeZdV7S2DUwc13l9+ZfLrg0Le/PndhnZs+G7c6vB30O5H+IP
    3ur+ILGr+zSfpucPi3Njb39oNLtT+XuvCdy8Cxvdm21eNmqgyL2J+4vF7Yc0YEzYnGphH13T0L/RBmt/x7t9Ahn+5lkshnzk8rCF
    N+UtEwvmNKtuE3cyStxo8vhKTwv2GetF1ikM/9gOtrduHuj2ellfkwUbu7b0bTDALWvAYuNCkw17J9zxjYnOLs6ZNOF5s6dvb2eN
    yDu53/7ZbwVW9GsNWdLGmxvMMm53gOsaapI57QK9X5fmI9paXe5T3DauScuO7aYOjTPdNGDTfEdfs2u5ztGv0tsVNc5K6cBYZenY
    eJNrhfX08MWi5b3cenh8fG6XnFy4YB1VMdv7XXKJZ1lBRd7hszZXbhSMfTDrwOzZSR2r5j09nPv2gyTveeDuc0z3msSILUNGJ2+6
    1yvgxWCGY0Zy5Y0enXv4TDd/Per1M06P/KclPe7aBP54lukpezLrfiE7ss2oRNtH4ZMXX3s4d+Z8x2VtEA51UuW+jVHHnu6If1g1
    gONYPmhk3MmA6rTOccFF24Z7UmjrBkhyfrYZI10wfFpcl71zTdA18wN/OzOwM7tjWw7dbWFS/pePa0tymooT0+bHLqNVzlpz7rq7
    f9DwD4dXvdq/6+iHN0EtFtDNLg8xnbmnoY0YmeFa1Ke3T1xwm7mnHx9OrWn49MKCZjkFTf4wMjLqrfAboeCyFagrR4aCLyMmnelI
    pTtTGc6hDEdXOsuVZW9Lp7vS6ebZS99rVRBJuAJeYu0VgkspQ0EFa3UFhUCEyhVskVS3jourvTNe586uPcmgjiWsI+BwXDkSaaJM
    EMNXGAVJEPUDBVHKUYQnQ1FhYuaqj9NAjW6qGlxUzpEJpAqBRGwkD+nfD4lWCoQKqkD8dOa5M6CgUYBfkO+ufmNhpT/zodnR7PoG
    sxMGoGwuKjP6l3zo+Ke2bzrd3kHzG6Yz6EwG0whJMPo3fJQAdTLQvdH/5ofpjIggAXswnJyd6I4uDJY9jcmkM5gMl+ZGf3/+6z/x
    aLRcoECjFKhIKgTSzC5GIokRogx7lgvbhY462XOYPK5TNI2vEAn/Ef53dMB4nOHEopO/8Y8jy4jBYjowGHSGvQPDiM6k27NYRgj9
    38n/MqVYXJcA/FX+f+gHRzcV0gA1DpUJeAIOGyoZV+RfQAh/f/7Pff7W/3/rf7X+d3BhAWTQnBhMlpOL/d/6/39R/wvEXDThnyvl
    69b/9kxHuiOm/8EPBgPKAgZIYfyt//8dH3dz3yE+oSOD/RCIcs/m7vALEbLFMR4WqNgCJgDN4NkcQdxFqIKNcPhsmRxVeFgMD/Wn
    OlsgdposMRuIEYs4ARovlcgUFsCvFStQMSgaL+Aq+B5cNE7AQanYAwURiAUKAVtIlXPYQiB8aHSsKdiWQqAQop4DgHcrESMDgF+M
    DOPwJRIhEoLCduXIRMSbw0HlcmSYQskFHYAMuVKokLvb4VV1BkTynEljIppgi7kIVxIvFkrYXETCA9YPGBRioHMZ3jkFkROdSlEZ
    TyITscUcFGRyJDIuyITNsTlAk4oEHJCKDQuRiIUCMUozAKtYNDEeViSNq/Z5U2ofFsJhY92rxkb0TNGMhTRaUAyvC2uyhQaGJZNE
    SxTkQWEygYLwJEKhJN5ABbZSwZfI6pyFqhYARSwYn9DDgsMWS8TA2gRZfBnK87DgKxRSuaudHR+rzAd1aWJUIRTwEmlsqdROTSBY
    x1KZBMxJkehhIYlxxfBeLyDWSjxas9Jq/P8k/WhPP1FKnj0h0msvr5QJScX/FNhxfCviBQoFKnOFZEdqSa4UidiyRAMEoqrwz0aU
    duv/VExhFVTIqhMncDR4xwjEBGAHqVRIuFF2Qq7tOLkEClIESQL/EMSiLza0BIWFK6IGPmBJVMSmSWQxFhSiFIZVUMSPq8TbYguH
    yGLYYsEE7ElVDoIAFjPAckQJMlgMFiTJDbU00cAKFxOA+WWIlC0DxeQ0VcuQkEiTqI2CiNJsLheATw5q4IAgTzJYAtSw0JsoQVHl
    E1UGS4CYECgSYclAZaxELNEr4iNRihUyrMTw/hZYZjLRrxxAyBt2GxEJEpIhuuxwiHj+i1Hn7z00mB2DqtJFbIHYT6zAJxJBTEAF
    C3K9oUpUTsKxNp4l8QA3SACgSYyuRYlAJwuEXCt57cTsRW4I1pMqUK63WB6PykjY0B4DkU0hZ6qmPgxlyzh8jCQSJUqZegBwjIA1
    EAVfICfohgIeUDECyAYlCqu0DxsMXMJFEYUEwQIPiXi2AFIiABENGQJFIx6TQLkUmAsoU4xAA0Obm9lCIcKOYwuE7GghWhdLwwEr
    +GyFiuBpFurZJRO/kil/ES/hsF0wbbbuBP/ZwB9pCIigYwBnRCkWjFeiiFgpAmDj4FlQ46JcCGYtXBEtgAFBzqEhAQokni2HiiIO
    oEBVHolONATQeIhTTXNYTVQsA/YByv1nALVeEA9M1CY8rgQFvyRAWkulgECBiYnBRI4RKw3BECTnS5RgvAGg8D8bLwE8MkQEGoAg
    oF8DMIxWKuD4BDKDowcWMkZOQkEsKkxEuEqMUQD5S5QKICrFXIE4BuGhoE40WwiNCBoSLETZchRTfGyOAp873hWm7rAGgOkONDCW
    R9TDWCkGSCaVQAGlYHbtbGQIwdi3vnjVtvgEHKiYcUErEAHBaCcVx6isv2g2NwalYQm6lqJckShE5XwUVagK4yk0jlyuX1oKrSng
    q3EUuoYlTwKVFx7XZEsFoL5E9BfqA/ArBBy8MkcmkcslMkEMgB+poV/3awfGzvTisUUCYaJHAJSOrvExfEVfBzrdjQX+OYJ/TnS6
    JVcgB355ooc8ni210AcH6NPdDvfQ3KMl3EQc5HwsmIdwhGy5HJSHsV08CdNloABXEKfKxTOoAujdEtnaBXAaogolMRJ1PighEMUg
    chlHC3FsYd023WDYhqpd2CAVq0zgoNauMYb3JPGcu1wKOEG/EBVqWAsDviOgSFCjXk3IldEWnviAER9oCAVjeky3BXc7ME41uDQP
    6p84WlAZYTxDBiN6g7+phHlqgWXD4QB6A3JOgxUZgBRbJmBTgVZDoRiJTsSTMVQC7tcgi88AWlM3sw40qGbEZ6ibkJI7xmFA2M31
    MZW1DQEaQtgHQHdAkFJU+h2Y6XyDNgBFW5/jUpFkc2usUoEYCj+FMJHmbidVQZyAnS4oIUyi2ZCuqUSSmvZ0kzWA9rAg5gZKAMXI
    QUUQSwhfwAUGgxpckPhh8xhr4Y1ZaLGbKq02jtBu3a4eMwFGLoAcB9Wbi36GHtmoi6ioQ3c+JJ7TLUuNlwHJQyJ8MuOoC8cKOLFQ
    fqgJA0+X67ENn6mqKxDzNOSqPUU1EQeIMCsSuCYq4gAt8pmk9qR6QxEAPwBIKl8SNSEwsilXQDoV4Y4dULRAIwLKlUlEBihbQ126
    nA4hpTVWoUCu0McFluqpJRdqwa2KajCWMUyRISQGwMupxIaOuFSgUiruk2PECR/V5Q3KV1gCF8Gew8BvhKE1WS10YWVx/93TB8AI
    WhowB/eTDTkEtWEKawk6pRae/gIZQAsHb07LPMXKEs4DtBjIphLeoxCAChpJGAoxY0cZE4Nbi3Iy/qBOhnCEo/OwwNql4rXVaAMO
    hpCLQxtMDqa522FPZLInOuTW0gaBQIFYCuw6rRJaakeC+eFIHFuoBFYQADveLAfvFM+to8IwVCwAQB4ihhhT/65/xdB4iboi+P0n
    KgKTRtMn9lT/yv4As+q68OFPVBXEabqFD/WvOkyQoK4JfutXhDwJoW+YOTBM4tIPAYY8qkKxNj5JpEU4lSqWUxOGjCoRCxMhh+FR
    FkjDevQFlQrskJSAEKYy0Z5WDs7fWv1p5RsgSu0CwKjkoHyJEFgoHhYoLQYobYGUL0F8hcAkFQu0C7OVCgmwWaVCVAHGAywCnWwo
    q7TLYPJPvxC0egAs4cg0zEqtpTCaAFQHcEUBf7KFclS7BLCLobfPJSXaaUGTLOYIKHDUghVXfx4WCplS27KEjB4Xg+ArFhYMZ6Ap
    Ubg1D/8NQw/9JAkeFnSEjjAdwF8LhCcQAhEtlgCORORgcrFg9hylDAbKfCRCGBXHU6lEm0x1AowgcthSGG9XirlayeMkArEq3dOd
    I5BxgNPGAR0zGICuEvFvgDhnCztP6HGgSAIDNA0SE4nvBCYo5EhzZIEk9U9Q2A5MTwtO2gIfJCiFOHXpIkgNTb0MnDvg72hJgo4x
    RaJ5soBWWyDudkphLVa1NhLFEioefKWK5IS9oJOEt0h0DyQFcOUkQhgI17b7AXr1msRJ4/8u1pkE1pkY1hl0LbTDVIh2ZwLrTALl
    TIOlGI6qYjQ6Q0UcBikDs9zc4fjEMRrhJYFmFBghULFEFvInIx8Usi4XAReAFLMB7geKRAMAwZ2wEh4MeKhGUFvgQz2OfygCogp+
    8FGyv0Grj9tXT0sMd4LqRaZ12mjM+tloYaSYqo67VQ/DzA+L15IAqI40QhuPQAk+N0yFolzceybCFdD0VVGHAu5wFvMEMpFOgFfb
    RoO2OdYgDicqfLYApAF0ugBus65N64KBUeHACCWnr3V1515/xYtF9vTUrk6H/4DaxYGsA+A/pXmxHkSgFpBFeMzXolaFGK1UQEYk
    bApltEiggVa0QkyVygTY6p2nD4EtS7ZI6oaEwXA7ETlwt8Mb0WpXTT6oTCaRaaQz6VGf5EEFVKbAbLsAMUcCBCfgVx1IqPlZIUtE
    2DBEqUMydpBGSM/EDElzgsE4fDTwKZrNibXwtOzpzHJhuqmCFFjDuhOrB2erlgJ12VymWqasD59rWFfVmsqow2JeKk8Ty8GjYHW4
    v1gpwvsN5aPEcj0Ul+RdAmBAKGm1BLKzyrd21QYvMATUjcP6hLbHx0RKIOt7b3XD5D4tdNX8rzAlRuPVnqsOvuCA2UBj8CHr4MCq
    C3u/8rt57PF1BE5grn64rY4QhlYFfxk6XonCaBXiLYdrMKrVDK0gBpm6YHU8dEDSBagCwFROLgKoSKRj1eBL7+RC4y08/8paodwL
    gAxvzTCbw7bZ0LGqYw1QKYdkp1l9QYBhiLCjJXEoDA5iHi/JBjDkwf/z1g3/8uogsAdhVdwwYcvVbIK1A+lVN8xgR6Dqn4O8f3R5
    q35o/CurVnWORLP5CFsFBn9FErni37SOBfpWSKB801rKIouifynKal8Drh8y/iVLu4YRWdeCLwnioC5GEDBOTSMMfYTPjgN9Q6zi
    3WIRfCmBNchueij710Ndb4cOJAfyHisoq8FPTPZjBEzapAD5Jr6eOAolVSNgV/9FEgGgd1BDyVbvuZHTED82xjU4jUKiF4hxlBNF
    IEaV0eOg0ASji4WchYIqcIFPIoKIITaOAekqw4SdgU1/hHCDvCBE42BAHchVEQKFN+DaX6DnT4SxMY1oMIg9ABVKsWGSEPFLneoZ
    hALK5IOqXrWpTKx8jEzAtagtjAfyMRuttvAQVuLPB4eYTE2YAP7+Xw0T6AWQ+PZ1CkKQrcVfnv/2DS0k8VTnRhZdrqgtNPXfTmJS
    toKPAOM2kMlAGKw4BzYTYSJwCAwq+DWARX6mMuOoDpBOgLOTiFGVFLSlADBwQhigBBO0gDDgb8Nkx9KmTvt6kxy0dQGKNZvCyGrf
    ANXpWIkcoYATC/euscWJWivOmO1FWDjQNKchGL1GyyTxwIXFbUSJFFARyZHCCnMlHCWUzogACGmugK1AhYl/k5QeSQGCcKS5MOPs
    yVREYziDB4YLzcmF+B/PcKY52lPtaXQnb5DKQrD/sBzEgeYIitqz9evA4rCik7e6A8Se5ojATvikXmGHNCcmh8ZgOtFcHGn2joDa
    aS7wPBQo68wgjY/mwILVGYNBQQbiAv5jMxzBPLBsOuKIOAoZNKYTFf5H7gBUgXU5oFUnmr09GCfNmUVjOdljHdDII1TBZUK9GcBH
    z142QPVa5n69PAsKgnt40JGEZWLFgNGwNjS7N3RtPwNWEG4HakcYtGIMhLlPBaI+GpXJdUlYKPB0ZxP7qgAnudoyWY5OzkxnB0em
    o6OFJ3xEwDMCEhCQ4m7HBmADlerRCovJojMc6A6qVlhMBCQgIOXPtOLEdKC72DPUrTgxQQsuCEj5c604spztHexJrYAEBKQYaqWO
    xZNabTZ3O7j/CP/Jk0gUOvvE8CQD+8TwDN19YlKdbHiXhYWnJfxyI7aKYGERPBfamlhoGMbUDTol3oCCsLsw5NCQR2Vx2mpYrz9A
    kai2TS7ALW38SIpADhQ/sWsd3+ioBDJVwBZjO/gTaQh5awWXrWDD6sBiUGAhbWgrGAo9kPceqXd74QPSOg2AbY/Df9PGYUEw1b5I
    dzt8sx7gUOzU1d/nf/8+//t/5Pwvy4Fmz7K3d2Y4/H3+93/x/K9aXv1z+b/2878OTMDu6vO/TIYj4H+WE8Pp7/O//46PNU8pxiM4
    1jbYoQcruPIO7HYBR2HlBrWZnR3yZOkU8BcZpgAEQjz8F/zF56aKw0XguxltMZM4EvFAkhBiMTsI28OMZRM/gQ0gp6gDe/DIQxxb
    vRolh4YBqJ7spkonVsgHoYkgXawUCnVzsC2HIM9KF+K+QwJBNzw58l8Aa3zC0KUIwPZI4mBSecm0GFThJ8T2Q/dLDOBaW2nvbrOy
    UcEMwwOxhvSLJsi7MDUNkBaZBgN/vs4x6Gx1go2okY1KiSWwumeh3oNLGgFIJLZn/LIuvsSlXZc4CfvLukQonDxqguAgHfvi1F1n
    K4b2dmjGgo/NH27cUH/qaI20sUPTBrCpfQA7aUii7ja0N19omsE2GwTKY8geVh3NqPcmkGCDQwtCps6q5OV4Tf+4LMDoqe7K6nVz
    TV38rMAwIpBeR23tUwskmgAKvB+WVycyNecSyOPGN6zXo3fdcwb6jfyKnbR2x5MgH60Q92NzYn+JNdUeDhLoFOIgNJ7gw1/U1Owp
    ILqutTTJW7WyocFdQD74aR0ovtF4xBeoQWsbWM0fCPORoJi1jUpwD5HiR7URuMNFge+24KJCQTQqwyKAUKlw4EoNvsIKrApofdMQ
    uGFDCI85wyAM3hR+cyUWB4fXQgqEKO614kFrGGdhc+FeXaCs2QqJhrlhloiCyOEyK94SO1oORw+bQIQoOw4L9yDA4lMkknck4Ud3
    UUSqjBYKOBhZ0UADagsBhjZV8/PGpkfYDGQ69oFrulxs4B5IhBUOBjs8D54Ss6IgOonjpAYT0RirSDescYUsUdWttW4nFM3wrJUy
    oY364KaGKWjADwdjAbluRJ4Wx9GIjToeCLbNGC+TjCMU0RzTRmDI0lU9TLjsohSywZS4PDB6bJEL5BLRLlUuyCGdvgf5Iaj2iRS5
    dnmaFemsrE6PMIgM6AOTnzq9BuJZ2PkWvS61DsKISCWxOAmxuQzCwUp1jJMGMuCiJMk2VLGuBr48VAFKqNJpcKwUMGYRquBLuKDb
    AX7evmAsHLhSCR7FMJwNfoJObNQRKxrc46Hdi43WyVsBD7E2B6k0SawNGL9MEo8xoB8U39ZW6r2ouEBQ75njckOIYWnGrSlA7p/D
    VmhNE/SebKOhAHiclcwAZELkkEhQIvaHw4C32nDRBNUUsAdAVvj3xIkIHW8ZTgpP8/RANM3QhKg4RsG3AeJUoZSJ8bI4lDWFIrCK
    kX8S0r+A81+AMjFj/ZGpy2g6NwDkekISWOMMGzeDqDCIZJIwYpPVAS5JCRlvbcVWzYRNU1v10Pom6SdxrJWqCIzQglwtSlflAdcR
    aAFYNypayCZVkqFCmCqWwKUhIFnFEtAKsDqA7lUVUS8W4XY/noiFVweEBg5GPAjwWRk+cFfryo8VYqvGkpXWyg+DtPLD+P+/8sMA
    XQ5wVC+P0OFaYhyDlAC+mXwGk5xAZYY5TzC4ughbQ8B/zgiTjmDHIjQdgQxGnCM5TeQCy9rDP1TNEqMGdFZEjFp7DzzIR4BoZUtR
    iCGN7MPEsA3ItFLvenaXi9hCIalFQ/VIclpVG6tF7lwP8WzAlfEGMf9k1hKiJkFOZLsM3sSCirk+cN3Hmm2jXaB2TZis7Y8OhvQq
    YosFPKC8aBx5HGINd/oBswJvCG7ZB4OBUuxf50hq2SPeQqHaCMGEZSgw19SGNrbly448YKvaZCI088hCUe3Ne8AFBDkaSDSCF6yX
    lDPUFh4WqE3BaGZAsmVwfUCoApAOhb8hYV+nlP/TMp7oFVaBM7bWV4sasgiGANKmi/+I2I8/thHJFTOMsZPqBKKi8AfMNosCLmcU
    tsyJH0nFKSsKO2WBt4IgVHULrurrqqBSiSLfI4AnAF0nEQVha5zQiNO0oNObK1CFKL7fZUgIIgIupwDYaYhUIIUOGKBINrT0sUgU
    Ys2AcVTGROyLaUMmKAOkS1KSuC+LkMgSpkK5Ksc8MVieBpx+gcLabrTMa7TYzkZjwphj5XSsFqJJwniOw070wIUlSPpYQz7DwgaD
    etZY5Qh6pA1NxJaSiJavIVmiRT5NIROIgL+lkAyWxEPbX45qUyPekcAfujdw2ESXNMyGGMIDDh2BHStycdzTN1ScTAVaVWCkopYq
    2tjTroVF9wzVwihKq2woJCuDZcl0Z6XyTiANA40BnWJAMEKNbMbvWsC8Rc3sAbkIlSKxBtFoDJuTGAgn5QFMUtwM8vBAqAxoreKz
    xZ9V/UGPwRobKKjBcANf7giZDECKra0Gg1i8TiKUE469AQIQRBLYVcsdjLRgJaJJOBTsGRILtj9SIFaibs1JxTXzIAtd6IwDiOCb
    DZEnU5Yi45QAMGxsFw9UVCq4yN2gN0wEdTUb0lXZ6gbhbFSJEGDYoHCqA4Y8HfEiSNAVoUfCUVvB4AE2OTdtU1vVCHk65E4wpuWq
    FI8/UdpaXU2nOby4ocaw1mKxsLOV5rT4RGjHiCFYACRQa7w6jSD6IAMd4DwdARqK1HbPEFKObsTcFdFvmBRFdyWPyIqitTVC54Nh
    Bt7epg65wwck2U3vGiKtAdFU1xVJlXK+dRIiJg8Kz8PHpJHdauwmaxQ8ggJTiDRrMhUwEW06qBXtmLBlxxMiR1MpCFtrqLsSIXQ0
    lWBCrZXUtXChA91P/CdBoqo2YJpOI4g2TvTbJIQTbBP/qd0mlmagTStdcsIBYZBg1SP30PyEDWqNSwtCuAr0UEGKREheqjRChVlN
    tNLVNhzoixJ6hqOCJTSlAIYVqMy6HxDqKFtsQ2oV0h5pCCoWY5JHDGx5HSZTzblWzmLWzlpMA7xFtKfFUKru62QmFTdpr2HpMZQu
    1WPiNBAF3q7K5mCrj8nAI7yqUzLYfVgac0WtjuBuSTIUsUYMxJk4OlDAwEQGBQ2vqVKLoDyhoRADpTDO52gFh5LJkaDmWthQc7X2
    GDQSmVmrSGZq9aFaESDiCmrOAYRsrWrJS9UmSRQhrnViTvNRdwsrw/gtMM+gNdnLTkCBzKc1Gi24GJSJxFANSUKmDsDI38nNtf0E
    zO7T+AYqIBEa+C9YxtpWrL4eJFmyoATmXqnyVCwfZUWyWLFCKrPCHbFXG62qNVmtdVwCe1glYHtoSVrNko9qvQhvWy4UAGQwKJAg
    aTDugY1AqyrG9+QuoH2jacvAmIiEJPLAXBEttaqp70oeV7IBXw3YX5jJhltjcPufNR+AHi4ujFdKoFOB3YYjt0H+pd6XGrd6RqE6
    fKryMdQgiYh003VdPDCzgKOUYVE0GDYcCich147mq81WuA5AdwNfuNmqtlo5elYrXE6CRSI4kWTLFKYDeWNlYaUbQFX3bGlJVITR
    y0hScWyYttgT1qMbyW7BxS15+Oaq35pihFQmj4NiBftTlyUPCocQIQOVMqDXVGCq1bYhBsjhuxlkdJ326mB/NXY1GlCukOkENECK
    tk9HkmZyW7sYIMoQLebBmiat/9PYXK5fHGACGOCCYVbgUQFajgH2JKIfiNHdbEFuCLs4yI0opt6iQFNdMgPRoVVdvyjWAgm6fAEX
    HabZPqDifrEEX7wPlMdoom0wfKcRDlod2Wh1wpNwlERTyTY6W0W8STcFIP+p+0PYHHgnQACxcEJlqJeISXCra51X60YYwksmA1Cf
    YPB9BIboBXY7XonKoPekh2ktka7afFAbRs2xZqAE0KeKOihCtRCkslOwAQXhuy08SJyFt06wLCwkgqFITIgMwY4N0oDul1urgpBq
    41YzZZBPkhwkHQinrtofBX0qtSmgZmH8HmXSMoqHhw6nWVrqWDVkmYBXJ5m2NmrTDpuqDWLugROC1qosBlVimuoIlGoCtbBeLcxH
    UhOI1tJb8q8YFstXDUIOdC4JomwKEq0DUvYwLayp4cqOrHP+EKDQyaIDz4PhRm4wupYGo/9CgxCgYIAQ2NHD1AoXpFDBs05IuJaR
    07DtC6gPkEFsWe3D0cGkDAWjk5HRRUCUgtM6WdipwatfCVI4BVHxGS6dtfd5kZfWNAtuhuQNAv05lS1HpyAMuo0BT0WLZ+rBMnhk
    tY5VSaFAE3wXCrRXJjVzwU5FW5HKyVGFtwJIo2ilArW2gldxwc0U+FV3Wg1qF4SHHahgfKAwnInao4XD5Kp3hhlgUC0SFELlrSOO
    iLI25IkLuAlEDBCroSFJsujCzShQ0pzw6DSGCRi+/sqozrIa0S9NroyG21fFMRB3oDkbG9KKKFzag4eoddfy9CuDmlh1UA4bpUrM
    ECt1qkbqOxZDLdVqh4Hpau86ItrTmGUatOqrNJFECRxLSbxYS61pxUZRmlSGwlq+KI8NpBs5goRLb8JZs8aJQ9d31uUt8uKiUKDL
    47qlDclesobVplStuwAhbUPpa2Vg8UxP7tciCAxo6frIij8xQmxe6k0GejKGGHjdhgkAfO04xHgKCAJsWUVn7BiN4XatRAaXR61o
    uuKD7IpizRjcewKzURoWPoauhjdceIabiUiuT+10pD3rQLaCTxMBd5icDNwjCkLuHigbhroBpRRuCfHGy4NC1lhJFWlpPCHdEQ6X
    /uXxsRO0xkeF46P/IwPCLvjCPDRyuzBqWp8hajMi1lsEqZ1IaP0akug2dY8Jk1JWvzSXkg3o3VpAoHLYMVTqa0qYjktizWo0SMI1
    HEayCklMjBAyEtayFS534WhJ09VZ8cMMFj2gGgCSnANP8weIFRJ4k5l1EhIN7JRYuHQO2A/M2Yq8pK3WzQY8S3hA3DA34iEdEjer
    bvEAcMd3CNlgXroup+oX+zNYIZ3LwHfDqwyQ/yy3TxMI0tM7Oo47fn4CZOiF6mqxuQz557VaNAYA3hwhd0/aM6+jnQ20CY0ExBoa
    GfouEsiyIXQKefe7XgihNrdStfI7QMDFb0BSbYrHDgfAK7MEMrjTGJAl6RYZhM+Wq9c5EeyaJK4m5qAaJBY2J6Qx5iiQNtTyJfHE
    yQNVeRvDfhOONfUBCdX4tTWu+gwEKZtsD2jOORhS2FqAI4dFtDkjjPRG4//MoIj6iIUBkYRf4mhYJtWmVuqKVdTCUST+IwVgsZvZ
    iFVqXSJWR0dUpOqNvV4HkpqqGkavqkAFvHNDwtO98wgjUZxe1b1it4GSWA4nWLkELgsYXtXDyJjo1U07goA1Vh8C1zHQ9WGoFUSo
    k6lrJ13DKtfQoNRhVNWBlV/LI1WwVnVKhWThIpqDZ9qrE6rFIrhCoZNUi4jA6AI7VVA/Nxcrq+3pYhE8rA8shKddUnuaVkGSum62
    TEQVxK1XArl6/FYaiaUGBdl5wfoxjHWMkOExDfWiJmg3XgbXUMRwLx1IofKw9xkQVwypj3PQkEEoEMzwbkFyWxIZPHMBb1iLg1FG
    FecJwaMQkbGxGy5BHUADYKTwmjy2UMpnR6MKuINfmEjT3V6oRo6h7fswx0ZnG4NQgNQTUao91Uh9dlVrJAQH31VN2oqJud1iyLXD
    QwJgqEgihpXxMjQDO1yI/dewY1CPnF7L3mvN/mucsmvfgl3bjmtDu66xKdRrz7Xevuv/I9eJ/9v2XWvudHIk7nSyx+90ciauEbM3
    XMpJu5ST4e7oiAv2xxlxMbBvm7x3WzfVAELxK3X1tnVjtIivMNvWqznVVv46NmfXTSKOJBJx/N++56sWZLqRg2OG9rLXIdTVESmN
    +0gyUvWsUG0j9Rc2rJ6NSjIFSAq6dhMVO/loh912o/iPNFGJ05v1cpnVLp0e3OsGe62hQv0VUfKR0D83JoPH9A15kXX4aEg9KOQX
    9FXvmWpoCN6tCfdY/xfcEqDGC0kiG9w6oDp3od4zYAl3DCBW2K31mlMe6mx3IluoMJTrSeTGGMy1IHLhBhk3K4NcLJEo/qtuxjB0
    1teNSMfO3AASTLaBX3/fofOf/Pn7/q+/7//Svv/LiWbPZNg7M+z/5uz/yfu/1K9W/jfd/2XPZDIdNfd/MRmA/x0ZDIe/7//6d3z6
    UJA+rq7RKE8iQ7GfbB7cEwytrWhJAlUumCAQx7iC3zBYRQVJ0AgQsWUxArErfqGBFBjZWBnwBKwiVxk0hWB9KjVGhqKgWE8G2zHa
    nu2mSaNy2bJYmGHvwGNyyBlC6AWDHNSZx0JRco5IwAXpTK4zm8Uk0iVCmMRxYTOZTpokdRs8Ls8RtcczYOgQGwmDzUQ1SVSREpj8
    IIPFZvGciCHikwWJXAcuB3UmEmNgi048Fx4x3ng+4BuYhn3wNMwdgEOi27swo0lpRHUuj44SReV8NlcST5WLXKFDLU1AHMA/WUw0
    25pOwf7Q6E42WkVFXFgUFmPSdcsy6NplhRAfiDNsVr+sPVFWxuYKlHJX4KxLE8gp2KCcVWnw7eWuiBX2jnLgOMkTgdcioioFFIQK
    nGwhSsVTQA5bDG/bAl4dRgjwZlOMDvAVWGo0ymfHCSB45CJAInysELwGFSsEe6Hir0N3hdFDa7xjbKTweGMMFqRQZUXHYBkcGPhQ
    panPZcNgAxUPp4C50RwxihWINUl0ehy2w5mIj7siPCGKTRZ+U7kCGX4c3pU4L4kNlEa8UAwbrFQCZCZWgh0tB2UUGEVhgRdXiEv4
    pO4NfyRziYaBqEQuvBCIJ5TEuxJvjcImJxRI4W5+jkKFOWx2GNlR5VLgmLgiYgl8PykGI4JmcSa066NxDbFXwRNPfezARDSvg8fZ
    XA+6GL/hkIciAKMoSE6ATDBa0SEnZ6yoBiZyBXCxsUVahURKTHgCFdvyhQEfhyf5vfPYOETsBCoBQhc6XaotaLAXoWnBkeEIGcEB
    lsPQo3lLPdaaHm7ZQCrg21/g6SdsQQgmx7DBEBnqZjQvpcdaIcbDYmrjVPVMSEUVHzkSyMQ23VJ5AoWr6n0eauKS82UCcSwBFH3I
    Y8i10ZonU2eG2JGapD9NvaQ33GPvolfzHJDwgJAYBLtjSfHENJ3odD0u04xQh82Y+h3JldF6/RCQI9rECInJYlFU/+g0J/wsldZA
    HPCBCFF4XgWjfZyVaHR7VKRD8IFwdiRyh7PF3jsL11aSCDARO1//JMnRMXpDHImSKn6n03vrMZ1MQh4DHz7DvjGVg1EimQbVHbAg
    TdMxka1iElCRz9CFIkfIFkmtmaA0BWHFxVOwGjb1wp9G/daCRBUEgBZUKCQiMnPA0RhEqpMWUjXiGNeuNjqgZjnWAmotGKovUsMv
    68Ju71a/BJ4MXO1b8QhRQohXZwye9hopobkDT5uHsE1BekjFh60CD5NFN8TgcTgSdWQBDgT8SU+U4plqtW5DFuDQFAA6RcBVqTos
    3QYfv+4dfFqzhZF5B5Vk0hfqasuolu6AycV1cWHXZy4a4aSGrHpkxKt7sBdnk4anISdHvSqxQF8QCCFIqKdztCOPwXLToTNG7UJK
    VzbQGEwoGgiWU8iAbQLPIboiSqkUlXHYMMRqEKI08guItOBrjxOT3vCxVy6SR2+QAcjz0IcadieBFkXCVxphCgp8U1XeGRWX6XLs
    lB/KVlhDxoGESIEWDtzAyIS2IQVh8GT4NkBcvzENdCiO/QuKklkLp9v8Sp/9VXoDtqg2yeHAI/DKRTkSGRtXd3BpDMuAuCZMETXe
    AT2w5Ah8SwWFxIkIjalJxIaATUydbABmrnzI8Lgw1xAV9hMiaKQ1lUmIYnKLZCasWxpg09UmLw7B5/py3MaQXaGPZvxOKopeMrxs
    yqAYrK0JPcHvYJiuNC3XlyOYZI1AxexGfZrFVlfrBIWmRaaWjqEKUR6QFQbUDPZKZR9s+wjZRsZf4EuWsvU2twhWoRtQpSy17NO8
    RjbpTzOOjl74E+La3pFwCOujkMh8JIFiVZEILC4mi8wampdTa8NIIMYMC7VWrcXHIOmjuoyUv6gCgHno/EsdQIIOtO0YdIOWvQ41
    qXUZUwuf2N08ujyiqvpLy0wtRnU6cdbqA94Zp8eGrPrYX1qNqhU3iRPwW3Txm5hJvICdkNSoc42fJ0OBxBPEqUiB9M5rsu+ksqS0
    BDjigAEb+8+RBB9iPuSk2mICBFsw/xJbqLSKRKmAhKpRHnVxomFskdlES4kAXpFrKRuYgIVU4tHoWIGCit+ZAd/OqOlfPw3CFjuR
    j79IsRazVWWjYCClqzxRwyMmQZpOc2bKcB7Rp0/YM/kKdaxr8gBVDhJHKZPDfrAdINBa0KUHV/VBal1p1xNFeUwei9yKWAJdJHhO
    jGugJWwZnAhPGtCvtUQu6LjhRvibjhQG3YnCcobBC6aNgU5cyZcik81SNjuaznHQrqFSz4aDQjLCtVJZLVCzsXCOqMV+ALlkclP3
    iajgS8X29co1RFM/E06b2TU73bX0ns6F93VMDJsJhy3kwNvneiO2UHziLiWmaumkydPrbRzSWDrsTEboL7m5Vo0mxMOF6iAUk1AY
    YIZUbL1DA0sth0/lqKrCc9REV/2oAOazEyShA74IPJoXqc20Gr7WOZeFQ1stJ5m1SEeVtNflunpJKDUWAO8Dk9jtT3sA6tiEzuBx
    w5iil07DT+PUFWqslxlgqEti3YJgUSzEo7pvitA+TtqBOyeD2p3gx7oCofrBO8IkgjKUZRib8Iio3rRVdPZLe0dXFGt4l3h1BPbe
    BQ3v9iS9XuEvGawaYx7eli+VCbBXLSfpqW6mc60hBjW4agkZGqBirbk64nxZH72vVti/FgsGGKVWloAqW+M0AhYxrML14+LOhrQL
    k2VjKFBIBjDJn6yNO9R0od2rynjS79aebqPfD4kNSYpHDo/sW9NpLs6kKurYgCHC1UXAr10xle6rC60GcKSmOjyIRzfo9CvhRQDQ
    hFPnqlOoEh5PjioIR7JW54Q8aRIy6hZE6veHGPaIdapjpWuJkKhWB3WCHHTCPK635sOK4KrXXkeHqrrXGvivVJPWTSr1CBThAlLB
    lil0VIW2U69FCwZB1tOF6UCn6xnjcJk1Go3W9yRUGvLPQslBE4jisVxQerShpUOWAWDUH3Jq01Bfe+gFOjTz57o4OdEdDWEB23Ku
    tYRBlMGO/qrw9OfAbdBZ1Aa8ehX9nxfD05g2OjOoB3DVk8WOkOmtHYskYgm2Oqo7b3t9tWVouV0zSDBGlkHa0vedVW9kIpnSqvcE
    GQzyQDr7f+1dTW8TMRD9K7mgtKJE22Q36scFLkgckbghDkuzIZWqNCVFiAN3znCH/8Yv4Cew9q7tmfGM1xsFTvMOlapuvI7XHr95
    89ytP7ZxpW3RnH85X1Sr5sNZO9jF+nJdToonZ85YYLcNWVCk/Ljfa6B4ZrvhXqfly4dQolhmShRkPyKaQt8KXxd0/Qhi+ZEEjBl4
    p5Jtl2P1o7kQr/fML7ib3t1O6kNrzlFBu5DIVU62NE78EBX0QdmAJW9VN2HG6SGEbVX7XHEcDz9kUSNkAZlC9R6ZpMTPzoMZOCsG
    ZbDFBc5E3O/kOaW5NUiUcuea+Z/XJi3w6VH4U7wrpUmP+F1DTIH7Unw5PiaVEvJBcrWUahz4pGgWESMhBS29pNEom1x028aq3m8a
    ZlFy5X+0j7x88XriXsUG9pJ1/cBUfH1YCn4Bc2FcShwV9FwzXpD4d8WJQd3GqbSROwp28q053/lucO27jzyQlBZG3cz8lI+PTAqT
    tRXlLlxLap69bx4/N802qsh+2lsPhNFpkXJsvvHVlZOcV81jfXvXLph2VVMLRPSh4AQN2srT6XVGeaOUuIOgpWDKwD7eVKd+f/s+
    hT3PyeKGY6xrr8YTpujcl5wwlxFymATDT9uOscsrqe8CZPYoeLzaru+56GENDUz46NrxLDzyPRyziDUv0G2MqeE4doeSszvg7/Tf
    y6zzcqDC6nmM7R+lCq5Au/GrKZcUkB46FnsErjAiDQg6Yhj9zUKi/PnxVSqLhrvsDiEBsVvY1r2Mb/OmpTb2/5HvxcQCLSVpuAf2
    XT9b6U19YiGJaalxIw9JJPujkpZkVwU5/s+vHz+nrN9DbAh4bJKCn7TRA0JlX+Ua2Z+7N7zmqa4Jq2p1SueAd5jM+pfIHuxvNrmm
    jyQHTSnxUQo01HX55n73JWvaVWjaMQN0eYoa3rY/xxuCK/pM3/gME9cunaVmQENCqs2ubcTI4qGd5+YVqfXkBDyupXlc3QlkYtwh
    Qd/vk/Z0cLD9UgtO4Yso4MLe5ov33ZJe5v23bDR11xFrvVRV7C5nveGJ5kVJCXf9HH4mcmEjRmVJeOmH5KuePFQoFAqFQqFQKBQK
    hUKhUCgUCoVCoVAoFAqFQqFQKBSKQfwF1nK/pgAgAwA=
"""
)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def random_number(length=6):
    return "".join(random.choices(string.digits, k=length))


def safe_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]', "", name)
    return re.sub(r"\s+", " ", name).strip() or "Unknown"


def normalize_name(name: str) -> str:
    return re.sub(r"\s+", " ", name.lower().strip())


def is_numeric_code(val: str) -> bool:
    return bool(re.match(r"^\d{4,}$", val.strip()))


def fuzzy_match(target: str, candidates, threshold=0.78):
    norm  = normalize_name(target)
    cands = list(candidates)
    if norm in cands:
        return norm
    norm_sorted = " ".join(sorted(norm.split()))
    for key in cands:
        if " ".join(sorted(key.split())) == norm_sorted:
            return key
    matches = difflib.get_close_matches(norm, cands, n=1, cutoff=threshold)
    return matches[0] if matches else None


# ═══════════════════════════════════════════════════════════════════════════════
# FILE PARSERS
# ═══════════════════════════════════════════════════════════════════════════════

def _rows_to_pairs(rows):
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    if n_cols < 2:
        return []

    def num_score(col):
        vals = [r[col].strip() for r in rows if col < len(r) and r[col].strip()]
        return (sum(1 for v in vals if is_numeric_code(v)) / len(vals)) if vals else 0.0

    scores   = [num_score(i) for i in range(n_cols)]
    code_col = max(range(n_cols), key=lambda i: scores[i])
    if scores[code_col] < 0.4:
        return []

    name_col, best = None, -1
    for i in range(n_cols):
        if i == code_col:
            continue
        vals = [r[i].strip() for r in rows if i < len(r) and r[i].strip()]
        if not vals:
            continue
        ts = sum(1 for v in vals if not is_numeric_code(v) and len(v) > 2) / len(vals)
        if ts > best:
            best, name_col = ts, i

    if name_col is None:
        return []

    start = 0
    if code_col < len(rows[0]) and not is_numeric_code(rows[0][code_col].strip()):
        start = 1

    pairs = []
    for row in rows[start:]:
        name = row[name_col].strip() if name_col < len(row) else ""
        code = row[code_col].strip() if code_col < len(row) else ""
        if name and code and is_numeric_code(code):
            pairs.append((name, code))
    return pairs


def parse_paycode_file(path: str):
    ext = Path(path).suffix.lower()
    try:
        if ext == ".csv":
            return _parse_csv_codes(path)
        elif ext in (".xlsx", ".xls"):
            return _parse_excel_codes(path)
        elif ext in (".docx", ".doc"):
            return _parse_docx_codes(path)
        elif ext == ".pdf":
            return _parse_pdf_codes(path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    except (RuntimeError, ValueError):
        raise
    except Exception as exc:
        raise RuntimeError(f"Error reading {Path(path).name}: {exc}")


def _parse_csv_codes(path):
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(path, newline="", encoding=enc) as f:
                return _rows_to_pairs([r for r in csv.reader(f)])
        except UnicodeDecodeError:
            continue
    return []


def _parse_excel_codes(path):
    import openpyxl
    wb   = openpyxl.load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
        break
    wb.close()
    return _rows_to_pairs(rows)


def _parse_docx_codes(path):
    import docx as _docx
    doc  = _docx.Document(path)
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
    if not rows:
        for para in doc.paragraphs:
            text = para.text.strip()
            if "\t" in text:
                rows.append(text.split("\t"))
            elif "," in text:
                rows.append(text.split(","))
    return _rows_to_pairs(rows)


def _parse_pdf_codes(path):
    doc  = fitz.open(path)
    rows = []
    for page in doc:
        words = page.get_text("words")
        if not words:
            continue
        lines: dict = {}
        for w in words:
            y_key = round(w[1] / 5) * 5
            lines.setdefault(y_key, []).append(w)
        for y in sorted(lines):
            lw  = sorted(lines[y], key=lambda w: w[0])
            row, cur, prev_x1 = [], [], None
            for w in lw:
                if prev_x1 is not None and w[0] - prev_x1 > 20:
                    row.append(" ".join(cur)); cur = []
                cur.append(w[4]); prev_x1 = w[2]
            if cur:
                row.append(" ".join(cur))
            if len(row) >= 2:
                rows.append(row)
    doc.close()
    return _rows_to_pairs(rows)


def parse_names_file(path: str):
    ext   = Path(path).suffix.lower()
    names = []
    skip  = {"name", "student", "student name", "no", "no.", "sn", "#", "sr"}
    try:
        if ext == ".csv":
            for enc in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    with open(path, newline="", encoding=enc) as f:
                        for row in csv.reader(f):
                            for cell in row:
                                c = cell.strip()
                                if c and not is_numeric_code(c):
                                    names.append(c)
                    break
                except UnicodeDecodeError:
                    continue
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        val = str(cell).strip() if cell is not None else ""
                        if val and val.lower() != "none" and not is_numeric_code(val):
                            names.append(val)
                break
            wb.close()
        elif ext in (".docx", ".doc"):
            import docx as _docx
            doc = _docx.Document(path)
            for para in doc.paragraphs:
                t = para.text.strip()
                if t and not is_numeric_code(t):
                    names.append(t)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and not is_numeric_code(t):
                            names.append(t)
        elif ext == ".pdf":
            doc = fitz.open(path)
            for page in doc:
                for line in page.get_text("text").split("\n"):
                    line = line.strip()
                    if line and not is_numeric_code(line) and len(line) > 2:
                        names.append(line)
            doc.close()
    except (RuntimeError, ValueError):
        raise
    except Exception as exc:
        raise RuntimeError(f"Error reading {Path(path).name}: {exc}")

    return [n for n in names if normalize_name(n) not in skip and len(n) >= 3]


# ═══════════════════════════════════════════════════════════════════════════════
# DATA STORE
# ═══════════════════════════════════════════════════════════════════════════════

class DataStore:
    APP_DIR = Path(os.environ.get("APPDATA", Path.home())) / "HiltonHighCardSplitter"

    def __init__(self):
        self.APP_DIR.mkdir(parents=True, exist_ok=True)
        self._codes_path    = self.APP_DIR / "paycodes.json"
        self._default_path  = self.APP_DIR / "defaulters.json"
        self._settings_path = self.APP_DIR / "settings.json"
        self._history_path  = self.APP_DIR / "report_history.json"
        self.report_archive = self.APP_DIR / "report_archive"
        self.report_archive.mkdir(parents=True, exist_ok=True)
        self._load()

    def _load(self):
        # Pay codes
        if self._codes_path.exists():
            try:
                d = json.loads(self._codes_path.read_text("utf-8"))
                self.students = d.get("students", {})
            except Exception:
                self.students = {}
        else:
            self.students = {}

        # Defaulters
        if self._default_path.exists():
            try:
                d = json.loads(self._default_path.read_text("utf-8"))
                self._defaulters        = set(d.get("defaulters", []))
                self._defaulter_display = d.get("display", {})
            except Exception:
                self._defaulters, self._defaulter_display = set(), {}
        else:
            self._defaulters, self._defaulter_display = set(), {}

        try:
            self.settings = json.loads(self._settings_path.read_text("utf-8"))
        except Exception:
            self.settings = {}
        self.settings.setdefault("upload_password", DEFAULT_UPLOAD_PASSWORD)
        self.settings.setdefault("netlify_token", DEFAULT_NETLIFY_TOKEN)
        self.settings.setdefault("netlify_site_id", DEFAULT_NETLIFY_SITE_ID)
        self.settings.setdefault("copyright_year", str(datetime.now().year))
        self.settings.setdefault("theme", DEFAULT_THEME)

        try:
            self.report_history = json.loads(self._history_path.read_text("utf-8"))
            if not isinstance(self.report_history, list):
                self.report_history = []
        except Exception:
            self.report_history = []

    def _save_settings(self):
        self._settings_path.write_text(
            json.dumps(self.settings, indent=2, ensure_ascii=False), "utf-8")

    def _save_history(self):
        self._history_path.write_text(
            json.dumps(self.report_history, indent=2, ensure_ascii=False), "utf-8")

    def storage_bytes(self) -> int:
        total = 0
        for root, _, files in os.walk(self.report_archive):
            for name in files:
                try:
                    total += (Path(root) / name).stat().st_size
                except OSError:
                    pass
        return total

    @staticmethod
    def human_size(size: int) -> str:
        value = float(size)
        for unit in ("B", "KB", "MB", "GB"):
            if value < 1024 or unit == "GB":
                return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
            value /= 1024
        return f"{value:.1f} GB"

    def add_report_batch(self, title: str, source_dir: str, filenames, manifest_rows=None) -> dict:
        files = [Path(source_dir) / name for name in filenames]
        batch_bytes = sum(p.stat().st_size for p in files if p.exists())
        current = self.storage_bytes()
        if current + batch_bytes > MAX_REPORT_STORAGE:
            raise RuntimeError(
                "The saved report limit of 5 GB would be exceeded. "
                "Delete an older report batch before saving this one.")
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        batch_dir = self.report_archive / f"{stamp}_{safe_filename(title)}"
        batch_dir.mkdir(parents=True, exist_ok=True)
        copied = []
        for source in files:
            if source.exists():
                target = batch_dir / source.name
                shutil.copy2(source, target)
                copied.append(target.name)
        record = {
            "title": title,
            "date": datetime.now().isoformat(timespec="seconds"),
            "size": batch_bytes,
            "folder": str(batch_dir),
            "files": copied,
            "manifest_rows": manifest_rows or [],
        }
        self.report_history.insert(0, record)
        self._save_history()
        return record

    def delete_report_batch(self, record: dict):
        folder = Path(record.get("folder", ""))
        if folder.exists() and folder.is_dir():
            shutil.rmtree(folder, ignore_errors=True)
        self.report_history = [
            item for item in self.report_history
            if item.get("folder") != record.get("folder")]
        self._save_history()

    def save_settings(self, values: dict):
        for key in ("upload_password", "netlify_token", "netlify_site_id",
                    "copyright_year", "theme"):
            if values.get(key) is not None:
                self.settings[key] = str(values[key]).strip()
        self._save_settings()

    def export_archive(self, destination: str):
        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in self.APP_DIR.rglob("*"):
                if path.is_file() and path.resolve() != Path(destination).resolve():
                    archive.write(path, Path("app_data") / path.relative_to(self.APP_DIR))
            for index, record in enumerate(self.report_history):
                folder = Path(record.get("folder", ""))
                if folder.exists():
                    for path in folder.rglob("*"):
                        if path.is_file():
                            archive.write(path, Path("reports") / str(index) / path.name)

    def import_archive(self, source: str):
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            with zipfile.ZipFile(source) as archive:
                for member in archive.infolist():
                    target = (root / member.filename).resolve()
                    if not str(target).startswith(str(root.resolve())):
                        raise RuntimeError("The import archive contains an unsafe path.")
                    if not member.is_dir():
                        target.parent.mkdir(parents=True, exist_ok=True)
                        with archive.open(member) as src, open(target, "wb") as dst:
                            shutil.copyfileobj(src, dst)
            imported = root / "app_data"
            if imported.exists():
                for path in imported.rglob("*"):
                    if path.is_file():
                        target = self.APP_DIR / path.relative_to(imported)
                        target.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copy2(path, target)
        self._load()


    def _save_codes(self):
        self._codes_path.write_text(
            json.dumps({"students": self.students}, indent=2, ensure_ascii=False), "utf-8")

    def _save_defaulters(self):
        self._default_path.write_text(
            json.dumps({"defaulters": list(self._defaulters),
                        "display":    self._defaulter_display},
                       indent=2, ensure_ascii=False), "utf-8")

    # ── Pay codes ──────────────────────────────────────────────────────────────

    def import_pairs(self, pairs) -> int:
        added = 0
        for name, code in pairs:
            existing_key = fuzzy_match(name, self.students)
            key = existing_key if existing_key else normalize_name(name)
            if key not in self.students:
                self.students[key] = {"display": name, "codes": []}
            if code not in self.students[key]["codes"]:
                self.students[key]["codes"].append(code)
                added += 1
        self._save_codes()
        return added

    def add_code(self, display_name: str, code: str):
        key = normalize_name(display_name)
        if key not in self.students:
            self.students[key] = {"display": display_name, "codes": []}
        if code not in self.students[key]["codes"]:
            self.students[key]["codes"].append(code)
        self._save_codes()

    def get_codes(self, name: str):
        key = fuzzy_match(name, self.students)
        return list(self.students[key]["codes"]) if key else []

    def all_students(self):
        return sorted(
            [(v["display"], v["codes"]) for v in self.students.values()],
            key=lambda x: x[0].lower())

    def remove_student(self, display_name: str):
        key = normalize_name(display_name)
        if key in self.students:
            del self.students[key]
            self._save_codes()

    def clear_all_codes(self):
        self.students = {}
        self._save_codes()

    # ── Defaulters ─────────────────────────────────────────────────────────────

    def add_defaulter(self, name: str):
        key = normalize_name(name)
        self._defaulters.add(key)
        self._defaulter_display[key] = name
        self._save_defaulters()

    def import_defaulters(self, names):
        for name in names:
            self.add_defaulter(name)

    def remove_defaulter(self, name: str):
        key = normalize_name(name)
        self._defaulters.discard(key)
        self._defaulter_display.pop(key, None)
        self._save_defaulters()

    def is_defaulter(self, name: str) -> bool:
        return bool(fuzzy_match(name, self._defaulters))

    def all_defaulters(self):
        return sorted(
            [self._defaulter_display.get(k, k) for k in self._defaulters],
            key=str.lower)


# ═══════════════════════════════════════════════════════════════════════════════
# MISSING CODES DIALOG
# ═══════════════════════════════════════════════════════════════════════════════

class MissingCodesDialog(tk.Toplevel):
    def __init__(self, parent, missing_names, store: DataStore):
        super().__init__(parent)
        self.title("Missing School Pay Codes")
        self.geometry("580x520")
        self.minsize(480, 400)
        self.resizable(True, True)
        self.grab_set()
        self.configure(bg=BG)
        self._store     = store
        self._missing   = list(missing_names)
        self._code_vars: dict = {}
        self.result     = "cancel"
        self._mode      = tk.StringVar(value="manual")
        self._build()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg="#92400e", pady=10)
        hdr.pack(fill="x")
        tk.Label(hdr, text="Missing School Pay Codes",
                 bg="#92400e", fg="white", font=FONT_LG).pack(padx=14, anchor="w")
        n = len(self._missing)
        tk.Label(hdr,
                 text=f"{n} student{'s' if n != 1 else ''} have no pay code in stored data.",
                 bg="#92400e", fg="#fef3c7", font=FONT).pack(padx=14, anchor="w")

        # Mode radio buttons
        rr = tk.Frame(self, bg=BG, pady=8)
        rr.pack(fill="x", padx=14)
        tk.Label(rr, text="What would you like to do?",
                 bg=BG, font=FONT_BOLD, fg=TEXT).pack(anchor="w", pady=(0, 4))
        for val, label in [
            ("manual", "Enter codes manually"),
            ("upload", "Upload a file with the missing codes"),
            ("skip",   "Skip these students (no report cards for them)"),
        ]:
            tk.Radiobutton(rr, text=label, variable=self._mode, value=val,
                           command=self._switch_mode,
                           bg=BG, font=FONT, activebackground=BG,
                           cursor="hand2").pack(anchor="w", padx=16)

        self._content = tk.Frame(self, bg=BG)
        self._content.pack(fill="both", expand=True, padx=14, pady=(4, 4))

        # Buttons
        bar = tk.Frame(self, bg=BORDER, pady=8)
        bar.pack(fill="x", side="bottom")
        _btn(bar, "Cancel", self._on_cancel, bg=BTN_BG).pack(side="right", padx=(0, 14))
        _btn(bar, "Continue  >", self._on_confirm,
             bg=PRIMARY, fg="white", font=FONT_BOLD).pack(side="right", padx=(0, 6))

        self._switch_mode()

    def _clear(self):
        for w in self._content.winfo_children():
            w.destroy()

    def _switch_mode(self):
        self._clear()
        mode = self._mode.get()
        if mode == "manual":
            self._build_manual()
        elif mode == "upload":
            self._build_upload()
        else:
            self._build_skip()

    def _build_manual(self):
        tk.Label(self._content, text="Enter a pay code next to each student name:",
                 bg=BG, font=FONT, fg=TEXT_MUTED).pack(anchor="w", pady=(4, 6))
        outer = tk.Frame(self._content, bg=BG)
        outer.pack(fill="both", expand=True)
        cv  = tk.Canvas(outer, bg=BG, highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient="vertical", command=cv.yview)
        cv.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        cv.pack(fill="both", expand=True)
        inner = tk.Frame(cv, bg=BG)
        cv.create_window((0, 0), window=inner, anchor="nw")
        self._code_vars = {}
        for name in self._missing:
            row = tk.Frame(inner, bg=BG)
            row.pack(fill="x", pady=3)
            tk.Label(row, text=name, bg=BG, font=FONT_BOLD, fg=TEXT,
                     width=30, anchor="w").pack(side="left")
            var = tk.StringVar()
            self._code_vars[name] = var
            tk.Entry(row, textvariable=var, font=FONT,
                     bd=1, relief="solid", width=14, bg=ENTRY_BG).pack(side="left", padx=(8, 0))
        inner.update_idletasks()
        cv.configure(scrollregion=cv.bbox("all"))

    def _build_upload(self):
        tk.Label(self._content,
                 text="Choose a CSV, Excel, Word or PDF file containing\n"
                      "student names and their school pay codes.",
                 bg=BG, font=FONT, fg=TEXT_MUTED, justify="left").pack(anchor="w", pady=(8, 10))
        self._upload_status = tk.Label(self._content, text="No file selected.",
                                       bg=BG, font=FONT_SM, fg=TEXT_MUTED,
                                       wraplength=440, justify="left")
        self._upload_status.pack(anchor="w", pady=(0, 8))
        _btn(self._content, "Browse for File", self._do_upload,
             bg=ACCENT, fg="white").pack(anchor="w")

    def _do_upload(self):
        path = filedialog.askopenfilename(
            parent=self, title="Select file with school pay codes",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("All", "*.*")])
        if not path:
            return
        try:
            pairs = parse_paycode_file(path)
            added = self._store.import_pairs(pairs)
            self._upload_status.config(
                text=f"OK  {Path(path).name} — {len(pairs)} pairs read, {added} new codes stored.",
                fg=PRIMARY)
        except Exception as exc:
            self._upload_status.config(text=f"Error: {exc}", fg=RED)

    def _build_skip(self):
        n = len(self._missing)
        tk.Label(self._content,
                 text=f"These {n} student{'s' if n != 1 else ''} will be skipped:\n"
                      "No report card will be generated for them.",
                 bg=BG, font=FONT, fg=TEXT_MUTED, justify="left").pack(anchor="w", pady=(8, 8))
        frm = tk.Frame(self._content, bg=BG)
        frm.pack(fill="both", expand=True)
        sb = tk.Scrollbar(frm); sb.pack(side="right", fill="y")
        lb = tk.Listbox(frm, font=FONT, bg=CARD, bd=1, relief="solid",
                        yscrollcommand=sb.set, highlightthickness=0)
        lb.pack(fill="both", expand=True)
        sb.config(command=lb.yview)
        for name in self._missing:
            lb.insert("end", name)

    def _on_confirm(self):
        mode = self._mode.get()
        if mode == "manual":
            for name, var in self._code_vars.items():
                code = var.get().strip()
                if code:
                    self._store.add_code(name, code)
        self.result = "skip" if mode == "skip" else "continue"
        self.destroy()

    def _on_cancel(self):
        self.result = "cancel"
        self.destroy()


class NameReviewDialog(tk.Toplevel):
    """Let the operator verify OCR results before any PDFs are generated."""

    def __init__(self, parent, cards):
        super().__init__(parent)
        self.title("Review identified student names")
        self.geometry("760x560")
        self.minsize(600, 400)
        self.configure(bg=BG)
        self.result = None
        self._vars = []
        self.transient(parent)
        self.grab_set()

        tk.Label(
            self, text="Check the identified names before generating reports.",
            bg=BG, fg=TEXT, font=FONT_LG
        ).pack(anchor="w", padx=16, pady=(14, 2))
        tk.Label(
            self,
            text="Only the name text from the selected region is used. Edit any "
                 "uncertain entry, then click Save names.",
            bg=BG, fg=TEXT_MUTED, font=FONT_SM, wraplength=700, justify="left"
        ).pack(anchor="w", padx=16, pady=(0, 10))

        frame = tk.Frame(self, bg=BG)
        frame.pack(fill="both", expand=True, padx=16)
        canvas = tk.Canvas(frame, bg=BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        body = tk.Frame(canvas, bg=BG)
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=body, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        for index, (path, name, pages) in enumerate(cards, 1):
            row = tk.Frame(body, bg=CARD, bd=1, relief="solid")
            row.pack(fill="x", pady=3)
            page_text = ", ".join(str(p + 1) for p in pages[:4])
            if len(pages) > 4:
                page_text += "..."
            tk.Label(
                row, text=f"{index}. {Path(path).name}  •  page(s) {page_text}",
                bg=CARD, fg=TEXT_MUTED, font=FONT_SM, anchor="w"
            ).pack(fill="x", padx=8, pady=(5, 1))
            var = tk.StringVar(value=name)
            self._vars.append(var)
            tk.Entry(
                row, textvariable=var, font=FONT, bg=ENTRY_BG,
                relief="solid", bd=1
            ).pack(fill="x", padx=8, pady=(1, 6))

        actions = tk.Frame(self, bg=BG)
        actions.pack(fill="x", padx=16, pady=12)
        _btn(actions, "Cancel", self._cancel, bg=BTN_BG).pack(side="right")
        _btn(actions, "Save names", self._save, bg=PRIMARY, fg="white").pack(
            side="right", padx=6
        )

    def _save(self):
        names = [re.sub(r"\s+", " ", var.get()).strip() for var in self._vars]
        if any(not name for name in names):
            messagebox.showerror("Name required", "Every report must have a student name.")
            return
        self.result = names
        self.destroy()

    def _cancel(self):
        self.result = "cancel"
        self.destroy()


# ═══════════════════════════════════════════════════════════════════════════════
# STANDALONE BUTTON HELPER (outside class, used in dialogs)
# ═══════════════════════════════════════════════════════════════════════════════

def _btn(parent, text, command, bg=None, fg=None, font=None, padx=10, pady=5):
    bg = BTN_BG if bg is None else bg
    fg = TEXT if fg is None else fg
    return tk.Button(parent, text=text, command=command, bg=bg, fg=fg,
                     font=font or FONT, relief="flat", padx=padx, pady=pady,
                     cursor="hand2", activebackground=bg, activeforeground=fg)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN APPLICATION
# ═══════════════════════════════════════════════════════════════════════════════

class ReportSplitter(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title("Hilton High School")
        self.geometry("1400x850")
        self.minsize(1040, 680)
        self.configure(bg=BG)

        self._store = DataStore()
        self._theme_name = self._store.settings.get("theme", DEFAULT_THEME)
        if self._theme_name not in THEMES:
            self._theme_name = DEFAULT_THEME
        self._app_icon = self._load_embedded_icon()
        if self._app_icon is not None:
            try:
                self.iconphoto(True, self._app_icon)
            except tk.TclError:
                pass

        # Generate-tab state
        self.pdf_doc         = None
        self.current_path    = ""
        self.current_page    = 0
        self.page_scale      = 1.0
        self.page_photo      = None
        self.sel_start       = None
        self.sel_rect_id     = None
        self.name_rects:dict = {}
        self.same_region_var = tk.BooleanVar(value=True)
        self.circular_doc    = None
        self.circular_path   = ""
        self.medical_path    = ""
        self.banner_path     = ""
        self.processing      = False
        self._secret_key_presses = []
        self._upload_unlocked = False
        self._settings_tab = None
        self._settings_hide_after_id = None
        self._prepared_cards = []

        # Per-file report card titles
        self.file_report_names: dict = {}
        self.file_report_classes: dict = {}

        self._apply_theme(self._theme_name)
        self._build_ui()
        self.after(100, self._fit_canvas)

    # ── Theme ──────────────────────────────────────────────────────────────────

    @staticmethod
    def _load_embedded_icon(master=None):
        try:
            return tk.PhotoImage(master=master, data=APP_ICON_BASE64, format="png")
        except tk.TclError:
            return None

    def _apply_theme(self, theme_name=None):
        old = {key: globals()[key] for key in (
            "PRIMARY", "PRIMARY_DARK", "PRIMARY_PALE", "ACCENT", "RED",
            "RED_PALE", "AMBER", "BG", "CARD", "BORDER", "TEXT",
            "TEXT_MUTED", "BTN_BG", "ENTRY_BG", "PANEL_TINT", "CANVAS_BG"
        )}
        theme_name = theme_name if theme_name in THEMES else DEFAULT_THEME
        self._theme_name = theme_name
        for key, value in THEMES[theme_name].items():
            globals()[key] = value
        try:
            s = ttk.Style(self)
            try:
                s.theme_use("clam")
            except Exception:
                try:
                    s.theme_use("alt")
                except Exception:
                    pass
            s.configure("TNotebook",
                        background=BG, borderwidth=0,
                        tabmargins=[0, 0, 0, 0])
            s.configure("TNotebook.Tab",
                        background=BG, foreground=TEXT_MUTED,
                        padding=[18, 10], font=("Segoe UI", 10, "bold"),
                        borderwidth=0)
            s.map("TNotebook.Tab",
                  background=[("selected", CARD), ("active", PRIMARY_PALE)],
                  foreground=[("selected", PRIMARY), ("active", PRIMARY)])
            s.configure("TProgressbar",
                        troughcolor=BORDER, background=PRIMARY,
                        borderwidth=0, thickness=8)
            s.configure("Treeview",
                        background=CARD, fieldbackground=ENTRY_BG,
                        foreground=TEXT,
                        rowheight=24, font=FONT)
            s.configure("Treeview.Heading",
                        background=PRIMARY_PALE, foreground=PRIMARY,
                        font=FONT_BOLD, relief="flat")
            s.map("Treeview",
                  background=[("selected", PRIMARY_PALE)],
                  foreground=[("selected", PRIMARY_DARK)])
            s.configure("TCombobox",
                        fieldbackground=ENTRY_BG, background=BTN_BG,
                        foreground=TEXT, selectbackground=PRIMARY_PALE,
                        selectforeground=TEXT)
        except Exception:
            pass  # Theme failures are non-fatal
        if hasattr(self, "nb"):
            self._restyle_widgets(old)
        for attr in ("theme_value_label", "secret_theme_value_label"):
            label = getattr(self, attr, None)
            if label is not None:
                label.config(text=self._theme_name.title())

    def _restyle_widgets(self, old):
        """Repaint existing Tk widgets without rebuilding application state."""
        replacement = {
            old["PRIMARY"]: PRIMARY, old["PRIMARY_DARK"]: PRIMARY_DARK,
            old["PRIMARY_PALE"]: PRIMARY_PALE, old["ACCENT"]: ACCENT,
            old["RED"]: RED, old["RED_PALE"]: RED_PALE, old["AMBER"]: AMBER,
            old["BG"]: BG, old["CARD"]: CARD, old["BORDER"]: BORDER,
            old["TEXT"]: TEXT, old["TEXT_MUTED"]: TEXT_MUTED,
            old["BTN_BG"]: BTN_BG, old["ENTRY_BG"]: ENTRY_BG,
            old["PANEL_TINT"]: PANEL_TINT, old["CANVAS_BG"]: CANVAS_BG,
            "white": ENTRY_BG, "#ffffff": ENTRY_BG,
            "#dde8f0": PANEL_TINT, "#c8d4e0": CANVAS_BG,
        }

        def recolour(value):
            return replacement.get(value, value)

        def visit(widget):
            try:
                if hasattr(widget, "refresh_theme"):
                    widget.refresh_theme(BG, CARD, BORDER, TEXT)
                cls = widget.winfo_class()
                if cls == "Frame":
                    widget.configure(background=recolour(widget.cget("background")))
                elif cls == "Labelframe":
                    widget.configure(
                        background=recolour(widget.cget("background")),
                        foreground=recolour(widget.cget("foreground")),
                    )
                elif cls in ("Label", "Checkbutton", "Radiobutton"):
                    widget.configure(
                        background=recolour(widget.cget("background")),
                        foreground=recolour(widget.cget("foreground")),
                        activebackground=recolour(widget.cget("activebackground")),
                    )
                elif cls in ("Button", "Menubutton"):
                    widget.configure(
                        background=recolour(widget.cget("background")),
                        foreground=recolour(widget.cget("foreground")),
                        activebackground=recolour(widget.cget("activebackground")),
                    )
                elif cls in ("Entry", "Listbox", "Text"):
                    widget.configure(
                        background=recolour(widget.cget("background")),
                        foreground=recolour(widget.cget("foreground")),
                        insertbackground=TEXT,
                        selectbackground=PRIMARY_PALE,
                    )
                elif cls == "Canvas":
                    widget.configure(background=recolour(widget.cget("background")))
            except (tk.TclError, TypeError):
                pass
            for child in widget.winfo_children():
                visit(child)

        visit(self)

    # ── Root ───────────────────────────────────────────────────────────────────

    def _build_ui(self):
        hdr = tk.Frame(self, bg=CARD, height=76)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        brand = tk.Frame(hdr, bg=PRIMARY_DARK, width=76, height=76)
        brand.pack(side="left", fill="y")
        brand.pack_propagate(False)
        if self._app_icon is not None:
            self._header_icon = self._app_icon.subsample(2, 2)
            tk.Label(brand, image=self._header_icon, bg=PRIMARY_DARK).pack(
                expand=True)
        tk.Label(
            hdr,
            text="Hilton High School  -  Report Card Splitter",
            bg=CARD,
            fg=TEXT,
            font=("Segoe UI", 20, "bold")
        ).pack(side="left", padx=26)

        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True)

        self._build_tab_generate()
        self._build_tab_paycodes()
        self._build_tab_defaulters()
        self._build_tab_upload()
        self.bind_all("<Shift-F5>", self._secret_keypress)

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 — GENERATE
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_generate(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  ▤  Generate  ")
        shell = tk.Frame(tab, bg=BG)
        shell.pack(fill="both", expand=True, padx=14, pady=12)

        # Fluent/Office-style command ribbon.  The buttons are wired to the
        # existing application actions; this changes only their presentation.
        ribbon = self._card(shell)
        ribbon.pack(fill="x", pady=(0, 12))
        ribbon_tabs = tk.Frame(ribbon, bg=CARD, height=42)
        ribbon_tabs.pack(fill="x")
        ribbon_tabs.pack_propagate(False)
        tk.Label(ribbon_tabs, text="Home", bg=CARD, fg=PRIMARY,
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=20)
        for label in ("Export", "About"):
            tk.Label(ribbon_tabs, text=label, bg=CARD, fg=TEXT_MUTED,
                     font=("Segoe UI", 10)).pack(side="left", padx=18)
        ttk.Separator(ribbon, orient="horizontal").pack(fill="x")

        actions = tk.Frame(ribbon, bg=CARD, height=82)
        actions.pack(fill="x", padx=10, pady=7)
        actions.pack_propagate(False)

        def command_tile(label, icon, command):
            tile = tk.Frame(actions, bg=CARD)
            tile.pack(side="left", padx=7)
            button = tk.Button(
                tile, text=icon, command=command, bg=CARD, fg=PRIMARY,
                activebackground=PRIMARY_PALE, activeforeground=PRIMARY,
                relief="flat", bd=0, cursor="hand2",
                font=("Segoe UI Symbol", 20), padx=12, pady=0)
            button.pack()
            tk.Label(tile, text=label, bg=CARD, fg=TEXT,
                     font=("Segoe UI", 8)).pack(pady=(1, 0))

        command_tile("Add Files", "▣", self._add_file)
        command_tile("Remove", "−", self._remove_file)
        command_tile("Clear", "×", self._clear_files)
        command_tile("Prev", "‹", self._prev_page)
        command_tile("Next", "›", self._next_page)
        command_tile("Generate Report Cards", "✓", self._start_processing)
        command_tile("Save", "▣", self._save_current_title)

        outer = tk.Frame(shell, bg=BG)
        outer.pack(fill="both", expand=True)
        self._gen_left(outer)
        self._gen_middle(outer)
        right = tk.Frame(outer, bg=BG)
        right.pack(side="right", fill="y", padx=(6, 0))
        generate_bar = tk.Frame(right, bg=BG)
        generate_bar.pack(fill="x", pady=(0, 8))
        self.start_btn = self._b(generate_bar, "GENERATE REPORT CARDS",
                                 self._start_processing, bg=PRIMARY, fg="white",
                                 font=FONT_LG, pady=12)
        self.start_btn.pack(side="left", fill="x", expand=True)
        menu_button = tk.Menubutton(
            generate_bar, text="▼", bg=PRIMARY_DARK, fg="white",
            font=FONT_LG, relief="flat", cursor="hand2", padx=10, pady=12,
            activebackground=PRIMARY_DARK, activeforeground="white"
        )
        menu_button.pack(side="right", fill="y")
        menu = tk.Menu(menu_button, tearoff=False)
        menu.add_command(label="Review prepared names",
                         command=self._review_prepared_names)
        menu_button.configure(menu=menu)
        self._gen_right(right)

    def _gen_left(self, parent):
        col = self._card(parent, "  Source Files  ", width=270)
        col.pack(side="left", fill="y", padx=(0, 6))
        self._b(col, "+ Add File(s)", self._add_file, bg=PRIMARY, fg="white").pack(fill="x", pady=(0, 4))
        self._b(col, "Remove Selected", self._remove_file, bg=RED, fg="white").pack(fill="x", pady=(0, 10))
        tk.Label(col, text="Loaded files:", bg=CARD, font=FONT, fg=TEXT_MUTED).pack(anchor="w")
        frm = tk.Frame(col, bg=CARD); frm.pack(fill="both", expand=True)
        sb  = tk.Scrollbar(frm); sb.pack(side="right", fill="y")
        self.file_list = tk.Listbox(frm, font=("Segoe UI", 8),
                                    selectmode="single", activestyle="none",
                                    bg=ENTRY_BG, bd=1, relief="solid",
                                    highlightthickness=0, yscrollcommand=sb.set)
        self.file_list.pack(fill="both", expand=True)
        sb.config(command=self.file_list.yview)
        self.file_list.bind("<<ListboxSelect>>", self._on_file_select)

    def _gen_middle(self, parent):
        col = self._card(parent, "  Preview  -  drag on page to mark the student name area  ")
        col.pack(side="left", fill="both", expand=True, padx=6)

        nav = tk.Frame(col, bg=CARD); nav.pack(fill="x", pady=(0, 2))
        self._b(nav, "< Prev", self._prev_page, bg=BTN_BG).pack(side="left")
        self.page_label = tk.Label(nav, text="No file loaded", bg=CARD, font=FONT, fg=TEXT_MUTED)
        self.page_label.pack(side="left", padx=10)
        self._b(nav, "Next >", self._next_page, bg=BTN_BG).pack(side="left")
        self.region_label = tk.Label(nav, text="", bg=CARD, font=FONT, fg=PRIMARY)
        self.region_label.pack(side="right", padx=6)

        tog = tk.Frame(col, bg=PANEL_TINT, pady=5); tog.pack(fill="x", pady=(0, 4))
        tk.Checkbutton(tog, text="Same name region for all files",
                       variable=self.same_region_var,
                       command=self._on_same_region_toggle,
                        bg=PANEL_TINT, font=FONT, activebackground=PANEL_TINT,
                       cursor="hand2").pack(side="left", padx=8)
        self.region_status_lbl = tk.Label(tog, text="", bg=PANEL_TINT,
                                          font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED)
        self.region_status_lbl.pack(side="right", padx=8)

        cf  = tk.Frame(col, bg=CANVAS_BG); cf.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(cf, orient="vertical")
        hsb = ttk.Scrollbar(cf, orient="horizontal")
        self.canvas = tk.Canvas(cf, bg=CANVAS_BG, cursor="crosshair",
                                highlightthickness=0,
                                yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self.canvas.yview)
        hsb.config(command=self.canvas.xview)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.canvas.pack(fill="both", expand=True)
        self.canvas.bind("<ButtonPress-1>",   self._sel_start)
        self.canvas.bind("<B1-Motion>",       self._sel_drag)
        self.canvas.bind("<ButtonRelease-1>", self._sel_end)

        self.hint_lbl = tk.Label(col,
            text="Load a file, then drag on the preview to mark the student name area.",
            bg=CARD, font=("Segoe UI", 9, "italic"), fg=TEXT_MUTED)
        self.hint_lbl.pack(pady=4)

    def _gen_right(self, parent):
        col = self._card(parent, "  Settings  ", width=304)
        col.pack(side="right", fill="y", padx=(6, 0))

        self._lbl(col, "Report Card Title (for selected file):")
        self.report_name_var = tk.StringVar(value="")
        self.report_name_var.trace_add("write", self._on_report_name_changed)
        title_row = tk.Frame(col, bg=CARD)
        title_row.pack(fill="x", pady=(2, 2))
        tk.Entry(title_row, textvariable=self.report_name_var,
                 font=("Segoe UI", 10), bd=1, relief="solid",
                 bg=ENTRY_BG).pack(side="left", fill="x", expand=True)
        self.report_class_var = tk.StringVar(value=CLASS_OPTIONS[0])
        self.report_class_combo = ttk.Combobox(
            title_row, textvariable=self.report_class_var,
            values=CLASS_OPTIONS, state="readonly", width=14)
        self.report_class_combo.pack(side="left", padx=(5, 0))
        self.report_class_combo.bind("<<ComboboxSelected>>",
                                     self._on_report_class_changed)
        tk.Label(col, text="Class (select the arrow)  |  Report title",
                 bg=CARD, fg=TEXT_MUTED, font=FONT_SM).pack(anchor="w")
        self.report_name_hint = tk.Label(col,
            text="Select a file on the left to set its title.",
            bg=CARD, font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED,
            wraplength=268, justify="left")
        self.report_name_hint.pack(anchor="w", pady=(0, 12))

        self._lbl(col, "Split Mode:")
        self.split_mode = tk.StringVar(value="fixed")
        fr = tk.Frame(col, bg=CARD); fr.pack(fill="x", pady=(2, 4))
        tk.Radiobutton(fr, text="Fixed pages per report:",
                       variable=self.split_mode, value="fixed",
                       bg=CARD, font=FONT, activebackground=CARD).pack(side="left")
        self.pages_spin = tk.Spinbox(fr, from_=1, to=100, width=5,
                                     font=("Segoe UI", 10), bd=1, relief="solid")
        self.pages_spin.delete(0, "end"); self.pages_spin.insert(0, "1")
        self.pages_spin.pack(side="left", padx=6)
        tk.Radiobutton(col, text="Auto-detect: new report when student name changes",
                       variable=self.split_mode, value="auto",
                       bg=CARD, font=FONT, justify="left",
                       activebackground=CARD, wraplength=260).pack(anchor="w", pady=(0, 14))

        tk.Label(col, text="Reports are saved automatically in this app's internal storage.",
                 bg=CARD, fg=PRIMARY, font=FONT_SM, wraplength=268,
                 justify="left").pack(anchor="w", pady=(0, 12))

        ttk.Separator(col, orient="horizontal").pack(fill="x", pady=(0, 10))
        ch = tk.Frame(col, bg=CARD); ch.pack(fill="x")
        self._lbl(ch, "Append Circular")
        tk.Label(ch, text="optional", bg=CARD,
                 font=("Segoe UI", 8, "italic"), fg=TEXT_MUTED).pack(side="left", padx=(6, 0))
        tk.Label(col, text="These pages are added after each student card.",
                 bg=CARD, font=FONT_SM, fg=TEXT_MUTED).pack(anchor="w", pady=(2, 6))
        self._b(col, "Choose Circular PDF", self._browse_circular,
                bg="#5b4fcf", fg="white").pack(fill="x", pady=(0, 4))
        self.circular_label = tk.Label(col, text="No circular selected",
                                       bg=CARD, font=("Segoe UI", 8, "italic"),
                                       fg=TEXT_MUTED, wraplength=250, justify="left")
        self.circular_label.pack(anchor="w")
        self._b(col, "Remove Circular", self._clear_circular, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        ttk.Separator(col, orient="horizontal").pack(fill="x", pady=(0, 10))
        self._lbl(col, "Optional Medical Form")
        self._b(col, "Choose Medical Form PDF", self._browse_medical,
                bg="#0f766e", fg="white").pack(fill="x", pady=(0, 4))
        self.medical_label = tk.Label(col, text="No medical form selected",
                                       bg=CARD, font=("Segoe UI", 8, "italic"),
                                       fg=TEXT_MUTED, wraplength=250, justify="left")
        self.medical_label.pack(anchor="w")
        self._b(col, "Remove Medical Form", self._clear_medical, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        self._lbl(col, "Optional Website Banner")
        self._b(col, "Choose Banner Image", self._browse_banner,
                bg="#0369a1", fg="white").pack(fill="x", pady=(0, 4))
        self.banner_label = tk.Label(col, text="No banner selected",
                                     bg=CARD, font=("Segoe UI", 8, "italic"),
                                     fg=TEXT_MUTED, wraplength=250, justify="left")
        self.banner_label.pack(anchor="w")
        self._b(col, "Remove Banner", self._clear_banner, bg=BTN_BG).pack(anchor="w", pady=(4, 10))

        self._lbl(col, "Progress:")
        self.progress_var = tk.DoubleVar()
        self.pct_label = tk.Label(col, text="0%", bg=CARD, font=FONT, fg=TEXT_MUTED)
        self.pct_label.pack(anchor="e")
        ttk.Progressbar(col, variable=self.progress_var, maximum=100).pack(fill="x", pady=(0, 6))
        self.status_label = tk.Label(col, text="Ready.", bg=CARD, font=FONT_SM,
                                     fg=TEXT_MUTED, wraplength=268, justify="left")
        self.status_label.pack(anchor="w")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 — PAY CODES
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_paycodes(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  ▣  Pay Codes  ")

        shell = tk.Frame(tab, bg=BG)
        shell.pack(fill="both", expand=True, padx=14, pady=12)

        header = self._card(shell)
        header.pack(fill="x", pady=(0, 12))
        top = tk.Frame(header, bg=CARD, height=62)
        top.pack(fill="x")
        top.pack_propagate(False)
        tk.Label(top, text="School Pay Codes", bg=CARD, fg=TEXT,
                 font=("Segoe UI", 14, "bold")).pack(side="left", padx=20)
        actions = tk.Frame(top, bg=CARD)
        actions.pack(side="right", padx=10)
        self._b(actions, "▣  Import File", self._pc_import,
                bg=ACCENT, fg="white").pack(side="left", padx=4)
        self._b(actions, "▥  Remove Selected", self._pc_remove_selected,
                bg=RED, fg="white").pack(side="left", padx=4)
        self._b(actions, "×  Clear All", self._pc_clear_all,
                bg="#7f1d1d", fg="white").pack(side="left", padx=4)

        meta = tk.Frame(header, bg=PRIMARY_PALE, pady=6, padx=12)
        meta.pack(fill="x")
        self.pc_stats = tk.Label(meta, text="", bg=PRIMARY_PALE,
                                 font=FONT_SM, fg=PRIMARY)
        self.pc_stats.pack(side="left")
        tk.Label(meta, text="Search:", bg=PRIMARY_PALE,
                 font=FONT).pack(side="right", padx=(0, 6))
        self.pc_search_var = tk.StringVar()
        self.pc_search_var.trace_add("write", lambda *_: self._pc_refresh())
        tk.Entry(meta, textvariable=self.pc_search_var, font=FONT,
                 bd=1, relief="solid", bg=ENTRY_BG, width=28,
                 highlightbackground=BORDER, highlightthickness=1
                 ).pack(side="right", padx=(0, 6), ipady=4)

        tbl = self._card(shell)
        tbl.pack(fill="both", expand=True)
        cols = ("student_name", "codes")
        self.pc_tree = ttk.Treeview(tbl, columns=cols, show="headings", selectmode="extended")
        self.pc_tree.heading("student_name", text="Student Name")
        self.pc_tree.heading("codes",        text="School Pay Code(s)")
        self.pc_tree.column("student_name", width=450, minwidth=180, anchor="w")
        self.pc_tree.column("codes",        width=720, minwidth=180, anchor="w")
        vsb = ttk.Scrollbar(tbl, orient="vertical", command=self.pc_tree.yview,
                            style="Modern.Vertical.TScrollbar")
        self.pc_tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self.pc_tree.pack(side="left", fill="both", expand=True, padx=(1, 0), pady=1)
        self.pc_empty = tk.Frame(tbl, bg=CARD)
        self.pc_empty.place(relx=.5, rely=.52, anchor="center")
        tk.Label(self.pc_empty, text="▤", bg=CARD, fg="#AEB6C0",
                 font=("Segoe UI Symbol", 38)).pack()
        tk.Label(self.pc_empty, text="No school pay codes loaded yet.",
                 bg=CARD, fg=TEXT_MUTED, font=("Segoe UI", 11)).pack(pady=(3, 0))
        self._pc_refresh()

    def _pc_refresh(self):
        query = normalize_name(self.pc_search_var.get()) if hasattr(self, "pc_search_var") else ""
        self.pc_tree.delete(*self.pc_tree.get_children())
        students = self._store.all_students()
        shown = 0
        for name, codes in students:
            if query and query not in normalize_name(name):
                continue
            self.pc_tree.insert("", "end", values=(name, "  |  ".join(codes)))
            shown += 1
        n = len(students)
        if hasattr(self, "pc_stats"):
            self.pc_stats.config(text=f"{n} student{'s' if n != 1 else ''}  •  showing {shown}")
        if hasattr(self, "pc_empty"):
            if shown:
                self.pc_empty.place_forget()
            else:
                self.pc_empty.place(relx=.5, rely=.52, anchor="center")

    def _pc_import(self):
        paths = filedialog.askopenfilenames(
            title="Import school pay codes",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("CSV", "*.csv"), ("Excel", "*.xlsx *.xls"),
                       ("Word", "*.docx *.doc"), ("PDF", "*.pdf"), ("All", "*.*")])
        if not paths:
            return
        total, errors = 0, []
        for path in paths:
            try:
                pairs = parse_paycode_file(path)
                total += self._store.import_pairs(pairs)
            except Exception as exc:
                errors.append(str(exc))
        self._pc_refresh()
        msg = f"Done. {total} new code(s) added."
        if errors:
            msg += "\n\nErrors:\n" + "\n".join(errors)
        messagebox.showinfo("Import complete", msg)

    def _pc_remove_selected(self):
        sel = self.pc_tree.selection()
        if not sel:
            messagebox.showinfo("Nothing selected", "Select one or more rows first.")
            return
        names = [self.pc_tree.item(i)["values"][0] for i in sel]
        if not messagebox.askyesno("Remove", f"Remove {len(names)} student(s)?"):
            return
        for name in names:
            self._store.remove_student(name)
        self._pc_refresh()

    def _pc_clear_all(self):
        if not messagebox.askyesno("Clear all", "Delete ALL stored pay codes?\nThis cannot be undone."):
            return
        self._store.clear_all_codes()
        self._pc_refresh()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 3 — DEFAULTERS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_defaulters(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  ♙  Defaulters  ")

        shell = tk.Frame(tab, bg=BG)
        shell.pack(fill="both", expand=True, padx=14, pady=12)

        warning = self._card(shell)
        warning.pack(fill="x", pady=(0, 12))
        top = tk.Frame(warning, bg=CARD, height=62)
        top.pack(fill="x")
        top.pack_propagate(False)
        title_area = tk.Frame(top, bg=CARD)
        title_area.pack(side="left", fill="y")
        tk.Label(title_area, text="⚠", bg=CARD, fg=RED,
                 font=("Segoe UI Symbol", 20)).pack(side="left", padx=(18, 8))
        tk.Label(title_area, text="Fee Defaulters", bg=CARD, fg=TEXT,
                 font=("Segoe UI", 14, "bold")).pack(side="left")
        tk.Label(title_area, text="These students will NOT receive report cards",
                 bg=CARD, fg=TEXT_MUTED, font=("Segoe UI", 9)
                 ).pack(side="left", padx=12)
        actions = tk.Frame(top, bg=CARD)
        actions.pack(side="right", padx=10)
        self._b(actions, "♙  Import Names", self._df_import,
                bg=AMBER, fg="white").pack(side="left", padx=4)
        self._b(actions, "▥  Remove Selected", self._df_remove_selected,
                bg=RED, fg="white").pack(side="left", padx=4)

        self.df_stats = tk.Label(warning, bg=RED_PALE, font=FONT_SM,
                                 fg=RED, pady=6, padx=20, anchor="w")
        self.df_stats.pack(fill="x")

        lst = self._card(shell)
        lst.pack(fill="both", expand=True, pady=(0, 10))
        list_area = tk.Frame(lst, bg=CARD)
        list_area.pack(fill="both", expand=True)
        sb = ttk.Scrollbar(list_area, orient="vertical",
                           style="Modern.Vertical.TScrollbar")
        sb.pack(side="right", fill="y")
        self.df_listbox = tk.Listbox(
            list_area, font=FONT, bg=ENTRY_BG, fg=TEXT,
            selectbackground=RED_PALE, selectforeground=TEXT,
            bd=0, relief="flat", highlightthickness=0,
            selectmode="extended", activestyle="none",
            yscrollcommand=sb.set)
        self.df_listbox.pack(side="left", fill="both", expand=True, padx=(1, 0), pady=1)
        sb.config(command=self.df_listbox.yview)
        self.df_empty = tk.Frame(lst, bg=CARD)
        self.df_empty.place(relx=.5, rely=.5, anchor="center")
        tk.Label(self.df_empty, text="♙", bg=CARD, fg="#B1BAC5",
                 font=("Segoe UI Symbol", 40)).pack()
        tk.Label(self.df_empty, text="No students on the defaulters list.",
                 bg=CARD, fg=TEXT_MUTED, font=("Segoe UI", 11)).pack(pady=(3, 0))

        ar = tk.Frame(shell, bg=BG, height=58)
        ar.pack(fill="x")
        ar.pack_propagate(False)
        tk.Label(ar, text="Add name manually:", bg=BG,
                 font=FONT).pack(side="left", padx=(0, 12))
        self.df_manual_var = tk.StringVar()
        tk.Entry(ar, textvariable=self.df_manual_var, font=FONT,
                 bd=1, relief="solid", bg=ENTRY_BG, width=32,
                 highlightbackground=BORDER, highlightthickness=1
                 ).pack(side="left", padx=(0, 4), ipady=5)
        self._b(ar, "＋  Add", self._df_add_manual,
                bg=RED, fg="white").pack(side="left", padx=8)
        self._df_refresh()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 4 — WEBSITE UPLOAD AND SAVED REPORTS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_tab_upload(self):
        tab = tk.Frame(self.nb, bg=BG)
        self.nb.add(tab, text="  ⇧  Website Upload  ")
        self.upload_tab = tab

        shell = tk.Frame(tab, bg=BG)
        shell.pack(fill="both", expand=True, padx=14, pady=12)
        main = self._card(shell)
        main.pack(fill="both", expand=True)

        gate = tk.Frame(main, bg=BTN_BG, pady=10, padx=14,
                        bd=1, relief="solid")
        gate.pack(fill="x", padx=18, pady=(18, 12))
        tk.Label(gate, text="♙", bg=BTN_BG, fg=TEXT,
                 font=("Segoe UI Symbol", 18)).pack(side="left", padx=(4, 8))
        tk.Label(gate, text="Website upload password:",
                 bg=BTN_BG, fg=TEXT, font=FONT_BOLD).pack(side="left")
        self.upload_password_var = tk.StringVar()
        tk.Entry(gate, textvariable=self.upload_password_var, show="•",
                 font=FONT, width=24, bd=1, relief="solid",
                 bg=ENTRY_BG).pack(side="left", padx=12, ipady=5)
        self._b(gate, "Unlock Upload", self._unlock_upload,
                bg=PRIMARY_PALE, fg=PRIMARY_DARK).pack(side="left")
        self.upload_gate_label = tk.Label(gate, text="Locked",
                                          bg="#fff4df", fg=AMBER,
                                          font=FONT_BOLD, padx=10, pady=5)
        self.upload_gate_label.pack(side="left", padx=10)

        self.upload_body = tk.Frame(main, bg=CARD, padx=18, pady=0)
        self.upload_body.pack(fill="both", expand=True)
        self._build_history_table(self.upload_body)

    def _build_history_table(self, parent):
        toolbar = tk.Frame(parent, bg=CARD, height=68)
        toolbar.pack(fill="x", pady=(0, 4))
        toolbar.pack_propagate(False)
        self._b(toolbar, "Compile Website Folder", self._compile_site,
                bg=ACCENT, fg="white").pack(side="left", pady=10)
        self._b(toolbar, "Upload Folder to Netlify", self._upload_site,
                bg=PRIMARY, fg="white").pack(side="left", padx=6, pady=10)
        self._b(toolbar, "Delete Selected Batch", self._delete_selected_batch,
                bg=RED, fg="white").pack(side="left", padx=6, pady=10)
        self._b(toolbar, "Refresh", self._refresh_history,
                bg=BTN_BG).pack(side="left", pady=10)
        self.history_storage_label = tk.Label(
            toolbar, text="", bg=BTN_BG, fg=TEXT_MUTED, font=FONT_SM,
            bd=1, relief="solid", padx=12, pady=7)
        self.history_storage_label.pack(side="right", pady=10)
        ttk.Separator(parent, orient="horizontal").pack(fill="x", pady=(0, 6))
        self.upload_status_label = tk.Label(parent, text="No website folder compiled yet.",
                                             bg=CARD, fg=TEXT_MUTED, font=FONT_SM,
                                             anchor="w")
        self.upload_status_label.pack(fill="x", pady=(0, 8))
        website_options = tk.Frame(parent, bg=CARD)
        website_options.pack(fill="x", pady=(0, 8))
        self._b(website_options, "Choose Website Banner", self._browse_banner,
                bg="#0369a1", fg="white").pack(side="left")
        self.website_banner_status = tk.Label(
            website_options, text="No banner selected",
            bg=CARD, fg=TEXT_MUTED, font=FONT_SM)
        self.website_banner_status.pack(side="left", padx=10)

        frame = tk.Frame(parent, bg=CARD)
        frame.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(frame, orient="vertical")
        vsb.pack(side="right", fill="y")
        self.history_tree = ttk.Treeview(
            frame, columns=("serial", "title", "date", "size"),
            show="headings", selectmode="extended",
            yscrollcommand=vsb.set)
        vsb.config(command=self.history_tree.yview)
        headings = (("serial", "No.", 60), ("title", "Title of reports", 350),
                    ("date", "Date", 190), ("size", "Size of all reports", 160))
        for key, text, width in headings:
            self.history_tree.heading(key, text=text)
            self.history_tree.column(key, width=width, anchor="w")
        self.history_tree.pack(fill="both", expand=True)
        self._refresh_history()

    def _refresh_history(self):
        if not hasattr(self, "history_tree"):
            return
        self.history_tree.delete(*self.history_tree.get_children())
        for number, record in enumerate(self._store.report_history, 1):
            self.history_tree.insert("", "end", iid=str(number - 1),
                                     values=(number, record.get("title", ""),
                                             record.get("date", "").replace("T", " "),
                                             DataStore.human_size(int(record.get("size", 0)))))
        used = self._store.storage_bytes()
        colour = RED if used >= MAX_REPORT_STORAGE else (
            AMBER if used >= STORAGE_WARNING_BYTES else PRIMARY)
        self.history_storage_label.config(
            text=f"Saved report storage: {DataStore.human_size(used)} / 5.0 GB",
            fg=colour)

    def _unlock_upload(self):
        if self.upload_password_var.get() == self._store.settings.get("upload_password"):
            self._upload_unlocked = True
            self.upload_gate_label.config(text="Unlocked", fg="#86efac")
            self._refresh_history()
        else:
            self._upload_unlocked = False
            self.upload_gate_label.config(text="Incorrect password", fg="#fca5a5")
            messagebox.showerror("Upload locked", "The upload password is incorrect.")

    def _delete_selected_batch(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        selected = self.history_tree.selection()
        if not selected:
            messagebox.showinfo("Nothing selected", "Select one or more saved report batches.")
            return
        records = [self._store.report_history[int(item)] for item in selected]
        if not messagebox.askyesno("Delete saved reports",
                                   f"Delete {len(records)} saved report batch(es)?"):
            return
        for record in records:
            self._store.delete_report_batch(record)
        self._refresh_history()

    def _secret_keypress(self, _event=None):
        now = time.monotonic()
        self._secret_key_presses = [stamp for stamp in self._secret_key_presses
                                    if now - stamp < 5]
        self._secret_key_presses.append(now)
        if len(self._secret_key_presses) >= 5:
            self._secret_key_presses.clear()
            self._open_secret_settings()

    def _schedule_settings_hide(self):
        if self._settings_hide_after_id is not None:
            try:
                self.after_cancel(self._settings_hide_after_id)
            except tk.TclError:
                pass
        self._settings_hide_after_id = self.after(
            180000, self._hide_secret_settings)

    def _hide_secret_settings(self):
        """Remove the administrator tab after its temporary access window."""
        self._settings_hide_after_id = None
        if self._settings_tab is None:
            return
        try:
            if str(self._settings_tab) in self.nb.tabs():
                self.nb.forget(self._settings_tab)
        except tk.TclError:
            pass

    def _open_theme_picker(self):
        dialog = tk.Toplevel(self)
        dialog.title("Choose App Theme")
        dialog.transient(self)
        dialog.resizable(False, False)
        dialog.configure(bg=BG)
        dialog.geometry("360x330")
        try:
            dialog.iconphoto(True, self._app_icon)
        except tk.TclError:
            pass

        tk.Label(dialog, text="Choose the look for Hilton High School",
                 bg=BG, fg=TEXT, font=FONT_LG).pack(anchor="w", padx=20, pady=(18, 4))
        tk.Label(dialog, text="Your choice is saved locally for the next launch.",
                 bg=BG, fg=TEXT_MUTED, font=FONT_SM).pack(anchor="w", padx=20, pady=(0, 12))
        selected = tk.StringVar(value=self._theme_name)
        choices = tk.Frame(dialog, bg=CARD, bd=1, relief="solid")
        choices.pack(fill="both", expand=True, padx=20, pady=(0, 14))
        for name in THEME_NAMES:
            tk.Radiobutton(
                choices, text=name.title(), variable=selected, value=name,
                bg=CARD, fg=TEXT, activebackground=CARD,
                activeforeground=TEXT, selectcolor=PRIMARY_PALE,
                font=FONT, anchor="w", padx=12, pady=5,
                cursor="hand2"
            ).pack(fill="x")

        actions = tk.Frame(dialog, bg=BG)
        actions.pack(fill="x", padx=20, pady=(0, 16))
        self._b(actions, "Cancel", dialog.destroy, bg=BTN_BG).pack(side="right")
        self._b(
            actions, "Apply Theme",
            lambda: (self._change_theme(selected.get()), dialog.destroy()),
            bg=PRIMARY, fg="white"
        ).pack(side="right", padx=(0, 8))
        dialog.grab_set()
        dialog.focus_set()

    def _change_theme(self, theme_name):
        if theme_name not in THEMES:
            return
        self._apply_theme(theme_name)
        self._store.save_settings({"theme": theme_name})

    def _open_secret_settings(self):
        if self._settings_tab is None:
            tab = tk.Frame(self.nb, bg=BG)
            self.nb.add(tab, text="  Settings  ")
            self._settings_tab = tab
            self._build_secret_settings(tab)
        elif str(self._settings_tab) not in self.nb.tabs():
            self.nb.add(self._settings_tab, text="  Settings  ")
        self.nb.select(self._settings_tab)
        self._schedule_settings_hide()

    def _build_secret_settings(self, tab):
        panel = self._card(tab, "  Administrator Settings  ")
        panel.pack(fill="x", padx=20, pady=20)
        tk.Label(panel, text="These settings are stored locally on this computer.",
                 bg=CARD, fg=TEXT_MUTED, font=FONT_SM).pack(anchor="w", pady=(0, 12))
        appearance = tk.Frame(panel, bg=CARD)
        appearance.pack(fill="x", pady=(0, 8))
        tk.Label(appearance, text="App theme", width=30, anchor="w",
                 bg=CARD, fg=TEXT, font=FONT).pack(side="left")
        self.secret_theme_value_label = tk.Label(
            appearance, text=self._theme_name.title(), bg=CARD,
            fg=PRIMARY, font=FONT_BOLD)
        self.secret_theme_value_label.pack(side="left", padx=(0, 10))
        self._b(appearance, "Change Theme", self._open_theme_picker,
                bg=BTN_BG).pack(side="left")
        self.secret_vars = {
            "upload_password": tk.StringVar(value=self._store.settings.get("upload_password", "")),
            "netlify_token": tk.StringVar(value=self._store.settings.get("netlify_token", "")),
            "netlify_site_id": tk.StringVar(value=self._store.settings.get("netlify_site_id", "")),
            "copyright_year": tk.StringVar(value=self._store.settings.get("copyright_year", "")),
        }
        fields = [
            ("Upload password", "upload_password", True),
            ("Netlify personal access token", "netlify_token", True),
            ("Netlify site ID", "netlify_site_id", False),
            ("Website copyright year", "copyright_year", False),
        ]
        for label, key, masked in fields:
            row = tk.Frame(panel, bg=CARD)
            row.pack(fill="x", pady=4)
            tk.Label(row, text=label, width=30, anchor="w",
                     bg=CARD, fg=TEXT, font=FONT).pack(side="left")
            tk.Entry(row, textvariable=self.secret_vars[key],
                     show="•" if masked else "", font=FONT,
                     width=48).pack(side="left", fill="x", expand=True)
        actions = tk.Frame(panel, bg=CARD)
        actions.pack(fill="x", pady=(14, 0))
        self._b(actions, "Save Settings", self._save_secret_settings,
                bg=PRIMARY, fg="white").pack(side="left")
        self._b(actions, "Export All App Data", self._export_app_data,
                bg=ACCENT, fg="white").pack(side="left", padx=6)
        self._b(actions, "Import App Data ZIP", self._import_app_data,
                bg=BTN_BG).pack(side="left")

    def _save_secret_settings(self):
        self._store.save_settings({key: var.get() for key, var in self.secret_vars.items()})
        self.upload_password_var.set("")
        messagebox.showinfo("Saved", "Administrator settings saved.")

    def _export_app_data(self):
        path = filedialog.asksaveasfilename(
            title="Export all app data", defaultextension=".zip",
            filetypes=[("ZIP archive", "*.zip")])
        if not path:
            return
        try:
            self._store.export_archive(path)
            messagebox.showinfo("Export complete", f"All app data was exported to:\n{path}")
        except Exception as exc:
            messagebox.showerror("Export failed", str(exc))

    def _import_app_data(self):
        path = filedialog.askopenfilename(
            title="Import app data ZIP", filetypes=[("ZIP archive", "*.zip")])
        if not path:
            return
        if not messagebox.askyesno("Import app data",
                                   "Importing replaces the local saved app data. Continue?"):
            return
        try:
            self._store.import_archive(path)
            self.upload_password_var.set("")
            self._upload_unlocked = False
            for key, var in self.secret_vars.items():
                var.set(self._store.settings.get(key, ""))
            self._refresh_history()
            messagebox.showinfo("Import complete",
                                "App data imported. Click Save Settings to confirm settings.")
        except Exception as exc:
            messagebox.showerror("Import failed", str(exc))

    def _template_root(self) -> Path:
        """Extract the website bundled in this Python file, never from a sidecar."""
        target = self._store.APP_DIR / "website_template_embedded"
        marker = target / ".embedded-site"
        if marker.exists() and marker.read_text(encoding="utf-8") == "embedded-v2" \
                and (target / "index.html").exists():
            return target
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)
        target.parent.mkdir(parents=True, exist_ok=True)
        with tarfile.open(fileobj=io.BytesIO(
                base64.b64decode(EMBEDDED_WEBSITE_ARCHIVE)), mode="r:gz") as archive:
            archive.extractall(target.parent)
        extracted = target.parent / "website_template"
        if not extracted.exists():
            raise RuntimeError("The embedded website bundle could not be extracted.")
        extracted.rename(target)
        self._patch_website_template(target)
        marker.write_text("embedded-v2", encoding="utf-8")
        return target

    @staticmethod
    def _patch_website_template(target: Path):
        """Make the bundled portal data-driven instead of relying on fixed classes."""
        script = target / "script.js"
        content = script.read_text("utf-8")
        old = """        students = parseManifest(text);
      })
"""
        new = """        students = parseManifest(text);
        populateClasses();
      })
"""
        if old in content:
            content = content.replace(old, new, 1)
        old_error = """      .catch(function () {
        students = {};
      });
"""
        new_error = """      .catch(function () {
        students = {};
        populateClasses();
      });
"""
        if old_error in content:
            content = content.replace(old_error, new_error, 1)
        marker = "  function populateClasses() {"
        if marker not in content:
            insert_at = content.index("  function fetchText(url) {")
            function = """  function populateClasses() {
    var classes = [];
    Object.keys(students).forEach(function (key) {
      var value = students[key].className || '';
      if (value && classes.indexOf(value) === -1) classes.push(value);
    });
    classes.sort(function (a, b) {
      return a.localeCompare(b, undefined, { numeric: true, sensitivity: 'base' });
    });
    classSelect.innerHTML = '';
    var placeholder = document.createElement('option');
    placeholder.value = '';
    placeholder.textContent = classes.length ? 'Select class' : 'No classes available';
    classSelect.appendChild(placeholder);
    classes.forEach(function (value) {
      var option = document.createElement('option');
      option.value = value;
      option.textContent = value;
      classSelect.appendChild(option);
    });
    classSelect.disabled = classes.length === 0;
    searchInput.disabled = true;
  }

"""
            content = content[:insert_at] + function + content[insert_at:]
        script.write_text(content, "utf-8")
        index = target / "index.html"
        html = index.read_text("utf-8")
        html = re.sub(
            r'\s*<option value="Senior (?:One|Two|Three|Four|Five|Six)">'
            r'Senior (?:One|Two|Three|Four|Five|Six)</option>',
            "",
            html
        )
        index.write_text(html, "utf-8")

    def _compile_site(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        destination_base = filedialog.askdirectory(
            title="Choose where to create the website folder")
        if not destination_base:
            return
        try:
            template = self._template_root()
            folder_name = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            site_dir = Path(destination_base) / folder_name
            site_dir.mkdir(parents=True, exist_ok=False)
            for name in ("index.html", "styles.css", "script.js",
                         "badge.png", "google1359a90e73c2fd7b.html"):
                source = template / name
                if source.exists():
                    shutil.copy2(source, site_dir / name)
            (site_dir / "report-cards").mkdir()

            # Copy every saved batch, newest first, and rewrite the manifest
            # filenames if two batches contain the same PDF name.
            manifest = []
            used_names = set()
            for batch_index, record in enumerate(self._store.report_history):
                folder = Path(record.get("folder", ""))
                rows_by_file = {
                    str(row[0]): row for row in record.get("manifest_rows", []) if row}
                for original in record.get("files", []):
                    source = folder / original
                    if not source.exists():
                        continue
                    target_name = original
                    if target_name in used_names:
                        target_name = f"{batch_index + 1}_{original}"
                    used_names.add(target_name)
                    shutil.copy2(source, site_dir / "report-cards" / target_name)
                    row = rows_by_file.get(original)
                    if row:
                        # New rows include class and report title; retain
                        # compatibility with older locally archived batches.
                        manifest.append((
                            target_name, row[1], row[2],
                            row[3] if len(row) > 3 else "Unknown",
                            row[4] if len(row) > 4 else record.get("title", "")
                        ))
                    else:
                        stem = Path(original).stem.split("_")
                        student = stem[0] if stem else "Unknown"
                        manifest.append((target_name, student, "", "Unknown",
                                         record.get("title", "")))

            with open(site_dir / "report-cards" / "manifest.csv", "w",
                      newline="", encoding="utf-8") as handle:
                writer = csv.writer(handle)
                writer.writerow(["filename", "student_name", "school_pay_code",
                                 "class", "report_title"])
                writer.writerows(manifest)

            optional_assets = []
            if self.circular_path and Path(self.circular_path).exists():
                optional_assets.append((self.circular_path, "circular.pdf"))
            if self.medical_path and Path(self.medical_path).exists():
                optional_assets.append((self.medical_path, "medical-form.pdf"))
            if self.banner_path and Path(self.banner_path).exists():
                ext = Path(self.banner_path).suffix.lower()
                if ext not in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
                    raise RuntimeError("The banner must be a supported image file.")
                optional_assets.append((self.banner_path, f"banner{ext}"))
            if optional_assets:
                assets = site_dir / "assets"
                assets.mkdir()
                for source, name in optional_assets:
                    shutil.copy2(source, assets / name)

            script = site_dir / "script.js"
            year = self._store.settings.get("copyright_year", str(datetime.now().year))
            if script.exists():
                content = script.read_text("utf-8")
                content = content.replace(
                    "document.getElementById('footer-year').textContent = new Date().getFullYear();",
                    "document.getElementById('footer-year').textContent = " +
                    json.dumps(str(year)) + ";")
                script.write_text(content, "utf-8")

            self.compiled_site_dir = site_dir
            self.upload_status_label.config(
                text=f"Compiled {len(manifest)} report(s) into {site_dir}",
                fg=PRIMARY)
            messagebox.showinfo("Website compiled",
                                f"Website folder created:\n{site_dir}\n\n"
                                f"{len(manifest)} report(s) included.")
        except Exception as exc:
            messagebox.showerror("Compile failed", str(exc))

    def _zip_folder(self, folder: Path, destination: Path):
        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in folder.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(folder))

    def _upload_site(self):
        if not self._upload_unlocked:
            messagebox.showwarning("Upload locked", "Unlock the upload tab first.")
            return
        folder = getattr(self, "compiled_site_dir", None)
        if not folder or not Path(folder).exists():
            messagebox.showinfo("Compile first",
                                "Compile a website folder before uploading it.")
            return
        token = self._store.settings.get("netlify_token", "").strip()
        site_id = self._store.settings.get("netlify_site_id", "").strip()
        if not token or not site_id:
            messagebox.showerror("Missing Netlify settings",
                                 "Set the Netlify token and site ID in Settings.")
            return
        self.upload_status_label.config(text="Preparing upload...", fg=ACCENT)
        threading.Thread(target=self._upload_site_worker,
                         args=(Path(folder), token, site_id), daemon=True).start()

    def _upload_site_worker(self, folder: Path, token: str, site_id: str):
        archive_path = None
        try:
            fd, temp_name = tempfile.mkstemp(suffix=".zip")
            os.close(fd)
            archive_path = Path(temp_name)
            self.after(0, lambda: self.upload_status_label.config(
                text="Compressing website folder...", fg=ACCENT))
            self._zip_folder(folder, archive_path)
            payload = archive_path.read_bytes()
            self.after(0, lambda: self.upload_status_label.config(
                text=f"Uploading {DataStore.human_size(len(payload))} to Netlify...",
                fg=ACCENT))
            request = urllib.request.Request(
                f"https://api.netlify.com/api/v1/sites/{site_id}/deploys",
                data=payload,
                method="POST",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/zip",
                    "Content-Length": str(len(payload)),
                })
            with urllib.request.urlopen(request, timeout=180) as response:
                response.read()
            self.after(0, lambda: self.upload_status_label.config(
                text="Upload complete. The website is deploying on Netlify.",
                fg=PRIMARY))
            self.after(0, lambda: messagebox.showinfo(
                "Upload complete",
                "The complete website folder was uploaded to Netlify successfully."))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")[:500]
            code = exc.code
            self.after(0, lambda: messagebox.showerror(
                "Netlify upload failed", f"Netlify returned {code}.\n{detail}"))
            self.after(0, lambda: self.upload_status_label.config(
                text="Upload failed. Check Netlify settings and try again.", fg=RED))
        except Exception as exc:
            error_text = str(exc)
            self.after(0, lambda: messagebox.showerror("Upload failed", error_text))
            self.after(0, lambda: self.upload_status_label.config(
                text=f"Upload failed: {error_text}", fg=RED))
        finally:
            if archive_path:
                try:
                    archive_path.unlink()
                except OSError:
                    pass

    def _df_refresh(self):
        self.df_listbox.delete(0, "end")
        names = self._store.all_defaulters()
        for name in names:
            self.df_listbox.insert("end", name)
        n = len(names)
        self.df_stats.config(text=f"  {n} student{'s' if n != 1 else ''} on the defaulters list")
        if hasattr(self, "df_empty"):
            if names:
                self.df_empty.place_forget()
            else:
                self.df_empty.place(relx=.5, rely=.5, anchor="center")

    def _df_import(self):
        paths = filedialog.askopenfilenames(
            title="Import defaulter names",
            filetypes=[("All supported", "*.csv *.xlsx *.xls *.docx *.doc *.pdf"),
                       ("All", "*.*")])
        if not paths:
            return
        total, errors = 0, []
        for path in paths:
            try:
                names = parse_names_file(path)
                self._store.import_defaulters(names)
                total += len(names)
            except Exception as exc:
                errors.append(str(exc))
        self._df_refresh()
        msg = f"Done. {total} name(s) added to defaulters list."
        if errors:
            msg += "\n\nErrors:\n" + "\n".join(errors)
        messagebox.showinfo("Import complete", msg)

    def _df_remove_selected(self):
        sel = list(self.df_listbox.curselection())
        if not sel:
            messagebox.showinfo("Nothing selected", "Select one or more names first.")
            return
        names = [self.df_listbox.get(i) for i in sel]
        if not messagebox.askyesno("Remove", f"Remove {len(names)} name(s)?"):
            return
        for name in names:
            self._store.remove_defaulter(name)
        self._df_refresh()

    def _df_add_manual(self):
        name = self.df_manual_var.get().strip()
        if name:
            self._store.add_defaulter(name)
            self.df_manual_var.set("")
            self._df_refresh()

    # WIDGET HELPERS
    # ══════════════════════════════════════════════════════════════════════════

    def _card(self, parent, title="", **kw):
        if ModernCard is not None:
            return ModernCard(parent, title=title.strip(), bg=CARD, fg=TEXT,
                              border=BORDER, outside=BG, padx=10, pady=10, **kw)
        return tk.LabelFrame(parent, text=title, bg=CARD,
                             font=FONT_BOLD, padx=10, pady=10,
                             bd=1, relief="solid", **kw)

    def _b(self, parent, text, command, bg=BTN_BG, fg=TEXT,
           font=None, padx=10, pady=5):
        return tk.Button(parent, text=text, command=command,
                         bg=bg, fg=fg, font=font or FONT,
                         relief="flat", padx=padx, pady=pady,
                         cursor="hand2", activebackground=bg, activeforeground=fg)

    def _lbl(self, parent, text):
        tk.Label(parent, text=text, bg=CARD,
                 font=FONT_BOLD, fg=TEXT).pack(anchor="w")

    # ══════════════════════════════════════════════════════════════════════════
    # REGION HELPERS
    # ══════════════════════════════════════════════════════════════════════════

    def _cur_rect(self):
        return self.name_rects.get(self.current_path)

    def _on_same_region_toggle(self):
        if self.same_region_var.get():
            rect = self._cur_rect()
            if rect:
                for p in self.file_list.get(0, "end"):
                    self.name_rects[p] = rect
        self._update_region_status()

    def _update_region_status(self):
        files = list(self.file_list.get(0, "end"))
        n = len(files)
        if n == 0:
            self.region_status_lbl.config(text="")
            return
        if self.same_region_var.get():
            if self._cur_rect():
                self.region_status_lbl.config(text="Region set - applies to all files", fg=PRIMARY)
            else:
                self.region_status_lbl.config(text="No region set yet - drag on the preview", fg=TEXT_MUTED)
        else:
            have  = sum(1 for p in files if p in self.name_rects)
            color = PRIMARY if have == n else (RED if have == 0 else AMBER)
            self.region_status_lbl.config(
                text=f"{have} of {n} file{'s' if n != 1 else ''} have a region set", fg=color)

    # ══════════════════════════════════════════════════════════════════════════
    # FILE MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════

    def _clear_files(self):
        if not self.file_list.size():
            return
        if not messagebox.askyesno(
                "Clear source files",
                "Remove all loaded PDF files and their marked regions?"):
            return
        if self.pdf_doc is not None:
            try:
                self.pdf_doc.close()
            except Exception:
                pass
        self.pdf_doc = None
        self.current_path = ""
        self.current_page = 0
        self.name_rects.clear()
        self.file_report_names.clear()
        self.file_report_classes.clear()
        self.file_list.delete(0, "end")
        self.canvas.delete("all")
        self.page_label.config(text="No file loaded")
        self.region_label.config(text="")
        self.hint_lbl.config(
            text="Load a file, then drag on the preview to mark the student name area.",
            fg=TEXT_MUTED)
        self.report_name_var.set("")
        self.report_name_hint.config(text="Select a file on the left to set its title.")
        self._update_region_status()

    def _save_current_title(self):
        """Keep the ribbon's Save action meaningful without changing storage rules."""
        if self.current_path:
            self.file_report_names[self.current_path] = self.report_name_var.get()
            self.file_report_classes[self.current_path] = self.report_class_var.get()
            messagebox.showinfo(
                "Saved",
                "The report title and class for the selected file have been saved.")
        else:
            messagebox.showinfo(
                "Nothing to save",
                "Select a source PDF first, then set its report title and class.")

    def _add_file(self):
        paths = filedialog.askopenfilenames(
            title="Select report card PDF(s)",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        existing = list(self.file_list.get(0, "end"))
        added = False
        for p in paths:
            if p not in existing:
                self.file_list.insert("end", p)
                if self.same_region_var.get() and self.name_rects:
                    self.name_rects[p] = next(iter(self.name_rects.values()))
                # Initialise per-file title (inherit current entry value if set)
                if p not in self.file_report_names:
                    self.file_report_names[p] = self.report_name_var.get().strip()
                if p not in self.file_report_classes:
                    self.file_report_classes[p] = self.report_class_var.get()
                added = True
        if added:
            self._update_region_status()
            if self.pdf_doc is None:
                self.file_list.selection_set(0)
                self._load_selected()

    def _remove_file(self):
        sel = self.file_list.curselection()
        if not sel:
            return
        path = self.file_list.get(sel[0])
        self.file_list.delete(sel[0])
        self.name_rects.pop(path, None)
        self.file_report_names.pop(path, None)
        self.file_report_classes.pop(path, None)
        if path == self.current_path:
            self.pdf_doc = None; self.current_path = ""
            self.canvas.delete("all")
            self.page_label.config(text="No file loaded")
            self.region_label.config(text="")
            self.hint_lbl.config(
                text="Load a file, then drag on the preview to mark the student name area.",
                fg=TEXT_MUTED)
        self._update_region_status()

    def _on_file_select(self, _=None):
        self._load_selected()

    def _on_report_name_changed(self, *_):
        """Save the current entry value back to the per-file dict."""
        if self.current_path:
            self.file_report_names[self.current_path] = self.report_name_var.get()

    def _on_report_class_changed(self, _=None):
        if self.current_path:
            self.file_report_classes[self.current_path] = self.report_class_var.get()

    def _load_selected(self):
        sel = self.file_list.curselection()
        if sel:
            self._load_file(self.file_list.get(sel[0]))

    def _load_file(self, path):
        try:
            doc = fitz.open(path)
        except Exception as exc:
            messagebox.showerror("Error", f"Could not open file:\n{exc}"); return
        # Set current_path BEFORE updating the StringVar so the trace writes correctly
        self.pdf_doc, self.current_path = doc, path
        self.current_page = 0; self.sel_rect_id = None
        # Load this file's individual report card title into the entry
        self.report_name_var.set(self.file_report_names.get(path, ""))
        self.report_class_var.set(self.file_report_classes.get(path, CLASS_OPTIONS[0]))
        fname = Path(path).name
        self.report_name_hint.config(text=f"Title for: {fname}")
        if self._cur_rect():
            self.region_label.config(text="Region already set for this file", fg=PRIMARY)
            self.hint_lbl.config(
                text="Name region is marked (shown in orange). Drag again to change it.",
                fg=TEXT_MUTED)
        else:
            self.region_label.config(text="")
            self.hint_lbl.config(
                text="Drag on the preview to mark where the student name appears.",
                fg=TEXT_MUTED)
        self._render_page()
        self._update_region_status()

    # ══════════════════════════════════════════════════════════════════════════
    # PDF RENDERING
    # ══════════════════════════════════════════════════════════════════════════

    def _fit_canvas(self):
        if self.pdf_doc:
            self._render_page()

    def _render_page(self):
        if not self.pdf_doc:
            return
        page  = self.pdf_doc[self.current_page]
        total = len(self.pdf_doc)
        self.page_label.config(text=f"Page {self.current_page + 1} of {total}")
        self.canvas.update_idletasks()
        cw = max(self.canvas.winfo_width(), 400)
        self.page_scale = cw / page.rect.width
        mat = fitz.Matrix(self.page_scale, self.page_scale)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        self.page_photo = ImageTk.PhotoImage(img)
        self.canvas.delete("all")
        self.canvas.create_image(0, 0, anchor="nw", image=self.page_photo)
        self.canvas.config(scrollregion=(0, 0, pix.width, pix.height))
        self.sel_rect_id = None
        rect = self._cur_rect()
        if rect:
            self._draw_region(rect)

    def _draw_region(self, rect):
        s = self.page_scale
        self.sel_rect_id = self.canvas.create_rectangle(
            rect.x0*s, rect.y0*s, rect.x1*s, rect.y1*s,
            outline="#f59e0b", width=2, dash=(5, 3))

    def _prev_page(self):
        if self.pdf_doc and self.current_page > 0:
            self.current_page -= 1; self._render_page()

    def _next_page(self):
        if self.pdf_doc and self.current_page < len(self.pdf_doc) - 1:
            self.current_page += 1; self._render_page()

    # ══════════════════════════════════════════════════════════════════════════
    # REGION SELECTION
    # ══════════════════════════════════════════════════════════════════════════

    def _sel_start(self, event):
        self.sel_start = (self.canvas.canvasx(event.x), self.canvas.canvasy(event.y))
        if self.sel_rect_id:
            self.canvas.delete(self.sel_rect_id); self.sel_rect_id = None

    def _sel_drag(self, event):
        if not self.sel_start:
            return
        x0, y0 = self.sel_start
        x1, y1 = self.canvas.canvasx(event.x), self.canvas.canvasy(event.y)
        if self.sel_rect_id:
            self.canvas.coords(self.sel_rect_id, x0, y0, x1, y1)
        else:
            self.sel_rect_id = self.canvas.create_rectangle(
                x0, y0, x1, y1, outline="#f59e0b", width=2, dash=(5, 3))

    def _sel_end(self, event):
        if not self.sel_start:
            return
        x0, y0 = self.sel_start
        x1, y1 = self.canvas.canvasx(event.x), self.canvas.canvasy(event.y)
        self.sel_start = None
        rx0, rx1 = sorted([x0, x1]); ry0, ry1 = sorted([y0, y1])
        if (rx1 - rx0) < 8 or (ry1 - ry0) < 8:
            return
        s    = self.page_scale
        rect = fitz.Rect(rx0/s, ry0/s, rx1/s, ry1/s)
        if self.same_region_var.get():
            for p in self.file_list.get(0, "end"):
                self.name_rects[p] = rect
            scope = "applied to all files"
        else:
            self.name_rects[self.current_path] = rect
            scope = "set for this file"
        self.region_label.config(
            text=f"Region {scope}  ({int(rx0/s)},{int(ry0/s)}) > ({int(rx1/s)},{int(ry1/s)})",
            fg=PRIMARY)
        sample = self._extract_name(self.pdf_doc, self.current_page, rect)
        if sample:
            self.hint_lbl.config(text=f'Sample name from this page: "{sample}"', fg=PRIMARY)
        else:
            self.hint_lbl.config(
                text="No text found in that area. Try a larger or different region.", fg=RED)
        self._update_region_status()

    # ══════════════════════════════════════════════════════════════════════════
    # SETTINGS
    # ══════════════════════════════════════════════════════════════════════════

    def _browse_circular(self):
        path = filedialog.askopenfilename(
            title="Choose circular PDF",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        if not path:
            return
        try:
            doc   = fitz.open(path)
            pages = len(doc)
            self.circular_doc = doc
            self.circular_path = path
            self.circular_label.config(
                text=f"OK  {Path(path).name}  ({pages} page{'s' if pages != 1 else ''})",
                fg=PRIMARY)
        except Exception as exc:
            messagebox.showerror("Error", f"Could not open circular PDF:\n{exc}")

    def _clear_circular(self):
        self.circular_doc = None
        self.circular_path = ""
        self.circular_label.config(text="No circular selected", fg=TEXT_MUTED)

    def _browse_medical(self):
        path = filedialog.askopenfilename(
            title="Choose medical form PDF",
            filetypes=[("PDF", "*.pdf"), ("All", "*.*")])
        if path:
            self.medical_path = path
            self.medical_label.config(text=f"OK  {Path(path).name}", fg=PRIMARY)

    def _clear_medical(self):
        self.medical_path = ""
        self.medical_label.config(text="No medical form selected", fg=TEXT_MUTED)

    def _browse_banner(self):
        path = filedialog.askopenfilename(
            title="Choose website banner image",
            filetypes=[("Images", "*.png *.jpg *.jpeg *.gif *.webp *.bmp"),
                       ("All", "*.*")])
        if path:
            self.banner_path = path
            self.banner_label.config(text=f"OK  {Path(path).name}", fg=PRIMARY)
            if hasattr(self, "website_banner_status"):
                self.website_banner_status.config(
                    text=f"Selected: {Path(path).name}", fg=PRIMARY)

    def _clear_banner(self):
        self.banner_path = ""
        self.banner_label.config(text="No banner selected", fg=TEXT_MUTED)
        if hasattr(self, "website_banner_status"):
            self.website_banner_status.config(text="No banner selected", fg=TEXT_MUTED)

    # ══════════════════════════════════════════════════════════════════════════
    # PROCESSING
    # ══════════════════════════════════════════════════════════════════════════

    def _start_processing(self):
        if self.processing:
            return
        files = list(self.file_list.get(0, "end"))
        if not files:
            messagebox.showerror("No files", "Please add at least one PDF file."); return
        missing_regions = [p for p in files if p not in self.name_rects]
        if missing_regions:
            messagebox.showerror("Missing name region",
                "These files have no name region marked:\n\n" +
                "\n".join(f"  - {Path(p).name}" for p in missing_regions) +
                "\n\nClick each file and drag on the preview to mark the name area.")
            return
        # Collect per-file report card titles/classes; validate all are filled
        file_report_names_snap = {}
        file_report_classes_snap = {}
        missing_titles = []
        for p in files:
            title = self.file_report_names.get(p, "").strip()
            if not title:
                missing_titles.append(Path(p).name)
            file_report_names_snap[p] = title
            file_report_classes_snap[p] = self.file_report_classes.get(
                p, CLASS_OPTIONS[0])
        if missing_titles:
            messagebox.showerror("Missing Report Card Title",
                "Please set a Report Card Title for each file.\n\n"
                "Select each file in the list on the left and enter its title.\n\n"
                "Missing titles for:\n" +
                "\n".join(f"  - {n}" for n in missing_titles))
            return
        staging_path = Path(tempfile.mkdtemp(
            prefix="report_generation_", dir=str(self._store.APP_DIR)))
        mode = self.split_mode.get()
        pages_per = 1
        if mode == "fixed":
            try:
                pages_per = int(self.pages_spin.get()); assert pages_per >= 1
            except (ValueError, AssertionError):
                messagebox.showerror("Invalid", "Pages per report must be >= 1."); return
        self.processing = True
        self.start_btn.config(state="disabled", text="Processing...")
        self._set_status(0, "Starting...")
        threading.Thread(target=self._run_processing,
                         args=(files, mode, pages_per, file_report_names_snap,
                               file_report_classes_snap, str(staging_path)),
                         daemon=True).start()

    def _review_prepared_names(self):
        if not self._prepared_cards:
            messagebox.showinfo(
                "No prepared names",
                "Run Generate once to scan the selected reports. "
                "The identified names will then be available here for review."
            )
            return
        dlg = NameReviewDialog(self, self._prepared_cards)
        self.wait_window(dlg)
        if isinstance(dlg.result, list):
            self._prepared_cards = [
                (card[0], name, card[2])
                for card, name in zip(self._prepared_cards, dlg.result)
            ]
            messagebox.showinfo("Names saved", f"{len(dlg.result)} name(s) updated.")

    def _run_processing(self, files, mode, pages_per, file_report_names_snap,
                        file_report_classes_snap, save_path):
        try:
            # Phase 1: scan all files
            all_cards = []
            for fi, file_path in enumerate(files):
                self._set_status(int(fi / len(files) * 20),
                    f"Scanning {fi+1}/{len(files)}: {Path(file_path).name}")
                src  = fitz.open(file_path)
                rect = self.name_rects[file_path]
                tot  = len(src)
                if mode == "fixed":
                    for start in range(0, tot, pages_per):
                        end  = min(start + pages_per, tot)
                        name = self._extract_name(src, start, rect) or f"Student {start//pages_per+1}"
                        all_cards.append((file_path, name, list(range(start, end))))
                else:
                    for name, group in self._auto_detect(src, rect, tot):
                        all_cards.append((file_path, name, group))
                src.close()

            # Keep the scan result visible and editable before filtering or
            # generating anything. This prevents a bad OCR line from becoming
            # a filename, manifest entry, or pay-code lookup.
            self._prepared_cards = list(all_cards)
            decision_holder = [None]
            done_evt = threading.Event()

            def _review():
                dlg = NameReviewDialog(self, self._prepared_cards)
                self.wait_window(dlg)
                decision_holder[0] = dlg.result
                done_evt.set()

            self.after(0, _review)
            done_evt.wait()
            if not isinstance(decision_holder[0], list):
                self._set_status(0, "Cancelled.")
                return
            edited_names = decision_holder[0]
            all_cards = [
                (card[0], edited_names[i], card[2])
                for i, card in enumerate(all_cards)
            ]
            self._prepared_cards = list(all_cards)

            # Phase 2: filter defaulters
            to_process = [(fp, nm, grp) for fp, nm, grp in all_cards
                          if not self._store.is_defaulter(nm)]
            skipped    = len(all_cards) - len(to_process)

            # Phase 3: check pay codes
            missing_codes = list({nm for _, nm, _ in to_process
                                  if not self._store.get_codes(nm)})
            if missing_codes:
                decision_holder = [None]
                done_evt = threading.Event()

                def _show():
                    dlg = MissingCodesDialog(self, missing_codes, self._store)
                    self.wait_window(dlg)
                    decision_holder[0] = dlg.result
                    done_evt.set()

                self.after(0, _show)
                done_evt.wait()
                decision = decision_holder[0]
                if decision == "cancel":
                    self._set_status(0, "Cancelled."); return
                if decision == "skip":
                    to_process = [(fp, nm, grp) for fp, nm, grp in to_process
                                  if self._store.get_codes(nm)]

            # Phase 4: generate PDFs
            manifest_rows = []
            generated_by_title = {}
            total_cards   = len(to_process)
            for i, (file_path, student_name, group) in enumerate(to_process):
                self._set_status(20 + int(i / max(total_cards, 1) * 80),
                    f"Saving {i+1}/{total_cards}: {student_name}")
                src      = fitz.open(file_path)
                rpt_name = file_report_names_snap.get(file_path, "Report Card")
                filename = self._save_pdf(src, group, student_name, rpt_name, save_path)
                src.close()
                codes = self._store.get_codes(student_name)
                class_name = file_report_classes_snap.get(
                    file_path, CLASS_OPTIONS[0])
                manifest_rows.append(
                    (filename, student_name, "|".join(codes), class_name, rpt_name))
                generated_by_title.setdefault((class_name, rpt_name), []).append(filename)

            # Phase 5: manifest.csv
            with open(os.path.join(save_path, "manifest.csv"), "w",
                      newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(["filename", "student_name", "school_pay_code",
                            "class", "report_title"])
                for row in manifest_rows:
                    w.writerow(row)

            for (class_name, title), filenames in generated_by_title.items():
                rows_for_title = [row for row in manifest_rows if row[0] in filenames]
                self._store.add_report_batch(
                    f"{class_name} - {title}", save_path, filenames,
                    manifest_rows=rows_for_title)

            count = len(manifest_rows)
            self._set_status(100,
                f"Done! {count} report card(s) saved.\n"
                f"Defaulters skipped: {skipped}.\n"
                f"manifest.csv written to internal app storage.\n\n"
                f"Use Website Upload to compile and publish the site.")
            self.after(0, lambda: messagebox.showinfo("Complete",
                f"{count} report card(s) saved.\n"
                f"{skipped} student(s) skipped (defaulters).\n\n"
                "The reports are stored in the app's internal storage.\n\n"
                "Open Website Upload to compile and publish the site."))
            self.after(0, self._pc_refresh)
            self.after(0, self._refresh_history)

        except Exception as exc:
            self.after(0, lambda: messagebox.showerror("Error", str(exc)))
            self._set_status(0, f"Error: {exc}")
        finally:
            shutil.rmtree(save_path, ignore_errors=True)
            self.processing = False
            self.after(0, lambda: self.start_btn.config(
                state="normal", text="GENERATE REPORT CARDS"))

    def _auto_detect(self, source_doc, name_rect, total_pages):
        groups, cur_name, cur_pages = [], None, []
        for p in range(total_pages):
            name = self._extract_name(source_doc, p, name_rect) or cur_name or "Unknown"
            if name != cur_name:
                if cur_pages:
                    groups.append((cur_name or "Unknown", cur_pages))
                cur_name, cur_pages = name, [p]
            else:
                cur_pages.append(p)
        if cur_pages:
            groups.append((cur_name or "Unknown", cur_pages))
        return groups

    def _save_pdf(self, source_doc, page_nums, student_name, report_name, save_path):
        filename = (f"{safe_filename(student_name)}_"
                    f"{safe_filename(report_name)}_"
                    f"{random_number()}.pdf")
        new_doc = fitz.open()
        for p in page_nums:
            new_doc.insert_pdf(source_doc, from_page=p, to_page=p)
        if self.circular_doc:
            new_doc.insert_pdf(self.circular_doc)
        new_doc.save(os.path.join(save_path, filename), garbage=4, deflate=True, clean=True)
        new_doc.close()
        return filename

    @staticmethod
    def _extract_name(source_doc, page_num: int, rect) -> str:
        if not source_doc or rect is None:
            return ""
        page = source_doc[page_num]
        words = page.get_text("words", clip=rect)
        if not words:
            return ""

        # PyMuPDF words are (x0, y0, x1, y1, text, block, line, word).
        # Keep line geometry: plain get_text("text") loses the meaningful
        # horizontal/vertical gaps which separate a name from nearby labels.
        lines = {}
        for word in words:
            token = re.sub(r"\s+", " ", str(word[4])).strip()
            if not token or not re.search(r"[A-Za-z]", token):
                continue
            key = round(float(word[1]) / 2.5) * 2.5
            lines.setdefault(key, []).append((float(word[0]), float(word[1]),
                                               float(word[2]), float(word[3]), token))
        if not lines:
            return ""

        ordered = []
        for y, row in sorted(lines.items()):
            row.sort(key=lambda item: item[0])
            # Split on unusually large horizontal gaps. The selected box often
            # contains a label before the name and another field after it;
            # keeping the densest word cluster avoids both.
            gaps = [row[i + 1][0] - row[i][2] for i in range(len(row) - 1)]
            positive_gaps = sorted(g for g in gaps if g > 0)
            # The median is distorted when the row contains one or two
            # deliberately wide field gaps. The smallest gap is a safer
            # estimate of the spacing between words in the same name.
            normal_gap = positive_gaps[0] if positive_gaps else 0
            gap_limit = max(normal_gap * 2.2, 12)
            clusters, current = [], [row[0]]
            for index, gap in enumerate(gaps):
                if gap > gap_limit:
                    clusters.append(current)
                    current = []
                current.append(row[index + 1])
            clusters.append(current)
            # Prefer the cluster with the most name-like words; ties retain
            # reading order so a name below a label is not accidentally lost.
            cluster = max(
                clusters,
                key=lambda candidate: (
                    sum(1 for item in candidate
                        if re.search(r"[A-Za-z]", item[4]) and len(item[4]) > 1),
                    -candidate[0][0]
                )
            )
            text = " ".join(item[4] for item in cluster)
            # Ignore obvious field labels, punctuation/noise and long prose.
            letters = re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?", text)
            if not letters or len(letters) > 8:
                continue
            score = sum(1 for token in letters if len(token) > 1)
            if score:
                ordered.append((y, cluster, text))
        if not ordered:
            return ""

        # The first meaningful line is normally the name. Include a following
        # line only when it is close enough to be a wrapped name. A larger
        # vertical gap means it is another field/description and is ignored.
        first_y, first_row, first_text = ordered[0]
        selected = [first_text]
        first_height = max(item[3] - item[1] for item in first_row)
        for y, row, text in ordered[1:]:
            gap = y - first_y
            if gap <= max(first_height * 1.75, 7):
                selected.append(text)
            else:
                break

        result = re.sub(r"[^A-Za-z0-9 .'\-]", " ", " ".join(selected))
        result = re.sub(r"\s+", " ", result).strip(" .-'")
        return result

    def _set_status(self, pct: float, text: str):
        self.after(0, lambda: self.progress_var.set(pct))
        self.after(0, lambda: self.pct_label.config(text=f"{int(pct)}%"))
        self.after(0, lambda: self.status_label.config(text=text))


# ═══════════════════════════════════════════════════════════════════════════════
# STARTUP SPLASH
# ═══════════════════════════════════════════════════════════════════════════════

class _SplashWindow(tk.Tk):
    """Small Office-style loading screen shown before the main window."""

    def __init__(self):
        super().__init__()
        self.title("Hilton High School")
        self.overrideredirect(True)
        self.configure(bg=PRIMARY_DARK)
        self.resizable(False, False)
        self._icon = ReportSplitter._load_embedded_icon(self)
        if self._icon is not None:
            try:
                self.iconphoto(True, self._icon)
            except tk.TclError:
                pass

        width, height = 520, 300
        self.update_idletasks()
        x = max((self.winfo_screenwidth() - width) // 2, 0)
        y = max((self.winfo_screenheight() - height) // 2, 0)
        self.geometry(f"{width}x{height}+{x}+{y}")

        body = tk.Frame(self, bg=PRIMARY_DARK)
        body.pack(fill="both", expand=True, padx=32, pady=28)
        if self._icon is not None:
            tk.Label(body, image=self._icon, bg=PRIMARY_DARK).pack(pady=(0, 10))
        tk.Label(body, text="Hilton High School", bg=PRIMARY_DARK,
                 fg="white", font=("Segoe UI", 22, "bold")).pack()
        tk.Label(body, text="Preparing your workspace", bg=PRIMARY_DARK,
                 fg="#b7d9c5", font=("Segoe UI", 10)).pack(pady=(6, 18))

        self._progress = ttk.Progressbar(
            body, orient="horizontal", mode="indeterminate",
            length=330, style="Splash.Horizontal.TProgressbar")
        self._progress.pack()
        tk.Label(body, text="Loading...", bg=PRIMARY_DARK,
                 fg="#d7eee0", font=("Segoe UI", 9)).pack(pady=(10, 0))
        try:
            ttk.Style(self).configure(
                "Splash.Horizontal.TProgressbar",
                troughcolor=PRIMARY_DARK, background=PRIMARY,
                borderwidth=0, thickness=5)
            self._progress.start(12)
        except tk.TclError:
            pass


# ═══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    try:
        splash = _SplashWindow()
        # Keep the splash visible long enough to feel intentional while the
        # main window is prepared immediately afterward.
        splash.after(1100, splash.destroy)
        splash.mainloop()
        app = ReportSplitter()
        app.mainloop()
    except Exception as exc:
        import traceback
        try:
            root = tk.Tk(); root.withdraw()
            messagebox.showerror("Startup Error",
                f"The app failed to start:\n\n{exc}\n\n"
                f"{traceback.format_exc()[-800:]}")
            root.destroy()
        except Exception:
            pass


if __name__ == "__main__":
    main()
